from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd

from pq_constants import (
    MEASURED_CLOSE,
    MEASURED_OPEN,
    SEVERITY_LABEL,
    SEVERITY_ORDER,
    SEVERITY_SIGNIFICANT_PERSISTENCE,
    measured as _m,
    measured_pct as _mp,
    pct_text as _pct,
    VOLTAGE_BAND_LABEL,
    VOLTAGE_BAND_ORDER,
    strip_marks,
    is_single_phase_208,
    ll_factor,
    __version__,
    Thresholds,
    _H519_ORDERS,
    _SERVICE_TYPE_LABEL,
    _lookup_isc,
    _tdd_class,
    _tdd_limit,
)
from pq_adapter import PQDataset
from pq_analysis import (check_ride_through, check_frequency_ride_through,
                         _IMPEDANCE_MIN_CONSISTENCY, _IMPEDANCE_STEP_MIN_A,
                         _MIN_LOADED_AMPS, _VOLTAGE_RESOLUTION_V,
                         applicable_current_standard, check_trd, exports_power,
                         grade_finding, is_generation_only, standard_k_rating)

log = logging.getLogger(__name__)


def _word_current_standard(doc, report: dict, thresh) -> None:
    """State which current distortion standard governs, and how that was decided.

    Only printed where there is a decision to report -- a service with no
    generation is under 519 and always was, and saying so at every ordinary
    site would bury the cases where it matters.
    """
    std = report.get("current_standard") or {}
    if std.get("branch") == "no_der":
        return

    _body(doc, std.get("reason", ""))

    trd = report.get("trd_compliance") or {}
    if std.get("standard") != "1547":
        return

    if not trd.get("available"):
        _body(doc, trd.get("note", ""))
        return

    _body(doc,
          f"Total rated-current distortion (TRD) is "
          f"√(I_rms² − I₁²) ÷ I_rated, with I_rated the plant's rated current "
          f"capacity of {_m(trd['irated_amps'], '.0f', ' A')}. Measured TRD "
          f"reached {_m(trd.get('trd_pct'), '.2f', '%')} against the "
          f"{trd['trd_limit_pct']:.1f}% limit of IEEE 1547-2018 Table 26"
          + (f", on phase {trd['trd_phase'].upper()}."
             if trd.get("trd_phase") else "."))

    if trd.get("orders"):
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        _set_col_widths(tbl, [1.6, 3.0, 3.0, 2.0])
        for cell, text in zip(tbl.rows[0].cells,
                              ["Order", "Worst % of I_rated", "Limit (% of I_rated)",
                               "Result"]):
            _cell_shade(cell, _CHROME_BAND)
            cell.paragraphs[0].add_run(text).bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for h, od in sorted(trd["orders"].items()):
            cells = tbl.add_row().cells
            values = [
                f"H{h}" + (" (even)" if od["even"] else ""),
                f"{od['max_pct_irated']:.3f}%",
                f"{od['limit_pct']:.1f}%",
                "Within limit" if od["pass"] else "Outside limit",
            ]
            for cell, text in zip(cells, values):
                cell.paragraphs[0].add_run(text).font.size = Pt(9)

    for caveat in trd.get("caveats", []):
        _body(doc, caveat)


def _il_basis_phrase(tdd_info: dict) -> str:
    """One sentence saying what IL is, which is not the same thing every time.

    At a load service IL is the measured demand. At a plant it is the entered
    nameplate where there is one, and otherwise the largest export the
    recording happened to catch -- which is a weaker reference and has to say
    so, because a cloudy week shrinks it and inflates every percentage taken
    against it.
    """
    if not tdd_info or not tdd_info.get("il_amps"):
        return ""
    amps = _m(tdd_info["il_amps"], ".0f", " A")
    basis = tdd_info.get("il_basis", "demand")
    if basis == "rated_output":
        kw = tdd_info.get("rated_ac_kw")
        rating = f" ({kw:,.0f} kW AC nameplate)" if kw else ""
        return (f"This is a generating facility, so the reference current (IL) "
                f"is its rated output{rating}, {amps}, rather than a demand "
                f"load it does not have.")
    if basis == "measured_export":
        return (f"This is a generating facility with no demand load to take IL "
                f"from, and no nameplate rating was supplied, so IL is the "
                f"largest export measured in this recording, {amps}. The "
                f"percentages below are therefore relative to what the plant "
                f"did over these days rather than to what it is rated to do; "
                f"a recording taken in poorer conditions would raise them.")
    if basis == "billing":
        return (f"The maximum demand load current (IL) is {amps}, taken from "
                f"billing history as IEEE 519-2022 defines it: the twelve "
                f"previous months' 15- or 30-minute maximum demands, averaged.")
    return (f"IEEE 519-2022 defines the maximum demand load current (IL) as the "
            f"twelve previous months' 15- or 30-minute maximum demands averaged, "
            f"which is a billing quantity rather than a measurement. None was "
            f"supplied, so the largest fundamental current in this recording, "
            f"{amps}, is used in its place. Where the recording covers a period "
            f"that is not representative of the year, the percentages taken "
            f"against it move accordingly.")

# ─────────────────────────────────────────────────────────────────────────────
# Word report dependencies (optional)
# ─────────────────────────────────────────────────────────────────────────────

try:
    from docx import Document as _DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as _DOCX_REL
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ── Colour system ────────────────────────────────────────────────────────────
# One colour, one job. Xcel red is the brand: headings, titles and table header
# bands, and nothing else. Severity owns the green-amber-orange-maroon ramp,
# blue means "nothing is wrong yet", grey means "no reading to grade". Alarm was
# moved off #CC0000 to deep maroon because it sat six percent from brand red --
# on paper a heading and a severe finding printed as the same colour.
#
# Every text colour here clears 4.5:1 on white (WCAG AA for body text); the
# shades are backgrounds only and are never the sole carrier of a verdict, which
# always also has its own word.
_XE_RED    = RGBColor(0xDA, 0x10, 0x20) if _DOCX_AVAILABLE else None
# Compliance verdicts in prose and in the pass/fail column. These are the
# "compliant" and "severe" bands of _SEVERITY_STYLE below, named for the job
# they do here so a reader of either name lands on the same two colours.
_PASS_CLR  = RGBColor(0x1F, 0x7A, 0x1F) if _DOCX_AVAILABLE else None
_FAIL_CLR  = RGBColor(0x8C, 0x1D, 0x1D) if _DOCX_AVAILABLE else None
_GRAY_CLR  = RGBColor(0xF2, 0xF2, 0xF2) if _DOCX_AVAILABLE else None

#: Structural fills. Grey, so that a tinted cell anywhere in either document
#: means a severity band and never decoration.
_CHROME_HDR   = "DA1020"    # header band, white text on brand red
_CHROME_BAND  = "E4E4E4"    # group divider inside a table
_CHROME_LABEL = "EFEFEF"    # label column of a two-column table

#: Body font. Arial at 10 pt is what this group writes in, so a report that
#: gets pasted into other correspondence matches it without restyling.
_BASE_FONT = "Arial"
_BASE_PT   = 10


def _cell_shade(cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _write_in_field(doc, label: str, *, lines: int = 1, indent_cm: float = 0.0,
                    width_cm: float = 16.0, label_size_pt: int = 10):
    """A box the engineer types into, rather than a rule they would write on.

    Both documents are normally filled in and sent as files, not printed, and a
    row of underscores is the wrong shape for that: typing on it inserts ahead
    of the rule instead of sitting on it, so the line walks off the margin. A
    single-cell table takes a click anywhere inside it, grows down as the notes
    get longer, and still reads as a place to write on a printed copy.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    _bold(p, label, size_pt=label_size_pt)

    box = doc.add_table(rows=1, cols=1)
    _set_col_widths(box, [width_cm])
    cell = box.rows[0].cells[0]
    _cell_shade(cell, "FCFCFD")

    # A light outline, not the report's table rule: this is a field to fill,
    # and it should not read as another data table.
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'BFBFBF')
        borders.append(el)
    tblPr = box._tbl.tblPr
    tblPr.append(borders)
    if indent_cm:
        ind = OxmlElement('w:tblInd')
        ind.set(qn('w:w'), str(int(indent_cm * 567)))   # twips
        ind.set(qn('w:type'), 'dxa')
        tblPr.append(ind)

    # "At least", so the box is visibly a space to write in when empty and
    # still expands to hold however much is typed into it.
    trPr = box.rows[0]._tr.get_or_add_trPr()
    height = OxmlElement('w:trHeight')
    height.set(qn('w:val'), str(int(lines * 0.55 * 567)))
    height.set(qn('w:hRule'), 'atLeast')
    trPr.append(height)
    return cell


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


# ── Measured values in prose ─────────────────────────────────────────────────
# Marked values (see pq_constants.measured) render bold here; every other
# output strips the markers instead.
_MEASURED_OPEN  = MEASURED_OPEN
_MEASURED_CLOSE = MEASURED_CLOSE
_MEASURED_RE = re.compile(f"{_MEASURED_OPEN}(.*?){_MEASURED_CLOSE}", re.S)


def _emit_text(para, text: str, *, bold: bool = False, color=None,
               size_pt: Optional[float] = _BASE_PT):
    """Add ``text`` to a paragraph, breaking measured values into bold runs."""
    def _add(chunk: str, measured: bool):
        if not chunk:
            return
        run = para.add_run(chunk)
        run.bold = bold or measured
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = color

    pos = 0
    for match in _MEASURED_RE.finditer(text):
        _add(text[pos:match.start()], False)
        _add(match.group(1), True)
        pos = match.end()
    _add(text[pos:], False)


def _bold(para, text: str, color=None, size_pt: int = _BASE_PT):
    _emit_text(para, text, bold=True, color=color, size_pt=size_pt)


def _normal(para, text: str, color=None, size_pt: int = _BASE_PT):
    _emit_text(para, text, color=color, size_pt=size_pt)


def _pf_sym(passes) -> str:
    # "Within" / "Exceeded" states the compliance fact without the alarm that
    # FAIL carries; severity is now a separate column and does the grading.
    if passes is True:  return "Within"
    if passes is False: return "Exceeded"
    return "N/A"


def _pf_color(passes):
    if passes is True:  return _PASS_CLR
    if passes is False: return _FAIL_CLR
    return None


#: The one severity scale. Band → (text colour, cell shade). Alarm is deep
#: maroon rather than red so it cannot be confused with a brand-red heading;
#: "watch" is blue because nothing is wrong yet, and blue does no other job.
_SEVERITY_STYLE = {
    "compliant":    ((0x1F, 0x7A, 0x1F), "E9F4E9"),
    "watch":        ((0x1F, 0x5C, 0x8B), "E7EFF7"),
    "minor":        ((0x8A, 0x6D, 0x00), "FAF2D8"),
    "significant":  ((0xA8, 0x50, 0x00), "FBE8D6"),
    "severe":       ((0x8C, 0x1D, 0x1D), "F5DCDC"),
    "not_assessed": ((0x6B, 0x6B, 0x6B), "F0F0F0"),
}

#: Three analyses grew their own severity words before there was a shared
#: scale: neutral integrity grades critical/warning/caution/normal, the
#: assessment findings grade critical/warning/info, and recommended actions
#: rank High/Medium/Low. The analysis code keeps those keys -- they are its
#: data, and tests and CSVs read them -- but every one of them renders through
#: the band it maps to here, so one seriousness prints as one word in one
#: colour wherever the reader meets it.
_SEVERITY_ALIAS = {
    "critical": "severe",
    "warning":  "significant",
    "caution":  "minor",
    "normal":   "compliant",
    "info":     "watch",
    "High":     "severe",
    "Medium":   "significant",
    "Low":      "watch",
}


def _sev_band(band) -> str:
    """Canonical band for any of the scales the analysis code still speaks."""
    band = _SEVERITY_ALIAS.get(band, band)
    return band if band in _SEVERITY_STYLE else "not_assessed"


def _sev_color(band):
    rgb, _ = _SEVERITY_STYLE[_sev_band(band)]
    return RGBColor(*rgb) if _DOCX_AVAILABLE else None


def _sev_shade(band) -> str:
    return _SEVERITY_STYLE[_sev_band(band)][1]


def _sev_label(band) -> str:
    """The word the reader sees, from the shared vocabulary."""
    return SEVERITY_LABEL[_sev_band(band)]


# ─────────────────────────────────────────────────────────────────────────────
# 7. REPORT GENERATION & EXPORT
# ─────────────────────────────────────────────────────────────────────────────

def generate_report(
    ds: PQDataset,
    volt_result: dict,
    thd_result: dict,
    pf_result: dict,
    volt_imb_result: dict,
    curr_imb_result: dict,
    demand_result: dict,
    harm_result: dict,
    volt_harm_result: dict,
    neutral_harm_result: dict,
    source_harm_result: dict,
    stat_result: dict,
    event_result: dict,
    thresh: Thresholds,
    neutral_health_result: Optional[dict] = None,
    spectral_shape_result: Optional[dict] = None,
    itic_result: Optional[dict] = None,
    direction_result: Optional[dict] = None,
    impedance_result: Optional[dict] = None,
    ll_volt_result: Optional[dict] = None,
    frequency_result: Optional[dict] = None,
    flicker_result: Optional[dict] = None,
    kfactor_result: Optional[dict] = None,
) -> dict:
    """Compile all analysis results into a structured summary dictionary."""
    df = ds.df

    transformer_pass: Optional[bool] = None
    if "transformer" in demand_result:
        # None means the transformer is shared and this service's demand did
        # not on its own exceed the nameplate: not determinable, which is not
        # a pass. `not None` would have been True and read as an all-clear.
        _over = demand_result["transformer"]["overloaded"]
        transformer_pass = None if _over is None else not _over

    report = {
        "file_summary": {
            "start_time":       df.index[0].strftime("%Y-%m-%d %H:%M"),
            "end_time":         df.index[-1].strftime("%Y-%m-%d %H:%M"),
            "duration_hours":   round(ds.duration_hours, 3),
            "sample_count":     len(df),
            "channels":         sorted(df.columns.tolist()),
            "interval_minutes": ds.meta.get("interval_minutes", 5),
            "topology":         ds.meta.get("topology", "unknown"),
            "data_quality":     ds.meta.get("data_quality", {}),
            "channel_map":      ds.meta.get("channel_map", {}),
            "device_channels":  ds.meta.get("device_channels", 0),
            "sessions":         ds.meta.get("sessions", []),
            "session_index":    ds.meta.get("session_index", 0),
            "has_maxmin":       ds.has_maxmin,
            "has_adaptive":     ds.has_adaptive,
            "catalog":          ds.catalog(),
        },
        "voltage_compliance":    volt_result,
        "voltage_ll_compliance": ll_volt_result or {"available": False,
                                                    "error": "not evaluated"},
        "frequency":             frequency_result or {"available": False,
                                                      "error": "not evaluated"},
        "flicker":               flicker_result or {"available": False,
                                                    "error": "not evaluated"},
        "kfactor":               kfactor_result or {"available": False,
                                                    "note": "not evaluated"},
        "thd_compliance":        thd_result,
        # Which standard governs current distortion here, and the IEEE 1547
        # assessment where one applies. Both are derived rather than passed in:
        # they depend only on the thresholds and the frame, and threading two
        # more results through every caller would buy nothing.
        "current_standard":      applicable_current_standard(thresh),
        "trd_compliance":        (check_trd(ds.df, thresh) if exports_power(thresh)
                                  else {"available": False,
                                        "note": "No generation at this service."}),
        "power_factor":          pf_result,
        "voltage_imbalance":     volt_imb_result,
        "current_imbalance":     curr_imb_result,
        "demand":                demand_result,
        "individual_harmonics":         harm_result,
        "individual_voltage_harmonics": volt_harm_result,
        "neutral_harmonics":            neutral_harm_result,
        "harmonic_sources":             source_harm_result,
        "harmonic_direction":           direction_result or {"available": False,
                                                             "note": "not evaluated"},
        "spectral_shape":               spectral_shape_result or {"available": False},
        "harmonic_statistics":          stat_result,
        "events":                       event_result,
        "itic":                         itic_result or {"available": False, "note": "not evaluated"},
        # The plant-side counterpart of ITIC. Derived here rather than passed
        # in: it needs only the events and the thresholds, both already to hand.
        "ride_through":                 check_ride_through(event_result, thresh),
        "frequency_ride_through":       check_frequency_ride_through(ds, thresh),
        "neutral_health":               neutral_health_result or {"available": False, "reason": "not run"},
        "service_impedance":            impedance_result or {"available": False,
                                                             "reason": "not evaluated"},
        "pass_fail": {
            "transformer_loading":    transformer_pass,
            "voltage":                volt_result["total_pct_out_of_bounds"] == 0
                                      if volt_result["available"] else None,
            "voltage_line_to_line":   (ll_volt_result or {}).get("overall_pass")
                                      if (ll_volt_result or {}).get("available") else None,
            "frequency":              (frequency_result or {}).get("overall_pass")
                                      if (frequency_result or {}).get("available") else None,
            "flicker":                (flicker_result or {}).get("overall_pass")
                                      if (flicker_result or {}).get("available") else None,
            # IEEE 519-2022 Clause 5 judges voltage THD on the 95th percentile of
            # short-time values (and P99 against 1.5x the limit), not on every
            # sample.  Failing a site on a single artifact interval overstates
            # what the standard actually requires.
            "thd_voltage":            (thd_result["voltage"].get("p95_pass", True)
                                       and thd_result["voltage"].get("p99_pass", True))
                                      if thd_result["voltage"]["available"] else None,
            "thd_current":            thd_result["current"]["pct_exceeding"] == 0
                                      if thd_result["current"]["available"] else None,
            "individual_harmonics":   harm_result.get("overall_pass", None)
                                      if harm_result.get("available") else None,
            "individual_voltage_harmonics": volt_harm_result.get("overall_pass", None)
                                      if volt_harm_result.get("available") else None,
            "power_factor":           pf_result["pct_below_limit"] == 0
                                      if pf_result["available"] else None,
            "voltage_imbalance":      volt_imb_result["pct_exceeding"] == 0
                                      if volt_imb_result["available"] else None,
            # No limit means not evaluated, which is None -- not a pass. On a
            # two-leg service the leg difference is reported as a measurement
            # and nothing here is being tested against a threshold.
            "current_imbalance":      curr_imb_result["pct_exceeding"] == 0
                                      if (curr_imb_result["available"]
                                          and curr_imb_result.get("limit_pct") is not None)
                                      else None,
            "harmonic_statistics":    stat_result.get("overall_pass")
                                      if stat_result.get("available") else None,
            "neutral_health":         (
                (neutral_health_result or {}).get("severity") in ("normal", "caution")
                if (neutral_health_result or {}).get("available") else None
            ),
            "itic_transients":        (
                (itic_result or {}).get("overall_pass")
                if (itic_result or {}).get("available") else None
            ),
        },
    }
    return report


def print_report(report: dict) -> None:
    """Print a human-readable summary to stdout."""
    sep = "─" * 60
    print(f"\n{'═'*60}")
    print(f"  POWER QUALITY ANALYSIS SUMMARY  (pq-analyzer v{__version__})")
    print(f"{'═'*60}")
    fs = report["file_summary"]
    print(f"  Period : {fs['start_time']}  →  {fs['end_time']}")
    print(f"  Duration: {fs['duration_hours']:.2f} h   |   Samples: {fs['sample_count']:,}")
    print(f"  Channels in use: {', '.join(fs['channels'])}")

    # ── Demand ────────────────────────────────────────────────────────────────
    print(f"\n{sep}")
    print("  DEMAND")
    dem = report["demand"]
    if not dem["available"]:
        print(f"  {dem['error']}")
    else:
        if "apparent_power" in dem:
            ap = dem["apparent_power"]
            lf = f"{ap['load_factor']:.3f}" if ap["load_factor"] is not None else "n/a"
            print(f"  Apparent: peak={ap['peak_kva']:.1f} kVA  mean={ap['mean_kva']:.1f} kVA  "
                  f"8-hr peak={ap['peak_8h_kva']:.1f} kVA  load factor={lf}")
        if "real_power" in dem:
            rp = dem["real_power"]
            print(f"  Real:     peak={rp['peak_kw']:.1f} kW   mean={rp['mean_kw']:.1f} kW")
        if "reactive_power" in dem:
            qp = dem["reactive_power"]
            print(f"  Reactive: peak={qp['peak_kvar']:.1f} kVAR  mean={qp['mean_kvar']:.1f} kVAR")
        if "peak_current" in dem:
            pc = dem["peak_current"]
            ph_str = "  ".join(f"{ph.upper()}={a:.0f} A" for ph, a in pc["phases"].items())
            print(f"  Peak current (interval max): {pc['max_a']:.0f} A worst  [{ph_str}]")
        if "transformer" in dem:
            tx = dem["transformer"]
            if tx["overloaded"] is True:
                sym = "FAIL — OVERLOADED"
            elif tx["overloaded"] is False:
                sym = "PASS"
            else:
                sym = "NOT DETERMINABLE — shared transformer"
            print(f"  Transformer: {tx['nameplate_kva']:.0f} kVA nameplate  "
                  f"8-hr peak={tx['peak_8h_kva']:.1f} kVA ({tx['pct_nameplate']:.1f}%)  [{sym}]")
            if tx["overloaded"] is None and tx.get("note"):
                print(f"  {tx['note']}")
            elif tx["overloaded"] is True and not tx.get("dedicated"):
                # Determined, and determined from this service alone. The
                # shared-transformer caveat cuts the other way here.
                print("  This service alone exceeds the nameplate; other "
                      "customers on this transformer add to that.")
        else:
            print("  (Pass --transformer-kva to check transformer loading)")

    # ── Sessions ──────────────────────────────────────────────────────────────
    _sessions = report["file_summary"].get("sessions") or []
    if len(_sessions) > 1:
        _current = report["file_summary"].get("session_index", 0)
        print(f"\n{sep}")
        print(f"  ⚠  {len(_sessions)} RECORDING SESSIONS IN THIS FILE — "
              f"SESSION {_current + 1} ANALYSED")
        for s in _sessions:
            mark = "→" if s["index"] == _current else " "
            print(f"  {mark} Session {s['index'] + 1}: "
                  f"{(s['start_time'] or '')[:16].replace('T', ' ')} → "
                  f"{(s['end_time'] or '')[:16].replace('T', ' ')}  "
                  f"({s['hours']:.1f} h, {s['intervals']} intervals)")
        print("  Re-run with --session to analyse another one.")

    # ── File integrity ────────────────────────────────────────────────────────
    dq = (report["file_summary"].get("data_quality") or {})
    if dq.get("missing_bytes") or dq.get("unreadable_observations"):
        print(f"\n{sep}")
        print("  ⚠  SOURCE FILE INCOMPLETE")
        n_bad   = dq.get("unreadable_observations") or 0
        n_total = dq.get("total_observations") or 0
        if n_bad:
            scope = f"{n_bad} of {n_total}" if n_total else str(n_bad)
            print(f"  {scope} observation records could not be decoded and were "
                  "skipped.  An observation")
            print("  record is one block of measurements the meter wrote to the "
                  "file.")
            for d in (dq.get("unreadable_detail") or [])[:3]:
                label = d.get("name") or f"offset {d.get('offset', '?')}"
                print(f"    - {label}: {strip_marks(d.get('reason', ''))}")
        if dq.get("missing_bytes"):
            print(f"  The file is {dq['missing_bytes']:,} bytes shorter than its own "
                  "record headers")
            print("  declare, so it was cut short — an export or copy that ended "
                  "early.")
        if n_total and n_bad and n_bad < n_total:
            print(f"  Everything below was computed from the {n_total - n_bad} "
                  "records that read cleanly.")
        print("  Measurements in the damaged portion are missing entirely — not "
              "zero, just absent —")
        print("  so counts, minimums and maximums may understate what the meter "
              "saw.  Re-export")
        print("  the file and re-run before relying on these results.")

    # ── Voltage ───────────────────────────────────────────────────────────────
    print(f"\n{sep}")
    print("  VOLTAGE (ANSI C84.1)")
    vc = report["voltage_compliance"]
    if not vc["available"]:
        print(f"  {vc['error']}")
    else:
        _any_extremes = any(s.get("used_interval_extremes") for s in vc["phases"].values())
        for ph, s in vc["phases"].items():
            sym = {"range_a": "PASS", "range_b": "RANGE B"}.get(s["band"], "FAIL")
            print(f"  {ph:12s}: {s['min_v']:6.1f} / {s['mean_v']:6.1f} / {s['max_v']:6.1f} V  "
                  f"  {s['pct_out_of_bounds']:5.2f}% OOB  [{sym}]")
        grp = ("2.4–34.5 kV group" if vc.get("nominal_group") == "over_600v"
               else "120–600 V group")
        print(f"  Range A: {vc['range_v'][0]:.1f} – {vc['range_v'][1]:.1f} V  "
              f"(nominal {vc['nominal_v']:.1f} V, C84.1 Table 1 {grp})")
        if vc.get("range_b_v"):
            print(f"  Range B: {vc['range_b_v'][0]:.1f} – {vc['range_b_v'][1]:.1f} V")
        print("  (verdict from interval averages — C84.1 rates sustained voltage)")
        if _any_extremes:
            print("  (within-interval max/min also read; sags and swells are graded "
                  "against ITIC)")

    # ── Line-to-line voltage ──────────────────────────────────────────────────
    llv = report.get("voltage_ll_compliance") or {"available": False}
    if llv.get("available"):
        print(f"\n{sep}")
        print("  LINE-TO-LINE VOLTAGE (ANSI C84.1)")
        for pair, st in llv["pairs"].items():
            sym = "PASS" if st["pct_out_of_bounds"] == 0 else "FAIL"
            print(f"  {pair:12s}: {st['min_v']:6.1f} / {st['mean_v']:6.1f} / "
                  f"{st['max_v']:6.1f} V    {st['pct_out_of_bounds']:5.2f}% OOB  [{sym}]")
        print(f"  Allowed range: {llv['range_v'][0]:.1f} – {llv['range_v'][1]:.1f} V  "
              f"(nominal {llv['nominal_v']:.0f} V ± "
              f"{(llv['range_v'][1]/llv['nominal_v']-1)*100:.0f}%, {llv['configuration']})")

    # ── Frequency ─────────────────────────────────────────────────────────────
    freq = report.get("frequency") or {"available": False}
    if freq.get("available"):
        print(f"\n{sep}")
        print("  FREQUENCY")
        sym = "PASS" if freq["overall_pass"] else "FAIL"
        print(f"  {freq['min_hz']:.3f} / {freq['mean_hz']:.3f} / {freq['max_hz']:.3f} Hz"
              f"   max deviation {freq['max_deviation_hz']:.3f} Hz"
              f"   out of band {freq['pct_out_of_band']:.2f}%  [{sym}]")
        print(f"  Allowed range: {freq['range_hz'][0]:.2f} – {freq['range_hz'][1]:.2f} Hz "
              f"(nominal {freq['nominal_hz']:.0f} Hz)")

    # ── Flicker ───────────────────────────────────────────────────────────────
    fl = report.get("flicker") or {"available": False}
    if fl.get("available"):
        print(f"\n{sep}")
        print("  FLICKER (IEC 61000-3-3)")
        for kind, label, limit in (("pst", "Pst", fl["pst_limit"]),
                                   ("plt", "Plt", fl["plt_limit"])):
            if not fl[kind]:
                continue
            per_phase = "  ".join(
                f"{p}: med {v['median']:.2f} p95 {v['p95']:.2f} max {v['max']:.2f}"
                for p, v in sorted(fl[kind].items()))
            worst = max(fl[kind].values(), key=lambda v: v["max"])
            sym = "PASS" if worst["max"] <= limit else "FAIL"
            window = "10 min" if kind == "pst" else "2 h"
            print(f"  {label} ({window}) by phase — {per_phase}"
                  f"   limit={limit:.2f}  [{sym}]")
        print(f"  Governing phase: {fl['worst_phase']} "
              f"({fl['worst_ratio_of_limit']:.2f}× its limit)")

    # ── THD / TDD ─────────────────────────────────────────────────────────────
    print(f"\n{sep}")
    tdd_info = report["thd_compliance"].get("tdd_info", {})
    if tdd_info and tdd_info.get("isc_provided"):
        print(f"  THD / TDD (IEEE 519-2022)")
        print(f"  ISC={tdd_info['isc_amps']:.0f} A  IL={tdd_info['il_amps']:.0f} A  "
              f"ISC/IL={tdd_info['isc_il_ratio']:.1f}  "
              f"class {tdd_info['tdd_class']}  → TDD limit {tdd_info['tdd_limit_pct']:.1f}%")
        if tdd_info.get("isc_source"):
            print(f"  ISC source: {tdd_info['isc_source']}")
    elif tdd_info:
        print(f"  THD / TDD (IEEE 519-2022)")
        print(f"  ISC not provided  IL={tdd_info['il_amps']:.0f} A  "
              f"class {tdd_info['tdd_class']} → conservative TDD limit "
              f"{tdd_info['tdd_limit_pct']:.1f}%  [pass --isc for true class]")
    else:
        print("  THD (IEEE 519)  [no RMS current channels — TDD unavailable]")

    for key, label in [("voltage", "Voltage THD"), ("current", "Current TDD" if tdd_info else "Current THD")]:
        td = report["thd_compliance"][key]
        if not td["available"]:
            print(f"  {label}: no data")
            continue
        sym = "PASS" if td["pct_exceeding"] == 0 else "FAIL"
        ll_tag = "  [light-load intervals excluded]" if td.get("light_load_filtered") else ""
        print(f"  {label}: max={td['max_thd_pct']:.2f}%  mean={td['mean_thd_pct']:.2f}%  "
              f"limit={td['limit_pct']:.1f}%  exceed={td['pct_exceeding']:.2f}%  [{sym}]{ll_tag}")
        if key == "current" and td.get("harmonic_rms_source"):
            print(f"  Harmonic RMS source: {td['harmonic_rms_source']}"
                  + (f"   IL={td['il_amps']:.1f} A (fundamental)"
                     if td.get("il_amps") else ""))
        if key == "current" and "peak_max_tdd_pct" in td:
            pk_sym = "PASS" if td["peak_pct_exceeding"] == 0 else "FAIL"
            print(f"  {label} (peak within interval): max={td['peak_max_tdd_pct']:.2f}%  "
                  f"exceed={td['peak_pct_exceeding']:.2f}%  [{pk_sym}]")

    # ── Power factor ──────────────────────────────────────────────────────────
    print(f"\n{sep}")
    print("  POWER FACTOR")
    pfr = report["power_factor"]
    if not pfr["available"]:
        print(f"  {pfr['error']}")
    else:
        sym = "PASS" if pfr["pct_below_limit"] == 0 else "FAIL"
        print(f"  Min={pfr['min_pf']:.4f}  Mean={pfr['mean_pf']:.4f}  "
              f"Limit={pfr['limit']:.2f}  Below limit={pfr['pct_below_limit']:.2f}%  [{sym}]")

    # ── Voltage imbalance ─────────────────────────────────────────────────────
    print(f"\n{sep}")
    print("  VOLTAGE IMBALANCE (NEMA MG1 / ANSI C84.1)")
    imb = report["voltage_imbalance"]
    if not imb["available"]:
        print(f"  {imb['error']}")
    else:
        sym = "PASS" if imb["pct_exceeding"] == 0 else "FAIL"
        print(f"  Max={imb['max_imbalance_pct']:.2f}%  Mean={imb['mean_imbalance_pct']:.2f}%  "
              f"Limit={imb['limit_pct']:.1f}%  Exceed={imb['pct_exceeding']:.2f}%  [{sym}]")

    # ── Current imbalance ─────────────────────────────────────────────────────
    print(f"\n{sep}")
    print("  CURRENT IMBALANCE")
    ci = report["current_imbalance"]
    if not ci["available"]:
        print(f"  {ci['error']}")
    elif ci.get("limit_pct") is None:
        # Two legs: reported, not tested. Printing PASS here would imply a
        # threshold was cleared when none was applied.
        print(f"  {ci.get('metric_label', 'Leg current difference')}: "
              f"Max={ci['max_imbalance_pct']:.2f}%  "
              f"Mean={ci['mean_imbalance_pct']:.2f}%  [MEASUREMENT — no limit]")
        if ci.get("note"):
            print(f"  {ci['note']}")
    else:
        sym = "PASS" if ci["pct_exceeding"] == 0 else "FAIL"
        print(f"  Max={ci['max_imbalance_pct']:.2f}%  Mean={ci['mean_imbalance_pct']:.2f}%  "
              f"Limit={ci['limit_pct']:.1f}%  Exceed={ci['pct_exceeding']:.2f}%  [{sym}]")
        if "neutral_current" in ci:
            nc = ci["neutral_current"]
            print(f"  Neutral current: mean={nc['mean_amps']:.1f} A ({nc['mean_pct_of_phase']:.1f}% of phase avg)  "
                  f"max={nc['max_amps']:.1f} A ({nc['max_pct_of_phase']:.1f}%)")

    # ── Individual harmonics ──────────────────────────────────────────────────
    print(f"\n{sep}")
    ih = report["individual_harmonics"]
    if not ih.get("available"):
        note = ih.get("note", "Pass --isc to enable per-order IEEE 519-2022 check")
        print(f"  INDIVIDUAL CURRENT HARMONICS (IEEE 519-2022 Table 2)  [{note}]")
    else:
        sym = "PASS" if ih["overall_pass"] else "FAIL"
        print(f"  INDIVIDUAL CURRENT HARMONICS (IEEE 519-2022 Table 2)  [{sym}]")
        print(f"  IL={ih['il_amps']:.0f} A  ISC/IL={ih['isc_il_ratio']:.1f}")
        # Print header + one row per harmonic order showing worst phase
        order_rows = []
        for h in _H519_ORDERS:
            worst_max = 0.0; worst_ph = None; limit = 0.0
            for ph in ("a", "b", "c"):
                r = ih["phases"].get(ph, {}).get(h)
                if r and r["max_pct_il"] > worst_max:
                    worst_max = r["max_pct_il"]; worst_ph = ph; limit = r["limit_pct_il"]
            if worst_ph:
                sym_h = "PASS" if worst_max <= limit else "FAIL"
                order_rows.append((h, worst_max, limit, worst_ph, sym_h))
        if order_rows:
            print(f"  {'H':>3}  {'worst%IL':>9}  {'limit%IL':>9}  {'ph':>3}  status")
            for h, wmax, lim, wph, sym_h in order_rows:
                marker = " ←" if sym_h == "FAIL" else ""
                print(f"  H{h:<2}  {wmax:>9.2f}  {lim:>9.1f}  {wph:>3}  {sym_h}{marker}")

    # ── Neutral harmonics ─────────────────────────────────────────────────────
    nh = report.get("neutral_harmonics", {})
    if nh.get("available"):
        print(f"\n{sep}")
        print("  NEUTRAL HARMONICS (informational)")
        t_pct = nh.get("triplen_pct", 0.0)
        acc   = nh.get("accumulation_factor")
        if nh.get("accumulation_note"):
            # Split-phase: the factor is not withheld for lack of data, it does
            # not apply. Saying "n/a" would read as a gap in the recording.
            print(f"  Triplen content: {t_pct:.0f}% of neutral harmonic current")
            print(f"  {nh['accumulation_note']}")
        else:
            acc_s = f"{acc:.1f}×" if acc is not None else "n/a"
            print(f"  Triplen content: {t_pct:.0f}%  |  Accumulation factor: {acc_s}")
        for h, od in sorted(nh["orders"].items()):
            tag = " [triplen]" if od["is_triplen"] else "          "
            print(f"  H{h:<3}{tag}  mean={od['mean_a']:.3f} A  max={od['max_a']:.3f} A")

    # ── Harmonic source indication ────────────────────────────────────────────
    sh = report.get("harmonic_sources", {})
    if sh.get("available"):
        print(f"\n{sep}")
        print("  HARMONIC SOURCE ATTRIBUTION (indicative)")
        overall_labels = {
            "customer":          "Customer-side injection",
            "resonance_suspect": "Resonance suspected",
            "mixed":             "Mixed / indeterminate",
            "indeterminate":     "Indeterminate",
        }
        print(f"  Overall: {overall_labels.get(sh.get('overall'), sh.get('overall'))}")
        resonant = sh.get("resonant_orders", [])
        if resonant:
            print(f"  Resonance suspects: {', '.join('H'+str(h) for h in sorted(resonant))}")
        print(f"  {'Order':<6}  {'Z_h(Ω)':>8}  {'Z_ratio':>8}  {'Correlation':>10}  Indication")
        for h, od in sorted(sh["orders"].items()):
            ratio_s = f"{od['z_ratio']:.2f}×" if od["z_ratio"] is not None else "   n/a"
            corr_s  = f"{od['corr']:.2f}"     if od["corr"]    is not None else "   n/a"
            print(f"  H{h:<5}  {od['z_ohm']:>8.4f}  {ratio_s:>8}  {corr_s:>10}  {od['attribution']}")

    # ── Service impedance ─────────────────────────────────────────────────────
    si = report.get("service_impedance", {})
    if si.get("available"):
        print(f"\n{sep}")
        print(f"  SERVICE IMPEDANCE — {_IMPEDANCE_HEADLINE.get(si.get('overall'), si.get('overall'))}")
        for ph, fit in sorted(si.get("phases", {}).items()):
            if fit.get("identifiable"):
                r = f"{fit['r_ohm']:.4f}" if fit.get("r_ohm") is not None else "   —  "
                x = f"{fit['x_ohm']:.4f}" if fit.get("x_ohm") is not None else "   —  "
                print(f"  Phase {ph.upper()}:  R={r} Ω  X={x} Ω  |Z|={fit['z_ohm']:.4f} Ω  "
                      f"({fit['steps']} load steps, {fit['consistency']:.0%} consistent)")
            else:
                print(f"  Phase {ph.upper()}:  not measurable — "
                      f"{strip_marks(fit.get('reason', ''))}")
        asym = si.get("asymmetry") or {}
        if asym.get("ratio"):
            mark = "⚠ " if asym["flagged"] else "  "
            print(f"  {mark}Between phases: {asym['worst_phase']} is "
                  f"{asym['ratio']:.1f}× {asym['best_phase']} "
                  f"({asym['excess_v_at_peak']:.1f} V extra at peak load)")
        neutral = si.get("neutral") or {}
        if neutral.get("identifiable"):
            mark = "⚠ " if neutral.get("elevated") else "  "
            print(f"  {mark}Neutral: R={neutral['r_ohm']:.4f} Ω → "
                  f"{neutral['rise_at_peak_v']:.1f} V at {neutral['i_peak_a']:.0f} A")
        elif neutral.get("at_resolution"):
            print("    Neutral: no rise above the meter's resolution — a sound "
                  "neutral reads this way")
        cmp_ = si.get("comparison") or {}
        if cmp_:
            mark = "⚠ " if cmp_["verdict"] in ("high", "elevated") else "  "
            print(f"  {mark}Measured {cmp_['measured_ohm']:.4f} Ω vs expected "
                  f"{cmp_['expected_ohm']:.4f} Ω = {cmp_['ratio']:.1f}×  "
                  f"({cmp_['excess_v_at_peak']:+.1f} V at {cmp_['i_peak_a']:.0f} A peak)")
        elif not (si.get("expected") or {}).get("available"):
            if (si.get("expected") or {}).get("primary_metered"):
                print("    No expected impedance: this service is metered on "
                      "the primary — enter the primary line R1 and X1 to "
                      "compare against one.")
            else:
                print("    No expected impedance: pick the service conductor and run "
                      "length to compare against one.")
        exp_ = si.get("expected") or {}
        if exp_.get("shared_secondary_z_ohm") is not None:
            print(f"    Includes {exp_.get('shared_secondary_ft', 0):.0f} ft of "
                  f"shared secondary — neighbours' current on it widens the fit")

    # ── Harmonic source direction ─────────────────────────────────────────────
    hd = report.get("harmonic_direction", {})
    if hd.get("available"):
        iv = hd.get("interval") or {}
        wf = hd.get("waveform") or {}
        print(f"\n{sep}")
        print("  HARMONIC SOURCE DIRECTION (which side of the meter)")
        print(f"  Over the recording: {_DIRECTION_LABEL.get(iv.get('overall'), '—')}"
              f"   |   At the captures: {_DIRECTION_LABEL.get(wf.get('overall'), '—')}")
        if iv.get("available"):
            print(f"  {'Order':<6}  {'Z_h(Ω)':>8}  {'FromLoad':>9}  {'Backgnd':>8}  "
                  f"{'r':>6}  Indication")
            for h, od in sorted((iv.get("orders") or {}).items()):
                corr_s = f"{od['corr']:.2f}" if od["corr"] is not None else "   n/a"
                print(f"  H{h:<5}  {od['slope_ohm']:>8.3f}  {od['v_from_load_v']:>9.2f}  "
                      f"{od['v_background_v']:>8.2f}  {corr_s:>6}  "
                      f"{_DIRECTION_LABEL.get(od['indication'], od['indication'])}")
        if wf.get("available"):
            print(f"  Captures used: {wf['captures_used']} of {wf['captures_total']}"
                  f"{_direction_exclusions(wf)}")
            print(f"  {'Order':<6}  {'Readings':>8}  {'OutBound':>9}  {'MedianP(W)':>11}  "
                  f"{'Angle':>7}  Indication")
            for h, od in sorted((wf.get("orders") or {}).items()):
                print(f"  H{h:<5}  {od['samples']:>8}  {od['toward_system']:>9}  "
                      f"{_signed_watts(od['median_p_w']):>11}  "
                      f"{od['median_angle_deg']:>6.0f}°  "
                      f"{_DIRECTION_LABEL.get(od['indication'], od['indication'])}")
        elif wf.get("note"):
            print(f"  {wf['note']}")
        for h, verdict in sorted((hd.get("agreement") or {}).items()):
            if verdict == "disagree":
                print(f"  ⚠  H{h}: the two methods disagree — see the report section.")

    # ── Spectral shape (broadband vs. resonance classification) ───────────────
    ss = report.get("spectral_shape", {})
    if ss.get("available"):
        print(f"\n{sep}")
        print("  SPECTRAL SHAPE (single-visit classification, not a trend)")
        class_labels = {
            "broadband_consistent": "Broadband-consistent (elevated, flat across orders)",
            "resonance_present":    "Resonance present (see source indication above)",
            "elevated_uneven":      "Elevated but concentrated (no order flagged resonant)",
            "not_elevated":         "Not meaningfully elevated",
        }
        cls = ss.get("classification")
        print(f"  Classification: {class_labels.get(cls, cls)}")
        print(f"  Mean VTHD: {ss['mean_vthd_pct']:.2f}%  |  Elevation ratio: {ss['elevation_ratio']}  |  "
              f"Flatness CV: {ss['flatness_cv']}")

    # ── Neutral health ────────────────────────────────────────────────────────
    nh = report.get("neutral_health", {})
    if nh.get("available"):
        print(f"\n{sep}")
        print("  NEUTRAL HEALTH (split-phase)")
        sev = nh["severity"].upper()
        print(
            f"  Severity: {sev}  |  "
            f"L1+L2 sum={nh['sum_mean_v']:.1f} V (std {nh['sum_std_v']:.2f} V)  |  "
            f"Leg corr r=" + (f"{nh['leg_correlation']:.3f}" if
                                 nh.get("leg_correlation") is not None
                                 else "n/a") + "  |  "
            f"Asym={nh['asym_mean_v']:.1f} V ({nh['asym_pct']:.1f}%)"
        )
        if nh.get("vne_available"):
            print(f"  Vne: max={nh['vne_max_v']:.2f} V  mean={nh['vne_mean_v']:.2f} V")
        if nh.get("coincident_events"):
            print(f"  Coincident opposing sag/swell events: {nh['coincident_events']}")
        for f_txt in nh.get("findings", []):
            print(f"  • {strip_marks(f_txt)}")

    # ── Engineering assessment (likely causes) ────────────────────────────────
    print(f"\n{sep}")
    print("  ENGINEERING ASSESSMENT (LIKELY CAUSES)")
    rca = report.get("root_causes", [])
    if not rca:
        print("  No assessment findings generated.")
    else:
        _sev_rank = {"critical": 0, "warning": 1, "info": 2}
        for finding in sorted(rca, key=lambda f: _sev_rank.get(f["severity"], 9)):
            sev   = finding["severity"].upper()
            conf  = finding["confidence"].upper()
            title = finding["title"]
            print(f"\n  [{sev}] [{conf} confidence]  {title}")
            print(f"    Finding:       {strip_marks(finding['finding'])}")
            print(f"    Likely cause:  {strip_marks(finding['cause'])}")
            if finding.get("origin_evidence"):
                print(f"    Bearing on origin: "
                      f"{strip_marks(finding['origin_evidence'])}")
            print(f"    Candidate action:  "
                  f"{strip_marks(finding['recommendation'])}")

    # ── Events ────────────────────────────────────────────────────────────────
    print(f"\n{sep}")
    print("  EVENTS")
    ev = report["events"]
    src_label = " (adaptive cycle-level)" if ev.get("data_source") == "adaptive" else " (interval avg)"
    print(f"  Total events detected: {ev['event_count']}{src_label}")
    if ev.get("waveform_captures"):
        print(f"  Point-on-wave waveform captures: {ev['waveform_captures']}")
    if ev["event_count"] and len(ev["events"]) > 0:
        summary = ev["events"]["type"].value_counts()
        for etype, cnt in summary.items():
            print(f"    {etype:20s}: {cnt}")

    print(f"\n{'═'*60}\n")


def export_results(
    ds: PQDataset,
    report: dict,
    outdir: Path,
    stem: str = "pq_analysis",
) -> None:
    """Write CSVs: interval data, adaptive events, violations, and events."""
    df = ds.df
    outdir.mkdir(parents=True, exist_ok=True)

    # 1. Full interval dataset (avg + peak/min columns)
    data_path = outdir / f"{stem}_data.csv"
    df.to_csv(data_path)
    log.info("Saved interval data → %s", data_path)

    # 2. Adaptive (cycle-level) data when present
    if ds.has_adaptive:
        assert ds.adaptive_df is not None
        adap_path = outdir / f"{stem}_adaptive.csv"
        ds.adaptive_df.to_csv(adap_path)
        log.info("Saved adaptive data → %s  (%d rows)", adap_path, len(ds.adaptive_df))

    # 3. Violations CSV — union of all violation timestamps
    viol_sets = []
    vc = report["voltage_compliance"]
    if "violation_timestamps" in vc and len(vc["violation_timestamps"]) > 0:
        viol_sets.append(pd.Series("voltage_oob", index=vc["violation_timestamps"]))

    for key in ("voltage", "current"):
        td = report["thd_compliance"][key]
        if td and td.get("violation_timestamps"):
            idx = pd.DatetimeIndex(td["violation_timestamps"])
            viol_sets.append(pd.Series(f"thd_{key}", index=idx))

    pf_viol = report["power_factor"].get("violation_timestamps")
    if pf_viol is not None and len(pf_viol) > 0:
        viol_sets.append(pd.Series("power_factor", index=pf_viol))

    imb_viol = report["voltage_imbalance"].get("violation_timestamps")
    if imb_viol is not None and len(imb_viol) > 0:
        viol_sets.append(pd.Series("imbalance", index=imb_viol))

    if viol_sets:
        all_viols = pd.concat(viol_sets).rename("violation_type").sort_index()
        all_viols = all_viols[~all_viols.index.duplicated(keep="first")]
        viol_df = df.loc[df.index.intersection(all_viols.index)].copy()
        viol_df.insert(0, "violation_type", all_viols.reindex(viol_df.index))
        viol_path = outdir / f"{stem}_violations.csv"
        viol_df.to_csv(viol_path)
        log.info("Saved violations → %s  (%d rows)", viol_path, len(viol_df))

    # 4. Events CSV
    ev_df = report["events"]["events"]
    if len(ev_df) > 0:
        ev_path = outdir / f"{stem}_events.csv"
        ev_df.to_csv(ev_path, index=False)
        log.info("Saved events → %s  (%d rows)", ev_path, len(ev_df))


# ─────────────────────────────────────────────────────────────────────────────
# 8c. WORD REPORT GENERATOR — private section helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_font(style_or_run, name: str = _BASE_FONT) -> None:
    """Name a font on every script slot, not just the Latin one.

    Setting ``font.name`` alone leaves the theme font in ``w:eastAsia`` and
    ``w:cs``, and Word will fall back to it for characters this report does
    use -- the em dash and the ohm and degree signs among them.
    """
    style_or_run.font.name = name
    rpr = style_or_run.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for slot in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        fonts.set(qn(slot), name)


def _apply_base_style(doc) -> None:
    """Put both documents in Arial: body at 10 pt, headings in Xcel red."""
    normal = doc.styles["Normal"]
    _set_font(normal)
    normal.font.size = Pt(_BASE_PT)

    for level, size in ((1, 14), (2, 12), (3, 11)):
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:      # a template without the built-in heading styles
            continue
        _set_font(style)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _XE_RED


def _section_heading(doc, title: str, level: int = 1):
    """Real Word heading (level 1 or 2) so Word's navigation pane and an
    auto-generated table of contents both work — not just a bolded paragraph."""
    heading = doc.add_heading(title, level=level)
    for run in heading.runs:
        run.font.color.rgb = _XE_RED
        _set_font(run)
    return heading


def _body(doc, text: str) -> None:
    # Size comes from the Normal style, so body prose stays one size even if the
    # house size changes; only measured values pick up their own run.
    _emit_text(doc.add_paragraph(), text, size_pt=None)


def _session_note(fs: dict, *, plain: bool = False) -> str:
    """What the file held, when it held more than this report covers.

    A "download all data" export carries every session the meter still had --
    a reset or a re-arm in the field starts a new one -- and only one is
    analysed, because the gap between two sessions is not recorded time.
    Saying which one, and what else was in the file, is the difference between
    a report a reviewer can check and a report that quietly covers half a
    download. Empty string when the file holds a single session, which is the
    ordinary case and needs no explanation.
    """
    sessions = fs.get("sessions") or []
    if len(sessions) < 2:
        return ""
    current = fs.get("session_index", 0)
    others = []
    for s in sessions:
        if s["index"] == current:
            continue
        start = (s["start_time"] or "")[:16].replace("T", " ")
        hours = f"{s['hours']:.0f}" if plain else _m(s["hours"], ".0f")
        others.append(f"{start} ({hours} hours)")
    listed = "; ".join(others)
    n = len(sessions)
    return (
        f"This file holds {n if plain else _m(n, 'd')} separate recording "
        f"sessions. A meter that is reset or re-armed in the field starts a "
        f"new one, and a download of everything on the meter carries them all. "
        f"This report covers session {current + 1} of {n} only. Also in the "
        f"file: {listed}. Analysing one session at a time is deliberate: the "
        f"time between two sessions was not recorded, so a figure quoted as a "
        f"share of the recording would be wrong if they were run together. "
        f"Re-run against the other session to have it assessed."
    )


def _plot_path(outdir: Optional[Path], stem: str, name: str) -> Optional[Path]:
    """Path of a stem-prefixed plot PNG (matches pq_plots naming)."""
    if outdir is None:
        return None
    return outdir / (f"{stem}_{name}" if stem else name)


def _embed_plot(doc, outdir: Optional[Path], stem: str, name: str,
                caption: str = "", width_cm: float = 15.5) -> bool:
    """Embed a generated plot image if it exists; returns True when embedded."""
    p = _plot_path(outdir, stem, name)
    if p is None or not p.exists():
        return False
    doc.add_picture(str(p), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        r = cap.add_run(caption)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return True


#: Outlook's hyperlink colour. The signature is read beside emails from the
#: same person, so the address is styled the way their mail client styles one.
_LINK_TEAL = RGBColor(0x46, 0x78, 0x86) if _DOCX_AVAILABLE else None


def _add_hyperlink(paragraph, url: str, text: str, *,
                   size_pt: float = 10, color=None, bold: bool = False) -> None:
    """Append a real hyperlink to ``paragraph``.

    python-docx has no API for this: a link is a w:hyperlink element pointing at
    an external relationship on the document part, so both have to be built by
    hand. Worth the code rather than styling plain text to look like a link --
    the address in a customer letter is there to be clicked, and a blue-and-
    underlined run that does nothing when clicked is a small broken promise.
    """
    part  = paragraph.part
    r_id  = part.relate_to(url, _DOCX_REL.HYPERLINK, is_external=True)
    link  = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run  = OxmlElement("w:r")
    rPr  = OxmlElement("w:rPr")
    font = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        font.set(qn(attr), _BASE_FONT)
    rPr.append(font)
    if bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")                      # half-points
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    if color is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), str(color))
        rPr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _signature_block(doc, engineer_name: str, engineer_title: str,
                     engineer_email: str = "") -> None:
    """The sign-off, identical wherever a document is signed.

    Four lines, in the house format:

        Jacob Whitaker              bold, 11 pt, black
        Xcel Energy                 bold, 10 pt, brand red
        Manager, Electric Area Engineering    10 pt, black
        jacob.b.whitaker@…          10 pt, Outlook link teal, live mailto

    Each line is its own paragraph with the space before and after removed, so
    the block reads as one address rather than four stacked paragraphs -- Word's
    Normal style puts 8 pt between paragraphs, which is what made the old
    sign-off sprawl.

    No telephone number. That is a policy decision, not an omission: the letter
    invites a reply instead.
    """
    def line(text, *, size_pt, bold=False, color=None):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size_pt)
        if color is not None:
            run.font.color.rgb = color
        _set_font(run)
        return p

    # Black is left unset rather than set to 000000. It is already the default,
    # and a document whose only explicit colours are the ones that mean
    # something stays checkable -- there is a test asserting the letter paints
    # from the brand palette and nothing else, and an explicit black is a
    # colour that means nothing arriving in front of it.
    line(engineer_name or "[Engineer Name]", size_pt=11, bold=True)
    line("Xcel Energy", size_pt=10, bold=True, color=_XE_RED)
    if engineer_title:
        line(engineer_title, size_pt=10)
    if engineer_email:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _add_hyperlink(p, f"mailto:{engineer_email}", engineer_email,
                       size_pt=10, color=_LINK_TEAL)


def _add_page_field(paragraph, size_pt: float = 8) -> None:
    """Append a PAGE field so Word numbers the page as it paginates.

    A footer that ends in the word "page" followed by literal text cannot
    number anything -- the count only exists once Word lays the document out,
    so the number has to be a field it fills in rather than a string we write.
    """
    run = paragraph.add_run()
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "1"
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(el)


def _add_toc(doc) -> None:
    """Insert a TOC field (levels 1-2) that Word populates automatically on open."""
    _section_heading(doc, "Table of Contents", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "Right-click and choose \"Update Field\" to generate the table of contents."
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(el)

    # Force Word to prompt/auto-refresh all fields (including this TOC) on open,
    # so the reader doesn't have to manually right-click → Update Field.
    settings_el = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings_el.append(update_fields)

    doc.add_page_break()


def _integrity_summary(dq: dict, fs: dict) -> str:
    """The header table's version: enough to judge the report, not to debug it.

    The full diagnosis still has to exist somewhere, because a customer's .pqd
    cannot be sent anywhere and a failure that does not carry its own evidence
    cannot be diagnosed at all. It does not have to sit in the header, where a
    paragraph of byte offsets and zlib state buries the six rows the reader
    came for. Scope, what it costs, and what to do; the evidence is in
    Appendix B.
    """
    parts = []

    n_bad   = dq.get("unreadable_observations") or 0
    n_total = dq.get("total_observations") or 0
    missing = dq.get("missing_bytes") or 0

    if n_bad:
        # "1 of 87 observation records ... was skipped": the noun agrees with
        # the 87, the verb with the 1.
        scope = (f"{n_bad} of {n_total} observation records" if n_total else
                 f"{n_bad} observation record{'' if n_bad == 1 else 's'}")
        line = (f"{scope} could not be decoded and "
                f"{'was' if n_bad == 1 else 'were'} skipped")
        if n_total and n_bad < n_total:
            n_good = n_total - n_bad
            line += ("; the other one reads cleanly and covers the "
                     "recording period above" if n_good == 1 else
                     f"; the other {n_good} read cleanly and cover the "
                     "recording period above")
        parts.append(line + ".")

        # Which record was lost decides which checks are affected -- losing the
        # max/min observation removes the peak and minimum voltage check, and
        # no count conveys that. The names are short; the reasons are not.
        names = [d.get("name", "") for d in (dq.get("unreadable_detail") or [])]
        names = [n for n in names if n]
        if names:
            parts.append("Skipped: " + "; ".join(f"“{n}”" for n in names) + ".")

    if missing:
        parts.append(
            f"The file is {'also ' if n_bad else ''}{missing:,} bytes shorter "
            "than its own record headers declare, so it was cut short.")

    parts.append(
        "Counts, minimums and maximums may understate what the meter saw. "
        "Re-export from the meter and re-run before relying on these results; "
        "the full diagnosis is in Appendix B.")

    return "INCOMPLETE — " + " ".join(parts)


def _integrity_note(dq: dict, fs: dict) -> str:
    """Explain a damaged source file to someone who has never heard of PQDIF.

    The reader has to answer three questions before weighing anything else in
    the report: what was lost, how much of the file that is, and whether the
    findings can still be relied on.  A bare count of "observation records"
    answers none of them.
    """
    parts = []

    n_bad   = dq.get("unreadable_observations") or 0
    n_total = dq.get("total_observations") or 0
    missing = dq.get("missing_bytes") or 0

    if n_bad:
        # Name the unit before counting it — "observation record" is PQDIF
        # jargon that means nothing to a field engineer reading the report.
        scope = (f"{n_bad} of {n_total} observation records" if n_total else
                 f"{n_bad} observation record{'' if n_bad == 1 else 's'}")
        parts.append(
            f"{scope} could not be decoded and "
            f"{'was' if n_bad == 1 else 'were'} skipped. "
            "An observation record is one block of measurements the meter wrote "
            "to the file — typically one channel group over one span of time."
        )
        detail = dq.get("unreadable_detail") or []

        # Name what was lost.  Which record it was determines which checks are
        # affected — losing the max/min observation removes the peak/min voltage
        # check, and no count conveys that.
        names = [d.get("name", "") for d in detail]
        names = [n for n in names if n]
        if names:
            parts.append("Skipped: " + "; ".join(f"“{n}”" for n in names) + ".")

        reasons = [d.get("reason", "") for d in detail if d.get("reason")]
        if reasons:
            # One reason is worth quoting; a list of them is noise in a table.
            # The reader's reason carries its own structural evidence and ends
            # in its own full stop, which is what makes this paragraph enough
            # to diagnose the file without sending the file anywhere.
            first = reasons[0].rstrip().rstrip(".")
            parts.append(f"Reason reported by the reader: {first}"
                         + (f" (and {len(reasons) - 1} more)" if len(reasons) > 1 else "")
                         + ".")

    if missing:
        parts.append(
            f"The file is {'also ' if n_bad else ''}{missing:,} bytes shorter "
            "than its own record headers declare, which means it was cut short — "
            "an export or copy that ended early, not a format this reader fails "
            "to understand."
        )

    # What it means for the findings, stated plainly.
    if n_total and n_bad and n_bad < n_total:
        n_good = n_total - n_bad
        parts.append(
            f"Everything in this report was computed from the {n_good} "
            f"record{'' if n_good == 1 else 's'} that read cleanly, covering the "
            "recording period shown above. Measurements held only in the "
            "skipped records are absent "
            "from every statistic here — they are not counted as zero, but any "
            "interval or event they contained is simply missing, so counts, "
            "minimums and maximums may understate what the meter actually saw."
        )
    else:
        parts.append(
            "Measurements in the damaged portion are absent from every statistic "
            "in this report — not counted as zero, but missing, so counts, "
            "minimums and maximums may understate what the meter actually saw."
        )

    parts.append(
        "Re-export the file from the meter and re-run before relying on these "
        "results. If a re-export reads clean, use it instead of this report."
    )

    return "INCOMPLETE — " + " ".join(parts)


def _word_site_info_table(doc, site_name, stem, site_address,
                          fs, nominal_v, nominal_ll, prepared_by="") -> None:
    rows_data = [
        ("Customer / Site", site_name or stem),
    ]
    # Who ran the analysis, as a header field rather than a signature: an
    # internal document still has to say whose work it is.
    if prepared_by:
        rows_data.append(("Prepared by", prepared_by))
    if site_address:
        rows_data.append(("Address", site_address))
    rows_data += [
        ("Recording period", f"{fs['start_time']}  →  {fs['end_time']}"),
        ("Duration",         f"{fs['duration_hours']:.2f} hours  |  {fs['sample_count']:,} intervals"),
        ("Service voltage",  f"{nominal_v:.0f} V L-N  /  {nominal_ll} V L-L"),
        ("Topology",         fs.get("topology", "unknown")),
        ("Data sources",     (
            "Interval avg"
            + (", interval max/min" if fs.get("has_maxmin") else "")
            + (", adaptive events" if fs.get("has_adaptive") else "")
        )),
    ]
    # An incomplete source file belongs in the report header, not only in a log:
    # the findings below are drawn from partial data and a reader has to know.
    # The header carries the summary; Appendix B carries the evidence.
    _dq = fs.get("data_quality") or {}
    if _dq.get("missing_bytes") or _dq.get("unreadable_observations"):
        rows_data.append((
            "Source file integrity",
            _integrity_summary(_dq, fs),
        ))
    info_tbl = doc.add_table(rows=len(rows_data), cols=2)
    info_tbl.style = 'Table Grid'
    _set_col_widths(info_tbl, [5.0, 11.5])
    for i, (label, value) in enumerate(rows_data):
        cell_l, cell_r = info_tbl.rows[i].cells
        _cell_shade(cell_l, _CHROME_LABEL)
        cell_l.paragraphs[0].add_run(label).bold = True
        cell_r.paragraphs[0].add_run(value)
    doc.add_paragraph()


#: The standard cell for the steady-state voltage row. It names the basis as
#: well as the range: an engineer checking the row against the meter's own
#: max/min display has to be able to see, from the row, that the two are
#: answering different questions.
_C841_STANDARD = ("Steady-state voltage within ANSI C84.1-2016 Range A "
                  "(service voltage, interval averages)")


def _c841_row_title(res: dict) -> str:
    """Row title, naming the Table 1 group where it is not the familiar one.

    The two groups share their upper limits and differ only below nominal, so a
    primary service judged at 97.5% next to a row labelled the same way as every
    secondary report reads as an error in the tool rather than as the standard.
    """
    if res.get("nominal_group") == "over_600v":
        return (_C841_STANDARD[:-1] + ", 2.4–34.5 kV group: "
                "Range A 97.5–105% of nominal)")
    return _C841_STANDARD

#: Compliance-column wording for each voltage band. Range A is the line the
#: standard asks the utility to hold, so leaving it is never rendered as a pass;
#: "Range B" says which side of it the service landed on without borrowing the
#: word Exceeded, which the reader would take to mean outside C84.1 altogether.
_VOLTAGE_BAND_VERDICT = {
    "range_a":   "Within",
    "range_b":   "Range B",
    "outside_a": "Exceeded",
    "outside_b": "Exceeded",
}


def _grade_voltage_band(res: dict, missing: Optional[List[str]] = None) -> dict:
    """Severity for a band-classified voltage result.

    C84.1 does not present its two ranges as degrees of one limit, so this does
    not grade on a margin against one. Range B is a band the supply is permitted
    to enter, on the condition that the excursions are "limited in extent,
    frequency and duration" and are corrected within a reasonable time. That
    condition is the grade: entering Range B briefly is a minor finding, sitting
    in it is the thing the standard asks to be fixed, and leaving Range B is
    outside what C84.1 contemplates at all.
    """
    notes = ([f"no usable data on {', '.join(_phase_label(m) for m in missing)}"]
             if missing else [])
    band = res.get("band", "range_a")
    if band == "range_a":
        return grade_finding(True, confidence_notes=notes or None)

    pct_b   = res.get("total_pct_range_b", 0.0) or 0.0
    pct_out = res.get("total_pct_outside_b", 0.0) or 0.0
    sustained = SEVERITY_SIGNIFICANT_PERSISTENCE

    if band == "outside_b":
        sev_band = "severe" if pct_out >= sustained else "significant"
        reason = (f"{_pct(pct_out, '.1f')} of the recording outside Range B "
                  f"({res['range_b_v'][0]:.1f}–{res['range_b_v'][1]:.1f} V)")
    elif band == "range_b":
        sev_band = "significant" if pct_b >= sustained else "minor"
        reason = (f"{_pct(pct_b, '.1f')} of the recording in Range B; C84.1 asks "
                  "that Range B excursions be limited in extent, frequency and "
                  "duration and corrected within a reasonable time")
    else:
        # Range B was not evaluated, so how far past Range A this went is not
        # something the tool can say. Grading it on persistence alone is the
        # honest ceiling, and the reason says why it stops there.
        sev_band = "significant" if (res.get("total_pct_out_of_bounds") or 0.0) >= sustained else "minor"
        reason = (f"{_pct(res.get('total_pct_out_of_bounds') or 0.0, '.1f')} of "
                  "the recording outside Range A; Range B not evaluated at this "
                  "nominal, so the depth past Range A is not graded")

    downgraded = False
    if notes and sev_band != "minor":
        sev_band = SEVERITY_ORDER[SEVERITY_ORDER.index(sev_band) - 1]
        downgraded = True
        reason += "; " + "; ".join(notes) + " (severity reduced one band)"
    elif notes:
        reason += "; " + "; ".join(notes)

    return {"band": sev_band, "label": SEVERITY_LABEL[sev_band],
            "reason": reason, "margin": None, "downgraded": downgraded}


def _voltage_band_cell(st: dict, res: dict, lead: str) -> str:
    """The Measured cell for one band-classified voltage result.

    Shared by the line-to-neutral and line-to-line rows so the two cannot drift
    into describing the same classification differently.

    The readings quoted are interval averages, because that is what the verdict
    rests on. The meter's within-interval extremes follow on their own clause,
    named as such: a cell that opened with 258.6 V and closed with "all
    intervals in Range A" would read as a contradiction, and the reader would be
    right to distrust it -- the two numbers answer different questions.
    """
    rng  = res["range_v"]
    band = st["band"]

    if band == "range_a":
        band_note = "All intervals in Range A"
    else:
        # Quote the edge that was actually crossed, and split the share against
        # that same edge. An interval outside Range B described as "below
        # 263.1 V" invites the reader to check it against the Range A limit and
        # conclude the classification is one band out; a low/high split taken
        # from Range A under an outside-B headline does the same thing to the
        # arithmetic, since the two shares do not add up to the one quoted.
        outside_b = band == "outside_b" and res.get("range_b_v")
        edge = res["range_b_v"] if outside_b else rng
        lo_pct = st["pct_under_b"] if outside_b else st["pct_under"]
        hi_pct = st["pct_over_b"]  if outside_b else st["pct_over"]
        if lo_pct and hi_pct:
            direction = (f"{_mp(lo_pct, '.1f')} low / {_mp(hi_pct, '.1f')} high")
        elif lo_pct:
            direction = f"all below {edge[0]:.1f} V"
        else:
            direction = f"all above {edge[1]:.1f} V"
        if band == "outside_b":
            share = _mp(st["pct_outside_b"], '.1f')
            band_note = f"{share} of intervals outside Range B, {direction}"
        elif band == "range_b":
            share = _mp(st["pct_range_b"], '.1f')
            band_note = f"{share} of intervals in Range B, {direction}"
        else:
            share = _mp(st["pct_out_of_bounds"], '.1f')
            band_note = f"{share} of intervals outside Range A, {direction}"

    limits = f"Range A {rng[0]:.1f}–{rng[1]:.1f} V"
    if res.get("range_b_v"):
        limits += f"  ·  Range B {res['range_b_v'][0]:.1f}–{res['range_b_v'][1]:.1f} V"

    cell = (f"{lead}{_m(st['min_v'], '.1f')}–{_m(st['max_v'], '.1f', ' V')} "
            f"(mean {_m(st['mean_v'], '.1f')})  |  {band_note}  |  {limits}")

    # Only worth saying where the extremes actually left the band the averages
    # stayed inside -- that is the case a reader would otherwise mistake for a
    # missed violation, and it is the one pointing at the events section.
    if st.get("used_interval_extremes"):
        lo, hi = st.get("min_interval_v"), st.get("max_interval_v")
        if lo is not None and hi is not None and (lo < rng[0] or hi > rng[1]):
            cell += (f"  |  Within-interval extremes {_m(lo, '.1f')}–"
                     f"{_m(hi, '.1f', ' V')}, graded as events against ITIC, "
                     "not against C84.1")
    return cell


def _word_compliance_table(doc, report, thresh, df) -> None:
    pf   = report["pass_fail"]
    volt = report["voltage_compliance"]
    thd  = report["thd_compliance"]
    pfr  = report["power_factor"]
    imb  = report["voltage_imbalance"]
    ci   = report["current_imbalance"]
    dem  = report["demand"]
    ih   = report["individual_harmonics"]
    ivh  = report.get("individual_voltage_harmonics", {})
    hs   = report.get("harmonic_statistics", {})
    tdd_info = thd.get("tdd_info", {})
    sev = compute_severities(report, thresh)

    _bold(doc.add_paragraph(), "Compliance Status by Standard", size_pt=11)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    # Measured carries the most text now, so it gets the room; the total stays
    # inside the 16.5 cm the previous three-column layout used.
    _set_col_widths(tbl, [5.6, 6.0, 1.9, 3.0])

    # Header row
    hdr_cells = tbl.rows[0].cells
    for cell, text in zip(hdr_cells,
                          ["Standard", "Measured", "Compliance", "Severity"]):
        _cell_shade(cell, _CHROME_HDR)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    # Rows are buffered rather than written straight out, so they can be emitted
    # grouped by the quantity being measured.  Interleaving volts, amps, kVA and
    # hertz made the table read as a list of unrelated tests; grouped, a reader
    # can take in "the voltage picture" or "the current picture" at once.
    buffered: List[tuple] = []

    def add_row(standard, measured, passes, severity=None, group="other",
                verdict=None):
        """One finding: the compliance fact, then how much it matters.

        Callers that pass no severity fall back to grading on pass/fail alone,
        which yields Compliant / Minor / Not assessed — never a red row without
        a margin to justify it.

        ``verdict`` overrides the Compliance word for the findings whose
        standard is not binary — C84.1 has two named ranges, and collapsing
        Range B into "Exceeded" would lose the distinction the standard draws.
        """
        buffered.append((group, standard, measured, passes, severity, verdict))

    def _emit(standard, measured, passes, severity, verdict=None):
        row   = tbl.add_row()
        cells = row.cells
        cells[0].paragraphs[0].add_run(standard).font.size = Pt(10)
        # The Measured column is the case the marking exists for: one cell puts
        # a reading and the limit it is judged against in the same breath --
        # "P95 6.80% (limit 8.00%)" -- and only the first came off the meter.
        _emit_text(cells[1].paragraphs[0], measured, size_pt=10)

        # Compliance column — the standard's own binary verdict, unshaded so it
        # reads as a fact rather than an alarm.
        r = cells[2].paragraphs[0].add_run(verdict or _pf_sym(passes))
        r.bold = True
        r.font.size = Pt(10)
        if passes is False:
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif passes is True:
            r.font.color.rgb = _PASS_CLR

        # Severity column — carries the colour, because this is the axis that
        # actually says whether to worry.
        s = severity or grade_finding(passes)
        band = s["band"]
        _cell_shade(cells[3], _sev_shade(band))
        p3 = cells[3].paragraphs[0]
        r3 = p3.add_run(s["label"])
        r3.bold = True
        r3.font.size = Pt(10)
        clr = _sev_color(band)
        if clr:
            r3.font.color.rgb = clr
        if s.get("reason"):
            # Tables are not marked -- a column of readings is already read as
            # readings -- so the severity reason drops its markers here.
            note = p3.add_run("\n" + strip_marks(s["reason"]))
            note.font.size = Pt(7.5)
            note.italic = True
            note.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def _emit_group_heading(text):
        row = tbl.add_row()
        cell = row.cells[0].merge(row.cells[-1])
        _cell_shade(cell, _CHROME_BAND)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Demand / transformer loading. On a shared transformer the row states
    # what was measured -- this service's contribution -- rather than claiming
    # the transformer was inside its nameplate, which one meter cannot show.
    if "transformer" in dem:
        tx   = dem["transformer"]
        meas = (f"{_m(tx['peak_8h_kva'], '.0f', ' kVA')} 8-hr peak  /  "
                f"{tx['nameplate_kva']:.0f} kVA nameplate  "
                f"({_m(tx['pct_nameplate'], '.0f', '%')})")
        if tx["overloaded"] is None:
            add_row("This service's demand against transformer nameplate "
                    "(8-hr peak; shared transformer, total loading not measured)",
                    meas + "  |  this service only — other customers on this "
                           "transformer were not measured",
                    None, group="power")
        else:
            add_row("Steady-state demand ≤ transformer nameplate (8-hr peak)", meas,
                    not tx["overloaded"], sev.get("demand"), group="power")
    else:
        add_row("Steady-state demand ≤ transformer nameplate (8-hr peak)", "No nameplate provided", None, group="power")

    # Power factor
    if pfr["available"]:
        if thresh.customer_class == "r":
            meas = (f"Min {_m(pfr['min_pf'], '.3f')}  /  "
                    f"Mean {_m(pfr['mean_pf'], '.3f')}  "
                    f"(residential — tariff PF clause not applicable)")
            add_row("Power factor ≥ 0.90 lagging (Xcel tariff)", meas, None, group="power")
        else:
            meas = (f"Min {_m(pfr['min_pf'], '.3f')}  /  "
                    f"Mean {_m(pfr['mean_pf'], '.3f')}  "
                    f"(limit ≥ {pfr['limit']:.2f})")
            # Graded on the mean, not the minimum: the tariff clause is about
            # sustained operation, and a single low interval is not a billing
            # condition.  The minimum still shows in the Measured column.
            add_row("Power factor ≥ 0.90 lagging (Xcel tariff)", meas,
                    pf["power_factor"], sev.get("power_factor"), group="power")
    else:
        add_row("Power factor ≥ 0.90 lagging (Xcel tariff)", "No data", None, group="power")

    # Voltage compliance
    if volt["available"]:
        phases = volt["phases"]
        rng = volt["range_v"]
        missing = volt.get("phases_missing_data") or []
        missing_note = (
            f"  |  No usable data: {', '.join(_phase_label(m) for m in missing)}"
            if missing else ""
        )
        # Name the worst phase and quote what it actually measured.  The old row
        # printed range_v -- the *allowed* band -- under the bare label "Range",
        # so the only voltage in the cell was the limit dressed as a measurement,
        # and "worst phase" was never followed by which phase or what it read.
        # Ranked on the band reached first: a phase that leaves Range B for one
        # interval is a worse finding than one sitting in Range B for a tenth of
        # the recording, and sorting on share alone would print the wrong phase.
        worst_col, worst_st = max(
            phases.items(),
            key=lambda kv: (VOLTAGE_BAND_ORDER.index(kv[1]["band"]),
                            kv[1]["pct_out_of_bounds"]))
        meas = _voltage_band_cell(
            worst_st, volt, f"Worst phase {_phase_label(worst_col)}: ") + missing_note
        if volt.get("range_b_note"):
            meas += f"  |  {volt['range_b_note']}"
        # Voltage severity is driven by how much of the recording sat out of
        # band; the excursion depth now shows in the Measured column.
        add_row(_c841_row_title(volt), meas, pf["voltage"], sev.get("voltage"),
                group="voltage",
                verdict=_VOLTAGE_BAND_VERDICT.get(worst_st["band"]))
    else:
        add_row(_C841_STANDARD, volt.get("error", "No data"), None, group="voltage")

    # Voltage transients / ITIC
    itic = report.get("itic", {})
    if itic.get("available"):
        if itic["n_events"] == 0:
            it_meas = "No voltage sag/swell events detected"
        else:
            it_meas = (f"{_m(itic['n_events'])} sag/swell event(s) evaluated; "
                       f"{_m(itic['n_violations'])} outside the ITIC envelope")
            # A count of events says nothing about whether equipment would have
            # ridden them through.  Depth and duration are what the ITIC curve
            # is actually plotted against.
            w = itic.get("worst")
            if w:
                it_meas += (f"  |  Worst: {_m(w['value_v'], '.1f', ' V')} "
                            f"({_m(w['pct_nominal'], '.0f', '%')} of nominal) "
                            f"for {_m(w['duration_ms'], '.0f', ' ms')}")
                if w.get("phase"):
                    it_meas += f" on {_phase_label(w['phase'])}"
        add_row("Voltage sags/swells within ITIC voltage tolerance curve",
                it_meas, pf.get("itic_transients"), group="voltage")
    else:
        add_row("Voltage sags/swells within ITIC voltage tolerance curve",
                itic.get("note", "Event duration data not available"), None, group="voltage")

    # Voltage THD
    v_thd = thd["voltage"]
    if v_thd["available"]:
        # Lead with the percentile the standard actually judges on; the maximum
        # follows as context, flagged when it is a spike rather than a level.
        meas = (f"P95 {_m(v_thd.get('p95_thd_pct', v_thd['max_thd_pct']), '.2f', '%')}  "
                f"(limit {v_thd['limit_pct']:.1f}%)  /  "
                f"Mean {_m(v_thd['mean_thd_pct'], '.2f', '%')}  /  "
                f"Max {_m(v_thd['max_thd_pct'], '.2f', '%')}")
        if v_thd.get("max_is_outlier"):
            meas += " (isolated spike)"
        add_row("Voltage THD < 8.0% (IEEE 519-2022 P95, secondary)", meas,
                pf["thd_voltage"], sev.get("thd_voltage"), group="voltage")
    else:
        add_row("Voltage THD < 8.0% (IEEE 519-2022, secondary)", "No voltage THD channel", None, group="voltage")

    # Current TDD / THD
    c_thd = thd["current"]
    if c_thd["available"]:
        metric = "TDD" if tdd_info else "THD"
        lim    = c_thd["limit_pct"]
        if tdd_info and tdd_info.get("isc_provided"):
            cls = f"  [ISC/IL={tdd_info['isc_il_ratio']:.0f}, class {tdd_info['tdd_class']}]"
        elif tdd_info:
            cls = "  [most restrictive class assumed — conservative]"
        else:
            cls = ""
        meas   = (f"Max {_m(c_thd['max_thd_pct'], '.2f', '%')}  /  "
                  f"Mean {_m(c_thd['mean_thd_pct'], '.2f', '%')}  "
                  f"(limit {lim:.1f}%{cls})")
        add_row(f"Current {metric} within IEEE 519-2022 Table 2", meas,
                pf["thd_current"], sev.get("thd_current"), group="current")
    else:
        add_row("Current TDD within IEEE 519-2022 Table 2", "No current THD channel", None, group="current")

    # Individual current harmonics
    if ih.get("available"):
        # Report the order with the tightest margin, not the largest current.
        # The Table 2 limits fall steeply with h, so the biggest harmonic is
        # often comfortably legal while a small high order is the binding one.
        mo   = ih.get("worst_margin_order")
        mpct = ih.get("worst_margin_pct_of_il")
        mlim = ih.get("worst_limit_pct")
        # Spell out IL the first time it appears -- "% of IL" is IEEE shorthand
        # for percent of maximum demand load current and is not self-evident to
        # everyone the report goes to.
        il_note = (f", IL = {_m(ih['il_amps'], '.0f', ' A')} max demand current"
                   if ih.get("il_amps") else " (IL = max demand current)")
        if mo and mpct is not None and mlim:
            tight = (f"tightest: H{mo[0]} phase {mo[1].upper()} at "
                     f"{_m(mpct, '.2f', '%')} of IL "
                     f"vs {mlim:.1f}% limit ({_m(mpct / mlim * 100, '.0f', '%')} of it)")
        else:
            wo = ih.get("worst_order")
            tight = (f"worst: H{wo[0]} at {_m(ih['worst_pct_of_il'], '.2f', '%')} of IL"
                     if wo else "")
        head = ("All current harmonic orders within limits" if ih["overall_pass"]
                else "One or more current harmonic orders exceed limit")
        meas = (f"{head}  ({tight}{il_note})" if tight
                else f"{head}{il_note}")
        add_row("Individual current harmonics within IEEE 519-2022 Table 2", meas,
                pf["individual_harmonics"], sev.get("individual_harmonics"), group="current")
    else:
        note = ih.get("note", "Pass --isc to enable per-order check")
        add_row("Individual current harmonics within IEEE 519-2022 Table 2", note, None, group="current")

    # Individual voltage harmonics
    if ivh.get("available"):
        vwo = ivh.get("worst_order")
        # Mirror the current row's wording so the two are impossible to confuse
        # at a glance: name the quantity, then what the percentage is of.
        vhead = ("All voltage harmonic orders within 5% limit"
                 if ivh["overall_pass"]
                 else "One or more voltage harmonic orders exceed 5% limit")
        meas = (f"{vhead}  (worst: H{vwo[0]} phase {vwo[1].upper()} at "
                f"{_m(ivh['worst_pct_nom'], '.2f', '%')} of nominal voltage)"
                if vwo else vhead)
        add_row("Individual voltage harmonics within IEEE 519-2022 Table 1 (5% of nominal)", meas,
                pf.get("individual_voltage_harmonics"),
                sev.get("individual_voltage_harmonics"), group="voltage")
    else:
        add_row("Individual voltage harmonics within IEEE 519-2022 Table 1 (5% of nominal)",
                ivh.get("note", "Per-order voltage harmonics not available in this meter format"), None, group="voltage")

    # Statistical compliance (IEEE 519-2022 Clause 5)
    if hs.get("available"):
        period_note = f"{_m(hs['period_days'], '.1f')}-day recording"
        if hs["overall_pass"]:
            hs_meas = f"P95 ≤ 1.0× and P99 ≤ 1.5× limits for all orders ({period_note})"
        else:
            hs_meas = f"One or more orders exceed P95 or P99 statistical limits ({period_note})"
        # Name the order that binds and quote it.  "Within limits for all orders"
        # gives no sense of how close the service came, and the tightest margin
        # is what a follow-up recording should be aimed at.
        b = hs.get("binding")
        if b:
            label = "TDD" if b["order"] == "thd" else f"H{b['order']}"
            hs_meas += (f"  |  Tightest: {label} phase {b['phase'].upper()} "
                        f"P95 {_m(b['p95'], '.2f', '%')} vs {b['limit']:.2f}% "
                        f"limit ({_m(b['ratio'] * 100, '.0f', '%')} of it)")
        # Already a percentile test, so persistence is baked in; grade on the
        # verdict and note when the recording was too short to be a true week.
        add_row(
            "Current harmonic P95 / P99 within IEEE 519-2022 Clause 5 statistical limits",
            hs_meas, pf.get("harmonic_statistics"),
            sev.get("harmonic_statistics"), group="current",
        )
    else:
        add_row(
            "Current harmonic P95 / P99 within IEEE 519-2022 Clause 5 statistical limits",
            hs.get("note", "Pass --isc to enable statistical check"), None,
            group="current",
        )

    # Voltage imbalance — the metric itself differs by service configuration,
    # so the row must not claim NEMA MG1 on a single-phase service where that
    # definition does not apply.
    if imb["available"]:
        meas = (f"Max {_m(imb['max_imbalance_pct'], '.2f', '%')}  /  "
                f"Mean {_m(imb['mean_imbalance_pct'], '.2f', '%')}  "
                f"(limit {imb['limit_pct']:.1f}%)")
        if imb.get("metric") == "nema_mg1":
            imb_label = "Voltage imbalance < 3% (ANSI C84.1 / NEMA MG1)"
        else:
            imb_label = (f"Leg-to-leg voltage difference < "
                         f"{imb['limit_pct']:.0f}% of nominal (single-phase service)")
            meas += "  |  NEMA MG1 unbalance is a three-phase definition"
            if imb.get("note"):
                meas += "; only two of three phases measured"
        add_row(imb_label, meas,
                pf["voltage_imbalance"], sev.get("voltage_imbalance"), group="voltage")
    else:
        add_row("Voltage imbalance < 3% (ANSI C84.1 / NEMA MG1)", "No data", None, group="voltage")

    # Current imbalance — a limit on three-phase service, a reported
    # measurement on two legs, where no PSCo or standard limit exists.
    if thresh.customer_class in ("c", "sg", "pg"):
        # Sheet R123 is a billing-demand provision, not a limit: above 15%
        # between phases the Company "may take as the Billing Demand" the
        # three-phase equivalent of the worst phase. Calling it a limit told a
        # customer they were in breach of something they were not.
        ci_label = ("Current imbalance < 10% (NEMA MG1); above 15% between "
                    "phases PSCo Sheet R123 allows billing demand to be taken "
                    "from the worst phase")
    else:
        ci_label = "Current imbalance < 10% (NEMA MG1)"
    if ci["available"] and ci.get("limit_pct") is None:
        add_row("Leg current difference (reported, no limit applies)",
                f"Max {_m(ci['max_imbalance_pct'], '.2f', '%')}  /  "
                f"Mean {_m(ci['mean_imbalance_pct'], '.2f', '%')}  |  "
                f"measurement, not a violation — no limit is set for a "
                f"two-leg service",
                None, group="current")
    elif ci["available"]:
        meas = (f"Max {_m(ci['max_imbalance_pct'], '.2f', '%')}  /  "
                f"Mean {_m(ci['mean_imbalance_pct'], '.2f', '%')}  "
                f"(limit {ci['limit_pct']:.1f}%)")
        add_row(ci_label, meas, pf["current_imbalance"],
                sev.get("current_imbalance"), group="current")
    else:
        add_row(ci_label, "No data", None, group="current")

    # Flicker — every phase the meter recorded, not just phase A. The row makes
    # an unqualified statement about the service, so it has to be based on the
    # worst phase or it can read as a pass while another phase is over limit.
    fl = report.get("flicker") or {"available": False}
    if fl.get("available"):
        pst_max = fl.get("pst_max")
        plt_max = fl.get("plt_max")
        phases = ", ".join(fl.get("phases_read", []))
        worst = fl.get("worst_phase")
        measured = (
            f"Pst max {_m(pst_max, '.2f')} (limit {fl['pst_limit']:.2f})"
            if pst_max is not None else "Pst n/a"
        ) + "  /  " + (
            f"Plt max {_m(plt_max, '.2f')} (limit {fl['plt_limit']:.2f})"
            if plt_max is not None else "Plt n/a"
        ) + f"  — worst phase {worst} of {phases}"
        add_row(
            f"Flicker within IEC 61000-3-3 limits, all phases "
            f"(Pst ≤ {fl['pst_limit']:.1f}, Plt ≤ {fl['plt_limit']:.2f})",
            measured,
            fl.get("overall_pass"),
            sev.get("flicker"), group="voltage",
        )
    else:
        add_row("Flicker within IEC 61000-3-3 limits (Pst ≤ 1.0, Plt ≤ 0.65)",
                "Not measured in this recording", None, group="voltage")

    # Line-to-line voltage and frequency were being counted in the standards
    # tally without appearing here, so the summary claimed more standards than
    # the table showed.
    llv = report.get("voltage_ll_compliance") or {}
    if llv.get("available"):
        worst_pair, worst = max(
            llv["pairs"].items(),
            key=lambda kv: (VOLTAGE_BAND_ORDER.index(kv[1]["band"]),
                            kv[1]["pct_out_of_bounds"]))
        meas = _voltage_band_cell(worst, llv, f"Worst pair {worst_pair}: ")
        if llv.get("range_b_note"):
            meas += f"  |  {llv['range_b_note']}"
        # Where the engineer stated the nominal, say so: on a primary service
        # the band is built from what they entered, and a reader checking the
        # limit needs to know it came from the form rather than from the file.
        nominal_note = ("" if llv.get("nominal_source") != "entered"
                        else ", entered")
        add_row(f"Line-to-line voltage within ANSI C84.1 Range A "
                f"({llv['nominal_v']:.0f} V nominal{nominal_note})",
                meas, llv.get("overall_pass"),
                sev.get("voltage_line_to_line"), group="voltage",
                verdict=_VOLTAGE_BAND_VERDICT.get(worst["band"]))
    else:
        add_row("Line-to-line voltage within ANSI C84.1 Range A",
                llv.get("error", "Not evaluated"), None, group="voltage")

    frq = report.get("frequency") or {}
    if frq.get("available"):
        meas = (f"{_m(frq['min_hz'], '.3f')}–{_m(frq['max_hz'], '.3f', ' Hz')}  "
                f"(allowed {frq['range_hz'][0]:.2f}–{frq['range_hz'][1]:.2f} Hz)")
        add_row(f"System frequency within ±{(frq['range_hz'][1] - frq['nominal_hz']):.1f} Hz "
                f"of {frq['nominal_hz']:.0f} Hz", meas, frq.get("overall_pass"),
                sev.get("frequency"), group="frequency")
    else:
        add_row("System frequency within tolerance", "No frequency channel", None, group="frequency")

    # ── Emit, grouped by the quantity measured ────────────────────────────────
    # Voltage first: it is what the customer experiences and what most calls are
    # about.  Current second, since distortion and imbalance there are usually
    # the cause of what the voltage rows show.  Power and frequency last.
    # Within a group, magnitude checks come before distortion checks: a reader
    # wants "is the voltage right" answered before "is its shape right".  Code
    # order does not give that, so rank explicitly on a distinctive fragment of
    # each label.  Anything unmatched keeps insertion order, after the ranked.
    order_within = {
        "voltage": ["Steady-state voltage", "Line-to-line voltage",
                    "Voltage imbalance", "Voltage sags/swells", "Voltage THD",
                    "Individual voltage harmonics", "Flicker"],
        "current": ["Current imbalance", "Current TDD", "Current THD",
                    "Individual current harmonics", "Current harmonic P95"],
        "power":   ["Steady-state demand", "Power factor"],
    }

    def _rank(group_key, standard):
        for i, frag in enumerate(order_within.get(group_key, [])):
            if standard.startswith(frag):
                return i
        return len(order_within.get(group_key, []))

    for key, heading in (
        ("voltage",   "VOLTAGE  —  volts, and quantities derived from voltage"),
        ("current",   "CURRENT  —  amps, and quantities derived from current"),
        ("power",     "DEMAND AND POWER FACTOR  —  kVA, power factor"),
        ("frequency", "FREQUENCY  —  hertz"),
        ("other",     "OTHER"),
    ):
        rows = [r for r in buffered if r[0] == key]
        if not rows:
            continue
        rows.sort(key=lambda r: _rank(key, r[1]))
        _emit_group_heading(heading)
        for _, standard, measured, passes, severity, verdict in rows:
            _emit(standard, measured, passes, severity, verdict)

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Narrative helpers — executive summary, key findings, structured actions
# ─────────────────────────────────────────────────────────────────────────────

_PF_FRIENDLY = {
    "transformer_loading":          "transformer loading",
    "voltage":                      "steady-state voltage (ANSI C84.1)",
    "thd_voltage":                  "voltage THD (IEEE 519)",
    "thd_current":                  "current TDD (IEEE 519)",
    "individual_harmonics":         "individual harmonic currents (IEEE 519)",
    "individual_voltage_harmonics": "individual harmonic voltages (IEEE 519)",
    "power_factor":                 "power factor (PSCo tariff)",
    "voltage_imbalance":            "voltage imbalance",
    "current_imbalance":            "current imbalance",
    "harmonic_statistics":          "current harmonic statistical limits (IEEE 519 Clause 5)",
    "neutral_health":               "neutral integrity",
    "itic_transients":              "voltage sags/swells (ITIC curve)",
    "voltage_line_to_line":         "line-to-line voltage (ANSI C84.1)",
    "frequency":                    "system frequency",
    "flicker":                      "voltage flicker (IEC 61000-3-3)",
}

_PRIORITY_FROM_SEV = {"critical": "High", "warning": "Medium", "info": "Low"}
# Actions are no longer grouped by who owns them: the tool reports evidence and
# the reviewing engineer writes the attribution.


def _phase_label(ph: str) -> str:
    """Customer-facing phase label from an internal phase/channel key."""
    return ph.replace("voltage_", "").replace("current_", "").strip("_").upper() or ph.upper()


def _flicker_status(report: dict) -> Optional[dict]:
    """Worst-phase flicker severity and pass/fail, or None when not measured.

    Read off the analysis result rather than off phase A's columns. On a real
    split-phase file phase A reached Pst 1.43 while phase B reached 4.98, so a
    status taken from phase A alone understated the service by 3.5x and could
    report a pass with another phase over limit.
    """
    fl = report.get("flicker") or {}
    if not fl.get("available"):
        return None
    return {
        "pst_max":   fl.get("pst_max"),
        "plt_max":   fl.get("plt_max"),
        "pst_limit": fl.get("pst_limit"),
        "plt_limit": fl.get("plt_limit"),
        "pst_phase": fl.get("pst_worst_phase"),
        "plt_phase": fl.get("plt_worst_phase"),
        "passes":    fl.get("overall_pass"),
    }


def _collect_key_findings(report: dict, thresh: Thresholds, df) -> List[str]:
    """Build the Key Findings list: compliance failures first, then significant
    warnings/observations, each as 1–2 customer-facing sentences."""
    pf    = report["pass_fail"]
    thd   = report["thd_compliance"]
    items: List[Tuple[int, str]] = []          # (severity rank, text)

    c_thd = thd["current"]
    if pf.get("thd_current") is False and c_thd.get("available"):
        tdd_info = thd.get("tdd_info", {})
        metric   = "TDD" if tdd_info else "THD"
        assumed  = ("" if not tdd_info or tdd_info.get("isc_provided")
                    else " (most restrictive class limit assumed)")
        items.append((0,
            f"Current {metric} exceeded the IEEE 519-2022 limit of {c_thd['limit_pct']:.1f}%"
            f"{assumed} during {_pct(c_thd['pct_exceeding'])} of the recording "
            f"(maximum {c_thd['max_thd_pct']:.2f}%)."))

    v_thd = thd["voltage"]
    if pf.get("thd_voltage") is False and v_thd.get("available"):
        # State the percentile that failed, since that is the criterion, and say
        # how much of the recording was actually above the limit.
        if not v_thd.get("p95_pass", True):
            basis = (f"the 95th-percentile voltage THD was "
                     f"{v_thd.get('p95_thd_pct', 0):.2f}%, above the "
                     f"{v_thd['limit_pct']:.1f}% IEEE 519-2022 limit")
        else:
            basis = (f"the 99th-percentile voltage THD was "
                     f"{v_thd.get('p99_thd_pct', 0):.2f}%, above the "
                     f"{v_thd.get('p99_limit_pct', 0):.1f}% short-time limit "
                     f"(1.5 x the {v_thd['limit_pct']:.1f}% limit)")
        items.append((0,
            f"Voltage distortion exceeded IEEE 519-2022: {basis}, sustained "
            f"across {_pct(v_thd['pct_exceeding'])} of the recording "
            f"(maximum {v_thd['max_thd_pct']:.2f}%)."))

    ih = report["individual_harmonics"]
    if pf.get("individual_harmonics") is False and ih.get("worst_order"):
        h, ph = ih["worst_order"]
        items.append((0,
            f"Individual harmonic currents exceeded IEEE 519-2022 per-order limits; "
            f"the worst order was H{h} at {ih['worst_pct_of_il']:.2f}% of IL "
            f"(phase {ph.upper()})."))

    ivh = report.get("individual_voltage_harmonics", {})
    if pf.get("individual_voltage_harmonics") is False and ivh.get("worst_order"):
        vh = ivh["worst_order"]
        items.append((0,
            f"The H{vh[0]} voltage harmonic exceeded the recommended limit of 5% of "
            f"nominal voltage ({ivh['worst_pct_nom']:.2f}% measured)."))

    if pf.get("harmonic_statistics") is False:
        items.append((0,
            "One or more current harmonic orders exceeded the IEEE 519-2022 "
            "statistical (95th/99th percentile) limits over the recording period."))

    volt = report["voltage_compliance"]
    if pf.get("voltage") is False and volt.get("available"):
        worst_ph, worst = max(volt["phases"].items(),
                              key=lambda kv: kv[1]["pct_out_of_bounds"])
        ph_label = _phase_label(worst_ph)
        items.append((0,
            f"Steady-state voltage was outside ANSI C84.1 Range A during "
            f"{_pct(worst['pct_out_of_bounds'], '.2f')} of intervals on the worst phase "
            f"(phase {ph_label}: {worst['min_v']:.1f}–{worst['max_v']:.1f} V)."))

    if pf.get("transformer_loading") is False:
        tx = report["demand"]["transformer"]
        items.append((0,
            f"The serving transformer is overloaded: the 8-hour peak demand of "
            f"{tx['peak_8h_kva']:.0f} kVA is {tx['pct_nameplate']:.0f}% of the "
            f"{tx['nameplate_kva']:.0f} kVA nameplate."))

    pfr = report["power_factor"]
    if (pf.get("power_factor") is False and thresh.customer_class != "r"
            and pfr.get("available")):
        items.append((0,
            f"Power factor fell below the {pfr['limit']:.2f} tariff requirement during "
            f"{_pct(pfr['pct_below_limit'])} of the recording "
            f"(minimum {pfr['min_pf']:.3f})."))

    imb = report["voltage_imbalance"]
    if pf.get("voltage_imbalance") is False and imb.get("available"):
        items.append((0,
            f"Voltage imbalance exceeded the {imb['limit_pct']:.1f}% limit during "
            f"{_pct(imb['pct_exceeding'])} of the recording "
            f"(maximum {imb['max_imbalance_pct']:.2f}%)."))

    ci = report["current_imbalance"]
    if pf.get("current_imbalance") is False and ci.get("available"):
        items.append((0,
            f"Current imbalance exceeded the {ci['limit_pct']:.1f}% limit during "
            f"{_pct(ci['pct_exceeding'])} of the recording "
            f"(maximum {ci['max_imbalance_pct']:.2f}%)."))

    fl = _flicker_status(report)
    if fl and fl["passes"] is False:
        items.append((0,
            "Voltage flicker exceeded IEC 61000-3-3 limits "
            f"(Pst maximum {fl['pst_max']:.2f} on phase {fl['pst_phase']} "
            f"against a {fl['pst_limit']:.2f} limit; Plt maximum "
            f"{fl['plt_max']:.2f} on phase {fl['plt_phase']} against a "
            f"{fl['plt_limit']:.2f} limit)."))

    itic = report.get("itic", {})
    if pf.get("itic_transients") is False and itic.get("worst"):
        w = itic["worst"]
        w_kind = "sag to" if w["type"] == "voltage_sag" else "swell to"
        items.append((0,
            f"{itic['n_violations']} voltage event(s) fell outside the ITIC voltage "
            f"tolerance envelope (worst: {w_kind} {w['pct_nominal']:.0f}% of nominal "
            f"for {w['duration_ms']:.0f} ms) — sensitive electronic equipment may "
            f"misoperate or reset during these events."))

    nhh = report.get("neutral_health", {})
    if nhh.get("available") and nhh.get("severity") in ("critical", "warning"):
        if nhh["severity"] == "critical":
            items.append((0,
                "Indicators consistent with an open or high-resistance service neutral "
                "were detected. This is a safety concern requiring prompt investigation."))
        else:
            items.append((1,
                "Neutral integrity indicators are outside normal ranges and warrant "
                "investigation."))

    nc = ci.get("neutral_current") if ci.get("available") else None
    if nc and nc["mean_pct_of_phase"] > 15:
        items.append((1,
            f"Neutral current was elevated, averaging {nc['mean_amps']:.1f} A "
            f"({nc['mean_pct_of_phase']:.0f}% of the phase average), consistent with "
            f"load imbalance and/or triplen harmonic currents."))

    sh = report.get("harmonic_sources", {})
    resonant = sh.get("resonant_orders", []) if sh.get("available") else []
    if resonant:
        orders = ", ".join(f"H{h}" for h in sorted(resonant))
        items.append((1,
            f"Harmonic resonance is suspected near {orders} based on impedance and "
            f"voltage–current correlation diagnostics."))

    # Remaining warning-level assessment findings not already represented above
    _covered_kw = ("resonan", "neutral", "imbalance", "voltage", "power factor",
                   "flicker", "overload")
    for f in report.get("root_causes", []):
        if f.get("severity") != "warning":
            continue
        title = f.get("title", "")
        if any(k in title.lower() for k in _covered_kw):
            continue
        first = f.get("finding", "").split(". ")[0].rstrip(".")
        items.append((1, f"{title}. {first}." if first else f"{title}."))

    items.sort(key=lambda t: t[0])
    return [txt for _, txt in items[:8]]


def _purpose_from_title(title: str) -> str:
    words = title.split()
    if words and not (words[0].isupper() or words[0][0].isdigit()):
        title = title[0].lower() + title[1:]
    return f"Address {title}."


#: Most actions any one report should carry.  Past this the section stops being
#: a plan and becomes bulk the reader skims, which is worse than a short list
#: that names the few things actually worth doing.
_MAX_ACTIONS = 6

#: Distinct things an action can ask for.  Two recommendations that both amount
#: to "re-measure with the customer's load disconnected" are one action, however
#: differently they are worded.
_ACTION_INTENTS = (
    ("remeasure_isolated", ("loads disconnected", "with all customer loads")),
    ("neutral_inspect",    ("neutral connection", "tighten all neutral")),
    ("neutral_sizing",     ("neutral conductor sizing",)),
    ("pf_correction",      ("power factor correction",)),
    ("transformer_upgrade", ("upgrade of the overloaded", "transformer upgrade")),
    ("voltage_review",     ("ansi c84.1 range a",)),
    ("harmonic_study",     ("harmonic impedance frequency sweep", "detailed harmonic study")),
    ("harmonic_source",    ("identify the source", "inventory the nonlinear",
                            "mitigate harmonic current sources", "document existing harmonic")),
    ("load_balance",       ("redistribute single-phase", "balanced between the two")),
    ("capacitor_review",   ("feeder capacitor", "capacitor banks")),
)

_PRIORITY_ORDER = {"High": 0, "Medium": 1, "Low": 2}


def _action_intent(recommendation: str) -> Optional[str]:
    """Classify what an action is actually asking for, for de-duplication."""
    low = recommendation.lower()
    for intent, needles in _ACTION_INTENTS:
        if any(n in low for n in needles):
            return intent
    return None


def _prune_actions(actions: List[dict]) -> List[dict]:
    """Collapse duplicates and cap the list.

    Findings arrive from several independent checks, and more than one can
    conclude the same next step -- two separate findings both ended with
    "measure again with all customer loads disconnected", which reached the
    report as two actions. Keeping the higher-priority wording of each distinct
    intent removes the repetition without losing anything.
    """
    actions.sort(key=lambda a: _PRIORITY_ORDER.get(a.get("priority"), 3))

    kept: List[dict] = []
    seen_intents: set = set()
    for a in actions:
        intent = _action_intent(a["recommendation"])
        if intent is not None:
            if intent in seen_intents:
                continue
            seen_intents.add(intent)
        kept.append(a)
        if len(kept) >= _MAX_ACTIONS:
            break
    return kept


def _build_structured_actions(report: dict, thresh: Thresholds) -> List[dict]:
    """Assemble recommended actions as dicts with recommendation, purpose and
    priority (High/Medium/Low).

    Actions are deliberately not grouped by who owns them. The tool reports what
    the measurements support and the reviewing engineer assigns responsibility;
    grouping actions under 'Customer' and 'Utility' headings assigned it by
    implication, and nothing here can commit Xcel Energy to an action.
    """
    pf  = report["pass_fail"]
    rca = report.get("root_causes", [])
    sev_rank = {"critical": 0, "warning": 1, "info": 2}
    actions: List[dict] = []

    cls = thresh.customer_class
    residential = cls == "r"
    small = cls in ("r", "c")

    for f in sorted(rca, key=lambda f: sev_rank.get(f["severity"], 9)):
        if not f.get("recommendation"):
            continue
        # A finding that reports the absence of a result is not something to
        # act on. Without this the actions list carried entries like
        # "Address no recognised load signature."
        if f.get("no_action"):
            continue
        # Informational findings are hypotheses, not conclusions. The load
        # signature match is the main one: promoting its recommendation to a
        # Recommended Action told a homeowner to fit input reactors to VFDs
        # they do not own, on the strength of a shape match. Its advice stays
        # attached to the finding in the Engineering Assessment instead.
        if f.get("severity") == "info":
            continue
        actions.append({
            "recommendation": f["recommendation"],
            "purpose":        _purpose_from_title(f.get("title", "the identified condition")),
            "priority":       _PRIORITY_FROM_SEV.get(f.get("severity"), "Low"),
            "origin_evidence": f.get("origin_evidence"),
        })

    covered = " ".join(a["recommendation"].lower() for a in actions)

    # Compliance-driven actions not already covered by assessment findings
    if (not residential and pf["power_factor"] is False
            and "power factor" not in covered):
        # Both clauses bind a C&I customer at once, and neither is class
        # specific: Sheet R73 sits in the General rules and sets 90% lagging
        # for every customer; Sheet R121 sits in the Commercial and Industrial
        # rules and asks for near unity from all of C, SG and PG alike.
        rec = ("Install power factor correction capacitors to bring power factor "
               "above 0.90 lagging, which PSCo Electric Tariff Sheet R73 states "
               "the Company's rates contemplate. Sheet R121 additionally asks "
               "commercial and industrial customers to maintain a power factor "
               "as near unity as practicable at the point of delivery.")
        actions.append({"recommendation": rec,
                        "purpose":  "Meet the tariff power factor requirement and avoid "
                                    "penalty or discontinuance exposure.",
                        "priority": "High"})

    if not any(k in covered for k in ("harmonic", "vfd", "rectifier")):
        if pf.get("thd_current") is False:
            # Filters and 12-pulse topologies are industrial remedies. At a
            # house the useful step is finding which appliance is responsible.
            if residential:
                rec = ("Identify the source by switching major nonlinear loads "
                       "individually — EV charger, air conditioning or heat pump, "
                       "and any solar inverter — and noting which one changes the "
                       "distortion. Report the result to your Xcel Energy contact; "
                       "harmonic filters are not an appropriate remedy at a "
                       "residential service.")
                purpose = ("Locate the responsible equipment before considering any "
                           "corrective work.")
            elif cls == "c":
                rec = ("Inventory the nonlinear equipment on site (electronic "
                       "lighting drivers, small drives, EV charging, office "
                       "equipment) and identify the dominant contributor before "
                       "considering mitigation. At this service size, replacing or "
                       "relocating a single offending load is usually cheaper and "
                       "more effective than a filter.")
                purpose = ("Bring current distortion within IEEE 519-2022 limits "
                           "proportionately to the size of the service.")
            else:
                rec = ("Investigate and mitigate harmonic current sources "
                       "(VFDs, rectifiers, UPS). Consider passive or active "
                       "harmonic filters, or 12-pulse drive topologies.")
                purpose = "Bring current distortion within IEEE 519-2022 limits."
            actions.append({"recommendation": rec, "purpose": purpose,
                            "priority": "High"})

        if pf.get("individual_harmonics") is False and not small:
            # A commissioned harmonic study is not a proportionate ask of a
            # house or a corner store.
            actions.append({
                "recommendation": "Perform a detailed harmonic study with individual "
                                  "source identification for the harmonic orders "
                                  "exceeding IEEE 519-2022 per-order limits.",
                "purpose":  "Identify and correct the sources of the specific harmonic "
                            "orders that exceed their limits.",
                "priority": "High"})

    # Three-phase only. A two-leg service reports its leg difference as a
    # measurement with no limit, so `pf["current_imbalance"]` is None there and
    # this never fires -- leg imbalance is not a violation and must not raise a
    # High-priority action. The equivalent advice for a house is attached to
    # the "Elevated neutral current" finding, which keys off the measured
    # neutral current rather than off the imbalance itself.
    if (pf["current_imbalance"] is False
            and not any(k in covered for k in ("imbalance", "balance"))):
        actions.append({
            "recommendation": ("Redistribute single-phase loads across phases. "
                               "Investigate whether triplen harmonics are "
                               "contributing to elevated neutral current."),
            "purpose":  "Reduce current imbalance and neutral current.",
            "priority": "High"})

    if pf.get("transformer_loading") is False and "transformer upgrade" not in covered:
        actions.append({
            "recommendation": "Contact your Xcel Energy Area Engineer to discuss an "
                              "upgrade of the overloaded serving transformer.",
            "purpose":  "Prevent thermal overload and premature transformer failure.",
            "priority": "High"})

    if pf["voltage"] is False and "voltage excursion" not in covered:
        actions.append({
            "recommendation": "Steady-state voltage excursions outside ANSI C84.1 Range A "
                              "warrant a distribution system review of secondary voltage, "
                              "service conductor length and tap settings.",
            "purpose":  "Return service voltage to within ANSI C84.1 Range A.",
            "priority": "High"})

    return _prune_actions(actions)


def _flicker_severities(fl: dict, report: dict) -> dict:
    """Severity for Pst and Plt separately, and the worse of the two overall.

    Graded off the 95th percentile rather than the maximum, with the share of
    the recording spent over the limit as persistence -- the same way voltage
    THD is graded, and for the same reason. A single ten-minute Pst of 4.98 is
    five times the limit, and grading on the maximum called that "severe" on
    the strength of one interval out of two thousand while ignoring the phase
    that sat over its Plt limit for a quarter of the recording. The maximum is
    still what decides pass or fail: compliance is binary, severity is not.
    """
    duration_h = (report.get("file_summary") or {}).get("duration_hours") or 0.0
    notes = []
    if duration_h and duration_h < 24.0:
        # Both standards assess the 95th percentile over a week, and no survey
        # recording runs that long -- so "shorter than a week" cannot be the
        # trigger, or every flicker finding would be discounted a band forever,
        # which is a bias rather than a discount. Under a day is different: it
        # holds only a handful of the two-hour windows Plt is built from, and
        # a percentile over those is genuinely soft.
        notes.append(f"{duration_h:.0f}-hour recording holds only "
                     f"{duration_h / 2:.0f} of the two-hour windows Plt is "
                     "built from")

    out: dict = {}
    worst_band = None
    for kind, key in (("pst", "flicker_pst"), ("plt", "flicker_plt")):
        phases = fl.get(kind) or {}
        if not phases:
            continue
        limit = fl[f"{kind}_limit"]
        # The governing phase is the one furthest past the limit at its 95th
        # percentile, which is the phase the severity should describe.
        worst_phase = max(phases.values(), key=lambda s: s["p95"])
        graded = grade_finding(
            all(s["pass"] for s in phases.values()),
            measured=worst_phase["p95"], limit=limit,
            persistence_pct=max(s["pct_exceeding"] for s in phases.values()),
            confidence_notes=notes)
        out[key] = graded
        if (worst_band is None
                or SEVERITY_ORDER.index(graded["band"])
                > SEVERITY_ORDER.index(worst_band["band"])):
            worst_band = graded

    if worst_band is not None:
        out["flicker"] = worst_band
    return out


def compute_severities(report: dict, thresh: Thresholds) -> dict:
    """Grade every finding once, keyed the same way as ``pass_fail``.

    The executive summary is written before the compliance table is drawn, and
    both quote severity.  Computing it in one place is what keeps the summary's
    "1 significant, 2 minor" from disagreeing with the rows underneath it.
    """
    pf   = report["pass_fail"]
    thd  = report["thd_compliance"]
    volt = report["voltage_compliance"]
    pfr  = report["power_factor"]
    imb  = report["voltage_imbalance"]
    ci   = report["current_imbalance"]
    dem  = report["demand"]
    ih   = report["individual_harmonics"]
    ivh  = report.get("individual_voltage_harmonics", {})
    hs   = report.get("harmonic_statistics", {})
    llv  = report.get("voltage_ll_compliance") or {}
    frq  = report.get("frequency") or {}
    fl   = report.get("flicker") or {}
    tdd_info = thd.get("tdd_info", {})

    sev: dict = {}

    # Voltage THD — artifact gating is a confidence signal, not a silent edit.
    v_thd = thd["voltage"]
    if v_thd.get("available"):
        notes = []
        if v_thd.get("artifact_samples"):
            notes.append(f"{v_thd['artifact_samples']} interval(s) dropped as "
                         f"artifacts below {v_thd.get('artifact_floor_v', 0):.0f} V")
        sev["thd_voltage"] = grade_finding(
            pf.get("thd_voltage"), measured=v_thd.get("p95_thd_pct"),
            limit=v_thd.get("limit_pct"),
            persistence_pct=v_thd.get("pct_exceeding"), confidence_notes=notes)

    # Current TDD — an assumed ISC class is exactly the kind of soft input that
    # should hold a finding back from the top band.
    c_thd = thd["current"]
    if c_thd.get("available"):
        notes = []
        if tdd_info and not tdd_info.get("isc_provided"):
            notes.append("ISC not supplied — most restrictive class assumed")
        if c_thd.get("light_load_filtered"):
            notes.append("light-load intervals excluded")
        sev["thd_current"] = grade_finding(
            pf.get("thd_current"), measured=c_thd.get("mean_thd_pct"),
            limit=c_thd.get("limit_pct"),
            persistence_pct=c_thd.get("pct_exceeding"), confidence_notes=notes)

    if volt.get("available"):
        sev["voltage"] = _grade_voltage_band(
            volt, volt.get("phases_missing_data") or [])

    if pfr.get("available") and thresh.customer_class != "r":
        sev["power_factor"] = grade_finding(
            pf.get("power_factor"), measured=pfr.get("mean_pf"),
            limit=pfr.get("limit"), persistence_pct=pfr.get("pct_below_limit"),
            lower_is_worse=True)

    if imb.get("available"):
        sev["voltage_imbalance"] = grade_finding(
            pf.get("voltage_imbalance"), measured=imb.get("mean_imbalance_pct"),
            limit=imb.get("limit_pct"), persistence_pct=imb.get("pct_exceeding"))

    if ci.get("available"):
        sev["current_imbalance"] = grade_finding(
            pf.get("current_imbalance"), measured=ci.get("mean_imbalance_pct"),
            limit=ci.get("limit_pct"), persistence_pct=ci.get("pct_exceeding"))

    if "transformer" in dem:
        tx = dem["transformer"]
        # None passes straight through to grade_finding, which returns
        # "not_assessed" -- there is nothing to grade on a shared transformer.
        sev["demand"] = grade_finding(
            None if tx["overloaded"] is None else not tx["overloaded"],
            measured=tx["peak_8h_kva"], limit=tx["nameplate_kva"])

    if ih.get("available"):
        sev["individual_harmonics"] = grade_finding(
            pf.get("individual_harmonics"),
            measured=ih.get("worst_margin_pct_of_il"),
            limit=ih.get("worst_limit_pct"))

    if ivh.get("available"):
        sev["individual_voltage_harmonics"] = grade_finding(
            pf.get("individual_voltage_harmonics"),
            measured=ivh.get("worst_pct_nom"), limit=ivh.get("limit_pct"))

    if hs.get("available"):
        # Grade on the order that binds, so a P95 sitting at 91% of its limit
        # reads as Watch here rather than as a bare Compliant.
        b = hs.get("binding") or {}
        sev["harmonic_statistics"] = grade_finding(
            pf.get("harmonic_statistics"),
            measured=b.get("p95"), limit=b.get("limit"),
            confidence_notes=([f"{hs['period_days']:.1f}-day recording, short of "
                               "the 7 days Clause 5 assumes"]
                              if hs.get("period_days", 7) < 7 else None))

    if fl.get("available"):
        sev.update(_flicker_severities(fl, report))

    if llv.get("available"):
        sev["voltage_line_to_line"] = _grade_voltage_band(llv)

    if frq.get("available"):
        sev["frequency"] = grade_finding(
            frq.get("overall_pass"), persistence_pct=frq.get("pct_out_of_bounds"))

    if report.get("itic", {}).get("available"):
        sev["itic_transients"] = grade_finding(pf.get("itic_transients"))

    return sev


def _severity_rollup(sev: dict) -> str:
    """"1 significant, 2 minor" — the proportion a bare fail count hides."""
    counts: dict = {}
    for s in sev.values():
        if s["band"] in ("severe", "significant", "minor"):
            counts[s["band"]] = counts.get(s["band"], 0) + 1
    order = ["severe", "significant", "minor"]
    return ", ".join(f"{counts[b]} {SEVERITY_LABEL[b].lower()}"
                     for b in order if b in counts)


def _exec_summary_bullets(report: dict, thresh: Thresholds, df,
                          key_findings: List[str], actions: List[dict]) -> List[str]:
    """3–5 bullets: overall compliance, most significant finding, principal
    concern, action outlook, and overall assessment."""
    pf        = report["pass_fail"]
    evaluated = {k: v for k, v in pf.items() if v is not None}
    if thresh.customer_class == "r":
        # Residential services are not subject to the power factor tariff clause
        evaluated.pop("power_factor", None)
    fl        = _flicker_status(report)
    # pass_fail already carries a "flicker" entry, so counting fl separately
    # inflated the tally by one against the table.
    n_eval    = len(evaluated) + (0 if "flicker" in evaluated else (1 if fl else 0))
    fails     = [_PF_FRIENDLY.get(k, k.replace("_", " "))
                 for k, v in evaluated.items() if v is False]
    if fl and fl["passes"] is False:
        fails.append("voltage flicker (IEC 61000-3-3)")

    sev = compute_severities(report, thresh)

    bullets: List[str] = []
    if fails:
        # A bare count of failed standards reads the same whether every one is a
        # marginal excursion or a sustained overload.  Leading with the severity
        # mix gives the reader proportion before they reach the table.
        rollup = _severity_rollup(sev)
        mix = f" ({rollup} by severity)" if rollup else ""
        bullets.append(
            f"{len(fails)} of the {n_eval} power quality standards evaluated were "
            f"not met{mix}: {', '.join(fails)}.")
    else:
        watch = sum(1 for s in sev.values() if s["band"] == "watch")
        watch_txt = (f" {watch} measurement(s) sit within the limit but close "
                     "enough to it to be worth tracking." if watch else "")
        bullets.append(
            f"All {n_eval} power quality standards evaluated for this service "
            f"were met.{watch_txt}")

    if key_findings and (fails or len(key_findings) > 0):
        bullets.append(key_findings[0])

    rca = report.get("root_causes", [])
    sev_rank = {"critical": 0, "warning": 1}
    risky = sorted((f for f in rca if f.get("severity") in sev_rank),
                   key=lambda f: sev_rank[f["severity"]])
    if risky:
        top = risky[0]
        bullets.append(
            f"Principal concern: {top['title']} "
            f"({top.get('confidence', 'low')} confidence). Attribution is left to "
            "the reviewing engineer; the evidence bearing on it is given with the "
            "finding.")

    if actions:
        n_high   = sum(1 for a in actions if a["priority"] == "High")
        high_txt = f", {n_high} of them high priority" if n_high else ""
        n_act    = len(actions)
        act_txt  = "action is" if n_act == 1 else "actions are"
        bullets.append(
            f"{n_act} candidate {act_txt} supported by the measurements"
            f"{high_txt} — see Recommended Actions.")
    else:
        bullets.append("No corrective actions are required at this time.")

    # Three marginal exceedances are not "significant deficiencies".  Escalate on
    # how bad the findings are, not on how many rows are red.
    n_serious = sum(1 for s in sev.values()
                    if s["band"] in ("severe", "significant"))
    has_critical = (any(f.get("severity") == "critical" for f in rca)
                    or report.get("neutral_health", {}).get("severity") == "critical"
                    or any(s["band"] == "severe" for s in sev.values())
                    or n_serious >= 3)
    has_warning  = any(f.get("severity") == "warning" for f in rca)
    if fails and has_critical:
        overall = ("Overall assessment: significant power quality deficiencies exist at "
                   "this service and warrant prompt attention.")
    elif fails:
        overall = ("Overall assessment: power quality at this service is generally "
                   "acceptable, but the standards listed above are not met and targeted "
                   "corrective action is recommended.")
    elif has_warning:
        overall = ("Overall assessment: all evaluated standards are met; the "
                   "observations noted in this report warrant monitoring but no "
                   "immediate corrective action.")
    else:
        overall = ("Overall assessment: power quality at this service is good and no "
                   "corrective action is required.")
    bullets.append(overall)

    return bullets[:5]


def _word_recording_overview(doc, report, outdir=None, stem="") -> None:
    """The measured series, whole and unprocessed, before any assessment.

    Everything further down is derived — filtered, averaged, compared against a
    limit. This section is the raw record those derivations came from, so a
    reader can confirm the file was read correctly before weighing any
    conclusion drawn from it.
    """
    p = _plot_path(outdir, stem, "overview.png")
    if p is None or not p.exists():
        return
    fs = report["file_summary"]
    dq = fs.get("data_quality") or {}
    _section_heading(doc, "Recording Overview", level=1)
    _body(doc,
        f"Measured RMS voltage and current over the full recording period, "
        f"{fs['start_time']} to {fs['end_time']} "
        f"({_m(fs['duration_hours'], '.1f', ' hours')}, {fs['sample_count']:,} intervals of "
        f"{_m(fs.get('interval_minutes', 5), 'g')} minutes). These are the recorded "
        "series as read from the meter file, with no filtering or scaling "
        "applied; every finding in this report derives from them. Any span the "
        "meter did not record is shaded and the trace is broken across it, so a "
        "gap cannot be mistaken for a flat reading."
        + ("  This file was incomplete when read — see Source file integrity in "
           "the header table above, and Appendix B for what could not be read."
           if (dq.get("missing_bytes") or dq.get("unreadable_observations"))
           else ""))
    note = _session_note(fs)
    if note:
        # Immediately under the period this report covers, where a reader
        # checking the dates is already looking.
        p = doc.add_paragraph()
        _bold(p, "More than one session in this file: ")
        _normal(p, note)
    _embed_plot(doc, outdir, stem, "overview.png",
                caption="Measured voltage and current over the recording period.")
    doc.add_paragraph()


def _word_exec_summary(doc, report, thresh, df,
                       key_findings: List[str], actions: List[dict]) -> None:
    _section_heading(doc, "Executive Summary and Compliance Status", level=1)
    for b in _exec_summary_bullets(report, thresh, df, key_findings, actions):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b).font.size = Pt(10)
    doc.add_paragraph()
    _word_compliance_table(doc, report, thresh, df)


def _word_key_findings(doc, key_findings: List[str]) -> None:
    _section_heading(doc, "Key Findings", level=1)
    if not key_findings:
        _body(doc,
            "No significant power quality deviations were identified during the "
            "recording period.")
    else:
        for txt in key_findings:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(txt)
    doc.add_paragraph()


def _word_demand(doc, report, thresh, outdir=None, stem="") -> Optional[str]:
    dem = report["demand"]

    if not dem["available"]:
        return "Demand"
    _section_heading(doc, "Demand", level=2)
    if dem["available"]:
        ap = dem.get("apparent_power", {})
        rp = dem.get("real_power", {})
        pc = dem.get("peak_current", {})
        _pk_str = (" Peak current within any 5-minute interval was "
                   f"{_m(pc['max_a'], '.0f', ' A')}." if pc else "")
        _body(doc,
            f"Peak apparent demand was {_m(ap.get('peak_kva', 0), '.1f', ' kVA')} "
            f"(mean {_m(ap.get('mean_kva', 0), '.1f', ' kVA')}, load factor {_m(ap.get('load_factor', 0), '.2f')}). "
            f"Peak real power was {_m(rp.get('peak_kw', 0), '.1f', ' kW')}.{_pk_str}"
        )
        if "transformer" in dem:
            tx = dem["transformer"]
            if tx["overloaded"] is True:
                # True is sound whether or not the transformer is shared: this
                # service's demand alone is a lower bound on its load, so
                # exceeding nameplate proves an overload regardless of who else
                # is connected.
                shared_note = ("" if tx.get("dedicated") else
                               " This service alone accounts for that, and the "
                               "transformer serves other customers who were not "
                               "measured, so the total is at least this much.")
                _body(doc,
                    f"The transformer is overloaded. The nameplate is {tx['nameplate_kva']:.0f} kVA; "
                    f"the 8-hour rolling peak demand was {_m(tx['peak_8h_kva'], '.1f', ' kVA')} "
                    f"({_m(tx['pct_nameplate'], '.0f', '%')} of nameplate).{shared_note} "
                    "Transformers can be loaded above nameplate on an 8-hour basis but not continuously. "
                    "A transformer upgrade should be evaluated."
                )
            elif tx["overloaded"] is False:
                _body(doc,
                    f"The transformer loading is within acceptable limits. "
                    f"The 8-hour peak was {_m(tx['peak_8h_kva'], '.1f', ' kVA')} "
                    f"({_m(tx['pct_nameplate'], '.0f', '%')} of the {tx['nameplate_kva']:.0f} kVA nameplate)."
                )
            else:
                # Shared transformer, this service below nameplate on its own.
                # Nothing here supports a statement about the transformer, and
                # this paragraph used to make one: "within acceptable limits",
                # about equipment whose load was never measured.
                _body(doc,
                    f"This service's 8-hour peak demand was {_m(tx['peak_8h_kva'], '.1f', ' kVA')} "
                    f"against a {tx['nameplate_kva']:.0f} kVA transformer "
                    f"({_m(tx['pct_nameplate'], '.0f', '%')} of nameplate). The transformer "
                    "serves other customers whose load was not measured, so this "
                    "figure is this service's contribution to its loading and not "
                    "that loading. Whether the transformer is adequately sized "
                    "cannot be determined from a recording at one meter; it would "
                    "need transformer-level load data or a coincident recording "
                    "across the services it feeds."
                )
    _embed_plot(doc, outdir, stem, "demand_profile.png",
                "Demand pattern over the recording period.")
    doc.add_paragraph()
    return None


def _word_power_factor(doc, report, thresh, outdir=None, stem="") -> Optional[str]:
    pfr = report["power_factor"]

    if not pfr["available"]:
        return "Power Factor"
    _section_heading(doc, "Power Factor", level=2)
    if pfr["available"]:
        direction = "lagging" if pfr["mean_pf"] > 0 else "leading"
        is_residential = thresh.customer_class == "r"
        if is_residential:
            _body(doc,
                f"Measured mean power factor was {_m(pfr['mean_pf'], '.3f')} {direction} "
                f"(minimum {_m(pfr['min_pf'], '.3f')}). "
                "Residential services (PSCo Schedule R) are not subject to the power factor "
                "tariff clause — values in the 0.85–0.95 range are normal for homes with HVAC, "
                "appliances, and lighting. No corrective action is required."
            )
        elif thresh.customer_class == "pg":
            # Schedule PG — C&I Primary: Sheet R121 requires "near unity"
            if pfr["pct_below_limit"] == 0:
                _body(doc,
                    f"Power factor was maintained near unity as required by PSCo Electric Tariff "
                    f"Sheet R121 (Schedule PG — C&I Primary service). "
                    f"Measured mean {_m(pfr['mean_pf'], '.3f')} {direction}, minimum {_m(pfr['min_pf'], '.3f')}."
                )
            else:
                _body(doc,
                    f"Power factor fell below {pfr['limit']:.2f} during "
                    f"{_mp(pfr['pct_below_limit'], '.1f')} of the recording "
                    f"(mean {_m(pfr['mean_pf'], '.3f')} {direction}, minimum {_m(pfr['min_pf'], '.3f')}). "
                    "PSCo Electric Tariff Sheet R121 requires commercial and industrial "
                    "customers to maintain power factor as near unity as practicable. "
                    "The customer should evaluate power factor correction equipment to comply "
                    "with tariff requirements."
                )
        else:
            # Sheet R73 sits in the General rules and sets 0.90 lagging for
            # every class, so it is not cited against a schedule.
            tariff_cite = "PSCo Electric Tariff Sheet R73"
            if pfr["pct_below_limit"] == 0:
                _body(doc,
                    f"Power factor was consistently above the 0.90 lagging requirement "
                    f"({tariff_cite}). "
                    f"Measured mean {_m(pfr['mean_pf'], '.3f')} {direction}, minimum {_m(pfr['min_pf'], '.3f')}."
                )
            else:
                _body(doc,
                    f"Power factor fell below the 0.90 lagging requirement during "
                    f"{_mp(pfr['pct_below_limit'], '.1f')} of the recording "
                    f"(mean {_m(pfr['mean_pf'], '.3f')} {direction}, minimum {_m(pfr['min_pf'], '.3f')}). "
                    f"{tariff_cite} requires Commercial and C&I Secondary customers to maintain "
                    f"power factor of not less than 0.90 lagging. The Company reserves the right "
                    f"to discontinue service to customers not complying with this requirement."
                )
    _embed_plot(doc, outdir, stem, "pf_load.png",
                "Power factor vs load. Low power factor at light load is common and "
                "usually benign; low power factor at high load is what the tariff "
                "addresses.", width_cm=12.5)
    doc.add_paragraph()
    return None


def _word_voltage(doc, report, outdir=None, stem="") -> Optional[str]:
    volt = report["voltage_compliance"]

    if not volt["available"]:
        return "Steady-State Voltage"
    _section_heading(doc, "Steady-State Voltage", level=2)
    if volt["available"]:
        rng = volt["range_v"]
        missing = volt.get("phases_missing_data") or []
        if missing:
            _body(doc,
                f"No usable voltage data for phase(s): "
                f"{', '.join(_phase_label(m) for m in missing)}. "
                "The compliance result below reflects only the phase(s) with valid data."
            )
        all_pass = all(v["pct_out_of_bounds"] == 0 for v in volt["phases"].values())
        # State the basis once, at the top, rather than appending a caveat to
        # every phase. What the verdict rests on is a property of the check, not
        # of phase B.
        _basis = ("These are interval averages. ANSI C84.1 rates sustained "
                  "service voltage, so the compliance result below is taken "
                  "from them")
        if any(v.get("used_interval_extremes") for v in volt["phases"].values()):
            _basis += (", and the meter's within-interval minima and maxima are "
                       "quoted separately where they left the band. Excursions "
                       "shorter than one interval are sags or swells and are "
                       "graded on depth and duration against the ITIC envelope "
                       "in Voltage Events, not against C84.1")
        _body(doc, _basis + ".")
        if volt.get("range_b_note"):
            _body(doc, volt["range_b_note"])

        def _extremes_clause(v) -> str:
            if not v.get("used_interval_extremes"):
                return ""
            lo, hi = v.get("min_interval_v"), v.get("max_interval_v")
            if lo is None or hi is None or (lo >= rng[0] and hi <= rng[1]):
                return ""
            return (f" Within-interval extremes reached {_m(lo, '.1f')}–"
                    f"{_m(hi, '.1f', ' V')}; see Voltage Events.")

        if all_pass:
            vals = {ph: v for ph, v in volt["phases"].items()}
            phase_str = "  ".join(
                f"Phase {_phase_label(ph)}: {_m(v['min_v'], '.1f')}–"
                f"{_m(v['max_v'], '.1f', ' V')} (mean {_m(v['mean_v'], '.1f', ' V')})"
                for ph, v in vals.items()
            )
            _body(doc,
                f"Voltage was within ANSI C84.1 Range A ({rng[0]:.1f}–{rng[1]:.1f} V) "
                f"for the entire recording period. {phase_str}."
                + "".join(_extremes_clause(v) for v in vals.values())
            )
        else:
            for ph, v in volt["phases"].items():
                pct_over  = v.get("pct_over",  0.0)
                pct_under = v.get("pct_under", 0.0)
                if v["pct_out_of_bounds"] == 0:
                    _body(doc,
                        f"Phase {_phase_label(ph)}: within ANSI C84.1 Range A ({rng[0]:.1f}–{rng[1]:.1f} V) "
                        f"for the entire recording. Min {_m(v['min_v'], '.1f', ' V')}, mean {_m(v['mean_v'], '.1f', ' V')}, "
                        f"max {_m(v['max_v'], '.1f', ' V')}.{_extremes_clause(v)}"
                    )
                else:
                    if pct_over > 0 and pct_under > 0:
                        direction = (
                            f"{_mp(pct_over, '.2f')} of intervals above the "
                            f"upper limit ({rng[1]:.1f} V) and "
                            f"{_mp(pct_under, '.2f')} below the lower limit "
                            f"({rng[0]:.1f} V)"
                        )
                    elif pct_over > 0:
                        direction = (f"{_mp(pct_over, '.2f')} of intervals above "
                                     f"the upper limit ({rng[1]:.1f} V)")
                    else:
                        direction = (f"{_mp(pct_under, '.2f')} of intervals below "
                                     f"the lower limit ({rng[0]:.1f} V)")
                    # Which band it reached is the finding, so it is named here
                    # rather than left for the reader to work out from the share.
                    band_clause = ""
                    if v["band"] == "range_b" and volt.get("range_b_v"):
                        band_clause = (
                            f" These stayed within Range B "
                            f"({volt['range_b_v'][0]:.1f}–{volt['range_b_v'][1]:.1f} V), "
                            "which C84.1 permits provided the excursions are "
                            "limited in extent, frequency and duration and are "
                            "corrected within a reasonable time.")
                    elif v["band"] == "outside_b" and volt.get("range_b_v"):
                        band_clause = (
                            f" {_mp(v['pct_outside_b'], '.2f')} of intervals were "
                            f"outside Range B as well "
                            f"({volt['range_b_v'][0]:.1f}–{volt['range_b_v'][1]:.1f} V), "
                            "which is beyond what C84.1 contemplates for a "
                            "service voltage.")
                    _body(doc,
                        f"Phase {_phase_label(ph)}: {direction}. "
                        f"Min {_m(v['min_v'], '.1f', ' V')}, mean {_m(v['mean_v'], '.1f', ' V')}, max {_m(v['max_v'], '.1f', ' V')} "
                        f"(ANSI C84.1 Range A: {rng[0]:.1f}–{rng[1]:.1f} V)."
                        f"{band_clause}{_extremes_clause(v)}"
                    )
    _embed_plot(doc, outdir, stem, "voltage.png",
                "Phase voltages against the ANSI C84.1 Range A and Range B "
                "limits; out-of-range periods are shaded.")
    doc.add_paragraph()
    return None


def _word_harmonics(doc, report, thresh, df, outdir, stem="") -> None:
    thd      = report["thd_compliance"]
    ih       = report["individual_harmonics"]
    ivh      = report.get("individual_voltage_harmonics", {})
    nh       = report.get("neutral_harmonics", {})
    sh       = report.get("harmonic_sources", {})
    hd       = report.get("harmonic_direction", {})
    ss       = report.get("spectral_shape", {})
    hs       = report.get("harmonic_statistics", {})
    dem      = report["demand"]
    tdd_info = thd.get("tdd_info", {})
    c_thd    = thd["current"]
    is_split = "voltage_c" not in report.get("file_summary", {}).get("channels", [])

    spec_img    = _plot_path(outdir, stem, "harmonic_spectrum.png")
    has_kfactor = (
        df is not None
        and "kfactor_meter" in df.columns
        and df["kfactor_meter"].notna().any()
    )
    any_harm = (
        c_thd.get("available") or thd["voltage"].get("available")
        or ih.get("available") or ivh.get("available") or nh.get("available")
        or sh.get("available") or ss.get("available") or hs.get("available")
        or hd.get("available") or spec_img.exists() or has_kfactor
    )
    if not any_harm:
        return   # no harmonic data in this recording — section suppressed

    _section_heading(doc, "Harmonic Evaluation", level=1)
    if tdd_info or c_thd.get("available") or ih.get("available"):
        _section_heading(doc, "Harmonic Compliance Evaluation", level=2)
    # Which standard applies comes before any limit is quoted: 519 and 1547
    # differ by three times in the aggregate limit and use different
    # denominators, so a limit stated without its standard is not a limit.
    _word_current_standard(doc, report, thresh)

    il_phrase = _il_basis_phrase(tdd_info)
    if tdd_info and tdd_info.get("isc_provided"):
        _body(doc,
            f"The available short-circuit current at the point of delivery is {tdd_info['isc_amps']:,.0f} A "
            f"(source: {tdd_info.get('isc_source', 'provided')}). "
            f"{il_phrase} "
            f"The resulting ISC/IL ratio is {tdd_info['isc_il_ratio']:.1f}, placing this service in the "
            f"IEEE 519-2022 {tdd_info['tdd_class']} class with a TDD limit of {tdd_info['tdd_limit_pct']:.1f}%."
        )
    elif tdd_info:
        _body(doc,
            f"Current distortion is evaluated as Total Demand Distortion (TDD), which "
            f"references harmonic current to a fixed reference current (IL) rather "
            f"than to the instantaneous fundamental — this avoids overstating distortion "
            f"when the fundamental is small. "
            f"{il_phrase} The available short-circuit current at the "
            f"point of delivery was not provided, so the most restrictive IEEE 519-2022 "
            f"class (ISC/IL < 20) is assumed, giving a conservative TDD limit of "
            f"{tdd_info['tdd_limit_pct']:.1f}%; the true limit for this service can only "
            f"be equal or higher."
        )
    if c_thd["available"]:
        metric = "TDD" if tdd_info else "THD"
        if c_thd["pct_exceeding"] == 0:
            _body(doc,
                f"Current {metric} was within the {c_thd['limit_pct']:.1f}% limit throughout the recording. "
                f"Maximum {metric} was {_m(c_thd['max_thd_pct'], '.2f', '%')}, mean {_m(c_thd['mean_thd_pct'], '.2f', '%')}."
            )
        else:
            ll_note = (
                " Light-load intervals are excluded from this evaluation "
                "(see Appendix B for the method)."
                if c_thd.get("light_load_filtered") else ""
            )
            _body(doc,
                f"Current {metric} exceeded the {c_thd['limit_pct']:.1f}% limit during "
                f"{_mp(c_thd['pct_exceeding'], '.1f')} of the recording "
                f"(max {_m(c_thd['max_thd_pct'], '.2f', '%')}, mean {_m(c_thd['mean_thd_pct'], '.2f', '%')}).{ll_note} "
                "Common sources include VFDs, UPS systems, switched-mode power supplies, and arc furnaces. "
                "Mitigation options include passive harmonic filters, active front-end drives, "
                "or 12-pulse converter topologies."
            )
        if "peak_max_tdd_pct" in c_thd:
            pk_pass = c_thd["peak_pct_exceeding"] == 0
            pk_verdict = "remained within" if pk_pass else "also exceeded"
            _body(doc,
                f"On a within-interval peak basis (using the meter's 5-minute max record), "
                f"current {metric} {pk_verdict} the {c_thd['limit_pct']:.1f}% limit "
                f"(peak max {_m(c_thd['peak_max_tdd_pct'], '.2f', '%')}, "
                f"peak exceedance {_mp(c_thd['peak_pct_exceeding'], '.1f')})."
            )

    if ih.get("available") and not ih["overall_pass"]:
        fail_orders = [
            f"H{h} (phase {ph})"
            for ph in ("a", "b", "c")
            for h, r in ih["phases"].get(ph, {}).items()
            if not r["pass"]
        ]
        _body(doc,
            f"The following individual current harmonic orders exceeded their IEEE 519-2022 "
            f"Table 2 per-order limits: "
            + ", ".join(fail_orders) + ". "
            "Individual harmonic limits are more restrictive than TDD for higher-order harmonics. "
            "Per-order magnitudes for every phase are tabulated below."
        )
    elif ih.get("available"):
        h_ord = ih.get("worst_order")
        if h_ord:
            worst_r = ih["phases"].get(h_ord[1], {}).get(h_ord[0], {})
            _body(doc,
                f"All individual current harmonic orders are within IEEE 519-2022 Table 2 limits. "
                f"The most constraining harmonic is H{h_ord[0]} (phase {h_ord[1].upper()}) "
                f"at {_m(ih['worst_pct_of_il'], '.2f', '%')} of IL "
                f"against a limit of {worst_r.get('limit_pct_il', '—')}%."
            )

    _embed_plot(doc, outdir, stem, "thd.png",
                "Voltage THD and current TDD over the recording, against IEEE 519 "
                "limits; exceedance periods are shaded.")

    # Individual harmonic table (if available)
    if ih.get("available"):
        _ph_keys   = ("a", "b") if is_split else ("a", "b", "c")
        _ph_names  = (("a", "L1"), ("b", "L2")) if is_split else (("a", "Phase A"), ("b", "Phase B"), ("c", "Phase C"))
        _n_cols    = 2 + len(_ph_keys)
        _col_w     = [2.0] + [4.5 if is_split else 3.5] * len(_ph_keys) + [3.5]
        doc.add_paragraph()
        ih_hdr = doc.add_paragraph()
        _bold(ih_hdr, "Individual Harmonic Current Summary (% of IL)", size_pt=10)
        harm_tbl = doc.add_table(rows=1, cols=_n_cols)
        harm_tbl.style = 'Table Grid'
        _set_col_widths(harm_tbl, _col_w)
        _ih_hdrs = ["Order"] + [f"{nm} (%IL)" for _, nm in _ph_names] + ["Limit (%IL)"]
        for cell, text in zip(harm_tbl.rows[0].cells, _ih_hdrs):
            _cell_shade(cell, _CHROME_BAND)
            cell.paragraphs[0].add_run(text).bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for h in _H519_ORDERS:
            if not any(ih["phases"].get(ph, {}).get(h) for ph in _ph_keys):
                continue
            row_cells = harm_tbl.add_row().cells
            row_cells[0].paragraphs[0].add_run(f"H{h}").font.size = Pt(9)
            limit_shown = False
            any_fail = False
            for j, ph in enumerate(_ph_keys):
                r = ih["phases"].get(ph, {}).get(h)
                if r:
                    txt = f"{r['max_pct_il']:.2f}"
                    run = row_cells[j+1].paragraphs[0].add_run(txt)
                    run.font.size = Pt(9)
                    if not r["pass"]:
                        run.bold = True
                        run.font.color.rgb = _FAIL_CLR
                        any_fail = True
                    if not limit_shown:
                        row_cells[_n_cols-1].paragraphs[0].add_run(f"{r['limit_pct_il']:.1f}").font.size = Pt(9)
                        limit_shown = True
            if any_fail:
                for cell in row_cells:
                    _cell_shade(cell, _sev_shade("severe"))

    # ── Individual voltage harmonic table ─────────────────────────────────────
    if ivh.get("available"):
        _vph_keys  = ("a", "b") if is_split else ("a", "b", "c")
        _vph_names = (("a", "L1"), ("b", "L2")) if is_split else (("a", "Phase A"), ("b", "Phase B"), ("c", "Phase C"))
        _vn_cols   = 2 + len(_vph_keys)
        _vcol_w    = [2.0] + [4.5 if is_split else 3.5] * len(_vph_keys) + [3.5]
        _section_heading(doc, "Individual Harmonic Voltage Summary", level=2)
        _body(doc,
            f"Limit: 5.0% of nominal ({thresh.nominal_voltage:.0f} V) per IEEE 519-2022 Table 1 "
            f"(bus voltage < 1 kV). Values are absolute Volts converted to % of nominal."
        )
        volt_harm_tbl = doc.add_table(rows=1, cols=_vn_cols)
        volt_harm_tbl.style = 'Table Grid'
        _set_col_widths(volt_harm_tbl, _vcol_w)
        _ivh_hdrs = ["Order"] + [f"{nm} (%nom)" for _, nm in _vph_names] + ["Limit (%nom)"]
        for cell, text in zip(volt_harm_tbl.rows[0].cells, _ivh_hdrs):
            _cell_shade(cell, _CHROME_BAND)
            cell.paragraphs[0].add_run(text).bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for h in (3, 5, 7, 11, 13):
            if not any(ivh["phases"].get(ph, {}).get(h) for ph in _vph_keys):
                continue
            row_cells = volt_harm_tbl.add_row().cells
            row_cells[0].paragraphs[0].add_run(f"H{h}").font.size = Pt(9)
            limit_shown = False
            any_fail = False
            for j, ph in enumerate(_vph_keys):
                r = ivh["phases"].get(ph, {}).get(h)
                if r:
                    txt = f"{r['max_pct_nom']:.2f}"
                    run = row_cells[j+1].paragraphs[0].add_run(txt)
                    run.font.size = Pt(9)
                    if not r["pass"]:
                        run.bold = True
                        run.font.color.rgb = _FAIL_CLR
                        any_fail = True
                    if not limit_shown:
                        row_cells[_vn_cols-1].paragraphs[0].add_run(f"{r['limit_pct']:.1f}").font.size = Pt(9)
                        limit_shown = True
            if any_fail:
                for cell in row_cells:
                    _cell_shade(cell, _sev_shade("severe"))

    if nh.get("available") or sh.get("available") or ss.get("available"):
        _section_heading(doc, "Harmonic Source and Resonance Diagnostics", level=2)
        _embed_plot(doc, outdir, stem, "harmonic_trend.png",
                    "Dominant harmonic orders vs load over time. Harmonics that track "
                    "load point to customer equipment; load-independent harmonics "
                    "point to background or system sources.")

    # ── Neutral harmonic content (informational) ──────────────────────────────
    if nh.get("available"):
        doc.add_paragraph()
        nh_hdr = doc.add_paragraph()
        _bold(nh_hdr, "Neutral Harmonic Content (Informational)", size_pt=10)

        acc = nh.get("accumulation_factor")
        t_pct = nh.get("triplen_pct", 0.0)
        acc_str = (f"{_m(acc, '.1f')}×" if acc is not None
                   else "n/a (phase harmonics not available)")
        if is_split:
            _body(doc,
                f"In a single-phase 3-wire (split-phase) service, the neutral carries the difference "
                f"current between L1 and L2, not the sum of zero-sequence currents from three phases. "
                f"The two legs are 180 degrees apart, so odd harmonics — triplens included — subtract "
                f"in the neutral as the fundamental does, rather than adding as they would on a "
                f"120-degree system. No accumulation factor is reported for this service: the "
                f"quantity is defined over three phases and does not apply here. Neutral harmonic "
                f"content reflects imbalance between the legs. Triplen content: {_m(t_pct, '.0f', '%')} of total "
                f"neutral harmonic current."
            )
        else:
            _body(doc,
                f"Triplens (H3, H9, H15) are zero-sequence harmonics that add arithmetically in a "
                f"4-wire wye neutral. Triplen content: {_m(t_pct, '.0f', '%')} of total neutral harmonic current. "
                f"Accumulation factor (H3-neutral ÷ mean H3-phase): {acc_str}. "
                "A factor near 3 indicates full accumulation from balanced "
                "single-phase loads on all three phases, and above 3 indicates "
                "resonance amplification. Well below 3 means the third harmonic "
                "is largely cancelling in the neutral rather than adding, which "
                "is the case when the load is predominantly three-phase — the "
                "usual reading on a service with few single-phase nonlinear loads."
            )

        nh_tbl = doc.add_table(rows=1, cols=4)
        nh_tbl.style = 'Table Grid'
        _set_col_widths(nh_tbl, [2.0, 3.5, 3.5, 3.5])
        for cell, text in zip(nh_tbl.rows[0].cells,
                               ["Order", "Mean (A)", "Max (A)", "Type"]):
            _cell_shade(cell, _CHROME_BAND)
            cell.paragraphs[0].add_run(text).bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for h, od in sorted(nh["orders"].items()):
            row_cells = nh_tbl.add_row().cells
            row_cells[0].paragraphs[0].add_run(f"H{h}").font.size = Pt(9)
            row_cells[1].paragraphs[0].add_run(f"{od['mean_a']:.3f}").font.size = Pt(9)
            row_cells[2].paragraphs[0].add_run(f"{od['max_a']:.3f}").font.size = Pt(9)
            label = "Triplen (zero-seq)" if od["is_triplen"] else "Non-triplen"
            run = row_cells[3].paragraphs[0].add_run(label)
            run.font.size = Pt(9)
            if od["is_triplen"]:
                run.bold = True

    # ── Harmonic source indication ────────────────────────────────────────────
    if sh.get("available"):
        doc.add_paragraph()
        sh_hdr = doc.add_paragraph()
        _bold(sh_hdr, "Harmonic Source Indication (Indicative)", size_pt=10)

        resonant  = sh.get("resonant_orders", [])
        overall   = sh.get("overall", "indeterminate")
        res_str   = f"H{', H'.join(str(h) for h in sorted(resonant))}" if resonant else "none detected"
        overall_labels = {
            "customer":         "Customer-side injection",
            "resonance_suspect": "Resonance suspected",
            "mixed":            "Mixed / indeterminate",
            "indeterminate":    "Indeterminate",
        }
        _body(doc,
            f"Overall indication: {overall_labels.get(overall, overall)}. "
            f"Resonance suspects: {res_str}. "
            "This is an indication only; see Appendix B for the method and its limitations. It concerns the direction distortion appears to come from, not responsibility for it."
        )

        sh_tbl = doc.add_table(rows=1, cols=5)
        sh_tbl.style = 'Table Grid'
        _set_col_widths(sh_tbl, [1.5, 2.5, 2.5, 2.5, 3.5])
        for cell, text in zip(sh_tbl.rows[0].cells,
                               ["Order", "Apparent Z (Ω)", "Z_ratio", "Correlation", "Indication"]):
            _cell_shade(cell, _CHROME_BAND)
            cell.paragraphs[0].add_run(text).bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for h, od in sorted(sh["orders"].items()):
            row_cells = sh_tbl.add_row().cells
            row_cells[0].paragraphs[0].add_run(f"H{h}").font.size = Pt(9)
            row_cells[1].paragraphs[0].add_run(
                f"{od['z_ohm']:.4f}").font.size = Pt(9)
            ratio_str = f"{od['z_ratio']:.2f}×" if od["z_ratio"] is not None else "—"
            row_cells[2].paragraphs[0].add_run(ratio_str).font.size = Pt(9)
            corr_str  = f"{od['corr']:.2f}" if od["corr"] is not None else "—"
            row_cells[3].paragraphs[0].add_run(corr_str).font.size = Pt(9)
            attr      = od.get("attribution", "indeterminate")
            attr_labels = {
                "customer":          "Customer",
                "resonance_suspect": "Resonance suspect",
                "indeterminate":     "Indeterminate",
            }
            attr_run = row_cells[4].paragraphs[0].add_run(attr_labels.get(attr, attr))
            attr_run.font.size = Pt(9)
            if attr == "resonance_suspect":
                attr_run.bold = True
                attr_run.font.color.rgb = _FAIL_CLR
                for cell in row_cells:
                    _cell_shade(cell, _sev_shade("severe"))

    # ── Spectral shape (broadband vs. resonance classification) ───────────────
    if ss.get("available"):
        doc.add_paragraph()
        ss_hdr = doc.add_paragraph()
        _bold(ss_hdr, "Spectral Shape Classification", size_pt=10)
        class_labels = {
            "broadband_consistent": "Broadband-consistent — elevated and flat across orders",
            "resonance_present":    "Resonance present — see Harmonic Source Indication above",
            "elevated_uneven":      "Elevated but concentrated — no order flagged as resonant",
            "not_elevated":         "Not meaningfully elevated",
        }
        cls = ss.get("classification")
        elev_ratio = ss.get("elevation_ratio")
        elev_str = f"{elev_ratio:.0%} of the {thresh.thd_voltage_limit:.0f}% IEEE 519 limit" if elev_ratio is not None else "limit not configured"
        _body(doc,
            f"{class_labels.get(cls, cls)}. Mean voltage THD {_m(ss['mean_vthd_pct'], '.2f', '%')} "
            f"({elev_str}), per-order spectrum coefficient of variation {_m(ss['flatness_cv'], '.2f')}."
        )

    # ── Which side of the meter the distortion comes from ─────────────────────
    _word_harmonic_direction(doc, report, thresh)

    # ── IEEE 519-2022 Clause 5 statistical compliance tables ──────────────────
    if hs.get("available"):
        _section_heading(doc, "Statistical Harmonic Evaluation (IEEE 519-2022 Clause 5)", level=2)
        _body(doc,
            f"Percentiles computed over the {_m(hs['period_days'], '.1f')}-day recording period "
            f"(ISC/IL = {hs['isc_il_ratio']:.0f}, class {hs['isc_class']}). "
            "See Appendix B for the statistical method and its limitations."
        )
        # The bracketed figure is headroom against the limit, so a negative
        # value marks an exceedance. Without saying so, the convention
        # inverts on the reader: "7.86% (-0.86)" is over a 7.0% limit while
        # "1.63% (+5.37)" is comfortably under it.
        _body(doc,
            "Each cell gives the measured percentile and, in brackets, its "
            "margin against the limit. A positive margin is headroom; a "
            "negative margin means the limit was exceeded by that amount.")

        ph_cols = [ph for ph in ("a", "b", "c")
                   if any(ph in hs["weekly"].get(k, {}) for k in hs["weekly"])]
        ph_labels = {"a": "L1", "b": "L2"} if is_split else {"a": "Phase A", "b": "Phase B", "c": "Phase C"}
        # The aggregate row is measured against the current-distortion limit, so
        # it is TDD wherever a TDD limit applies. Labelling it THD in a table of
        # "% of IL" named a different quantity than the one tabulated.
        _aggregate_label = "TDD" if tdd_info else "THD"
        stat_orders = [("h3", "H3"), ("h5", "H5"), ("h7", "H7"),
                       ("h9", "H9"), ("h11", "H11"), ("h13", "H13"),
                       ("h17", "H17"), ("h19", "H19"), ("h23", "H23"), ("h25", "H25"),
                       ("thd", _aggregate_label)]


        def _stat_table(title: str, val_key: str, lim_key: str,
                        pass_key: str, lim_label: str) -> None:
            _bold(doc.add_paragraph(), title, size_pt=9)
            n_cols = 2 + len(ph_cols)
            tbl = doc.add_table(rows=1, cols=n_cols)
            tbl.style = 'Table Grid'
            col_w = [2.0] + [4.0] * len(ph_cols) + [3.0]
            _set_col_widths(tbl, col_w[:n_cols])
            hdrs = ["Order"] + [ph_labels[p] + " (%IL)" for p in ph_cols] + [lim_label]
            for cell, txt in zip(tbl.rows[0].cells, hdrs):
                _cell_shade(cell, _CHROME_BAND)
                r = cell.paragraphs[0].add_run(txt)
                r.bold = True
                r.font.size = Pt(9)

            for key, label in stat_orders:
                ph_data = hs["weekly"].get(key, {})
                if not any(ph in ph_data for ph in ph_cols):
                    continue
                row = tbl.add_row()
                row.cells[0].paragraphs[0].add_run(label).font.size = Pt(9)
                any_fail_row = False
                limit_shown = False
                for j, ph in enumerate(ph_cols):
                    d = ph_data.get(ph, {})
                    if not d:
                        continue
                    val = d.get(val_key, 0.0)
                    passes = d.get(pass_key)
                    margin = d.get(
                        "p95_margin" if pass_key == "p95_pass" else "p99_margin", 0.0
                    )
                    margin_str = (
                        f" (+{margin:.2f})" if margin is not None and margin >= 0
                        else (f" ({margin:.2f})" if margin is not None else "")
                    )
                    txt = f"{val:.2f}%{margin_str}"
                    run = row.cells[j + 1].paragraphs[0].add_run(txt)
                    run.font.size = Pt(9)
                    if passes is False:
                        run.bold = True
                        run.font.color.rgb = _FAIL_CLR
                        any_fail_row = True
                    if not limit_shown:
                        lim_val = d.get(lim_key, 0.0)
                        row.cells[n_cols - 1].paragraphs[0].add_run(
                            f"{lim_val:.1f}%"
                        ).font.size = Pt(9)
                        limit_shown = True
                if any_fail_row:
                    for cell in row.cells:
                        _cell_shade(cell, _sev_shade("severe"))

        _stat_table(
            f"Weekly 95th Percentile vs 1.0× Limit (Short Time, {hs.get('period_note', '')})",
            "p95", "limit", "p95_pass", "Limit (1.0×)",
        )
        doc.add_paragraph()
        _stat_table(
            "Weekly 99th Percentile vs 1.5× Limit (Short Time)",
            "p99", "limit_1p5x", "p99_pass", "1.5× Limit",
        )
        doc.add_paragraph()

        # VST daily proxy table (separate — uses daily_vst data)
        _bold(doc.add_paragraph(),
              "Daily 99th Percentile vs 2.0× Limit (Very Short Time proxy)", size_pt=9)
        n_cols = 2 + len(ph_cols)
        vst_tbl = doc.add_table(rows=1, cols=n_cols)
        vst_tbl.style = 'Table Grid'
        col_w = [2.0] + [4.5] * len(ph_cols) + [3.0]
        _set_col_widths(vst_tbl, col_w[:n_cols])
        hdrs = ["Order"] + [ph_labels[p] + " worst-day P99" for p in ph_cols] + ["2.0× Limit"]
        for cell, txt in zip(vst_tbl.rows[0].cells, hdrs):
            _cell_shade(cell, _CHROME_BAND)
            r = cell.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(9)
        for key, label in stat_orders:
            ph_data = hs["daily_vst"].get(key, {})
            if not any(ph in ph_data for ph in ph_cols):
                continue
            row = vst_tbl.add_row()
            row.cells[0].paragraphs[0].add_run(label).font.size = Pt(9)
            any_fail_row = False
            limit_shown = False
            for j, ph in enumerate(ph_cols):
                d = ph_data.get(ph, {})
                if not d:
                    continue
                val = d.get("p99", 0.0)
                passes = d.get("pass", True)
                margin = d.get("margin", 0.0)
                margin_str = f" (+{margin:.2f})" if margin >= 0 else f" ({margin:.2f})"
                day = d.get("worst_day", "")
                txt = f"{val:.2f}%{margin_str}\n({day})"
                run = row.cells[j + 1].paragraphs[0].add_run(txt)
                run.font.size = Pt(9)
                if not passes:
                    run.bold = True
                    run.font.color.rgb = _FAIL_CLR
                    any_fail_row = True
                if not limit_shown:
                    lim_val = d.get("limit_2x", 0.0)
                    row.cells[n_cols - 1].paragraphs[0].add_run(
                        f"{lim_val:.1f}%"
                    ).font.size = Pt(9)
                    limit_shown = True
            if any_fail_row:
                for cell in row.cells:
                    _cell_shade(cell, _sev_shade("severe"))
    if spec_img.exists() or has_kfactor:
        _section_heading(doc, "Harmonic Spectrum and Transformer Loading Impact", level=2)

    # Harmonic spectrum chart
    if spec_img.exists():
        ih_chart_hdr = doc.add_paragraph()
        _bold(ih_chart_hdr, "Current Harmonic Spectrum (Median over Recording Period)", size_pt=10)
        doc.add_picture(str(spec_img), width=Cm(15))

    # K-factor section
    if has_kfactor:
        doc.add_paragraph()
        kf_stats = report.get("kfactor") or {"available": False}
        if kf_stats.get("available"):
            kf_med, kf_min, kf_max = (kf_stats["median"], kf_stats["min"],
                                      kf_stats["max"])
            kf_phase = kf_stats["worst_phase"]
            kf_detail = ", ".join(f"phase {p} {_m(v['median'], '.1f')}"
                                  for p, v in sorted(kf_stats["phases"].items()))
        else:
            kf_med  = float(df["kfactor_meter"].median())
            kf_max  = float(df["kfactor_meter"].max())
            kf_min  = float(df["kfactor_meter"].min())
            kf_phase, kf_detail = "A", ""
        kf_rate, kf_interp = standard_k_rating(kf_med)
        _body(doc,
            f"The meter-measured harmonic K-factor (IEEE C57.110) over the recording "
            f"period, taken from the worst phase ({kf_phase}) because harmonic heating "
            f"is per-winding: median {_m(kf_med, '.1f')}, minimum {_m(kf_min, '.1f')}, maximum "
            f"{_m(kf_max, '.1f')}."
            + (f" By phase: {kf_detail}. " if kf_detail else " ")
            + f"Standard distribution transformers are designed for K=1 (sinusoidal load). "
            f"Harmonic currents cause additional eddy-current and hysteresis losses in the "
            f"transformer core and windings, reducing rated capacity and accelerating insulation "
            f"aging. {kf_interp}"
        )
        if dem.get("transformer"):
            tx     = dem["transformer"]
            pct_tx = tx.get("pct_nameplate", 0)
            # Harmonic heating scales with load current, so a K-factor only puts
            # a transformer near its thermal limit when the transformer is
            # actually loaded. Stating otherwise recommended replacing a unit
            # sitting at 27% of nameplate, in a report that had already called
            # its loading acceptable two sections earlier.
            if not tx.get("dedicated"):
                # pct_nameplate is this service's contribution, not the
                # transformer's loading, so it cannot decide either branch
                # below. The "> 70" test would read a house at 30% as lightly
                # loaded on a transformer that may sit at 95%, and the else
                # branch would then assert thermal margin nobody measured.
                _body(doc,
                    f"This service contributes {_m(pct_tx, '.0f', '%')} of the "
                    f"{tx['nameplate_kva']:.0f} kVA nameplate at its 8-hour "
                    f"peak, and draws a K-factor of {_m(kf_med, '.1f')}. Harmonic "
                    "heating scales with load current, so whether this "
                    "K-factor is thermally significant depends on the "
                    "transformer's total loading — which includes other "
                    "customers on the same transformer and was not measured "
                    "here. This recording establishes the harmonic character "
                    "of this service's load, not the transformer's thermal "
                    "margin."
                )
            elif pct_tx > 70:
                _body(doc,
                    f"With the transformer loaded to {_m(pct_tx, '.0f', '%')} of its "
                    f"{tx['nameplate_kva']:.0f} kVA nameplate and a K-factor of "
                    f"{_m(kf_med, '.1f')}, the effective thermal load exceeds what the "
                    "nameplate rating assumes. "
                    + (f"A K-{kf_rate} rated unit is indicated before any "
                       "additional load is added."
                       if kf_rate else
                       "No standard K-rating covers a K-factor this high, so "
                       "re-assess it under representative load before specifying "
                       "a replacement.")
                )
            else:
                _body(doc,
                    f"The transformer is loaded to {_m(pct_tx, '.0f', '%')} of its "
                    f"{tx['nameplate_kva']:.0f} kVA nameplate, so despite a "
                    f"K-factor of {_m(kf_med, '.1f')} it retains substantial thermal "
                    "margin: harmonic losses scale with load current, and at this "
                    "loading the additional heating is small in absolute terms. "
                    + (f"The K-factor becomes the governing constraint if load "
                       f"grows — a replacement at that point should be K-{kf_rate} "
                       "rated rather than standard."
                       if kf_rate else
                       "Re-assess the K-factor under representative load before "
                       "it drives any sizing decision.")
                )

    doc.add_paragraph()


#: How the two methods' verdicts are said out loud. The wording names a side
#: of the meter and stops there: which side a harmonic source sits on is a
#: measurement, whose equipment it is and what follows from that is the
#: engineer's call, written under "Engineer's assessment".
_DIRECTION_LABEL = {
    "downstream":    "Customer side",
    "upstream":      "Utility side",
    "mixed":         "Split — no clear side",
    "conflicting":   "The two methods disagree",
    "indeterminate": "Indeterminate",
    "not_assessed":  "Not assessed",
}

#: The same verdicts inside a sentence, where a table label reads wrong.
_DIRECTION_PHRASE = {
    "downstream":    "the customer side",
    "upstream":      "the utility side",
    "mixed":         "no single side across the orders",
    "conflicting":   "opposite sides",
    "indeterminate": "no clear side",
    "not_assessed":  "no assessment — the spectrum was too close to the meter's "
                     "resolution to read",
}


def _word_harmonic_direction(doc, report, thresh) -> None:
    """Where the harmonic distortion at the meter appears to originate.

    Two methods are printed side by side rather than merged into one verdict.
    They answer the same question from different data -- the whole recording
    versus the captured instants -- and where they disagree that is itself the
    finding, since a customer source that only appears in triggered captures
    and a background source that only shows in the trend are different
    situations with different next steps.
    """
    hd = report.get("harmonic_direction") or {}
    if not hd.get("available") or thresh.customer_class == "r":
        return

    iv = hd.get("interval") or {}
    wf = hd.get("waveform") or {}

    _section_heading(doc, "Harmonic Source Direction", level=2)
    _body(doc,
        "Harmonic distortion measured at the meter can originate on either "
        "side of it: from equipment inside the premises drawing distorted "
        "current, or from the distribution system delivering an already "
        "distorted voltage. The two readings below answer that question from "
        "different data and are reported as evidence of direction only — they "
        "do not assign responsibility for the distortion or the cost of "
        "correcting it."
    )

    # ── Method 1: the whole recording, from magnitudes ───────────────────────
    if iv.get("available"):
        hdr = doc.add_paragraph()
        _bold(hdr, "Across the whole recording (harmonic voltage vs. harmonic current)",
              size_pt=10)
        _body(doc,
            "For each order the harmonic voltage at the meter is fitted "
            "against the harmonic current flowing through it: V_h = Z_h·I_h + "
            "V_background. The slope is the part of the distortion the "
            "premises' own current accounts for; the intercept is what remains "
            "when that current falls away, which is distortion arriving from "
            "the system. The last column before the indication is the same "
            "quantity measured rather than extrapolated — the average harmonic "
            "voltage over the tenth of intervals drawing the least harmonic "
            "current."
        )
        tbl = doc.add_table(rows=1, cols=7)
        tbl.style = 'Table Grid'
        _set_col_widths(tbl, [1.3, 2.1, 2.4, 2.4, 2.4, 1.6, 2.8])
        for cell, text in zip(tbl.rows[0].cells,
                              ["Order", "Z_h (Ω)", "From load (V)",
                               "Background (V)", "At quietest 10% (V)",
                               "r", "Indication"]):
            _cell_shade(cell, _CHROME_BAND)
            cell.paragraphs[0].add_run(text).bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for h, od in sorted((iv.get("orders") or {}).items()):
            cells = tbl.add_row().cells
            quiet = od.get("v_at_quiet_v")
            corr = od.get("corr")
            values = [
                f"H{h}",
                f"{od['slope_ohm']:.3f}",
                f"{od['v_from_load_v']:.2f}",
                f"{od['v_background_v']:.2f}",
                f"{quiet:.2f}" if quiet is not None else "—",
                f"{corr:.2f}" if corr is not None else "—",
                _DIRECTION_LABEL.get(od.get("indication"), od.get("indication", "")),
            ]
            for cell, text in zip(cells, values):
                cell.paragraphs[0].add_run(text).font.size = Pt(9)
        if iv.get("overall") == "not_assessed":
            _body(doc, iv.get("note", ""))

    # ── Method 2: the captures, from the sign of harmonic power ──────────────
    if wf.get("available"):
        doc.add_paragraph()
        hdr = doc.add_paragraph()
        _bold(hdr, "At the point-on-wave captures (sign of harmonic power)",
              size_pt=10)
        _body(doc,
            "The captures record voltage and current at the same instants, so "
            "each harmonic has a phase angle and its power P_h = ½·Re(V_h·I_h*) "
            "has a direction. Power flowing out of the premises at a harmonic "
            "order means the source of that order is inside; power flowing in "
            "means it is upstream. "
            f"{wf['captures_used']} of {wf['captures_total']} captures were "
            "usable"
            + _direction_exclusions(wf)
            + (f", measured against a fundamental of {_m(wf['fundamental_hz'], '.2f', ' Hz')}."
               if wf.get("fundamental_hz") else ".")
        )
        if wf.get("polarity_note"):
            _body(doc, wf["polarity_note"])

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = 'Table Grid'
        _set_col_widths(tbl, [1.3, 2.4, 2.6, 2.6, 2.4, 3.7])
        for cell, text in zip(tbl.rows[0].cells,
                              ["Order", "Readings", "Toward system",
                               "Median P_h (W)", "Median V–I angle", "Indication"]):
            _cell_shade(cell, _CHROME_BAND)
            cell.paragraphs[0].add_run(text).bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for h, od in sorted((wf.get("orders") or {}).items()):
            cells = tbl.add_row().cells
            values = [
                f"H{h}",
                f"{od['samples']}",
                f"{od['toward_system']} of {od['samples']}",
                _signed_watts(od['median_p_w']),
                f"{od['median_angle_deg']:.0f}°",
                _DIRECTION_LABEL.get(od.get("indication"), od.get("indication", "")),
            ]
            for cell, text in zip(cells, values):
                cell.paragraphs[0].add_run(text).font.size = Pt(9)

        # The exporting captures, on a generating service. They are a separate
        # table rather than extra rows because they are a different population
        # -- the site is running as a source in them -- and averaging the two
        # together would describe an operating state the service never sat in.
        exporting = (wf.get("export_split") or {}).get("exporting") or {}
        if exporting.get("orders"):
            doc.add_paragraph()
            _body(doc,
                  f"While exporting ({exporting['capture_phases']} phase-captures). "
                  "The service is generating in these, so an order flowing out "
                  "of the premises here is carried on the same path as the "
                  "exported fundamental and does not by itself place the source "
                  "inside the customer's system.")
            tbl2 = doc.add_table(rows=1, cols=6)
            tbl2.style = 'Table Grid'
            _set_col_widths(tbl2, [1.3, 2.4, 2.6, 2.6, 2.4, 3.7])
            for cell, text in zip(tbl2.rows[0].cells,
                                  ["Order", "Readings", "Toward system",
                                   "Median P_h (W)", "Median V–I angle",
                                   "Indication"]):
                _cell_shade(cell, _CHROME_BAND)
                cell.paragraphs[0].add_run(text).bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            for h, od in sorted(exporting["orders"].items()):
                cells = tbl2.add_row().cells
                values = [
                    f"H{h}",
                    f"{od['samples']}",
                    f"{od['toward_system']} of {od['samples']}",
                    _signed_watts(od['median_p_w']),
                    f"{od['median_angle_deg']:.0f}°",
                    _DIRECTION_LABEL.get(od.get("indication"),
                                         od.get("indication", "")),
                ]
                for cell, text in zip(cells, values):
                    cell.paragraphs[0].add_run(text).font.size = Pt(9)
    elif wf.get("note"):
        doc.add_paragraph()
        hdr = doc.add_paragraph()
        _bold(hdr, "At the point-on-wave captures (sign of harmonic power)",
              size_pt=10)
        _body(doc, wf["note"])

    # ── What the two together support ────────────────────────────────────────
    doc.add_paragraph()
    hdr = doc.add_paragraph()
    _bold(hdr, "Combined reading", size_pt=10)
    _body(doc, _direction_summary_sentence(hd))
    _body(doc,
        "Limitations: the sign of harmonic power identifies the side whose "
        "harmonic source dominates at the meter and is least decisive when the "
        "impedances either side are comparable (Xu, Liu & Liu, IEEE "
        "Transactions on Power Delivery 18(1), 2003); the captures are "
        "triggered snapshots rather than a continuous record; and the "
        "regression infers direction from how distortion co-varies with load "
        "rather than measuring it. Neither method establishes what equipment "
        "is responsible."
    )
    _write_in_field(doc, "Engineer's assessment:", lines=2, label_size_pt=11)
    doc.add_paragraph()


def _signed_watts(p: float) -> str:
    """Harmonic power with its sign, without printing a negative zero.

    The orders that matter most are often fractions of a watt, and "-0.00"
    reads as a direction when it is really the absence of one.
    """
    if abs(p) < 0.005:
        return "≈0"
    return f"{p:+.2f}"


def _direction_exclusions(wf: dict) -> str:
    """Why the unusable captures were unusable, in one clause."""
    reasons = [
        (wf.get("excluded_event", 0), "taken during a voltage event"),
        (wf.get("excluded_light_load", 0), "at too little load to measure"),
        (wf.get("excluded_short", 0), "too short to resolve the orders"),
        (wf.get("excluded_no_fundamental", 0), "without a steady fundamental"),
    ]
    named = [f"{n} {why}" for n, why in reasons if n]
    return f" (excluded: {', '.join(named)})" if named else ""


def _direction_summary_sentence(hd: dict) -> str:
    """One sentence on what both methods together do and do not support."""
    iv_overall = (hd.get("interval") or {}).get("overall")
    wf_overall = (hd.get("waveform") or {}).get("overall")
    agreement = hd.get("agreement") or {}
    agreed = sorted(h for h, v in agreement.items() if v == "agree")
    disagreed = sorted(h for h, v in agreement.items() if v == "disagree")

    if iv_overall == wf_overall:
        parts = [
            "Both readings — the harmonic voltage across the recording and the "
            "harmonic power at the captures — point to "
            f"{_DIRECTION_PHRASE.get(iv_overall, 'no clear side')}."
        ]
    else:
        parts = [
            "Across the recording the harmonic voltage points to "
            f"{_DIRECTION_PHRASE.get(iv_overall, 'no clear side')}; "
            "at the captures the harmonic power points to "
            f"{_DIRECTION_PHRASE.get(wf_overall, 'no clear side')}."
        ]
    if agreed:
        parts.append(
            "The two methods agree at "
            + ", ".join(f"H{h}" for h in agreed)
            + ", which is the strongest statement this recording supports: "
            "an inference drawn from the whole period and a direct measurement "
            "at the captured instants reaching the same answer."
        )
    if disagreed:
        parts.append(
            "They disagree at "
            + ", ".join(f"H{h}" for h in disagreed)
            + ". That is not resolved here: the two look at different windows, "
            "and a source that is present only during the captured instants — "
            "or only outside them — produces exactly this split. Repeat "
            "monitoring, or captures triggered on harmonic content rather than "
            "on voltage events, would separate the two."
        )
    if not agreed and not disagreed:
        parts.append(
            "Only one method returned a direction, so the reading rests on "
            "that method's assumptions alone."
        )
    return " ".join(parts)


_IMPEDANCE_HEADLINE = {
    "high_impedance_suspected": "High impedance suspected",
    "elevated":                 "Above what the service accounts for",
    "consistent_with_expected": "Consistent with the service as built",
    "measured_only":            "Measured; no expected value to compare against",
    "not_measurable":           "Not measurable from this recording",
}


def _primary_sequence_note(expected: dict) -> str:
    """Which sequence impedance the comparison used, and why that one.

    Named explicitly rather than left implied: a primary line's Z0 differs from
    its Z1 by a factor of two or three, so a reader who assumes the wrong one
    misreads every figure beside it.
    """
    primary = expected.get("primary") or {}
    note = (
        "This service is metered on the primary, so the transformer and the "
        "secondary conductors belong to the customer and sit downstream of the "
        "meter. Neither is in the path measured here; the primary line is, and "
        "its impedance is as entered from a planning model or fault study "
        "rather than from any table in this tool. The comparison uses positive "
        "sequence, because balanced load current flows in positive sequence "
        "and the measurement below is a per-phase voltage-against-current fit, "
        "which on balanced load is that same loop."
    )
    if primary.get("z0_ohm") is None:
        return note + (" No zero-sequence impedance was entered, which the "
                       "comparison does not need.")
    ratio = primary.get("z0_over_z1")
    extra = (
        f" Zero sequence was also entered, at {primary['z0_ohm']:.4f} Ω"
        + (f" — {ratio:.1f}× the positive-sequence value" if ratio else "")
        + ", the usual relationship for a line whose zero-sequence return is "
        "through earth. It is not what the comparison below uses. It is the "
        "right impedance for two things this recording may raise elsewhere: "
        "the triplen harmonics — the 3rd, 9th and 15th — which are zero-"
        "sequence on a balanced system, and the earth-return path taken by "
        "whatever part of an unbalanced load current does not cancel between "
        "phases."
    )
    loop = primary.get("single_phase_loop_ohm")
    if loop:
        extra += (f" A single-phase load tapped off this line would see "
                  f"{loop:.4f} Ω, which is (2·Z1 + Z0)/3 rather than Z1.")
    return note + extra


def _word_service_impedance(doc, report, thresh) -> Optional[str]:
    """Series impedance from the source to the meter, measured and expected.

    The section leads with the two readings that need no conductor table --
    one phase against the others, and the neutral against itself -- because
    they hold whatever the engineer typed into the picker, and a comparison
    against a generic conductor constant should never outrank a measurement
    the service made against itself.
    """
    si = report.get("service_impedance") or {}
    if not si.get("available"):
        return "Service Impedance"

    _section_heading(doc, "Service Impedance and High-Impedance Screening", level=2)
    _body(doc,
        "Voltage that falls when current rises is measuring the impedance "
        "between the source and the meter. It is fitted here on the steps in "
        "load rather than on the levels: the whole neighbourhood's load rises "
        "at the same hours as this customer's, and a fit against the levels "
        "attributes that shared droop to this service. Only intervals where "
        "the current stepped are used, so slow feeder movement drops out."
    )

    # ── per phase ────────────────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = 'Table Grid'
    _set_col_widths(tbl, [1.6, 2.2, 2.2, 2.2, 2.4, 4.9])
    for cell, text in zip(tbl.rows[0].cells,
                          ["Phase", "R (Ω)", "X (Ω)", "|Z| (Ω)",
                           "Load steps", "Basis"]):
        _cell_shade(cell, _CHROME_BAND)
        cell.paragraphs[0].add_run(text).bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for ph, fit in sorted(si.get("phases", {}).items()):
        cells = tbl.add_row().cells
        if fit.get("identifiable"):
            basis = (f"{fit['consistency']:.0%} of steps moved the voltage the "
                     "expected way"
                     + ("" if fit.get("separated") else
                        "; power factor too steady to separate R from X"))
            values = [ph.upper(),
                      f"{fit['r_ohm']:.4f}" if fit.get("r_ohm") is not None else "—",
                      f"{fit['x_ohm']:.4f}" if fit.get("x_ohm") is not None else "—",
                      f"{fit['z_ohm']:.4f}", str(fit.get("steps", "—")), basis]
        else:
            values = [ph.upper(), "—", "—", "—", str(fit.get("steps", "—")),
                      fit.get("reason", "Not measurable")]
        for cell, text in zip(cells, values):
            cell.paragraphs[0].add_run(strip_marks(text)).font.size = Pt(9)

    # ── one phase against the others ─────────────────────────────────────────
    asym = si.get("asymmetry") or {}
    if asym.get("ratio"):
        para = doc.add_paragraph()
        _bold(para, "Between phases: ")
        if asym["flagged"]:
            _normal(para,
                f"phase {asym['worst_phase']} measures {_m(asym['ratio'], '.1f')}× the "
                f"impedance of phase {asym['best_phase']}, worth "
                f"{_m(asym['excess_v_at_peak'], '.1f', ' V')} of extra drop at peak load. "
                "A single connection — a lug, a splice, a crimp — degrading on "
                "one phase is what produces this; the conductors and the "
                "transformer are common to all phases and cannot.",
                color=_FAIL_CLR)
        else:
            _normal(para,
                f"phase {asym['worst_phase']} measures {_m(asym['ratio'], '.1f')}× "
                f"phase {asym['best_phase']} ({_m(asym['excess_v_at_peak'], '.1f', ' V')} "
                "at peak load), which is within what unequal loading and "
                "measurement scatter account for. No single-phase connection "
                "problem is indicated.")

    # ── the neutral, measured against itself ─────────────────────────────────
    neutral = si.get("neutral") or {}
    if neutral:
        para = doc.add_paragraph()
        _bold(para, "Neutral connection: ")
        if neutral.get("identifiable"):
            text = (f"neutral-to-earth voltage rises with neutral current at "
                    f"{_m(neutral['r_ohm'], '.4f', ' Ω')}, which is "
                    f"{_m(neutral['rise_at_peak_v'], '.1f', ' V')} at the "
                    f"{_m(neutral['i_peak_a'], '.0f', ' A')} peak neutral current.")
            if neutral.get("elevated"):
                _normal(para, text + " A sound neutral holds this near zero; "
                        "this is the signature of resistance in the neutral "
                        "path — a loose connection, a corroded splice, or a "
                        "degraded ground.", color=_FAIL_CLR)
            else:
                _normal(para, text + " That is low enough to indicate a sound "
                        "neutral connection.")
        elif neutral.get("at_resolution"):
            _normal(para,
                "neutral-to-earth voltage never moved by as much as the "
                f"meter's {_VOLTAGE_RESOLUTION_V} V resolution across "
                f"{neutral.get('steps', 0)} load steps. Any resistance in the "
                "neutral is below what this recording can resolve, which is "
                "the reading a sound neutral gives.")
        else:
            _normal(para, neutral.get("reason", "Not measurable."))

    # ── measured against expected ────────────────────────────────────────────
    expected = si.get("expected") or {}
    comparison = si.get("comparison") or {}
    doc.add_paragraph()
    hdr = doc.add_paragraph()
    _bold(hdr, "Against the service as built", size_pt=10)
    if not expected.get("available"):
        _body(doc, expected.get("reason", "No expected impedance was built."))
    else:
        parts = []
        if expected.get("upstream_ohm") is not None:
            parts.append(f"{expected['upstream_ohm']:.4f} Ω from "
                         f"{expected.get('upstream_source', 'upstream')}")
        if expected.get("shared_secondary_z_ohm") is not None:
            parts.append(
                f"{expected['shared_secondary_z_ohm']:.4f} Ω from "
                f"{expected.get('shared_secondary_ft', 0):.0f} ft of "
                f"{expected.get('shared_secondary_label')} shared secondary")
        if expected.get("conductor_z_ohm") is not None:
            parts.append(
                f"{expected['conductor_z_ohm']:.4f} Ω from "
                f"{expected.get('run_length_ft', 0):.0f} ft of "
                f"{expected.get('conductor_label')} counted "
                f"{expected.get('conductor_path', '')}")
        _body(doc, "Expected impedance is " + ", plus ".join(parts)
              + f" — {expected.get('total_ohm', 0):.4f} Ω in total."
              + (f" This omits {expected['partial']}, so it is a floor."
                 if expected.get("partial") else ""))
        if expected.get("primary_metered"):
            _body(doc, _primary_sequence_note(expected))
        else:
            _body(doc,
                "The conductor constants are generic published values (NEC "
                "Chapter 9 Table 8 resistance at 75 °C, with a typical reactance "
                "for the construction), not PSCo Blue Book figures, and the run "
                "length is as entered. Treat the comparison as an order-of-"
                "magnitude check, not a tolerance.")
        if expected.get("shared_secondary_z_ohm") is not None:
            _body(doc,
                "Part of that path is a secondary main shared with other "
                "services. This customer's current flows through it, so its "
                "drop is in the measurement and belongs in the expected value. "
                "The neighbours' current flows through it too, and their share "
                "of the drop moves independently of this customer's load — it "
                "widens the scatter in the fit below rather than shifting it, "
                "so an impedance measured here is less precise than one "
                "measured on a dedicated run.")
        if comparison:
            # The measured impedance is marked; the expected one beside it is
            # not, which is the whole point of marking -- these two sentences
            # put a reading and a calculation in the same breath, and only one
            # of them came off the meter.
            verdict_text = {
                "high": (f"The measured {_m(comparison['measured_ohm'], '.4f', ' Ω')} is "
                         f"{_m(comparison['ratio'], '.1f')}× the expected "
                         f"{comparison['expected_ohm']:.4f} Ω — "
                         f"{_m(comparison['excess_v_at_peak'], '.1f', ' V')} of drop at "
                         f"the {_m(comparison['i_peak_a'], '.0f', ' A')} peak beyond what "
                         "the transformer and conductor account for."),
                "elevated": (f"The measured {_m(comparison['measured_ohm'], '.4f', ' Ω')} "
                             f"is {_m(comparison['ratio'], '.1f')}× the expected "
                             f"{comparison['expected_ohm']:.4f} Ω, worth "
                             f"{_m(comparison['excess_v_at_peak'], '.1f', ' V')} at peak "
                             "load."),
                "consistent": (f"The measured {_m(comparison['measured_ohm'], '.4f', ' Ω')} "
                               f"sits within the range the expected "
                               f"{comparison['expected_ohm']:.4f} Ω accounts "
                               "for."),
                "below_expected": (
                    f"The measured {_m(comparison['measured_ohm'], '.4f', ' Ω')} is below "
                    f"the expected {comparison['expected_ohm']:.4f} Ω, which "
                    "usually means the run is shorter or the conductor larger "
                    "than what was entered, or the source stiffer than the "
                    "Blue Book figure."),
            }.get(comparison["verdict"], "")
            para = doc.add_paragraph()
            _normal(para, verdict_text,
                    color=_FAIL_CLR if comparison["verdict"] == "high" else None)

    _write_in_field(doc, "Engineer's assessment:", lines=2, label_size_pt=11)
    doc.add_paragraph()
    return None


#: IEEE 1453-2015 Table 2 compatibility levels at low voltage. Not what the
#: pass/fail above is judged against -- that stays IEC 61000-3-3 -- but the
#: number a reader needs to know exists before treating a Plt of 0.7 as a
#: failure of the supply system rather than of an equipment emission limit.
_IEEE1453_LV_PST = 1.0
_IEEE1453_LV_PLT = 0.8


def _word_flicker(doc, report, df, outdir=None, stem="") -> Optional[str]:
    """Flicker severity, short-term and long-term, on every phase measured.

    The two are different measurements of the same phenomenon over different
    windows, and the section says which is which: an engineer reading it
    should not have to already know why there are two numbers, and the
    distinction decides what to do about them -- a high Pst with a low Plt is
    a short burst, the reverse is a load cycling for hours.
    """
    fl = report.get("flicker") or {}
    if not fl.get("available"):
        return "Voltage Flicker"

    _section_heading(doc, "Voltage Flicker (IEC 61000-3-3)", level=2)
    _body(doc,
        "Flicker severity is reported as two numbers because it is measured "
        "over two windows. Pst, short-term severity, is computed over 10 "
        "minutes and scaled so that 1.0 is the level at which half of "
        "observers watching an incandescent lamp find the flicker "
        "objectionable — it catches motor starts, welders and other short "
        "bursts. Plt, long-term severity, aggregates twelve consecutive Pst "
        "values into one figure covering two hours (the cube root of their "
        "mean cube, so one bad ten minutes still shows), and it catches loads "
        "that cycle for hours: heat pumps, compressors, arc furnaces. A high "
        "Pst with a low Plt is a brief disturbance; a Plt near its limit with "
        "modest Pst values is something running repeatedly."
    )

    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = 'Table Grid'
    _set_col_widths(tbl, [2.6, 1.6, 2.0, 2.0, 2.0, 2.6, 2.7])
    for cell, text in zip(tbl.rows[0].cells,
                          ["Measure", "Phase", "Median", "95th pct", "Max",
                           "% of time over", "Result"]):
        _cell_shade(cell, _CHROME_BAND)
        cell.paragraphs[0].add_run(text).bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for kind, label, limit in (("pst", f"Pst (10 min)", fl["pst_limit"]),
                               ("plt", f"Plt (2 h)", fl["plt_limit"])):
        for phase, stats in sorted(fl.get(kind, {}).items()):
            cells = tbl.add_row().cells
            values = [f"{label}, limit {limit:.2f}", phase,
                      f"{stats['median']:.2f}", f"{stats['p95']:.2f}",
                      f"{stats['max']:.2f}", f"{stats['pct_exceeding']:.1f}%",
                      "PASS" if stats["pass"] else "FAIL"]
            for cell, text in zip(cells, values):
                run = cell.paragraphs[0].add_run(text)
                run.font.size = Pt(9)
            if not stats["pass"]:
                for cell in cells:
                    _cell_shade(cell, _sev_shade("severe"))
                cells[-1].paragraphs[0].runs[0].bold = True
                cells[-1].paragraphs[0].runs[0].font.color.rgb = _FAIL_CLR

    # Severity, separately for each measure. Pass/fail above is decided by the
    # maximum, which one bad ten-minute window is enough to fail; the band here
    # is graded at the 95th percentile with the share of time over the limit,
    # so a single anomaly does not read as a sustained condition.
    severities = _flicker_severities(fl, report)
    for kind, label in (("pst", "Pst"), ("plt", "Plt")):
        graded = severities.get(f"flicker_{kind}")
        if not graded or graded["band"] == "not_assessed":
            continue
        para = doc.add_paragraph()
        _bold(para, f"{label} severity: ")
        _normal(para, graded["label"] + (f" — {graded['reason']}"
                                         if graded["reason"] else ""),
                color=_sev_color(graded["band"]))

    failures = [(kind, phase, stats)
                for kind in ("pst", "plt")
                for phase, stats in sorted(fl.get(kind, {}).items())
                if not stats["pass"]]
    if not failures:
        _body(doc,
            f"Both measures stayed within their limits on every phase "
            f"measured ({', '.join(fl.get('phases_read', []))}). No "
            "objectionable lamp flicker from this service is expected."
        )
    else:
        worst = ", ".join(
            f"{'Pst' if k == 'pst' else 'Plt'} reached {_m(s['max'], '.2f')} on phase "
            f"{p} against a {fl[k + '_limit']:.2f} limit"
            for k, p, s in failures)
        _body(doc,
            f"Flicker severity exceeded IEC 61000-3-3 limits: {worst}. Flicker "
            "complaints are plausible under these conditions. Common causes "
            "are arc furnaces, large motor starting, welders and rapidly "
            "cycling loads. Separating a source inside the premises from one "
            "on the system requires voltage measurement at the service "
            "entrance with the customer's load disconnected."
        )
        if any(k == "plt" for k, _p, _s in failures) and not any(
                k == "pst" for k, _p, _s in failures):
            _body(doc,
                "Only the two-hour measure is over limit while every "
                "ten-minute value stays inside it, which is the signature of "
                "a load that cycles repeatedly rather than one large "
                "disturbance."
            )

    # The limits above are equipment emission limits; the numbers a supply
    # system is held to are different, and a reader comparing against the
    # wrong one draws the wrong conclusion.
    _body(doc,
        f"Limits applied are IEC 61000-3-3, which governs what a piece of "
        f"equipment may emit: Pst {fl['pst_limit']:.2f} and Plt "
        f"{fl['plt_limit']:.2f}. The compatibility levels IEEE 1453-2015 sets "
        f"for a low-voltage supply system are higher — Pst {_IEEE1453_LV_PST:.1f} "
        f"and Plt {_IEEE1453_LV_PLT:.1f} — so a Plt between "
        f"{fl['plt_limit']:.2f} and {_IEEE1453_LV_PLT:.1f} is over the "
        "equipment limit without exceeding what the supply system is expected "
        "to hold. Both standards assess against the 95th percentile over a "
        "week; this recording covers "
        f"{_m(report['file_summary']['duration_hours'] / 24, '.1f')} day(s), so the "
        "percentiles above describe the period recorded and not a week."
    )
    _embed_plot(doc, outdir, stem, "flicker.png",
                "Flicker severity over the recording. The bottom panel sorts "
                "every reading from worst to best against the share of the "
                "recording at or above it: a curve that falls below the limit "
                "line close to the left edge was over limit only briefly, "
                "while one still above it further right was over limit for "
                "that share of the time. The dotted line marks the 95th "
                "percentile, where both standards assess and where the "
                "severity above is graded.")
    doc.add_paragraph()
    return None


def _word_neutral_health(doc, report, thresh, outdir=None, stem="") -> Optional[str]:
    nh = report.get("neutral_health", {})
    if not nh.get("available"):
        return None   # split-phase-only diagnostic — suppress silently elsewhere

    _section_heading(doc, "Neutral Integrity Assessment", level=2)

    # The band word comes from the shared scale; only the finding after the dash
    # is particular to the neutral.
    sev = nh.get("severity", "normal")
    sev_findings = {
        "critical": "Open or High-Resistance Neutral Suspected",
        "warning":  "Neutral Integrity Concern",
        "caution":  "Neutral Anomaly Detected",
        "normal":   "Neutral Appears Healthy",
    }
    sev_p = doc.add_paragraph()
    _bold(sev_p,
          f"{_sev_label(sev).upper()} — {sev_findings.get(sev, 'Neutral Assessed')}",
          color=_sev_color(sev), size_pt=11)

    indicators = [
        ("L1 + L2 Sum (mean / std)",
         f"{nh['sum_mean_v']:.1f} V / {nh['sum_std_v']:.2f} V",
         (f"Healthy: ~{nh.get('healthy_sum_v', 240):.0f} V, std < 1 V"
          + (f"; open neutral → ~{nh.get('open_neutral_sum_v', 0):.0f} V"
             if nh.get("sum_is_diagnostic") else
             "; same either way on this service"))),
        ("L1–L2 Correlation (Pearson r)",
         (f"{nh['leg_correlation']:.3f}"
          if nh.get("leg_correlation") is not None else "not computable"),
         "Healthy: r > 0.80; open neutral → r ≈ −1"),
        ("Voltage Asymmetry |L1 − L2|",
         f"{nh['asym_mean_v']:.1f} V mean, {nh['asym_max_v']:.1f} V max ({nh['asym_pct']:.1f}%)",
         "Healthy: < 2% of nominal"),
        ("Coincident Opposing Events",
         str(nh["coincident_events"]),
         "Healthy: 0"),
    ]
    if nh.get("vne_available"):
        indicators.append((
            "Neutral-to-Earth Voltage (Vne)",
            f"{nh['vne_mean_v']:.2f} V mean, {nh['vne_max_v']:.2f} V max",
            "Healthy: < 0.5 V; > 5 V is safety hazard",
        ))

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [3.0, 3.0, 3.5])
    for cell, hdr_txt in zip(tbl.rows[0].cells, ["Indicator", "Measured", "Benchmark"]):
        _cell_shade(cell, _CHROME_BAND)
        cell.paragraphs[0].add_run(hdr_txt).bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for ind, val, bench in indicators:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].add_run(ind).font.size = Pt(9)
        cells[1].paragraphs[0].add_run(val).font.size = Pt(9)
        cells[2].paragraphs[0].add_run(bench).font.size = Pt(9)

    doc.add_paragraph()

    for finding in nh.get("findings", []):
        p = doc.add_paragraph(style="List Bullet")
        _normal(p, finding, size_pt=10)

    if sev in ("critical", "warning"):
        doc.add_paragraph()
        rec_p = doc.add_paragraph()
        rec_p.paragraph_format.left_indent = Cm(0.5)
        rec_run = rec_p.add_run("Recommendation:  ")
        rec_run.bold = True
        rec_run.font.size = Pt(10)
        if sev == "critical":
            rec_text = (
                "An open or high-resistance neutral is a safety emergency. This can cause "
                "severe overvoltage on the lightly-loaded leg (potentially exceeding 200 V on "
                "a 120 V circuit), damaging appliances and posing shock and fire hazards. "
                "Contact Xcel Energy immediately to inspect the service neutral from the "
                "transformer secondary to the meter socket. Schedule a same-day inspection."
            )
        else:
            rec_text = (
                "Investigate the service neutral for loose connections, corrosion, or "
                "undersized conductors. Check neutral bar connections in the main panel and "
                "the meter socket lug. Xcel Energy should inspect the transformer secondary "
                "neutral and service drop. Re-measure after any repairs to confirm resolution."
            )
        _normal(rec_p, rec_text, size_pt=10)

    _embed_plot(doc, outdir, stem, "neutral_health.png",
                "Neutral integrity indicators: leg voltages, voltage sum stability, "
                "and leg asymmetry over the recording.")
    doc.add_paragraph()
    return None


def _word_imbalance(doc, report, thresh, outdir=None, stem="") -> Optional[str]:
    imb = report["voltage_imbalance"]
    ci  = report["current_imbalance"]

    if not imb["available"] and not ci["available"]:
        return "Voltage and Current Imbalance"
    _section_heading(doc, "Voltage and Current Imbalance", level=2)
    if imb["available"]:
        if imb["pct_exceeding"] == 0:
            _body(doc,
                f"Voltage imbalance was within the 3% limit throughout the recording. "
                f"Maximum {_m(imb['max_imbalance_pct'], '.2f', '%')}, mean {_m(imb['mean_imbalance_pct'], '.2f', '%')}."
            )
        else:
            _body(doc,
                f"Voltage imbalance exceeded 3% during {_mp(imb['pct_exceeding'], '.1f')} of the recording "
                f"(max {_m(imb['max_imbalance_pct'], '.2f', '%')}, mean {_m(imb['mean_imbalance_pct'], '.2f', '%')}). "
                "Distinguishing a supply asymmetry from an unbalanced load requires "
                "repeating the measurement with all customer loads disconnected."
            )

    if ci["available"]:
        nc_text = ""
        if "neutral_current" in ci:
            nc = ci["neutral_current"]
            nc_text = (
                f" Neutral current averaged {_m(nc['mean_amps'], '.1f', ' A')} "
                f"({_m(nc['mean_pct_of_phase'], '.1f', '%')} of phase average) with a peak of "
                f"{_m(nc['max_amps'], '.1f', ' A')} ({_m(nc['max_pct_of_phase'], '.1f', '%')})."
            )
            if is_single_phase_208(thresh.service_type):
                # The two legs of a 120/208 service are 120 degrees apart, so
                # their currents do not subtract in the neutral the way the two
                # legs of a 120/240 service do. With balanced load the neutral
                # carries essentially full leg current, and calling that
                # "elevated" would send an engineer looking for a fault that
                # is not there.
                nc_text += (
                    " This is a single-phase 120/208 service — two legs of a "
                    "three-phase transformer — so the legs are 120° apart "
                    "rather than the 180° of a 120/240 service. Their currents "
                    "add vectorially instead of subtracting, and with balanced "
                    "load the neutral carries roughly the same current as each "
                    "leg. A neutral near 100% of leg current is therefore "
                    "normal here and is not evidence of imbalance. Triplen "
                    "harmonics (3rd, 9th, 15th) also add in phase on this "
                    "configuration, so neutral current above leg current points "
                    "to harmonics rather than to a wiring fault."
                )
            elif nc["mean_pct_of_phase"] > 15:
                nc_text += (
                    " Elevated neutral current is consistent with load imbalance and/or significant "
                    "triplen harmonic currents (3rd, 9th, 15th) from nonlinear single-phase loads "
                    "such as computers, lighting controls, and variable-speed drives."
                )

        if ci["pct_exceeding"] == 0:
            _body(doc,
                f"Current imbalance was within the 10% limit throughout the recording. "
                f"Maximum {_m(ci['max_imbalance_pct'], '.2f', '%')}, mean {_m(ci['mean_imbalance_pct'], '.2f', '%')}.{nc_text}"
            )
        else:
            _body(doc,
                f"Current imbalance exceeded 10% during {_mp(ci['pct_exceeding'], '.1f')} of the recording "
                f"(max {_m(ci['max_imbalance_pct'], '.2f', '%')}, mean {_m(ci['mean_imbalance_pct'], '.2f', '%')}). "
                f"Correcting current imbalance requires redistributing single-phase "
                f"load across the phases.{nc_text}"
            )
    _embed_plot(doc, outdir, stem, "imbalance.png",
                "Voltage and current imbalance against limits, with neutral current.")
    doc.add_paragraph()
    return None


def _word_events(doc, report, outdir=None, stem="") -> Optional[str]:
    ev = report["events"]

    _section_heading(doc, "Event Detection Summary", level=2)
    adap_note = (
        " Event detection used cycle-level adaptive records (~17 ms resolution),"
        " which capture within-interval sags/swells missed by 5-minute averages."
        if ev.get("data_source") == "adaptive"
        else " Event detection used 5-minute interval averages."
    )
    if ev["event_count"] == 0:
        _body(doc,
            "No significant voltage sag, swell, transient, or flicker events were detected"
            " during the recording period." + adap_note
        )
    else:
        edf = ev["events"]
        type_counts = edf["type"].value_counts().to_dict() if len(edf) > 0 else {}
        parts = [f"{cnt} {etype.replace('_', ' ')}" for etype, cnt in sorted(type_counts.items())]
        _body(doc,
            f"{ev['event_count']} event(s) detected: {', '.join(parts)}." + adap_note + " "
            "Voltage event causes may include faults on adjacent feeders, motor starting inrush, "
            "transformer energization, or switching operations. "
            "Flicker events (PST > 1.0 or PLT > 0.65) indicate arc-type or intermittent loads."
        )

    n_wf = ev.get("waveform_captures", 0)
    if n_wf:
        _body(doc,
            f"The meter also recorded {n_wf} point-on-wave waveform captures "
            f"(instantaneous voltage and current snapshots triggered by disturbances). "
            f"These are analyzed for sub-cycle events and the most severe capture is "
            f"shown below.")

    itic = report.get("itic", {})
    if itic.get("available") and itic.get("n_events", 0) > 0:
        if itic["overall_pass"]:
            _body(doc,
                f"All {itic['n_events']} detected sag/swell events fall within the ITIC "
                "voltage tolerance envelope — IT and electronic equipment is expected "
                "to ride through them without disruption.")
        else:
            w = itic.get("worst") or {}
            w_txt = (f" The most severe was a {w.get('type', '').replace('voltage_', '')} to "
                     f"{_m(w.get('pct_nominal', 0), '.0f', '%')} of nominal lasting "
                     f"{_m(w.get('duration_ms', 0), '.0f', ' ms')}." if w else "")
            _body(doc,
                f"{itic['n_violations']} of {itic['n_events']} detected sag/swell events "
                f"fall outside the ITIC voltage tolerance envelope — sensitive electronic "
                f"equipment may misoperate, reset, or drop out during these events.{w_txt}")
    if _embed_plot(doc, outdir, stem, "itic_curve.png",
                   "Detected sag/swell events plotted on the ITIC voltage tolerance "
                   "curve. Events inside the no-disruption zone are unlikely to "
                   "affect IT equipment.", width_cm=13.5):
        pass
    _embed_plot(doc, outdir, stem, "waveform_worst.png",
                "Most severe point-on-wave capture in the recording: instantaneous "
                "voltage and current waveforms around the disturbance.")
    doc.add_paragraph()
    return None


def _word_measurement_review(doc, report, thresh, df, outdir=None, stem="") -> None:
    """Detailed Measurement Review — supporting measurements, after conclusions.
    Sections without usable data are suppressed and summarized in one line."""
    _section_heading(doc, "Detailed Measurement Review", level=1)
    _body(doc,
        "The measurements below provide the supporting evidence for the findings "
        "and recommendations above.")

    unavailable: List[str] = []
    for label in (
        _word_demand(doc, report, thresh, outdir, stem),
        _word_power_factor(doc, report, thresh, outdir, stem),
        _word_voltage(doc, report, outdir, stem),
        # Directly after the voltage it explains: a service impedance is what
        # turns the load into the voltage the section above reports.
        _word_service_impedance(doc, report, thresh),
        _word_flicker(doc, report, df, outdir, stem),
        _word_imbalance(doc, report, thresh, outdir, stem),
        _word_neutral_health(doc, report, thresh, outdir, stem),
        _word_events(doc, report, outdir, stem),
    ):
        if label:
            unavailable.append(label)

    if unavailable:
        if len(unavailable) == 1:
            _body(doc,
                f"{unavailable[0]} data was not available for this recording period.")
        else:
            _body(doc,
                "The following measurements were not available for this recording "
                "period: " + ", ".join(unavailable) + ".")
        doc.add_paragraph()


_FINDING_CONF_RANK = {"high": 0, "medium": 1, "low": 2}
_FINDING_SEV_RANK  = {"critical": 0, "warning": 1, "info": 2}


def _word_finding(doc, finding: dict, show_recommendation: bool = False) -> None:
    """Render one assessment finding. Shared by the assessment and Appendix D
    so an experimental finding is laid out the same way, only somewhere else.

    The assessment section does not print recommendations -- those are gathered
    into Recommended Actions. Appendix D does, because its findings are
    informational and never reach that list, so their advice would otherwise be
    dropped entirely.
    """
    sev  = finding["severity"]
    conf = finding.get("confidence", "")
    p = doc.add_paragraph()
    _bold(p, f"{_sev_label(sev).upper()}: {finding['title']}",
          color=_sev_color(sev), size_pt=10)
    tag = p.add_run(f"  [{conf.capitalize()} confidence]")
    tag.font.size = Pt(9)
    tag.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    body_p = doc.add_paragraph()
    body_p.paragraph_format.left_indent = Cm(0.5)
    run_f = body_p.add_run("Finding:  ")
    run_f.bold = True
    run_f.font.size = Pt(10)
    _normal(body_p, finding["finding"], size_pt=10)

    body_p2 = doc.add_paragraph()
    body_p2.paragraph_format.left_indent = Cm(0.5)
    run_c = body_p2.add_run("Likely cause:  ")
    run_c.bold = True
    run_c.font.size = Pt(10)
    _normal(body_p2, finding["cause"], size_pt=10)

    if finding.get("origin_evidence"):
        body_p3 = doc.add_paragraph()
        body_p3.paragraph_format.left_indent = Cm(0.5)
        run_e = body_p3.add_run("Evidence bearing on origin:  ")
        run_e.bold = True
        run_e.font.size = Pt(10)
        _normal(body_p3, finding["origin_evidence"], size_pt=10)

    if show_recommendation and finding.get("recommendation"):
        body_p4 = doc.add_paragraph()
        body_p4.paragraph_format.left_indent = Cm(0.5)
        run_r = body_p4.add_run("Candidate action:  ")
        run_r.bold = True
        run_r.font.size = Pt(10)
        _normal(body_p4, finding["recommendation"], size_pt=10)

    # The tool states evidence; attribution is the reviewing engineer's to
    # write, so the document carries an explicit place for it rather than
    # pre-empting it.
    _write_in_field(doc, "Engineer's assessment:", lines=2, indent_cm=0.5,
                    width_cm=15.5)

    doc.add_paragraph()


def _word_load_signature_appendix(doc, report) -> None:
    """Appendix D: load-signature matching, held apart as experimental.

    This is the only analysis in the report that tries to say what equipment is
    behind the meter, and it is the weakest. Measured against synthetic
    mixtures of library loads it named a load that was not present in a third
    to a half of the cases where it named one at all, which is why the finding
    reports a load *family* and not a device. It sits last, under its own
    heading, so it is read as a hypothesis to check on site rather than as one
    of the measured findings.
    """
    sigs = [f for f in report.get("root_causes", []) if f.get("experimental")]
    if not sigs:
        return

    _section_heading(doc, "Appendix D: Load Signature Matching (Experimental)",
                     level=1)
    _body(doc,
        "Everything in this appendix is experimental and is not part of the "
        "compliance assessment. It compares the measured harmonic spectrum "
        "against a library of reference load types and reports the closest "
        "family — the converter topology the shape is consistent with — where "
        "one is close enough to name.")
    _body(doc,
        "What it can and cannot establish: a meter at the service entrance "
        "measures the sum of everything behind it, so a spectrum that sits "
        "nearest one reference entry is not evidence that entry is what draws "
        "the current. A blend of two loads can land nearer a third entry that "
        "neither resembles. Tested against synthetic mixtures of two or three "
        "library loads on one service, matching to a specific device named a "
        "load that was not present in roughly a third to a half of the cases "
        "where it named one, so no device is named here — only the family, and "
        "only above a similarity floor set from a measured null distribution. "
        "Treat anything below as a hypothesis to confirm on site.")
    _body(doc,
        "The reliable way to attribute distortion to a particular load is a "
        "recording taken while that load is switched on and off, which "
        "separates its contribution directly rather than inferring it from "
        "spectral shape.")
    doc.add_paragraph()

    for finding in sorted(sigs, key=lambda f: _FINDING_CONF_RANK.get(
            f.get("confidence", "low"), 3)):
        _word_finding(doc, finding, show_recommendation=True)


def _word_engineering_assessment(doc, report) -> None:
    # Experimental findings are held back for Appendix D. Load-signature
    # matching is the only one: it names a load *family* from spectral shape,
    # which is a hypothesis to check on site rather than a likely cause, and
    # mixing it in here gave it the standing of the measured findings.
    rca = [f for f in report.get("root_causes", []) if not f.get("experimental")]
    if not rca:
        return

    _section_heading(doc, "Engineering Assessment: Likely Causes and Contributing Conditions",
                     level=1)
    _body(doc,
        "The findings below describe the likely causes of the observed conditions. "
        "Each is graded by confidence level and states the evidence bearing on "
        "where the condition originates, leaving the attribution itself to the "
        "reviewing engineer. "
        "Supporting measurements are presented in the sections that follow.")

    ordered = sorted(rca, key=lambda f: (
        _FINDING_SEV_RANK.get(f.get("severity"), 9),
        _FINDING_CONF_RANK.get(f.get("confidence", "low"), 3),
    ))
    for finding in ordered:
        _word_finding(doc, finding)


def _word_recommended_actions(doc, actions: List[dict]) -> None:
    _section_heading(doc, "Recommended Actions", level=1)

    if not actions:
        _body(doc,
            "No corrective actions are required at this time. All measured parameters "
            "are within applicable standards. Continue to monitor power quality if "
            "issues recur.")
        doc.add_paragraph()
        return

    # Priority keeps its own words -- how soon to act is not the same question
    # as how bad the reading is -- but it is coloured from the one severity
    # ramp, so a High action and a Severe finding look like the same weight.
    _prio_rank = {"High": 0, "Medium": 1, "Low": 2}
    # One list, ordered by priority. Which side of the meter each action falls to
    # follows from the action itself and is the engineer's call to record.
    _body(doc,
        "Candidate actions supported by the measurements, in priority order. "
        "Assignment of these actions between the utility and the customer is left "
        "to the reviewing engineer.")
    for a in sorted(actions, key=lambda a: _prio_rank.get(a["priority"], 3)):
        p = doc.add_paragraph()
        _bold(p, "Recommendation:  ", size_pt=10)
        _normal(p, a["recommendation"], size_pt=10)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        _bold(p2, "Purpose:  ", size_pt=10)
        _normal(p2, a["purpose"], size_pt=10)
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Cm(0.5)
        _bold(p3, "Priority:  ", size_pt=10)
        prio_run = p3.add_run(a["priority"])
        prio_run.bold = True
        prio_run.font.size = Pt(10)
        prio_run.font.color.rgb = _sev_color(a["priority"])
        doc.add_paragraph()


#: Grouping order for the channel appendix: the measured quantities first, then
#: what is derived from them, so a reader checking the basics is not reading past
#: forty harmonic orders to find the voltage.
_CHANNEL_ORDER: Dict[str, int] = {
    "voltage": 0, "current": 1, "power": 2, "frequency": 3,
    "thd": 4, "hrms": 5, "flicker": 7, "kfactor": 8,
}


def _channel_sort_key(name: str) -> tuple:
    """Group, then harmonic order numerically — h11 after h3, not before it."""
    head = name.split("_")[0]
    m = re.fullmatch(r"h(\d+)", head)
    if m:
        return (6, int(m.group(1)), name)
    return (_CHANNEL_ORDER.get(head, 99), 0, name)


#: Canonical channel name → what it is, for names that carry no phase suffix.
_CHANNEL_PLAIN: Dict[str, str] = {
    "frequency":       "System frequency",
    "power_real":      "Real power, total",
    "power_reactive":  "Reactive power, total",
    "power_apparent":  "Apparent power, total",
    "power_factor":    "Power factor, as reported by the meter",
    "kfactor_meter":   "Transformer K-factor, as reported by the meter",
    "voltage_neutral": "RMS voltage, neutral to earth",
    "current_neutral": "RMS current in the neutral conductor",
}

#: Ordinal words for harmonic orders, so "3rd" reads as an order and not a count.
def _ordinal(n: int) -> str:
    if 10 <= n % 100 <= 20:
        return f"{n}th"
    return f"{n}{ {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th') }"


def _phase_word(suffix: str, is_split: bool) -> str:
    """Name a phase the way the rest of the report names it for this service."""
    if suffix == "neutral":
        return "neutral"
    if is_split:
        return {"a": "L1", "b": "L2"}.get(suffix, suffix.upper())
    return f"phase {suffix.upper()}"


def _channel_description(name: str, is_split: bool) -> str:
    """Plain description of one canonical channel.

    Derived from the name rather than tabulated, because the harmonic channels
    alone run to several dozen and a hand-written table would fall out of date
    the first time a meter reports an order nobody listed.
    """
    base, _, stat = name.rpartition("_")
    if stat in ("peak", "min") and base:
        inner = _channel_description(base, is_split)
        return (f"{inner} — highest value within each interval" if stat == "peak"
                else f"{inner} — lowest value within each interval")

    if name in _CHANNEL_PLAIN:
        return _CHANNEL_PLAIN[name]

    m = re.fullmatch(r"h(\d+)_(voltage|current)_(a|b|c|neutral)", name)
    if m:
        order, kind, ph = int(m.group(1)), m.group(2), m.group(3)
        return (f"{_ordinal(order)}-order harmonic {kind}, "
                f"{_phase_word(ph, is_split)}")

    m = re.fullmatch(r"hrms_(voltage|current)_(a|b|c|neutral)", name)
    if m:
        kind, ph = m.group(1), m.group(2)
        return (f"Total harmonic {kind} (RMS of all orders combined), "
                f"{_phase_word(ph, is_split)}, as reported by the meter")

    m = re.fullmatch(r"thd_(voltage|current)_(a|b|c|neutral)", name)
    if m:
        kind, ph = m.group(1), m.group(2)
        return (f"Total harmonic distortion of the {kind}, "
                f"{_phase_word(ph, is_split)}")

    m = re.fullmatch(r"kfactor_current_(a|b|c|neutral)", name)
    if m:
        return f"Transformer K-factor from the {_phase_word(m.group(1), is_split)} current"

    m = re.fullmatch(r"flicker_(pst|plt)(?:_(a|b|c))?", name)
    if m:
        kind = ("Short-term flicker severity (Pst, 10-minute)" if m.group(1) == "pst"
                else "Long-term flicker severity (Plt, 2-hour)")
        return f"{kind}, {_phase_word(m.group(2) or 'a', is_split)}"

    m = re.fullmatch(r"voltage_(ab|bc|ca)", name)
    if m:
        p, q = m.group(1)[0], m.group(1)[1]
        return (f"RMS voltage, line to line "
                f"({_phase_word(p, is_split)} to {_phase_word(q, is_split)})")

    m = re.fullmatch(r"(voltage|current)_(a|b|c)", name)
    if m:
        kind, ph = m.group(1), m.group(2)
        what = ("RMS voltage, line to neutral" if kind == "voltage" else "RMS current")
        return f"{what}, {_phase_word(ph, is_split)}"

    return name.replace("_", " ").capitalize()


def _word_channel_appendix(doc, report, df) -> None:
    """Appendix C: every channel read out of the file, and what it holds.

    The channel-to-quantity match is the one step that fails silently: a label
    matched to the wrong quantity produces a full report of confident numbers
    about the wrong thing. Listing the match next to the device's own label is
    what makes that checkable.
    """
    fs  = report["file_summary"]
    cmap = fs.get("channel_map") or {}
    if not cmap:
        return

    doc.add_page_break()
    _section_heading(doc, "Appendix C: Channels Read From the Meter File", level=1)

    is_split  = "voltage_c" not in fs.get("channels", [])
    intervals = fs.get("sample_count", len(df))
    unmatched = max((fs.get("device_channels") or 0) - len(cmap), 0)
    _body(doc,
        f"{len(cmap):,} channels were read from the meter file and matched to the "
        "quantities below; every number in this report derives from these and "
        "nothing else."
        + (f" A further {unmatched:,} channel(s) found in the file were not "
           "matched to a quantity and were not read."
           if unmatched else
           " Every channel found in the file was matched, so none was dropped.")
        + " The device name is the label the meter itself gave the channel, so "
          "each row can be checked against the meter's own channel list. "
          "\"Intervals with data\" is how many of the "
          f"{intervals:,} recording intervals carried a value — a channel that "
          "matched but arrived empty shows here as 0.")

    # Within-interval max/min ride with the channel they qualify, so the table
    # stays one row per measured quantity rather than three.
    extras: Dict[str, List[str]] = {}
    for col in df.columns:
        base, _, stat = col.rpartition("_")
        if stat in ("peak", "min") and base in cmap:
            extras.setdefault(base, []).append("max" if stat == "peak" else "min")

    rows = sorted(cmap.items(), key=lambda kv: _channel_sort_key(kv[0]))

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    _set_col_widths(table, [3.6, 6.2, 1.4, 3.2, 2.2])
    for i, head in enumerate(("Channel", "What it is", "Unit",
                              "Device name in file", "Intervals with data")):
        cell = table.rows[0].cells[i]
        _cell_shade(cell, _CHROME_BAND)
        r = cell.paragraphs[0].add_run(head)
        r.bold = True
        r.font.size = Pt(9)

    for name, info in rows:
        cells = table.add_row().cells
        extra = extras.get(name)
        label = f"{name}  (+ {', '.join(sorted(extra))})" if extra else name
        n_valid = int(df[name].notna().sum()) if name in df.columns else 0
        for cell, text in zip(cells, (
            label,
            _channel_description(name, is_split),
            info.get("unit") or "—",
            info.get("device") or "—",
            f"{n_valid:,}",
        )):
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(8)
    doc.add_paragraph()


def _word_appendix(doc, report, thresh, df) -> None:
    """Appendix B: Standards, Methods, and Limitations — methodology disclosures
    moved out of the main body."""
    doc.add_page_break()
    _section_heading(doc, "Appendix B: Standards, Methods, and Limitations", level=1)
    _body(doc,
        "The notes below describe the measurement basis, statistical methods, and "
        "known limitations behind the findings in this report.")

    entries: List[Tuple[str, str]] = []
    fs       = report["file_summary"]
    interval = f"{fs.get('interval_minutes', 5):g}"

    # First, because it conditions everything below it: a reader weighing any
    # finding needs to know the file it came from was incomplete before they
    # weigh it. The header says so in three lines; this is where the evidence
    # that makes the failure diagnosable without the .pqd lives.
    _dq = fs.get("data_quality") or {}
    if _dq.get("missing_bytes") or _dq.get("unreadable_observations"):
        entries.append(("Source file integrity — what could not be read",
                        _integrity_note(_dq, fs)))

    entries.append(("Measurement basis",
        f"Results are computed from {interval}-minute interval averages recorded by "
        "the meter"
        + (", supplemented by within-interval maximum/minimum records"
           if fs.get("has_maxmin") else "")
        + (" and cycle-level adaptive event records (~17 ms resolution)"
           if fs.get("has_adaptive") else "")
        + ". Interval averaging can mask short-duration excursions; where "
          "maximum/minimum or adaptive records exist they are used to capture "
          "within-interval behavior."))

    volt = report["voltage_compliance"]
    if volt.get("available"):
        entry = (
            "ANSI C84.1 compliance is evaluated on interval averages. The "
            "standard rates sustained service voltage; an excursion shorter "
            "than one recording interval is a sag or a swell, which IEEE 1159 "
            "and the ITIC envelope grade on depth and duration and which this "
            "report evaluates there. Grading C84.1 on within-interval extremes "
            "would count one such event twice, once against a standard that "
            "does not cover it.")
        if any(v.get("used_interval_extremes") for v in volt["phases"].values()):
            entry += (" The meter's within-interval extreme records are read "
                      "and reported alongside the averages, and are labelled "
                      "as within-interval wherever they appear.")
        entries.append(("Voltage compliance basis", entry))

        rng_b = volt.get("range_b_v")
        entries.append(("ANSI C84.1 voltage ranges",
            (volt.get("band_basis") or "")
            + " Range A is the band the supply is expected to hold under normal "
            f"conditions ({volt['range_v'][0]:.1f}–{volt['range_v'][1]:.1f} V "
            "line-to-neutral here). "
            + (f"Range B ({rng_b[0]:.1f}–{rng_b[1]:.1f} V) is wider on the low "
               "side than the high side, and covers voltages that result from "
               "practical design and operating conditions; C84.1 asks that "
               "excursions into it be limited in extent, frequency and duration "
               "and be corrected within a reasonable time. Readings are "
               "classified into Range A, Range B, or outside both."
               if rng_b else (volt.get("range_b_note") or ""))))

    thd = report["thd_compliance"]
    if thd.get("tdd_info"):
        ti = thd["tdd_info"]
        ll = ("  Intervals below 10% of peak demand are excluded per IEEE 519-2022 "
              "§2.1 guidance, which evaluates distortion at maximum demand conditions."
              if thd["current"].get("light_load_filtered") else "")
        isc_note = (
            ""
            if ti.get("isc_provided") else
            "  The available short-circuit current (ISC) was not provided for this "
            "study, so the most restrictive class (ISC/IL < 20) is assumed; the "
            "resulting limit is conservative and the true limit for this service "
            "can only be equal or higher."
        )
        entries.append(("IEEE 519 TDD evaluation",
            "Total Demand Distortion (TDD) references harmonic current to the maximum "
            "demand load current (IL), taken as the highest interval current over the "
            "recording — unlike THD, TDD does not overstate distortion at light load. "
            "Per-interval TDD is derived as measured THD scaled by the ratio of "
            "interval current to IL. The applicable IEEE 519-2022 Table 2 limit is "
            "selected by the short-circuit ratio ISC/IL." + isc_note + ll))

    hs = report.get("harmonic_statistics", {})
    if hs.get("available"):
        entries.append(("Statistical evaluation method (IEEE 519-2022 Clause 5)",
            "Short Time (ST) statistics use 5-minute interval data as a proxy for "
            "IEC 61000-4-30 10-minute measurements. Very Short Time (VST) values are "
            "approximated from the daily 99th percentile of 5-minute data — a "
            "conservative approximation; true VST evaluation requires 3-second "
            "measurements and may not capture sub-minute harmonic peaks."))

    sh = report.get("harmonic_sources", {})
    if sh.get("available"):
        entries.append(("Harmonic source attribution",
            "Source attribution is indicative, based on harmonic impedance estimates "
            "and Pearson correlation between per-order voltage and current interval "
            "series. Direction measured from waveform phasors is reported separately "
            "under Harmonic Source Direction."))

    fl = report.get("flicker", {})
    if fl.get("available"):
        entries.append(("Flicker: what Pst and Plt each measure",
            "Both come from the same IEC 61000-4-15 flickermeter, over different "
            "windows. Pst is the severity over 10 minutes, normalised so 1.0 is "
            "the level half of observers watching a 60 W incandescent lamp find "
            "objectionable. Plt covers two hours and is built from twelve "
            "consecutive Pst values as the cube root of their mean cube, a "
            "weighting that keeps one bad ten-minute period visible instead of "
            "averaging it away. Pst therefore answers whether a disturbance is "
            "perceptible now, and Plt whether a repeatedly cycling load adds up "
            "to a nuisance over a working period; they fail independently and "
            "both are reported per phase. The meter carries each value forward "
            "across several recording intervals, so the interval count is not "
            "the number of measurement windows and the percentiles are "
            "time-weighted. The limits applied, Pst 1.0 and Plt 0.65, are IEC "
            "61000-3-3 equipment emission limits; IEEE 1453-2015 compatibility "
            "levels for a low-voltage supply system are Pst 1.0 and Plt 0.8, and "
            "both standards assess the 95th percentile over a week, which is "
            "longer than this recording."))

    si = report.get("service_impedance", {})
    if si.get("available"):
        expected = si.get("expected") or {}
        entries.append(("Service impedance — method and limits",
            "Impedance between the source and the meter is fitted on the steps in "
            "load, not on the levels: a fit against the levels attributes the "
            "whole feeder's coincident droop to this one service, which on a "
            "measured 150 ft residential drop read 0.29 Ω against an expected "
            "0.04 Ω. Only intervals whose current stepped by at least "
            f"{_IMPEDANCE_STEP_MIN_A} A (or a twentieth of peak) are used, and a "
            "fit is reported only where the voltage moved the opposite way in at "
            f"least {_IMPEDANCE_MIN_CONSISTENCY:.0%} of those steps. R and X "
            "separate only where the power factor varied enough to make the real "
            "and reactive parts of the current distinguishable; otherwise the "
            "effective magnitude is reported alone. Two readings need no expected "
            "value: one phase against the others, which no common conductor or "
            "transformer can explain, and neutral-to-earth voltage against "
            "neutral current."
            + (" This service is metered on the primary, so the expected value "
               "is the primary line impedance as entered, in positive sequence; "
               "the transformer and secondary conductors are the customer's and "
               "sit downstream of the meter."
               if expected.get("primary_metered") else
               " The expected impedance uses generic published conductor "
               "constants (NEC Chapter 9 Table 8 resistance at 75 °C with a "
               "typical reactance for the construction), not PSCo Blue Book "
               "figures, and the run length as entered, so it is an "
               "order-of-magnitude check rather than a tolerance."
               if expected.get("generic_conductor_constants") else "")
            + (" Part of the path is a secondary main shared with other "
               "services, whose current adds drop that does not move with this "
               "customer's load; that widens the scatter in the fit rather than "
               "biasing it, so the measurement is less precise here than on a "
               "dedicated run."
               if expected.get("shared_secondary_z_ohm") is not None else "")))

    hd = report.get("harmonic_direction", {})
    if hd.get("available"):
        wf = hd.get("waveform") or {}
        entries.append(("Harmonic source direction — method and limits",
            "Two independent readings. (1) Over the whole recording, each order's "
            "harmonic voltage is fitted against its harmonic current as V_h = "
            "Z_h·I_h + V_background; the slope is what the premises' own current "
            "accounts for and the intercept is what persists without it, checked "
            "against the measured harmonic voltage over the tenth of intervals "
            "drawing the least harmonic current. It uses magnitudes only and so "
            "infers direction from covariance rather than measuring it. (2) At each "
            "point-on-wave capture, harmonic phasors are projected onto the measured "
            "fundamental and its multiples, giving P_h = ½·Re(V_h·I_h*), whose sign "
            "is the direction of harmonic power flow; "
            + ("the service carries on-site generation, so the sign of the "
               "fundamental cannot establish CT orientation -- reversed clamps "
               "while importing and correct clamps while exporting read alike -- "
               "and the captures are instead separated by their own direction of "
               "flow and each half read on its own terms, with the CTs taken as "
               "installed arrow-toward-load. "
               if exports_power(thresh) else
               "the sign convention is fixed by the fundamental real power, so "
               "reversed CTs are detected and corrected rather than silently "
               "inverting the result. ")
            + "Captures taken "
            "during a voltage event, below "
            f"{_MIN_LOADED_AMPS:.0f} A of load, shorter than three cycles, or with "
            "no steady fundamental are excluded. The power-direction sign identifies "
            "the side whose harmonic source dominates at the meter and is least "
            "decisive when the impedances either side are comparable (Xu, Liu & Liu, "
            "IEEE Transactions on Power Delivery 18(1), 2003), and captures are "
            "triggered snapshots rather than a continuous record"
            + (f" — {wf.get('captures_used', 0)} of {wf.get('captures_total', 0)} "
               "were usable here." if wf else ".")
            + " Neither method establishes which equipment is responsible."))

    ss = report.get("spectral_shape", {})
    if ss.get("available"):
        note = ss.get("note", "")
        entries.append(("Spectral shape classification",
            "The spectral shape classification is a single-visit snapshot, not a "
            "trend. It distinguishes broadly elevated spectra from "
            "resonance-concentrated spectra but cannot by itself prove or exclude "
            "resonance." + (f"  {note}" if note else "")))

    if report.get("root_causes"):
        entries.append(("Confidence levels",
            "Engineering assessment findings are graded High, Medium, or Low "
            "confidence based on the strength and agreement of the supporting "
            "indicators. Lower-confidence findings identify plausible contributing "
            "conditions that require field verification before corrective investment."))

    if report.get("itic", {}).get("available"):
        entries.append(("ITIC evaluation",
            "Sag/swell events are evaluated against the ITI (CBEMA) Curve "
            "(ITIC 2000), the voltage tolerance envelope referenced by IEEE "
            "1159-2019 for information technology equipment on 120 V nominal "
            "systems. Event magnitude and duration come from cycle-level "
            "records; events shorter than one cycle may be under-resolved."))

    if report.get("events", {}).get("waveform_captures"):
        entries.append(("Waveform capture analysis",
            "Point-on-wave captures are analyzed with a sliding half-cycle RMS "
            "envelope per IEEE 1159 to characterize sag/swell magnitude and "
            "duration at sub-cycle resolution. Capture windows are typically "
            "0.1–1.5 seconds, so durations for events outlasting the capture "
            "are lower bounds. Waveform channels are identified by amplitude "
            "signature; events also present in the cycle-level record are "
            "counted once."))

    ev = report.get("events", {})
    if ev:
        entries.append(("Event detection",
            "Sag, swell, and flicker events are detected from "
            + ("cycle-level adaptive records."
               if ev.get("data_source") == "adaptive"
               else f"{interval}-minute interval averages, which can miss events "
                    "shorter than the recording interval.")
            + " Attributing events to specific causes generally requires "
              "time-correlated system operation records."))

    for title, text in entries:
        p = doc.add_paragraph()
        _bold(p, title + ".  ", size_pt=10)
        _normal(p, text, size_pt=10)
    doc.add_paragraph()


# The engineering report has no sign-off: it is not addressed to anyone, and
# who prepared it is a "Prepared by" row in the header table. The signature
# block that used to live here was unreachable, carried a phone number, and
# used a format nothing else in the tool used -- three ways for it to be wrong
# by the time anyone wired it back up. _signature_block is the one sign-off.


# ─────────────────────────────────────────────────────────────────────────────
# 8c. WORD REPORT GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def generate_word_report(
    report: dict,
    thresh: Thresholds,
    site_name: str,
    site_address: str,
    engineer_name: str,
    outdir: Path,
    stem: str,
    *,
    ds: Optional["PQDataset"] = None,
    engineer_title: str = "",
    engineer_email: str = "",
) -> Optional[Path]:
    """Generate a Word (.docx) power quality response letter matching the PSC template."""
    if not _DOCX_AVAILABLE:
        log.warning("python-docx not installed — skipping Word report. pip install python-docx")
        return None

    df: Optional[pd.DataFrame] = ds.df if ds is not None else None
    doc = _DocxDocument()
    _apply_base_style(doc)

    import datetime
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        footer_para = section.footer.paragraphs[0]
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run(
            f"pq-analyzer v{__version__}  |  Generated {datetime.date.today()}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    fs         = report["file_summary"]
    nominal_v  = thresh.nominal_voltage
    # A single-phase 120/208 service also lacks a C channel, so channel
    # presence alone reported it as 120/240 -- a 13% error that reads as a
    # severe undervoltage. The service-type picker settles it.
    nominal_ll = round(nominal_v * ll_factor(thresh.service_type, thresh.topology))

    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _bold(hdr, "Xcel Energy — Internal Engineering Report", color=_XE_RED, size_pt=14)
    sub = doc.add_paragraph()
    # Said on the document itself, not only in its filename: the two documents
    # for a service look similar at a glance, and the one that carries assumed
    # limits, assessment blanks and raw diagnostics must not be handed over by
    # mistake. The customer document is the one written to be sent.
    _normal(sub, "Power quality analysis — internal working document. "
                 "The customer document for this service is issued separately.",
            color=RGBColor(0x66, 0x66, 0x66), size_pt=9)
    # Without this line the bold is just emphasis, and the reader has to infer
    # the rule from examples.
    key = doc.add_paragraph()
    _normal(key, "Figures in ", color=RGBColor(0x66, 0x66, 0x66), size_pt=9)
    _bold(key, "bold", color=RGBColor(0x66, 0x66, 0x66), size_pt=9)
    _normal(key, " are values measured during this recording. Plain figures are "
                 "limits, ratings, published constants or entered inputs.",
            color=RGBColor(0x66, 0x66, 0x66), size_pt=9)
    doc.add_paragraph()

    _word_site_info_table(doc, site_name, stem, site_address,
                          fs, nominal_v, nominal_ll,
                          prepared_by=", ".join(
                              b for b in (engineer_name,
                                          engineer_title or "Electric Area Engineer")
                              if b) if engineer_name else "")

    opening = doc.add_paragraph()
    opening.add_run(
        "This report summarizes power quality measurements at the service listed above "
        "and evaluates them against the applicable standards. The executive summary and "
        "compliance status come first, followed by the detailed measurements and harmonic "
        "diagnostics they rest on; key findings, the engineering assessment and "
        "recommended actions are drawn from those measurements and appear after them."
    )
    doc.add_paragraph()

    _word_recording_overview(doc, report, outdir, stem)

    _add_toc(doc)

    key_findings = _collect_key_findings(report, thresh, df)
    actions      = _build_structured_actions(report, thresh)

    # Order: summary, then the measurements, then what is concluded from them.
    # This is the internal engineering document, and its reader wants the data
    # before the narrative -- they are checking the numbers, not being persuaded
    # by them. Key Findings, the assessment and the actions all cite the
    # measurement sections, so those sections now precede rather than follow
    # what refers to them.
    _word_exec_summary(doc, report, thresh, df, key_findings, actions)
    # Definitions live in Appendix A; the pointer sits ahead of the measurement
    # sections, which is where the terminology first appears.
    _word_terms_pointer(doc, report)
    _word_measurement_review(doc, report, thresh, df, outdir, stem)
    _word_harmonics(doc, report, thresh, df, outdir, stem)
    _word_key_findings(doc, key_findings)
    _word_engineering_assessment(doc, report)
    _word_recommended_actions(doc, actions)
    # No sign-off block: this document is not addressed to anyone. Who prepared
    # it is recorded in the header table, and the per-finding "Engineer's
    # assessment" blanks are where an engineer puts their name to a judgment.
    _word_terms_panel(doc, report)
    _word_appendix(doc, report, thresh, df)
    _word_channel_appendix(doc, report, df)
    # Last in the document, deliberately: experimental, and not part of the
    # compliance assessment.
    _word_load_signature_appendix(doc, report)

    # ── Save ──────────────────────────────────────────────────────────────────
    outdir.mkdir(parents=True, exist_ok=True)
    out_path = outdir / f"{stem}_internal_engineering_report.docx"
    try:
        doc.save(out_path)
    except PermissionError as exc:
        raise PermissionError(
            f"Could not write {out_path.name} -- it's likely still open in Word or another "
            "program. Close that document and run the analysis again."
        ) from exc
    log.info("Internal engineering report saved → %s", out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# Residential customer letter
# ─────────────────────────────────────────────────────────────────────────────
#
# A separate document for the customer, written to be read by someone with no
# electrical training. The governing rule is that no number appears without what
# it means and what follows from it -- otherwise the engineer rewrites it before
# it can be sent, which defeats the purpose.
#
# Deliberately omitted for residential: harmonics, THD, TDD, K-factor, resonance
# and load-signature matching. A homeowner owns no transformer and cannot act on
# distortion. Voltage level, dips and swells, neutral integrity and visible
# flicker are what a house actually experiences.
#
# Also deliberately absent: any statement of who is responsible, and any
# commitment on Xcel Energy's behalf. Both are the reviewing engineer's to add.

#: What a customer would have noticed, for each condition the analysis detects.
_SYMPTOMS = {
    "under_voltage": (
        "Lights dimmer than usual, especially in the evening. Motors in "
        "refrigerators, freezers and air conditioners laboring on start-up, "
        "running hot, or tripping their overload protection."),
    "over_voltage": (
        "Lights brighter than usual. Bulbs failing sooner than they should. "
        "Electronics running warm."),
    "sag_events": (
        "Lights dipping or blinking. Clocks on ovens and microwaves resetting. "
        "Computers, routers or televisions restarting for no apparent reason."),
    "swell_events": (
        "A brief flare in the lights. Occasionally a surge protector clamping, "
        "or an appliance shutting itself off to protect its electronics."),
    "flicker": (
        "A visible flutter or shimmer in the lights, often most noticeable in "
        "the corner of your eye or on a lit wall."),
    "neutral": (
        "Lights in one part of the house dimming while lights elsewhere "
        "brighten, often the moment a large appliance switches on. Bulbs "
        "failing repeatedly in some rooms but not others. Electronics failing "
        "with no obvious cause."),
    "imbalance": (
        "One part of the installation noticeably affected while the rest seems "
        "normal."),
}

#: The same conditions as they present at a plant, where nobody has noticed the
#: lights doing anything. What a producer sees is inverters coming off line,
#: output curtailed, or generation that does not match irradiance -- and unlike
#: a homeowner they can check any of it against SCADA and revenue metering,
#: so these are written to be verifiable rather than evocative.
_SYMPTOMS_GENERATION = {
    "under_voltage": (
        "Inverters running closer to their low-voltage limit than the design "
        "allows for, and reaching for reactive support they were not sized to "
        "provide. Check for units riding at the bottom of their operating band "
        "during the site's highest output."),
    "over_voltage": (
        "The most common cause of lost production on a generating service: as "
        "output rises it lifts the voltage at the point of interconnection, "
        "and the inverters curtail or disconnect to stay inside their limits. "
        "It shows in SCADA as clipping or trips that track the sunniest hours "
        "rather than the cloudiest."),
    "sag_events": (
        "Inverters tripping off line and going through their reconnect timers, "
        "so a momentary dip costs minutes of production rather than seconds. "
        "Cross-check the timestamps below against your trip logs."),
    "swell_events": (
        "Protective trips on overvoltage, and stress on DC-link capacitors and "
        "surge protection. Repeated events shorten the life of both."),
    "flicker": (
        "Rapid output fluctuation reaching the system, usually from cloud "
        "transients across a large array or from a battery cycling faster than "
        "the interconnection can absorb smoothly."),
    "neutral": (
        "A high-impedance or open neutral on the auxiliary supply. It affects "
        "trackers, controls and SCADA rather than the inverters themselves, so "
        "it shows as instrumentation faults rather than as lost generation."),
    "imbalance": (
        "Per-inverter production data with one unit or string group sitting "
        "consistently below its neighbours under the same irradiance, or a "
        "phase current trend in SCADA that stays split rather than moving "
        "together. Both point at the same unit."),
}


def _urgent_signs(register: dict) -> str:
    """What "call us now" looks like, in equipment the reader actually has.

    A plant has no outlets and no domestic panel, and telling an operator to
    watch for scorch marks around either reads as a letter sent to the wrong
    site -- which costs the trust every other sentence is spending.
    """
    if register.get("generating"):
        return ("If you find signs of heating or arcing at the point of "
                "interconnection, in the combiner boxes, or at the inverter or "
                "switchgear terminations, treat that as urgent: isolate the "
                "affected equipment and call the emergency number on your bill.")
    return ("If you ever smell burning or see scorch marks near outlets or your "
            "panel, treat that as urgent and call the emergency number on your "
            "bill.")


def _symptoms_for(register: dict) -> dict:
    """The symptom vocabulary matching who the letter is addressed to.

    A producer has not noticed the lights doing anything, so the load wording
    would read as a form letter sent to the wrong site -- and worse, would
    invite them to look for evidence that cannot exist there.
    """
    return _SYMPTOMS_GENERATION if register.get("generating") else _SYMPTOMS


def _event_counts(event_result: dict) -> Dict[str, int]:
    """Count detected events by type.

    detect_events returns a DataFrame when it has adaptive data to work from and
    a list of dicts otherwise, so both are handled.
    """
    events = (event_result or {}).get("events")
    if events is None:
        return {}
    if isinstance(events, pd.DataFrame):
        if events.empty or "type" not in events.columns:
            return {}
        return {str(k): int(v) for k, v in events["type"].value_counts().items()}
    counts: Dict[str, int] = {}
    for e in events:
        key = e.get("type", "unknown")
        counts[key] = counts.get(key, 0) + 1
    return counts


#: Every service class gets its own customer document, and they are not the
#: same document with the address changed. What varies is the register: a
#: homeowner needs the physics explained and cannot act on distortion; a plant
#: engineer at a primary-metered site owns the transformer, is billed under a
#: different tariff sheet, and is insulted by an explanation of what voltage
#: is. Writing down to that reader costs credibility, and writing up to a
#: homeowner costs comprehension, so the depth is a property of the class.
#:
#: Fields:
#:   site            what to call the premises
#:   reader          who this is written for, in one phrase
#:   explains_basics whether to say what voltage and current are before using them
#:   pf_sheet        the tariff sheet that bills power factor, None where it doesn't
#:   owns_transformer whether the customer owns the service transformer
#:   fix_agent       who they would call to act on something inside the premises
_LETTER_REGISTER: Dict[str, dict] = {
    "r": {
        "site": "your home",
        "reader": "a homeowner",
        "explains_basics": True,
        "pf_sheet": None,
        "owns_transformer": False,
        "fix_agent": "a licensed electrician",
    },
    "c": {
        "site": "your business",
        "reader": "a small business owner",
        "explains_basics": True,
        "pf_sheet": "Sheet R73",
        "owns_transformer": False,
        "fix_agent": "a licensed electrician",
    },
    "sg": {
        "site": "your facility",
        "reader": "whoever looks after the electrical system on site",
        "explains_basics": False,
        "pf_sheet": "Sheet R73",
        "owns_transformer": False,
        "fix_agent": "your electrical contractor or maintenance team",
        "detail": "full",
        "itic_curve": True,
    },
    "pg": {
        "site": "your facility",
        "reader": "the engineer responsible for the site's electrical system",
        "explains_basics": False,
        "pf_sheet": "Sheet R121",
        "owns_transformer": True,
        "fix_agent": "your electrical engineer or contractor",
        "detail": "full",
        "itic_curve": True,
    },
    # A producer's array is not a customer class -- it keeps whichever schedule
    # it takes service under -- but it is a different reader entirely. Nobody
    # there has noticed the lights flickering, because there are no lights.
    # What they have is a plant that trips, curtails or underproduces, and a
    # SCADA history to check a claim against.
    "generation": {
        "site": "the plant",
        "reader": "the operator responsible for the plant",
        "explains_basics": False,
        "pf_sheet": "Sheet R121",
        "owns_transformer": True,
        "fix_agent": "your O&M contractor or the inverter supplier",
        "detail": "full",
        # ITIC is a ride-through envelope for load equipment. A generating
        # plant's ride-through obligation is IEEE 1547 Clause 6, which is a
        # different curve and is not assessed here, so showing this one would
        # invite the reader to judge the plant against the wrong standard.
        "itic_curve": False,
        "symptom_label": "What this may look like on site:  ",
        "generating": True,
    },
}

#: A letter that lists only exceptions leaves its reader unable to tell a clean
#: check from one that was never run. A facility with maintenance staff can use
#: the whole picture -- and is likely to hand this letter to a contractor, who
#: will ask what else was looked at. A homeowner cannot, so they do not get it.
for _key, _reg in _LETTER_REGISTER.items():
    _reg.setdefault("detail", "brief")

#: Whether the letter shows the ITIC curve itself rather than only the verdict
#: taken from it. Off by default: it is a log-scaled scatter plot, and a reader
#: who is not going to read it is being shown something they did not ask for,
#: which costs the trust the rest of the letter is spending carefully.
for _key, _reg in _LETTER_REGISTER.items():
    _reg.setdefault("itic_curve", False)

#: Service classes that get the plain-language letter, and what to call the
#: site. Every class does; the depth is what differs, per _LETTER_REGISTER.
_LETTER_CLASSES = {cls: reg["site"] for cls, reg in _LETTER_REGISTER.items()
                   if not reg.get("generating")}


def _letter_register(thresh) -> dict:
    """The register for a service, defaulting to the most explanatory one.

    A plant with no load is addressed as a plant whatever schedule it takes
    service under, so the role is checked before the class. Everything else
    keys on the class: an unrecognised one gets the fullest explanation rather
    than the tersest, because over-explaining to an engineer wastes their time
    while under-explaining to a homeowner loses them entirely.
    """
    if getattr(thresh, "service_role", "load") == "generation":
        return _LETTER_REGISTER["generation"]
    return _LETTER_REGISTER.get(thresh.customer_class, _LETTER_REGISTER["r"])


def _discard_stale_letter(path: Path, reason: str) -> None:
    """Delete a letter left behind by an earlier run of the same file.

    Nothing on the letter itself says which run produced it, so one that
    survives a run that could not rewrite it sits in the output folder beside a
    fresh report and fresh plots, describing a different recording — a letter
    stating a two-hour recording next to a plot spanning a week. Removing it is
    the only way the folder can be trusted, so failing to remove it has to stop
    the run rather than pass quietly.
    """
    if not path.exists():
        return
    try:
        path.unlink()
    except OSError as exc:
        raise PermissionError(
            f"Could not replace {path.name} -- it's likely still open in Word or "
            "another program. Close that document and run the analysis again. "
            "Until then the letter in the output folder is from an earlier run "
            "and does not describe this recording."
        ) from exc
    log.info("Removed the previous customer letter at %s (%s).", path, reason)


def _customer_vocabulary(report: dict, thresh: Thresholds) -> dict:
    """Wording that has to change between a house and a small business.

    Two things vary: what to call the premises, and how to explain the neutral.
    A house is split-phase, so the neutral is the shared return between two
    120-volt halves; a three-phase business has a neutral shared by three phases,
    and describing it as two halves would simply be wrong.
    """
    register = _letter_register(thresh)
    site = register["site"]
    topology = (report.get("file_summary") or {}).get("topology", "")
    if register.get("generating"):
        # The neutral at a plant serves trackers, controls and SCADA, not the
        # inverters, which are line-to-line. Describing it as the return for
        # the site's supply would point an operator at the wrong cabinet.
        neutral_measured = (
            f"The auxiliary supply at {site} shares one return wire, called "
            "the neutral, between its phases. If that shared connection "
            "loosens or corrodes the phase voltages begin to move against each "
            "other instead of holding steady, so that is the pattern we look "
            "for. It affects trackers, controls and SCADA rather than the "
            "inverters, which do not use it.")
        neutral_symptom = _symptoms_for(register)["neutral"]
    elif topology == "3-phase":
        neutral_measured = (
            f"{site.capitalize()} is supplied by three separate live wires that "
            "share one return wire, called the neutral. If that shared connection "
            "loosens or corrodes, the voltages on those wires begin to move "
            "against each other instead of holding steady, so that is the pattern "
            "we look for.")
        neutral_symptom = (
            "Equipment on one part of the site behaving oddly while the rest "
            "seems normal, often changing when large equipment switches on. "
            "Lamps failing repeatedly in some areas but not others. Electronic "
            "equipment failing with no obvious cause.")
    else:
        neutral_measured = (
            f"{site.capitalize()} is supplied by two 120-volt halves that share "
            "one return wire, called the neutral. If that shared connection "
            "loosens or corrodes, the two halves begin to move in opposite "
            "directions -- one rising as the other falls -- so that is the "
            "pattern we look for.")
        neutral_symptom = _symptoms_for(register)["neutral"]
    return {"site": site, "neutral_measured": neutral_measured,
            "neutral_symptom": neutral_symptom,
            # Everything above residential is billed for power factor and can
            # trace distortion to its own equipment, so all three commercial
            # classes get those items -- what differs is the tariff sheet and
            # how much the wording explains.
            "is_business": thresh.customer_class != "r",
            "register": register}


def _neutral_indicator_sentence(nh: dict) -> str:
    """State the neutral indicators as data, in terms a reader can weigh.

    The severity alone does not tell an electrician what to look for, and a
    "caution" driven by a couple of volts of variation reads very differently
    from one driven by the two sides opposing each other. So the numbers go in
    the letter regardless of how it is worded around them.
    """
    bits: List[str] = []
    corr = nh.get("leg_correlation")
    if corr is not None:
        if corr >= 0.5:
            sense = ("they rose and fell together, which is what a sound shared "
                     "connection looks like")
        elif corr >= 0:
            sense = "they tracked each other only loosely"
        else:
            sense = ("they moved in opposite directions, which is the pattern a "
                     "failing connection produces")
        bits.append(f"the two sides moved with a correlation of {corr:+.2f}, meaning "
                    + sense)
    total, spread = nh.get("sum_mean_v"), nh.get("sum_std_v")
    if total is not None and spread is not None:
        bits.append(f"together they averaged {total:.0f} volts and varied by about "
                    f"{spread:.1f} volts")
    asym = nh.get("asym_mean_v")
    if asym is not None:
        bits.append(f"the two sides differed from each other by {asym:.1f} volts "
                    "on average")
    vne = nh.get("vne_max_v")
    if vne is not None:
        bits.append(f"the highest neutral-to-earth voltage was {vne:.1f} volts")
    if not bits:
        return ""
    return "In detail: " + "; ".join(bits) + "."


def _customer_checks(report: dict, thresh: Thresholds) -> List[dict]:
    """Every check this recording could run, and how the service did on each.

    The findings section describes only what went wrong, which leaves a reader
    unable to tell "we looked and it was fine" from "we never looked". A
    facility hands this letter to a contractor, and the contractor's first
    question is what else was examined.

    Each row names the standard the figure was judged against, because this
    reader can look one up and a homeowner cannot. It states measurements and
    limits and nothing else: no cause, no party, no undertaking.
    """
    rows: List[dict] = []

    def add(item, measured, against, ok):
        rows.append({"item": item, "measured": measured, "against": against,
                     "result": {True: "Within limits", False: "Outside limits"}
                                .get(ok, "Not measured")})

    vc = report.get("voltage_compliance") or {}
    if vc.get("available"):
        lo, hi = vc["range_v"]
        lows = [s["min_v"] for s in vc["phases"].values()]
        highs = [s["max_v"] for s in vc["phases"].values()]
        pct = vc.get("total_pct_out_of_bounds", 0.0)
        add("Supply voltage",
            f"{_m(min(lows), '.0f')} to {_m(max(highs), '.0f', ' V')}"
            + (f", outside the range for {_mp(pct, '.1f')} of the recording"
               if pct else ""),
            f"ANSI C84.1 Range A: {lo:.0f}–{hi:.0f} V",
            pct == 0)
    else:
        add("Supply voltage", "No usable voltage data", "ANSI C84.1 Range A", None)

    llv = report.get("voltage_ll_compliance") or {}
    if llv.get("available"):
        pairs = llv["pairs"].values()
        add("Voltage between phases",
            f"{_m(min(p['min_v'] for p in pairs), '.0f')} to "
            f"{_m(max(p['max_v'] for p in pairs), '.0f', ' V')}",
            f"ANSI C84.1 Range A: {llv['range_v'][0]:.0f}–"
            f"{llv['range_v'][1]:.0f} V",
            llv.get("overall_pass"))

    imb = report.get("voltage_imbalance") or {}
    if imb.get("available"):
        add("Voltage balance between phases",
            f"highest {_m(imb['max_imbalance_pct'], '.2f', '%')}, "
            f"average {_m(imb['mean_imbalance_pct'], '.2f', '%')}",
            f"{imb.get('metric_label') or 'NEMA MG1'}: "
            f"{imb['limit_pct']:.0f}% maximum",
            imb["max_imbalance_pct"] <= imb["limit_pct"])

    itic = report.get("itic") or {}
    if itic.get("available"):
        n_bad = itic.get("n_violations", 0)
        add("Brief dips and surges",
            (f"{_m(itic['n_events'])} recorded, {_m(n_bad)} beyond the curve"
             if itic.get("n_events") else "none recorded"),
            "ITIC (CBEMA) equipment tolerance curve",
            n_bad == 0)

    fl = report.get("flicker") or {}
    if fl.get("available") and fl.get("pst_max") is not None:
        add("Flicker (visible lamp flutter)",
            f"worst short-term reading {_m(fl['pst_max'], '.2f')}",
            f"IEC 61000-3-3: Pst {fl['pst_limit']:.2f}, "
            f"Plt {fl['plt_limit']:.2f}",
            fl.get("overall_pass"))

    v_thd = (report.get("thd_compliance") or {}).get("voltage") or {}
    if v_thd.get("available"):
        p95 = v_thd.get("p95_thd_pct", v_thd.get("max_thd_pct"))
        add("Waveform distortion on the voltage",
            f"{_m(p95, '.2f', '%')} at the 95th percentile",
            f"IEEE 519-2022: {v_thd['limit_pct']:.0f}% maximum",
            p95 <= v_thd["limit_pct"])

    pfr = report.get("power_factor") or {}
    register = _letter_register(thresh)
    if pfr.get("available") and register.get("pf_sheet"):
        add("Power factor",
            f"lowest {_m(pfr['min_pf'], '.2f')}, "
            f"average {_m(pfr['mean_pf'], '.2f')}",
            f"Xcel Energy tariff {register['pf_sheet']}: "
            f"{pfr['limit']:.2f} minimum",
            pfr["mean_pf"] >= pfr["limit"])

    dem = report.get("demand") or {}
    tx = dem.get("transformer") if isinstance(dem, dict) else None
    if tx and tx.get("overloaded") is not None:
        add("Demand against the transformer serving this service",
            f"{_m(tx['peak_8h_kva'], '.0f', ' kVA')} sustained peak, which is "
            f"{_m(tx['pct_nameplate'], '.0f', '%')} of the transformer",
            f"{tx['nameplate_kva']:.0f} kVA nameplate rating",
            not tx["overloaded"])

    return rows


def _customer_conditions(report: dict, thresh: Thresholds) -> List[dict]:
    """Conditions worth telling a residential or small-business customer about."""
    out: List[dict] = []
    vc = report.get("voltage_compliance") or {}
    fl = report.get("flicker") or {}
    nh = report.get("neutral_health") or {}
    ci = report.get("current_imbalance") or {}
    itic = report.get("itic") or {}
    pfr = report.get("power_factor") or {}
    thd = report.get("thd_compliance") or {}
    counts = _event_counts(report.get("events") or {})
    hours = (report.get("file_summary") or {}).get("duration_hours") or 0
    vocab = _customer_vocabulary(report, thresh)
    site = vocab["site"]
    symptoms = _symptoms_for(vocab["register"])
    #: Whether this reader gets extent and phase alongside the peak, or just
    #: the peak. A facility engineer acts on how long a condition held; a
    #: homeowner is served by the shortest true sentence.
    detailed = vocab["register"]["detail"] == "full"

    # ── Voltage outside the allowed range ─────────────────────────────────
    if vc.get("available") and vc.get("total_pct_out_of_bounds", 0) > 0:
        lo, hi = vc["range_v"]
        worst_low = min((s["min_v"] for s in vc["phases"].values()), default=None)
        worst_high = max((s["max_v"] for s in vc["phases"].values()), default=None)
        pct = vc["total_pct_out_of_bounds"]
        if worst_low is not None and worst_low < lo:
            out.append({
                "headline": f"The voltage at {site} dropped below the normal range",
                "measured": (
                    f"The lowest reading was {_m(worst_low, '.0f', ' volts')}. Normal "
                    f"service is {thresh.nominal_voltage:.0f} volts, and the allowed "
                    f"range is {lo:.0f} to {hi:.0f} volts. Readings fell outside that "
                    f"range during {_mp(pct, '.1f')} of the "
                    f"{_m(hours, '.0f', ' hours')} we recorded."),
                "means": (
                    "Low voltage makes motors work harder than they were designed to. "
                    "Over time that shortens the life of "
                    + ("refrigeration, compressors, pumps and air conditioning plant."
                       if vocab["is_business"] else
                       "refrigerators, freezers, air conditioners and well pumps.")),
                "symptom": symptoms["under_voltage"],
            })
        if worst_high is not None and worst_high > hi:
            out.append({
                "headline": f"The voltage at {site} rose above the normal range",
                "measured": (
                    f"The highest reading was {_m(worst_high, '.0f', ' volts')}, against "
                    f"an allowed maximum of {hi:.0f} volts."
                    # A peak on its own does not say whether this was a moment
                    # or a condition, and those call for different responses.
                    + (f" Readings sat outside the allowed range during "
                       f"{_mp(pct, '.1f')} of the "
                       f"{_m(hours, '.0f', ' hours')} we recorded."
                       if detailed and pct else "")),
                "means": (
                    "Sustained high voltage shortens the life of light bulbs and of "
                    "the electronics inside appliances."),
                "symptom": symptoms["over_voltage"],
            })

    # ── Voltage unbalance between the phases ──────────────────────────────
    # Only where NEMA MG1's definition applies. A split-phase service has two
    # legs 180 degrees apart, the leg difference is not unbalance in that
    # sense, and no limit is set for it -- reporting one would invent a
    # standard. Nothing here says where the unbalance comes from: it can
    # originate on either side of the meter, and this letter does not attribute.
    imb = report.get("voltage_imbalance") or {}
    if (imb.get("available") and imb.get("metric") == "nema_mg1"
            and imb.get("max_imbalance_pct", 0) > imb.get("limit_pct", 0)):
        out.append({
            "headline": "The three phases are not supplying equal voltage",
            "measured": (
                f"The largest difference between the phases was "
                f"{_m(imb['max_imbalance_pct'], '.2f', '%')}, against a limit "
                f"of {imb['limit_pct']:.0f}%. Across the recording it averaged "
                f"{_m(imb['mean_imbalance_pct'], '.2f', '%')}."),
            "means": (
                "Three-phase motors are built for phases that match. When they "
                "do not, the motor draws a current unbalance several times "
                "larger than the voltage difference — a few percent on the "
                "voltage can be ten times that on the current — and the extra "
                "current becomes heat in the windings. NEMA guidance is to "
                "derate a motor running on this much unbalance; left as it is, "
                "the winding insulation ages faster than it should."),
            "symptom": (
                "Three-phase motors running hot, humming louder than usual, or "
                "tripping their overload protection without an obvious load "
                "change. Motors and drives failing earlier than expected, often "
                "the same ones repeatedly."),
        })

    # ── Neutral integrity ─────────────────────────────────────────────────
    # Scaled to what was actually measured. These indicators move for ordinary
    # reasons as well as for a failing connection -- a couple of volts of
    # variation in the leg sum is normal loading -- so only the severities that
    # warrant it get urgent language, and the numbers are given either way so an
    # electrician can weigh them.
    severity = nh.get("severity") if nh.get("available") else None
    if severity in ("caution", "warning", "critical"):
        if severity == "critical":
            entry = {
                "headline": "We found signs of a problem with the neutral connection",
                "means": (
                    "This is the most important item in this letter. A failing "
                    "neutral lets the voltage on part of the installation climb "
                    "well above normal while the rest drops, which can damage "
                    "equipment on the high side. It is also a shock and fire "
                    "hazard, and it does not repair itself."),
                "safety": True,
            }
        elif severity == "warning":
            entry = {
                "headline": "The shared neutral connection is not behaving quite as it should",
                "means": (
                    "A connection that is loosening or corroding shows up in these "
                    "measurements before it causes obvious trouble, so this is "
                    "worth having looked at rather than left. On what we measured "
                    "here it is not an emergency."),
                "safety": False,
            }
        else:
            entry = {
                "headline": "We took a close look at the shared neutral connection",
                "means": (
                    "We check this because a weakening neutral is one of the few "
                    "faults that can damage equipment, and it shows in these "
                    "measurements before it becomes obvious. What we found sits "
                    "within the range that ordinary changes in load produce, so we "
                    "are reporting it as a baseline rather than as a problem. If "
                    "you contact us again, these are the numbers a later recording "
                    "would be compared against."),
                "safety": False,
            }
        entry["measured"] = (
            vocab["neutral_measured"] + " " + _neutral_indicator_sentence(nh)).strip()
        entry["symptom"] = (
            vocab["neutral_symptom"] if severity in ("warning", "critical") else
            "Nothing in particular. This is a check we run on every recording "
            "rather than something you would have seen.")
        # Promote above the other findings only when it warrants leading.
        if severity in ("warning", "critical"):
            out.insert(0, entry)
        else:
            out.append(entry)

    # ── Short dips and surges ─────────────────────────────────────────────
    n_sag = counts.get("voltage_sag", 0)
    n_swell = counts.get("voltage_swell", 0)
    if n_sag:
        worst = itic.get("worst") or {}
        detail = ""
        if worst.get("pct_nominal") and worst.get("duration_ms"):
            secs = worst["duration_ms"] / 1000.0
            detail = (f" The deepest fell to {_m(worst['pct_nominal'], '.0f', '%')} of "
                      f"normal voltage and lasted {_m(secs, '.1f', ' seconds')}.")
        out.append({
            "headline": "The voltage dipped briefly on several occasions",
            "measured": (f"We recorded {n_sag} short voltage dip"
                         f"{'s' if n_sag != 1 else ''}.{detail}"),
            "means": (
                "Brief dips are usually caused by a large load starting up, either "
                f"at {site} or nearby. Most equipment rides through them; clocks, "
                "controls and electronics without battery backup may not."),
            "symptom": symptoms["sag_events"],
        })
    if n_swell:
        out.append({
            "headline": "The voltage rose briefly on several occasions",
            "measured": (f"We recorded {n_swell} short voltage rise"
                         f"{'s' if n_swell != 1 else ''}."),
            "means": (
                "Brief rises often follow a large load switching off. Surge "
                "protection on sensitive electronics is worthwhile."),
            "symptom": symptoms["swell_events"],
        })

    # ── Visible flicker ───────────────────────────────────────────────────
    if fl.get("available") and fl.get("overall_pass") is False:
        pst = fl.get("pst_max")
        limit = fl["pst_limit"]
        # The flicker scale is meaningless to a customer without its basis: the
        # instrument is calibrated so that 1.0 is the conventional threshold of
        # irritability, the level at which about half of observers in laboratory
        # testing judged the flicker annoying. Quoting the number alone invites
        # "is 4.98 a lot?", so the multiple of that threshold is given too.
        if pst is not None:
            multiple = pst / limit if limit else 0
            scale = (
                f"Flicker measured {_m(pst, '.2f')} on the international scale used for "
                f"this. The scale is set so that {limit:.1f} is the point at which "
                "roughly half the people in laboratory testing judged the flicker "
                "in their lights annoying — it is a measure of irritation, not of "
                "damage. Your reading is about "
                + (f"{_m(multiple, '.0f')} times" if multiple >= 1.5
                   else "the same as")
                + " that level.")
        else:
            scale = ("Flicker exceeded the level at which roughly half of people "
                     "find it annoying.")
        out.append({
            "headline": "The lights were flickering enough to be noticeable",
            "measured": scale,
            "means": (
                "This measures how much the light output varies, not how much power "
                "you use, and it does not harm your equipment. It is measured "
                "because flicker at this level is a recognized nuisance: it is the "
                "kind of thing people notice as eye strain or restlessness in a "
                "room without necessarily realizing the lighting is the cause."),
            "symptom": symptoms["flicker"],
        })

    # ── Small business only: things that cost money or appear on a bill ───
    # A homeowner is not billed for power factor and cannot act on distortion, so
    # neither appears in a residential letter. A small business is billed for one
    # and can often trace the other to specific equipment.
    if vocab["is_business"]:
        register = vocab["register"]
        if pfr.get("available") and pfr.get("pct_below_limit", 0) > 0:
            if register.get("generating"):
                # A load PF clause is the wrong instrument here. The plant is
                # not drawing reactive power to serve a load; it is exporting
                # at whatever displacement its inverters are commanded to, and
                # what governs that is the interconnection agreement together
                # with the reactive capability IEEE 1547 Clause 5 requires of
                # the units. Quoting a tariff sheet written for load, and
                # recommending capacitors, would send them after the wrong fix.
                out.append({
                    "headline": "Displacement power factor at the point of interconnection",
                    "measured": (
                        f"Power factor averaged {_m(pfr['mean_pf'], '.2f')} and "
                        f"fell as low as {_m(pfr['min_pf'], '.2f')} over the "
                        "recording, measured at the point of interconnection "
                        "across both directions of flow."),
                    "means": (
                        "For a generating facility this is a question for the "
                        "interconnection agreement rather than for the load "
                        "power factor clauses of the tariff: what the plant is "
                        "required to hold, and whether the inverters are being "
                        "commanded to hold it, are set there and in the "
                        "reactive capability IEEE 1547 Clause 5 requires of the "
                        "units. It is reported here as measured, without a "
                        "compliance finding attached, because this recording "
                        "does not establish what was commanded. Capacitors are "
                        "not the remedy on an inverter-based plant; the "
                        "reactive capability is already in the units."),
                    "symptom": (
                        "Check the reactive setpoint or power factor mode the "
                        "inverters are running in against what the "
                        "interconnection agreement specifies. A plant left in "
                        "unity power factor mode where the agreement expects "
                        "voltage or reactive support is the usual explanation."),
                })
            else:
                sheet = register["pf_sheet"] or "the applicable tariff sheet"
                # Schedule PG asks for near unity rather than a stated 0.90
                # floor, so quoting a 0.90 requirement at a primary customer
                # would be quoting them the wrong tariff.
                requirement = ("near unity" if thresh.customer_class == "pg"
                               else "0.90 or better")
                explanation = (
                    "Power factor describes how much of the current you draw does "
                    "useful work. At a low power factor you draw more current for "
                    "the same output, which loads your wiring and ours without "
                    "producing anything extra. "
                    if register["explains_basics"] else
                    "Reactive demand raises the current for the same real load, "
                    "loading the service without producing output. ")
                out.append({
                    "headline": "Your power factor is below the level the tariff requires",
                    "measured": (
                        f"Power factor averaged {_m(pfr['mean_pf'], '.2f')} and fell as low "
                        f"as {_m(pfr['min_pf'], '.2f')}. PSCo Electric Tariff {sheet} requires "
                        f"{requirement}."),
                    "means": (
                        explanation
                        + f"Under {sheet} this can attract a billing adjustment, so "
                        "correcting it usually pays for itself. The normal remedy is "
                        "power factor correction capacitors, which "
                        + register["fix_agent"] + " can size and install."),
                    "symptom": (
                        "Nothing you would see or hear. This shows up on the bill "
                        "rather than in how equipment behaves."),
                })

        # Transformer loading, for the classes whose demand actually drives it.
        # A homeowner owns no transformer and a small commercial customer
        # cannot act on one; at SG and PG the loading is a planning number the
        # site's own growth decides, and at PG they own the transformer.
        dem = report.get("demand") or {}
        tx = dem.get("transformer") or {}
        if thresh.customer_class in ("sg", "pg") and tx.get("overloaded"):
            owner = ("your transformer" if register["owns_transformer"]
                     else "the transformer serving your facility")
            out.append({
                "headline": f"Demand is running above the rating of {owner}",
                "measured": (
                    f"The highest 8-hour average demand was "
                    f"{_m(tx['peak_8h_kva'], '.0f', ' kVA')} against a "
                    f"{tx['nameplate_kva']:.0f} kVA nameplate — "
                    f"{_m(tx['pct_nameplate'], '.0f', '%')} of rating."),
                "means": (
                    "Sustained loading above nameplate shortens transformer "
                    "insulation life and reduces the headroom available for "
                    "load growth. "
                    + ("As the transformer owner, planning its replacement or "
                       "uprating is yours to schedule; we can discuss the "
                       "loading figures behind this."
                       if register["owns_transformer"] else
                       "Tell us before adding significant load, so the service "
                       "can be reviewed against it.")),
                "symptom": (
                    "Voltage that sags at peak production and recovers when "
                    "large equipment stops. Transformer running hot or noisy."),
            })

        # Distortion, stated plainly and without per-order detail: enough to know
        # whether it matters and what causes it, not enough to invite questions
        # this letter cannot answer.
        #
        # The current-side limit depends on how stiff the supply is at this
        # service (the ISC/IL ratio). Without that figure the analysis falls back
        # to the most restrictive class, which on one test file turns a 0.35%
        # exceedance into 100%. That assumption is defensible in an engineering
        # report where it is labelled, but a customer letter must not tell a
        # business it breaches a standard on the strength of an assumed limit. So
        # an exceedance is only asserted for the quantity whose limit is known:
        # voltage distortion always (the limit is fixed), current distortion only
        # when the short-circuit current was supplied.
        v_thd = (thd.get("voltage") or {})
        i_thd = (thd.get("current") or {})
        isc_known = bool((thd.get("tdd_info") or {}).get("isc_provided"))
        v_exceeds = bool(v_thd.get("available") and v_thd.get("pct_exceeding", 0) > 0)
        i_exceeds = bool(i_thd.get("available") and i_thd.get("pct_exceeding", 0) > 0)
        i_undetermined = bool(i_thd.get("available")) and not isc_known

        which = []
        if v_exceeds:
            which.append("the voltage supplied to you")
        if i_exceeds and isc_known:
            which.append("the current your equipment draws")

        if which or i_undetermined:
            if which:
                claim = ("We measured distortion of that shape in "
                         + " and in ".join(which)
                         + ", beyond the level the applicable standard permits.")
            else:
                claim = ("We measured distortion of that shape in the current your "
                         "equipment draws.")
            if i_undetermined:
                claim += (
                    " Whether the current distortion exceeds the limit that applies "
                    "to you is not settled by this recording: that limit depends on "
                    "the strength of the supply at your service, which has not been "
                    "established yet. The engineer named below can complete that "
                    "assessment.")
            out.append({
                "headline": (
                    "The shape of the electrical waveform is more distorted than "
                    "the standard allows" if which else
                    "We measured distortion in the shape of the electrical waveform"),
                "measured": (
                    "Mains electricity should be a smooth repeating wave. " + claim),
                "means": (
                    "Distortion is produced by equipment that draws current in "
                    "pulses rather than smoothly — variable speed drives, battery "
                    "chargers, LED and fluorescent lighting, computer power "
                    "supplies, welders. It does not usually stop equipment "
                    "working, but it makes transformers and neutral conductors run "
                    "hotter than their rating assumes, which shortens their life. "
                    "Identifying which equipment is responsible needs a site "
                    "survey; the engineer named below can advise on whether that "
                    "is worthwhile here."),
                "symptom": (
                    "Transformers or conductors running hot. Nuisance tripping of "
                    "breakers that are not obviously overloaded. In some cases an "
                    "audible hum from a transformer or panel."),
            })

    # ── Ride-through, for a plant ─────────────────────────────────────────
    # The finding a producer actually wants: of the disturbances the system
    # handed the plant, which ones was it required to stay on through. 6.4.2.1
    # makes tripping inside a ride-through region the plant's non-compliance,
    # so this is evidence about their settings, not a complaint about ours --
    # and getting that backwards would blame the wrong party.
    rt = report.get("ride_through") or {}
    if vocab["register"].get("generating") and rt.get("available") and rt.get("n_events"):
        obliged = rt["n_required_to_ride_through"]
        beyond  = rt["n_beyond_requirement"]
        worst_o = rt.get("worst_obliged")
        if obliged:
            detail = (
                f"The deepest of them reached "
                f"{_m(worst_o['pct_nominal'], '.0f', '%')} of nominal for "
                f"{_m(worst_o['duration_s'], '.2f', ' s')}, inside the "
                f"{worst_o['region'].lower()} region for Category "
                f"{rt['category']}."
            )
        else:
            detail = ""
        out.append({
            "headline": "Voltage disturbances measured against your ride-through obligation",
            "measured": (
                f"{rt['n_events']} voltage event"
                f"{'s' if rt['n_events'] != 1 else ''} were recorded. "
                f"{obliged} fell inside the region IEEE 1547-2018 Clause 6.4.2 "
                f"requires a Category {rt['category']} resource to ride "
                f"through without tripping"
                + (f", and {beyond} fell outside it." if beyond else ".")
                + (" " + detail if detail else "")),
            "means": (
                "Clause 6.4.2.1 is explicit that tripping on a disturbance "
                "inside a ride-through region is non-compliance by the "
                "resource, not a fault in the supply. So if your plant came "
                "off line during any of the events counted above as inside "
                "the region, the place to look is the inverter protection "
                "settings and how they were commissioned against the "
                "category in your interconnection agreement"
                + (". Events outside the region are ones the standard permits "
                   "you to drop on, and no setting change would keep the plant "
                   "on through them." if beyond else ".")),
            "symptom": (
                "Cross-check the timestamps in the table below against your "
                "trip and reconnect logs. A trip that lines up with an event "
                "marked inside the region is worth raising with the inverter "
                "supplier; one that lines up with an event outside it is the "
                "plant behaving as the standard expects."),
        })

    # ── Frequency ride-through, for a plant ───────────────────────────────
    frt = report.get("frequency_ride_through") or {}
    if vocab["register"].get("generating") and frt.get("available"):
        if frt.get("assessable") and frt.get("n_excursions"):
            worst = frt["worst"]
            obliged = frt["n_required_to_ride_through"]
            power = frt.get("active_power_capability")
            out.append({
                "headline": "System frequency left the continuous operation band",
                "measured": (
                    f"{frt['n_excursions']} frequency excursion"
                    f"{'s' if frt['n_excursions'] != 1 else ''} outside "
                    f"58.8-61.2 Hz, the furthest reaching "
                    f"{_m(worst['extreme_hz'], '.2f', ' Hz')} for "
                    f"{_m(worst['duration_s'], '.0f', ' s')}. "
                    f"{obliged} fell in the mandatory operation region of IEEE "
                    f"1547-2018 Table 19."),
                "means": (
                    "Frequency is a property of the interconnection rather than "
                    "of your service or ours, so this is not something either "
                    "of us caused locally. What it bears on is your plant: in "
                    "the mandatory operation region the plant is required to "
                    "stay synchronised and keep exchanging power"
                    + (f", holding active power at {power}, per Table 20."
                       if power else ".")
                    + " Table 19 is the same for all three performance "
                    "categories, so the category does not change whether the "
                    "plant must ride through, only how much power it must hold."),
                "symptom": (
                    "If the plant came off line during one of these, the "
                    "under-frequency and over-frequency protection settings are "
                    "the place to look. Note that an excursion past 299 s "
                    "cumulative in a ten-minute window releases the plant from "
                    "the obligation, so a long event is not the same case as a "
                    "short one."),
            })
        elif frt.get("assessable") is False:
            out.append({
                "headline": "Frequency was not assessed against Clause 6.5.2",
                "measured": (
                    f"Frequency was recorded only as interval averages, ranging "
                    f"{_m(frt.get('min_hz'), '.2f')} to "
                    f"{_m(frt.get('max_hz'), '.2f', ' Hz')}."),
                "means": frt.get("note", ""),
                "symptom": (
                    "If you have reason to think the plant tripped on "
                    "frequency, a recording with the variable-rate record "
                    "enabled would settle it; this one cannot."),
            })

    # ── Load balance across the service ───────────────────────────────────
    if ci.get("available") and ci.get("pct_exceeding", 0) > 0:
        if vocab["register"].get("generating"):
            out.append({
                "headline": "Output is unevenly split across the phases",
                "measured": (
                    f"The phases differed in current by an average of "
                    f"{_m(ci['mean_imbalance_pct'], '.0f', '%')}, and at times "
                    f"by {_m(ci['max_imbalance_pct'], '.0f', '%')}."),
                "means": (
                    "On an inverter-based plant the phases are driven, not "
                    "loaded, so a persistent split usually points at one unit "
                    "or string group rather than at how something was "
                    "connected. It is worth tracing: an imbalance that follows "
                    "a single inverter is often the first sign of a fault on "
                    "that unit, and it costs production before it costs "
                    "anything else."),
                "symptom": symptoms["imbalance"],
            })
        else:
            out.append({
                "headline": f"The electrical load is unevenly split across {site}",
                "measured": (
                    f"The parts of your service differed in load by an average of "
                    f"{_m(ci['mean_imbalance_pct'], '.0f', '%')}, and at times by "
                    f"{_m(ci['max_imbalance_pct'], '.0f', '%')}."),
                "means": (
                    "An uneven split is common and is not a fault in itself. It does make "
                    "low voltage and neutral problems worse, so it is worth correcting if "
                    "an electrician is already working in your panel."),
                "symptom": symptoms["imbalance"],
            })

    return out


def generate_customer_letter(
    report: dict,
    thresh: Thresholds,
    site_address: str,
    engineer_name: str,
    outdir: Path,
    stem: str,
    *,
    engineer_title: str = "",
    engineer_email: str = "",
) -> Optional[Path]:
    """Write the customer document for this service class.

    Every class gets one, and the engineering report is never the customer
    document -- that stays internal. What changes between classes is the
    register: a homeowner is told what voltage is before a voltage is quoted;
    a primary-metered site is not, is billed under a different tariff sheet,
    and owns its own transformer. The depth comes from ``_LETTER_REGISTER``
    and the conditions from ``_customer_conditions``, so a class is added by
    describing its reader rather than by branching through the prose.
    """
    path = Path(outdir) / f"{stem}_customer_letter.docx"
    if not _DOCX_AVAILABLE:
        log.warning("python-docx not installed — skipping customer letter.")
        _discard_stale_letter(path, "python-docx is not installed")
        return None

    # Clear the previous letter before doing the work, not after: if this run
    # cannot write one, the folder must not still hold a letter describing a
    # different recording. Failing here also fails before the expensive part.
    _discard_stale_letter(path, "it is being replaced by this run")

    import datetime
    fs = report["file_summary"]
    vocab = _customer_vocabulary(report, thresh)
    conditions = _customer_conditions(report, thresh)
    doc = _DocxDocument()
    _apply_base_style(doc)

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(2.5)
        fp = section.footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Xcel Energy — Power Quality Review  |  page ")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        _add_page_field(fp)

    hdr = doc.add_paragraph()
    _bold(hdr, "Xcel Energy — Power Quality Review", color=_XE_RED, size_pt=15)
    sub = doc.add_paragraph()
    _normal(sub, f"What we measured at {vocab['site']}, and what it means",
            size_pt=11)

    # The date belongs at the top, where a letter's date goes. It used to sit
    # under the signature, which is not the house sign-off format -- and a
    # customer letter that carries no date at all is a filing problem for
    # whoever keeps it. Written out rather than 2026-08-13: this is the one
    # date in the document a reader reads as prose, and %-d is not portable.
    _today = datetime.date.today()
    dt = doc.add_paragraph()
    _normal(dt, f"{_today:%B} {_today.day}, {_today.year}", size_pt=10)

    intro = doc.add_table(rows=3, cols=2)
    intro.style = "Table Grid"
    _set_col_widths(intro, [4.5, 12.0])
    for i, (label, value) in enumerate([
        ("Service address", site_address or "—"),
        ("Dates measured", f"{fs['start_time']} to {fs['end_time']}"),
        ("Length of recording", f"{fs['duration_hours']:.0f} hours"),
    ]):
        cl, cr = intro.rows[i].cells
        _cell_shade(cl, _CHROME_LABEL)
        cl.paragraphs[0].add_run(label).bold = True
        cr.paragraphs[0].add_run(value)
    doc.add_paragraph()

    # ── Why you received this ─────────────────────────────────────────────
    register = vocab["register"]
    _section_heading(doc, "Why you received this", level=1)
    if register["explains_basics"]:
        _body(doc,
            f"You contacted us about the electricity supply at {vocab['site']}. "
            "We fitted a recording meter at your service for the period shown "
            "above. It measured the voltage and current many times a second and "
            "stored a summary every few minutes. This letter explains what those "
            "measurements show, in plain terms.")
    elif register.get("generating"):
        # A plant is metered at its point of interconnection, and what it wants
        # to know is whether the system it exports into is holding up its end.
        _body(doc,
            f"A power quality recorder was installed at the point of "
            f"interconnection of {vocab['site']} for the period shown above, "
            "logging voltage, current, distortion and disturbance events at "
            "interval resolution. This letter sets out what the recording "
            "shows, what each figure is measured against, and what follows "
            "from it. Where a condition would bear on the plant's own "
            "production, that is said with the finding.")
    else:
        # No explanation of what a meter does: this reader specifies them.
        _body(doc,
            f"A power quality recorder was installed at the service to "
            f"{vocab['site']} for the period shown above, logging voltage, "
            "current, distortion and disturbance events at interval resolution. "
            "This letter sets out what the recording shows, what each figure is "
            "measured against, and what follows from it.")

    # ── What we recorded ──────────────────────────────────────────────────
    # The whole recording on one page, before any finding is described. A
    # customer who has been told their supply was measured for a week should be
    # able to see that week.
    overview_img = _plot_path(outdir, stem, "overview.png")
    if overview_img is not None and overview_img.exists():
        _section_heading(doc, "What we recorded", level=1)
        _body(doc,
            "The chart below is everything the meter recorded, from the first "
            "reading to the last. The upper half is the voltage supplied to "
            f"{vocab['site']}; the lower half is the current, which is how much "
            "electricity was being used at the time. You do not need to read "
            "anything into the detail — it is here so you can see the whole "
            "period the findings below are drawn from."
            if register["explains_basics"] else
            ("Voltage and current across the whole recording, so the findings "
             "below can be read against the conditions they came from. The "
             "current is the plant's output, so it follows the resource rather "
             "than a load profile."
             if register.get("generating") else
             "Voltage and current across the whole recording, so the findings "
             "below can be read against the conditions they came from."))
        _embed_plot(doc, outdir, stem, "overview.png",
                    caption="Voltage and current recorded at your service.")
        doc.add_paragraph()

    # The customer is told the dates on the first page; if the meter recorded
    # more than one stretch, they are entitled to know this letter covers one
    # of them. Said without the file-format reasoning, which is ours to carry.
    sessions = fs.get("sessions") or []
    if len(sessions) > 1:
        current = fs.get("session_index", 0)
        other_periods = "; ".join(
            (s["start_time"] or "")[:10] + f" ({_m(s['hours'], '.0f')} hours)"
            for s in sessions if s["index"] != current)
        _body(doc,
            "The meter recorded in more than one stretch at this service — it "
            "was stopped and started again while it was installed. This letter "
            f"covers the period named above. The other recording{'s' if len(sessions) > 2 else ''} "
            f"{'were' if len(sessions) > 2 else 'was'} made on {other_periods}, "
            "and can be reviewed as well if you would like us to.")
        doc.add_paragraph()

    # ── What we checked ───────────────────────────────────────────────────
    # Before the exceptions, the whole list. A reader who sees only findings
    # cannot tell a check that passed from one that was never run, and this
    # reader is likely to be asked exactly that by a contractor.
    checks = _customer_checks(report, thresh) if register["detail"] == "full" else []
    if checks:
        _section_heading(doc, "What we checked", level=1)
        _body(doc,
            "Every measurement this recording supports, with the standard each "
            "was judged against. The findings that follow are drawn from the "
            "rows marked outside limits; the rest are here so the list is "
            "complete.")
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _set_col_widths(tbl, [4.4, 5.0, 4.6, 2.5])
        for cell, text in zip(tbl.rows[0].cells,
                              ["What we looked at", "What we measured",
                               "Measured against", "Result"]):
            _cell_shade(cell, _CHROME_HDR)
            r = cell.paragraphs[0].add_run(text)
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for chk in checks:
            cells = tbl.add_row().cells
            cells[0].paragraphs[0].add_run(chk["item"]).font.size = Pt(9.5)
            _emit_text(cells[1].paragraphs[0], chk["measured"], size_pt=9.5)
            cells[2].paragraphs[0].add_run(chk["against"]).font.size = Pt(9.5)
            rr = cells[3].paragraphs[0].add_run(chk["result"])
            rr.font.size = Pt(9.5)
            rr.bold = chk["result"] == "Outside limits"
        doc.add_paragraph()

    # ── What we found ─────────────────────────────────────────────────────
    _section_heading(doc, "What we found", level=1)
    if not conditions:
        _body(doc,
            f"We did not find a problem with the electricity supplied to "
            f"{vocab['site']} "
            "during this period. The voltage stayed within the normal range, and we "
            "did not record dips, surges or flicker beyond what is expected on a "
            "healthy supply.")
        _body(doc,
            "That does not mean nothing happened. A recording covers only the days "
            "it ran, and an intermittent fault can easily fall outside it. If the "
            "problem you reported continues, please contact us again and say when "
            "it happens, because a longer or repeated recording may be needed.")
    else:
        _body(doc,
            f"We found {len(conditions)} thing"
            f"{'s' if len(conditions) != 1 else ''} worth bringing to your "
            "attention. Each is explained below: what we measured, what it "
            "means, and "
            + ("what it would look like on site."
               if register.get("generating") else
               "what you may have noticed."))

    safety = [c for c in conditions if c.get("safety")]
    for idx, cond in enumerate(conditions, start=1):
        doc.add_paragraph()
        p = doc.add_paragraph()
        # Brand red, like every other heading in this letter. These are
        # headings, not severity statements: colouring them from the severity
        # palette put every ordinary finding in the Watch blue, which says
        # something about the finding that was never assessed. What is a safety
        # concern is answered in words, in its own section immediately below.
        _bold(p, f"{idx}. {cond['headline']}", color=_XE_RED, size_pt=11)
        # "What you may have noticed" assumes someone was standing there when
        # it happened. At a plant nobody was, and the question is instead what
        # this would look like in the trip logs and production data.
        for label, key in (("What we measured:  ", "measured"),
                           ("What this means:  ", "means"),
                           (register.get("symptom_label",
                                         "What you may have noticed:  "),
                            "symptom")):
            q = doc.add_paragraph()
            q.paragraph_format.left_indent = Cm(0.6)
            _bold(q, label, size_pt=10)
            _normal(q, cond[key], size_pt=10)

    # ── Ride-through, event by event ──────────────────────────────────────
    # A plant gets the list rather than a curve. The operator's next move is to
    # put these timestamps beside their trip logs, and a scatter plot does not
    # support that; the ITIC chart the other classes get would also be the
    # wrong envelope here, since 1547 Clause 6.4.2 is what binds them.
    rt_letter = report.get("ride_through") or {}
    if register.get("generating") and rt_letter.get("available") and rt_letter.get("events"):
        doc.add_paragraph()
        _section_heading(doc, "Every dip and surge we recorded", level=1)
        _body(doc,
            f"Each voltage event from the recording, with the IEEE 1547-2018 "
            f"Clause 6.4.2 region it falls in for a Category "
            f"{rt_letter['category']} resource. The last column is what the "
            f"standard required of the plant during that event.")
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        _set_col_widths(tbl, [3.6, 1.8, 1.9, 3.4, 3.6])
        for cell, text in zip(tbl.rows[0].cells,
                              ["When", "Depth", "Duration", "Clause 6.4.2 region",
                               "What was required"]):
            _cell_shade(cell, _CHROME_BAND)
            cell.paragraphs[0].add_run(text).bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for e in rt_letter["events"][:20]:
            ts = e.get("timestamp")
            cells = tbl.add_row().cells
            # Each region asks for something different, and the differences
            # are the point: permissive lets the plant stop exchanging current
            # so long as it does not trip, and Category III momentary cessation
            # requires it to stop.
            if not e["must_not_trip"]:
                required = "Tripping permitted"
            else:
                required = {
                    "continuous": "Remain in operation",
                    "mandatory":  "Ride through, do not trip",
                    "permissive": "Do not trip; may cease to energize",
                    "momentary":  "Do not trip; cease to energize",
                }.get(e["mode"], "Ride through, do not trip")
            values = [
                str(ts)[:19] if ts is not None else "—",
                f"{e['pct_nominal']:.0f}%",
                f"{e['duration_s']:.2f} s",
                e["region"],
                required,
            ]
            for cell, text in zip(cells, values):
                cell.paragraphs[0].add_run(text).font.size = Pt(9)
        for caveat in rt_letter.get("caveats", []):
            _body(doc, caveat)

    # ── The dips and surges, plotted ──────────────────────────────────────
    # "What we checked" already names the ITIC curve as the standard these
    # events were judged against, which leaves the reader referred to a curve
    # they cannot see and asked to take the verdict on trust. A facility with
    # maintenance staff can read the chart, and is likely to hand this letter
    # to a contractor who will want it. The register decides, not the data:
    # the same chart in front of a homeowner is a log-scaled scatter plot
    # nobody asked for.
    itic_letter = report.get("itic") or {}
    # n_events is checked as well as the image: pq_output keeps the previous
    # run's plots under the same stem, and a chart of last month's events under
    # this month's heading is worse than no chart.
    if (register.get("itic_curve") and itic_letter.get("available")
            and itic_letter.get("n_events")):
        img = _plot_path(outdir, stem, "itic_curve.png")
        if img is not None and img.exists():
            doc.add_paragraph()
            _section_heading(doc, "Every dip and surge we recorded", level=1)
            n_ev  = itic_letter.get("n_events", 0)
            n_bad = itic_letter.get("n_violations", 0)
            _body(doc,
                "The chart below plots each brief voltage dip and surge from the "
                "recording the way equipment manufacturers describe what their "
                "equipment will tolerate. Along the bottom is how long an event "
                "lasted, from a thousandth of a second at the left out to a "
                "hundred seconds at the right. Up the side is how far the "
                "voltage moved, as a percentage of normal.")
            _body(doc,
                "The green band is the range electronic equipment is built to "
                "ride through without misoperating. Anything inside it should "
                "have passed unnoticed. Points in the red bands went deeper or "
                "lasted longer than that, and those are the ones worth matching "
                "against equipment that tripped, reset or dropped out.")
            # Written out per count rather than with an inline plural: "Of the 1
            # event recorded, all fall inside" is the kind of sentence that
            # makes a reader wonder what else was generated rather than written.
            if n_ev == 1:
                _body(doc,
                    "One event was recorded. It falls "
                    + ("inside the green band, so it is unlikely to explain "
                       "equipment that has been misbehaving."
                       if not n_bad else
                       "outside the green band, so it is worth comparing its "
                       "time against anything that stopped working."))
            else:
                if n_bad:
                    one  = n_bad == 1
                    tail = (f"{_m(n_bad)} of them {'falls' if one else 'fall'} "
                            "outside the green band. Comparing "
                            f"{'its time' if one else 'their times'} against "
                            "anything that stopped working is the fastest way to "
                            f"tell whether {'it is' if one else 'they are'} behind "
                            "a problem you have been seeing.")
                else:
                    tail = ("all of them fall inside the green band, so none is "
                            "likely to explain equipment that has been "
                            "misbehaving.")
                _body(doc, f"Of the {_m(n_ev)} events recorded, {tail}")
            # No attribution: the chart says what reached the service, and a
            # motor start inside the facility and a fault on a neighbouring
            # feeder land in the same place on it.
            _body(doc,
                "A dip can begin on our system or inside the facility — a large "
                "motor starting is a common cause of one either way. The chart "
                "shows what reached your service, not where it started.")
            _embed_plot(doc, outdir, stem, "itic_curve.png",
                        caption="Each recorded dip and surge, by how far the "
                                "voltage moved and how long it lasted.",
                        width_cm=14.0)
            doc.add_paragraph()

    # ── Safety ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _section_heading(doc, "Is this a safety concern?", level=1)
    if safety:
        _body(doc,
            "Yes — one of the findings above needs prompt attention. A damaged or "
            "loose neutral connection can put well above normal voltage on part of "
            "your wiring, and it is a shock and fire risk as well as a risk to your "
            "appliances. Please arrange for a licensed electrician to inspect your "
            "main panel and service connections, and tell them this letter reports "
            "signs of a neutral problem.")
        _body(doc,
            "If you smell burning, see scorch marks around outlets or the panel, or "
            "notice lights surging brightly, treat it as urgent: stop using the "
            "affected circuits and call us immediately on the emergency number on "
            "your bill.")
    elif conditions:
        _body(doc,
            "We did not find anything in these measurements that suggests an "
            "immediate safety risk. The items above affect how well your equipment "
            "works and how long it lasts, rather than presenting a hazard. "
            + _urgent_signs(register))
    else:
        _body(doc,
            "We did not find anything in these measurements that suggests a safety "
            "risk. " + _urgent_signs(register))

    # ── What happens next ─────────────────────────────────────────────────
    doc.add_paragraph()
    _section_heading(doc, "What happens next", level=1)
    _body(doc,
        "These measurements have been reviewed by an engineer, whose name is at the "
        "end of this letter. Where the results point to something on our equipment, "
        f"we follow it up. Where they point to wiring or equipment inside "
        f"{vocab['site']}, "
        + register["fix_agent"] + " is the right party to look at it. The "
        "engineer's notes below say which applies in your case."
        + ("" if not register["owns_transformer"] else
           " The service transformer at this site is yours, so anything found "
           "on it falls to you to schedule."))

    _write_in_field(doc, "Engineer's notes for this service:", lines=4,
                    indent_cm=0.6, width_cm=15.5)

    # ── Questions ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    _section_heading(doc, "If you have questions", level=1)
    contact_bits = [b for b in (engineer_name, engineer_title) if b]
    _body(doc,
        "The engineer who reviewed your measurements is "
        + (", ".join(contact_bits) if contact_bits else "listed below")
        + ". "
        # No telephone number anywhere in the letter. The address in the
        # sign-off below is the route, and it is a live link.
        + ("Their email address is below, and you "
           if engineer_email else "You ")
        + "can reply to this letter with any questions.")
    _body(doc,
        "The measurements behind this letter were recorded and reviewed in detail. "
        "If you engage an electrician and they would find the underlying "
        "measurements useful, ask the engineer named above and they can discuss "
        "them directly.")

    doc.add_paragraph()
    _signature_block(doc, engineer_name, engineer_title, engineer_email)

    Path(outdir).mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(path))
    except PermissionError as exc:
        raise PermissionError(
            f"Could not write {path.name} -- it's likely still open in Word or "
            "another program. Close that document and run the analysis again."
        ) from exc
    log.info("Customer letter saved → %s  (%d condition(s) reported)",
             path, len(conditions))
    return path


# ─────────────────────────────────────────────────────────────────────────────
# Plain-language definitions for the engineering report
# ─────────────────────────────────────────────────────────────────────────────
#
# The engineering report keeps its technical terms -- an engineer needs them and
# vaguer language would cost precision -- but a report that goes to a customer,
# or to a colleague outside power quality, has to define them where they are
# actually read. So the recurring terms are glossed once, near the front, and
# only the ones this particular report uses are listed.

#: (term, plain definition, report key that has to be present for it to apply).
#: A report key of None means the term always applies.
_TERM_DEFINITIONS: List[Tuple[str, str, Optional[str]]] = [
    ("RMS voltage",
     "The effective voltage — what a meter reads and what equipment responds "
     "to, as distinct from the peak of the waveform.", None),
    ("ANSI C84.1 Range A",
     "The voltage band a distribution system is expected to hold in normal "
     "operation: within 5% either side of nominal. Range B is a wider band "
     "tolerated only briefly.", None),
    ("Sag / swell",
     "A short drop or rise in voltage, lasting from half a cycle up to about a "
     "minute. Distinct from a sustained voltage that is simply too low or high.",
     "events"),
    ("Total harmonic distortion (THD)",
     "How far the waveform departs from a clean sine wave, as a percentage of "
     "the fundamental 60 Hz component. Applied to voltage here.",
     "thd_compliance"),
    ("Total demand distortion (TDD)",
     "Harmonic current as a percentage of the site's maximum demand current. "
     "IEEE 519 limits current this way rather than as THD so that light-load "
     "periods, where harmonic current is small but the fundamental is smaller "
     "still, do not inflate the result.", "thd_compliance"),
    ("IL",
     "The maximum demand current during the recording, at the fundamental "
     "frequency. This is the reference TDD is measured against.",
     "thd_compliance"),
    ("ISC",
     "The short-circuit current available at the service point. The ratio "
     "ISC/IL selects which of the IEEE 519 limit classes applies — a stiffer "
     "supply tolerates more harmonic current.", "_isc"),
    ("Harmonic order (H3, H5, …)",
     "A component of the waveform at a whole multiple of 60 Hz. H3 is 180 Hz, "
     "H5 is 300 Hz, and so on.", "individual_harmonics"),
    ("Triplen harmonics",
     "The orders divisible by three (H3, H9, H15). On a four-wire system these "
     "add together in the neutral conductor instead of cancelling, which is why "
     "neutral current can exceed phase current.", "neutral_harmonics"),
    ("K-factor",
     "Harmonic current weighted by frequency, indicating the additional heating "
     "a transformer experiences beyond a sinusoidal load of the same size. A "
     "standard transformer is rated K=1.", "kfactor"),
    ("Pst and Plt",
     "Short-term (10-minute) and long-term (2-hour) flicker severity. The scale "
     "is set so that 1.0 is the level at which roughly half of observers judge "
     "light flicker annoying; it measures irritation, not damage.", "flicker"),
    ("Power factor",
     "Real power divided by apparent power. A low power factor means more "
     "current is drawn for the same useful work, loading conductors and "
     "transformers without doing more of it.", "power_factor"),
    ("ITIC (CBEMA) curve",
     "An industry reference curve describing how deep and how long a voltage "
     "dip can be before typical electronic equipment misoperates. Events "
     "plotted outside it are the ones likely to have caused trouble.", "itic"),
    ("Point-on-wave capture",
     "A short recording of the actual waveform shape, rather than a summary "
     "value over an interval. Used to see the shape of a disturbance.",
     "_waveforms"),
]


def _applicable_terms(report: dict) -> List[Tuple[str, str]]:
    """The defined terms this report actually uses.

    A residential report that never mentions TDD should not define it, and a
    report run without ISC should not define a ratio it never computed.
    """
    events = report.get("events") or {}
    thd = report.get("thd_compliance") or {}

    def present(key: Optional[str]) -> bool:
        if key is None:
            return True
        if key == "_isc":
            return bool((thd.get("tdd_info") or {}).get("isc_provided"))
        if key == "_waveforms":
            return bool(events.get("waveform_captures"))
        if key == "events":
            return bool(events.get("event_count"))
        if key == "thd_compliance":
            # Only if the section produced a usable result for either quantity.
            return bool((thd.get("voltage") or {}).get("available")
                        or (thd.get("current") or {}).get("available"))
        section = report.get(key)
        return isinstance(section, dict) and bool(section.get("available"))

    return [(t, d) for t, d, key in _TERM_DEFINITIONS if present(key)]


def _word_terms_pointer(doc, report: dict) -> None:
    """One line telling the reader where the definitions live.

    The glossary sits in Appendix A rather than inline, so a reader who meets an
    unfamiliar term in Key Findings needs somewhere to be sent; without this the
    definitions are only found by whoever happens to page to the back.
    """
    if not _applicable_terms(report):
        return
    p = doc.add_paragraph()
    r = p.add_run(
        "Terms used in this report — total demand distortion, percentiles, "
        "and the rest — are defined in Appendix A.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


def _word_terms_panel(doc, report: dict) -> None:
    """Define the report's recurring terms, in Appendix A."""
    terms = _applicable_terms(report)
    if not terms:
        return
    _section_heading(doc, "Appendix A: Terms Used in This Report", level=1)
    _body(doc,
        "This report is written for an engineering reader but is also shared with "
        "customers. The terms below are the ones it relies on; the measurement "
        "sections assume them rather than re-explaining them.")
    tbl = doc.add_table(rows=len(terms), cols=2)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [4.8, 11.7])
    for i, (term, definition) in enumerate(terms):
        cl, cr = tbl.rows[i].cells
        _cell_shade(cl, _CHROME_LABEL)
        run = cl.paragraphs[0].add_run(term)
        run.bold = True
        run.font.size = Pt(9)
        dr = cr.paragraphs[0].add_run(definition)
        dr.font.size = Pt(9)
    doc.add_paragraph()
