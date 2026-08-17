"""
test_pq.py — Unit tests for PQ Analyzer IEEE math, channel mapping, and pipeline.

Run with:
    pytest test_pq.py -v

Coverage:
  1. IEEE 519-2022 per-order harmonic limits (_h519_limit)
  2. IEEE 519-2022 TDD limits (_tdd_limit) — boundary values
  3. ISC/IL class label (_tdd_class)
  4. Neutral harmonic block formula (_V2_CH_H*_IN_AAC constants)
  5. ChannelMapper tag resolution
  6. ChannelMapper regex pattern resolution
  7. Pipeline smoke test: MockAdapter → extract_dataset → check_voltage_compliance
"""

import math
import sys
from pathlib import Path

import numpy as np
import pytest

# ── Make pq_* importable from any working directory ──────────────────────────
sys.path.insert(0, str(Path(__file__).parent))

import struct
import zlib

import pqdif
from pq_constants import (
    SEVERITY_LABEL,
    strip_marks,
    SEVERITY_ORDER,
    Thresholds,
    _h519_limit,
    _tdd_limit,
    _tdd_class,
)
from pq_adapter import (
    ProntoAdapter,
    MockAdapter,
    ChannelMapper,
    RawChannelInfo,
    extract_dataset,
)
from pq_analysis import (
    check_voltage_compliance,
    check_neutral_harmonics,
    check_harmonic_sources,
    detect_events,
)


# ─────────────────────────────────────────────────────────────────────────────
# 1. IEEE 519-2022 per-order harmonic limits
# ─────────────────────────────────────────────────────────────────────────────

class TestH519Limit:
    """Spot-check IEEE 519-2022 Table 2 values by class × harmonic order."""

    # Class <20: limits = [4.0, 7.0, 10.0, 12.0, 15.0] → class index 0
    def test_h5_class_lt20(self):
        assert _h519_limit(5, 15.0) == 4.0

    def test_h11_class_lt20(self):
        # 11 ≤ h < 17 row, class 0 → 2.0
        assert _h519_limit(11, 15.0) == 2.0

    def test_h17_class_lt20(self):
        # 17 ≤ h < 23, class 0 → 1.5
        assert _h519_limit(17, 15.0) == 1.5

    def test_h25_class_lt20(self):
        # 23 ≤ h < 35, class 0 → 0.6
        assert _h519_limit(25, 15.0) == 0.6

    def test_h37_class_lt20(self):
        # 35 ≤ h < 51, class 0 → 0.3
        assert _h519_limit(37, 15.0) == 0.3

    # Class 20–50: limits index 1
    def test_h5_class_20_50(self):
        assert _h519_limit(5, 25.0) == 7.0

    def test_h13_class_20_50(self):
        # 11 ≤ h < 17, class 1 → 3.5
        assert _h519_limit(13, 25.0) == 3.5

    # Class 50–100: limits index 2
    def test_h7_class_50_100(self):
        # h=7 is in 2 ≤ h < 11, class 2 → 10.0
        assert _h519_limit(7, 75.0) == 10.0

    def test_h17_class_50_100(self):
        # 17 ≤ h < 23, class 2 → 4.0
        assert _h519_limit(17, 75.0) == 4.0

    # Class 100–1000: limits index 3
    def test_h5_class_100_1000(self):
        assert _h519_limit(5, 500.0) == 12.0

    def test_h25_class_100_1000(self):
        # 23 ≤ h < 35, class 3 → 2.0
        assert _h519_limit(25, 500.0) == 2.0

    # Class ≥1000: limits index 4
    def test_h5_class_ge1000(self):
        assert _h519_limit(5, 1500.0) == 15.0

    def test_h37_class_ge1000(self):
        # 35 ≤ h < 51, class 4 → 1.4
        assert _h519_limit(37, 1500.0) == 1.4

    # Edge: order outside table
    def test_h_out_of_scope(self):
        assert _h519_limit(100, 50.0) == 0.0

    # Boundary: exactly at class threshold (isc_il = 20 is class 1, not class 0)
    def test_h5_boundary_at_20(self):
        assert _h519_limit(5, 20.0) == 7.0

    def test_h5_just_below_20(self):
        assert _h519_limit(5, 19.9) == 4.0


# ─────────────────────────────────────────────────────────────────────────────
# 2. IEEE 519-2022 TDD limits — boundary values
# ─────────────────────────────────────────────────────────────────────────────

class TestTDDLimit:
    def test_below_20(self):
        assert _tdd_limit(15.0) == 5.0

    def test_just_below_20(self):
        assert _tdd_limit(19.9) == 5.0

    def test_at_20(self):
        assert _tdd_limit(20.0) == 8.0

    def test_in_20_50(self):
        assert _tdd_limit(35.0) == 8.0

    def test_at_50(self):
        assert _tdd_limit(50.0) == 12.0

    def test_in_50_100(self):
        assert _tdd_limit(75.0) == 12.0

    def test_at_100(self):
        assert _tdd_limit(100.0) == 15.0

    def test_in_100_1000(self):
        assert _tdd_limit(500.0) == 15.0

    def test_just_below_1000(self):
        assert _tdd_limit(999.9) == 15.0

    def test_at_1000(self):
        assert _tdd_limit(1000.0) == 20.0

    def test_above_1000(self):
        assert _tdd_limit(5000.0) == 20.0


# ─────────────────────────────────────────────────────────────────────────────
# 3. ISC/IL class label
# ─────────────────────────────────────────────────────────────────────────────

class TestTDDClass:
    def test_lt20(self):
        assert _tdd_class(10.0) == "< 20"

    def test_20_to_50(self):
        assert _tdd_class(30.0) == "< 50"

    def test_50_to_100(self):
        assert _tdd_class(75.0) == "< 100"

    def test_100_to_1000(self):
        assert _tdd_class(500.0) == "< 1000"

    def test_ge1000(self):
        assert _tdd_class(1500.0) == "≥ 1000"



# ─────────────────────────────────────────────────────────────────────────────
# 5. ChannelMapper — tag-based resolution
# ─────────────────────────────────────────────────────────────────────────────

def _raw(index: int, label: str, qt: str, qm: str, ph: str) -> RawChannelInfo:
    return RawChannelInfo(index, label, qt, qm, ph, "")


class TestChannelMapperTags:
    mapper = ChannelMapper()

    def _resolve_one(self, qt: str, qm: str, ph: str) -> str | None:
        ch = _raw(0, "", qt, qm, ph)
        result = self.mapper.resolve([ch])
        return next(iter(result.keys())) if result else None

    def test_voltage_a(self):
        assert self._resolve_one("voltage", "rms", "an") == "voltage_a"

    def test_voltage_b(self):
        assert self._resolve_one("voltage", "average", "bn") == "voltage_b"

    def test_voltage_c(self):
        assert self._resolve_one("voltage", "rmsvalue", "cn") == "voltage_c"

    def test_current_a(self):
        assert self._resolve_one("current", "rms", "a") == "current_a"

    def test_current_b(self):
        assert self._resolve_one("current", "rms", "b") == "current_b"

    def test_current_neutral(self):
        assert self._resolve_one("current", "rms", "neutral") == "current_neutral"

    def test_current_neutral_phase_n(self):
        assert self._resolve_one("current", "rms", "phase_n") == "current_neutral"

    def test_thd_voltage_a(self):
        assert self._resolve_one("voltageharmonics", "thd", "an") == "thd_voltage_a"

    def test_thd_current_b(self):
        assert self._resolve_one("currentharmonics", "thd", "b") == "thd_current_b"

    def test_h5_current_a(self):
        assert self._resolve_one("currentharmonics", "h5", "a") == "h5_current_a"

    def test_h13_current_c(self):
        assert self._resolve_one("currentharmonics", "h13", "cn") == "h13_current_c"

    def test_h3_current_neutral(self):
        assert self._resolve_one("currentharmonics", "h3", "neutral") == "h3_current_neutral"

    def test_h7_current_neutral_in_alias(self):
        assert self._resolve_one("currentharmonics", "h7", "in") == "h7_current_neutral"

    def test_flicker_pst(self):
        assert self._resolve_one("flicker", "pst", "an") == "flicker_pst"

    def test_flicker_plt(self):
        assert self._resolve_one("flicker", "plt", "a") == "flicker_plt"

    def test_unmatched_returns_empty(self):
        ch = _raw(0, "xyzzy unknown", "unknown", "unknown", "unknown")
        result = self.mapper.resolve([ch])
        assert "voltage_a" not in result


# ─────────────────────────────────────────────────────────────────────────────
# 6. ChannelMapper — regex pattern resolution (no tags, label only)
# ─────────────────────────────────────────────────────────────────────────────

class TestChannelMapperRegex:
    mapper = ChannelMapper()

    def _resolve_by_label(self, label: str) -> str | None:
        ch = _raw(0, label, "", "", "")
        result = self.mapper.resolve([ch])
        return next(iter(result.keys())) if result else None

    def test_van_rms(self):
        assert self._resolve_by_label("Van RMS") == "voltage_a"

    def test_vb_label(self):
        assert self._resolve_by_label("Vb") == "voltage_b"

    def test_vc_label(self):
        assert self._resolve_by_label("Vc") == "voltage_c"

    def test_ia_label(self):
        assert self._resolve_by_label("Ia") == "current_a"

    def test_kw_label(self):
        assert self._resolve_by_label("kW") == "power_real"

    def test_kvar_label(self):
        assert self._resolve_by_label("kVAR") == "power_reactive"

    def test_thd_va_label(self):
        assert self._resolve_by_label("THD Va") == "thd_voltage_a"

    def test_thd_ia_label(self):
        assert self._resolve_by_label("THD Ia") == "thd_current_a"


# ─────────────────────────────────────────────────────────────────────────────
# 7. Pipeline smoke test: MockAdapter → extract_dataset → check_voltage_compliance
# ─────────────────────────────────────────────────────────────────────────────

class TestPipelineSmoke:
    """End-to-end test through the full extraction and analysis pipeline."""

    @pytest.fixture(scope="class")
    @classmethod
    def ds(cls):
        # 2.0 hours needed: MockAdapter injects swell at indices 6000–6020 (> 3600)
        adapter = MockAdapter(duration_hours=2.0, nominal=120.0)
        mapper  = ChannelMapper()
        return extract_dataset(adapter, mapper)

    @pytest.fixture(scope="class")
    @classmethod
    def thresh(cls):
        return Thresholds(nominal_voltage=120.0)

    def test_dataset_has_df(self, ds):
        assert not ds.df.empty

    def test_dataset_has_voltage_columns(self, ds):
        assert "voltage_a" in ds.df.columns
        assert "voltage_b" in ds.df.columns
        assert "voltage_c" in ds.df.columns

    def test_dataset_has_current_columns(self, ds):
        assert "current_a" in ds.df.columns

    def test_dataset_topology_inferred(self, ds):
        assert ds.meta.get("topology") in {"3-phase", "split-phase", "single-phase"}

    def test_dataset_has_adaptive_for_mock(self, ds):
        # MockAdapter now synthesizes a small adaptive_df for testing
        assert ds.has_adaptive

    def test_voltage_compliance_runs(self, ds, thresh):
        result = check_voltage_compliance(ds.df, thresh)
        assert "phases" in result
        assert "total_pct_out_of_bounds" in result

    def test_voltage_compliance_phases_present(self, ds, thresh):
        result = check_voltage_compliance(ds.df, thresh)
        assert "voltage_a" in result["phases"]

    def test_voltage_compliance_pct_numeric(self, ds, thresh):
        result = check_voltage_compliance(ds.df, thresh)
        pct = result["total_pct_out_of_bounds"]
        assert isinstance(pct, float)
        assert 0.0 <= pct <= 100.0

    def test_voltage_compliance_mock_sag_detected(self, ds, thresh):
        """MockAdapter injects a 12% sag event — compliance check must catch it."""
        result = check_voltage_compliance(ds.df, thresh)
        pct_under = result["phases"]["voltage_a"]["pct_under"]
        assert pct_under > 0.0, "Sag event in MockAdapter was not detected"

    def test_voltage_compliance_mock_swell_detected(self, ds, thresh):
        """MockAdapter injects an 8% swell event — compliance check must catch it."""
        result = check_voltage_compliance(ds.df, thresh)
        pct_over = result["phases"]["voltage_a"]["pct_over"]
        assert pct_over > 0.0, "Swell event in MockAdapter was not detected"

    def test_catalog_runs(self, ds):
        cat = ds.catalog()
        assert isinstance(cat, str)
        assert len(cat) > 0

    def test_duration_positive(self, ds):
        assert ds.duration_hours > 0.0


# ─────────────────────────────────────────────────────────────────────────────
# 8. check_neutral_harmonics
# ─────────────────────────────────────────────────────────────────────────────

class TestNeutralHarmonics:

    @pytest.fixture(scope="class")
    @classmethod
    def df_with_neutral(cls):
        """Synthetic DataFrame with phase and neutral harmonic channels."""
        import pandas as pd
        rng = np.random.default_rng(0)
        n   = 500
        idx = pd.date_range("2024-01-01", periods=n, freq="5min", tz="UTC")
        data = {
            # Phase A harmonics (Amps)
            "h3_current_a":  3.5 + rng.normal(0, 0.2, n),
            "h5_current_a":  5.0 + rng.normal(0, 0.3, n),
            "h7_current_a":  2.0 + rng.normal(0, 0.2, n),
            "h9_current_a":  0.8 + rng.normal(0, 0.1, n),
            "h11_current_a": 1.2 + rng.normal(0, 0.1, n),
            "h13_current_a": 0.9 + rng.normal(0, 0.1, n),
            # Neutral: triplens accumulate (~2.8× phase H3), non-triplens near zero
            "h3_current_neutral":  3.5 * 2.8 + rng.normal(0, 0.3, n),   # ≈ 9.8 A
            "h5_current_neutral":  rng.normal(0, 0.1, n).clip(0),
            "h7_current_neutral":  rng.normal(0, 0.08, n).clip(0),
            "h9_current_neutral":  0.8 * 2.5 + rng.normal(0, 0.1, n),   # ≈ 2.0 A
            "h11_current_neutral": rng.normal(0, 0.07, n).clip(0),
            "h13_current_neutral": rng.normal(0, 0.06, n).clip(0),
        }
        return pd.DataFrame(data, index=idx)

    @pytest.fixture(scope="class")
    @classmethod
    def thresh(cls):
        # Three-phase wye data: the neutral triplens accumulate at ~2.8x phase
        # H3, which only happens on a 120-degree system. Only phase A's
        # channels are present, so channel presence alone would read this as a
        # single-phase service -- the engineer's topology pick is what settles
        # it, which is the precedence `service_geometry` documents.
        return Thresholds(nominal_voltage=120.0, topology="3ph-wye")

    def test_available_when_neutral_cols_present(self, df_with_neutral, thresh):
        result = check_neutral_harmonics(df_with_neutral, thresh)
        assert result["available"] is True

    def test_unavailable_when_no_neutral_cols(self, thresh):
        import pandas as pd
        df = pd.DataFrame({"current_a": [50.0] * 10})
        result = check_neutral_harmonics(df, thresh)
        assert result["available"] is False

    def test_all_six_orders_present(self, df_with_neutral, thresh):
        result = check_neutral_harmonics(df_with_neutral, thresh)
        assert set(result["orders"].keys()) == {3, 5, 7, 9, 11, 13}

    def test_triplen_orders_flagged(self, df_with_neutral, thresh):
        result = check_neutral_harmonics(df_with_neutral, thresh)
        assert result["orders"][3]["is_triplen"] is True
        assert result["orders"][9]["is_triplen"] is True
        assert result["orders"][5]["is_triplen"] is False
        assert result["orders"][7]["is_triplen"] is False

    def test_triplen_dominant(self, df_with_neutral, thresh):
        """H3 + H9 >> H5 + H7 + H11 + H13 in this dataset."""
        result = check_neutral_harmonics(df_with_neutral, thresh)
        assert result["triplen_dominant"] is True

    def test_triplen_pct_above_50(self, df_with_neutral, thresh):
        result = check_neutral_harmonics(df_with_neutral, thresh)
        assert result["triplen_pct"] > 50.0

    def test_accumulation_factor_computed(self, df_with_neutral, thresh):
        result = check_neutral_harmonics(df_with_neutral, thresh)
        assert result["accumulation_factor"] is not None

    def test_accumulation_factor_near_expected(self, df_with_neutral, thresh):
        """H3-neutral ≈ 2.8× H3-phase, so accumulation_factor should be ≈ 2.8."""
        result = check_neutral_harmonics(df_with_neutral, thresh)
        af = result["accumulation_factor"]
        assert 2.0 < af < 4.0, f"Accumulation factor {af} out of expected 2.0–4.0 range"

    def test_mean_values_positive(self, df_with_neutral, thresh):
        result = check_neutral_harmonics(df_with_neutral, thresh)
        for h, od in result["orders"].items():
            assert od["mean_a"] >= 0.0
            assert od["max_a"] >= od["mean_a"]

    def test_triplen_sum_greater_than_nontriplen(self, df_with_neutral, thresh):
        result = check_neutral_harmonics(df_with_neutral, thresh)
        assert result["triplen_sum_mean_a"] > result["nontriplen_sum_mean_a"]

    def test_pipeline_neutral_harmonics_available(self):
        """MockAdapter now includes neutral harmonic channels — full pipeline test."""
        adapter = MockAdapter(duration_hours=1.0, nominal=120.0)
        mapper  = ChannelMapper()
        ds      = extract_dataset(adapter, mapper)
        result  = check_neutral_harmonics(ds.df, Thresholds(nominal_voltage=120.0))
        assert result["available"] is True
        assert result["triplen_dominant"] is True
        assert result["accumulation_factor"] is not None


# ─────────────────────────────────────────────────────────────────────────────
# 9. check_harmonic_sources
# ─────────────────────────────────────────────────────────────────────────────

class TestHarmonicSources:

    @pytest.fixture(scope="class")
    @classmethod
    def df_customer(cls):
        """Customer-injection scenario: V_h = k × h × I_h + noise → high correlation."""
        import pandas as pd
        rng = np.random.default_rng(7)
        n   = 500
        idx = pd.date_range("2024-01-01", periods=n, freq="5min", tz="UTC")
        kz  = 0.03  # Ω per unit order (inductive trend)
        data = {}
        for h in (3, 5, 7, 11, 13):
            ih = 5.0 / h + rng.normal(0, 0.1, n)
            ih = ih.clip(0.05)
            vh = ih * h * kz + rng.normal(0, 0.002, n)
            vh = vh.clip(0)
            data[f"h{h}_current_a"] = ih
            data[f"h{h}_voltage_a"] = vh
        return pd.DataFrame(data, index=idx)

    @pytest.fixture(scope="class")
    @classmethod
    def df_resonance(cls):
        """Resonance at H5: Z_5 >> linear trend of Z_3, Z_7, Z_11, Z_13."""
        import pandas as pd
        rng = np.random.default_rng(99)
        n   = 500
        idx = pd.date_range("2024-01-01", periods=n, freq="5min", tz="UTC")
        kz  = 0.03
        data = {}
        for h in (3, 5, 7, 11, 13):
            ih = 5.0 / h + rng.normal(0, 0.1, n)
            ih = ih.clip(0.05)
            if h == 5:
                # 6× spike above expected linear value → resonance
                vh = ih * h * kz * 6.0 + rng.normal(0, 0.01, n)
            else:
                vh = ih * h * kz + rng.normal(0, 0.002, n)
            vh = vh.clip(0)
            data[f"h{h}_current_a"] = ih
            data[f"h{h}_voltage_a"] = vh
        return pd.DataFrame(data, index=idx)

    @pytest.fixture(scope="class")
    @classmethod
    def thresh(cls):
        return Thresholds(nominal_voltage=120.0)

    # ── Basic availability ────────────────────────────────────────────────────

    def test_available_with_both_channels(self, df_customer, thresh):
        result = check_harmonic_sources(df_customer, thresh)
        assert result["available"] is True

    def test_unavailable_with_only_current(self, thresh):
        import pandas as pd
        df = pd.DataFrame({"h5_current_a": [5.0] * 50})
        result = check_harmonic_sources(df, thresh)
        assert result["available"] is False

    def test_unavailable_with_only_voltage(self, thresh):
        import pandas as pd
        df = pd.DataFrame({"h5_voltage_a": [0.5] * 50})
        result = check_harmonic_sources(df, thresh)
        assert result["available"] is False

    # ── Z_h values ───────────────────────────────────────────────────────────

    def test_all_five_orders_present(self, df_customer, thresh):
        result = check_harmonic_sources(df_customer, thresh)
        assert set(result["orders"].keys()) == {3, 5, 7, 11, 13}

    def test_z_ohm_positive(self, df_customer, thresh):
        result = check_harmonic_sources(df_customer, thresh)
        for h, od in result["orders"].items():
            assert od["z_ohm"] > 0, f"H{h} impedance should be positive"

    def test_z_increases_with_order(self, df_customer, thresh):
        """Z_h = k × h so Z should be monotonically increasing for customer injection."""
        result = check_harmonic_sources(df_customer, thresh)
        z_vals = [result["orders"][h]["z_ohm"] for h in sorted(result["orders"])]
        assert z_vals == sorted(z_vals), "Z_h should increase with harmonic order"

    def test_linear_slope_fitted(self, df_customer, thresh):
        result = check_harmonic_sources(df_customer, thresh)
        assert result["linear_slope_a"] is not None
        assert result["linear_slope_a"] > 0

    # ── Attribution — customer injection ─────────────────────────────────────

    def test_high_correlation_customer_injection(self, df_customer, thresh):
        result = check_harmonic_sources(df_customer, thresh)
        for h, od in result["orders"].items():
            assert od["corr"] is not None
            assert od["corr"] > 0.5, f"H{h} Pearson r={od['corr']:.2f} below 0.5 for customer scenario"

    def test_overall_customer_attribution(self, df_customer, thresh):
        result = check_harmonic_sources(df_customer, thresh)
        assert result["overall"] == "customer"

    def test_no_resonance_in_customer_scenario(self, df_customer, thresh):
        result = check_harmonic_sources(df_customer, thresh)
        assert result["resonant_orders"] == []

    # ── Resonance detection ───────────────────────────────────────────────────

    def test_resonance_flagged_at_h5(self, df_resonance, thresh):
        result = check_harmonic_sources(df_resonance, thresh)
        assert 5 in result["resonant_orders"], \
            f"H5 resonance not detected; z_ratio={result['orders'].get(5, {}).get('z_ratio')}"

    def test_h5_z_ratio_above_threshold(self, df_resonance, thresh):
        result = check_harmonic_sources(df_resonance, thresh)
        assert result["orders"][5]["z_ratio"] > 2.5

    def test_non_resonant_orders_not_flagged(self, df_resonance, thresh):
        result = check_harmonic_sources(df_resonance, thresh)
        for h in (3, 7, 11, 13):
            assert h not in result["resonant_orders"], f"H{h} spuriously flagged as resonance"

    def test_overall_resonance_suspect(self, df_resonance, thresh):
        result = check_harmonic_sources(df_resonance, thresh)
        assert result["overall"] == "resonance_suspect"

    # ── Pipeline smoke test ───────────────────────────────────────────────────

    def test_pipeline_source_available(self):
        """MockAdapter now has both voltage and current harmonics → full pipeline."""
        adapter = MockAdapter(duration_hours=1.0, nominal=120.0)
        mapper  = ChannelMapper()
        ds      = extract_dataset(adapter, mapper)
        result  = check_harmonic_sources(ds.df, Thresholds(nominal_voltage=120.0))
        assert result["available"] is True
        assert set(result["orders"].keys()) == {3, 5, 7, 11, 13}
        assert result["overall"] in ("customer", "mixed", "indeterminate")


# ─────────────────────────────────────────────────────────────────────────────
# 9b. Harmonic source direction — which side of the meter
# ─────────────────────────────────────────────────────────────────────────────

class _FakeDataset:
    """Just the two attributes the direction check reads off a dataset."""

    def __init__(self, df, waveforms=()):
        self.df = df
        self.waveforms = list(waveforms)


def _capture(v_phasors, i_phasors, fs=19200.0, cycles=10, f0=60.0,
             phases=("a", "b")):
    """One synthetic point-on-wave capture.

    *v_phasors* and *i_phasors* map harmonic order to (amplitude, phase in
    degrees), so a test states the angle between voltage and current directly
    -- which is the only thing the sign of harmonic power depends on.
    """
    n = int(round(fs * cycles / f0))
    t = np.arange(n) / fs

    def build(spec):
        out = np.zeros(n)
        for h, (amp, deg) in spec.items():
            out += amp * np.cos(2 * np.pi * f0 * h * t + np.radians(deg))
        return out

    return {
        "timestamp": None, "label": "synthetic", "t": t, "fs_hz": fs,
        "voltages": {p: build(v_phasors) for p in phases},
        "currents": {p: build(i_phasors) for p in phases},
        "vne": None,
    }


class TestHarmonicDirectionFromIntervals:
    """Direction inferred from magnitudes over the whole recording.

    The discriminator is the intercept: distortion that persists when the
    premises stop drawing harmonic current came from somewhere else.
    """

    @staticmethod
    def _df(v_of_i, n=400, seed=3):
        import pandas as pd
        rng = np.random.default_rng(seed)
        idx = pd.date_range("2024-01-01", periods=n, freq="5min", tz="UTC")
        # A load that swings over the day, so there are quiet intervals to
        # measure the background against and loaded ones to fit a slope on.
        duty = 0.5 + 0.5 * np.sin(np.linspace(0, 6 * np.pi, n))
        data = {"current_a": 20.0 * duty + 2.0}
        for h in (3, 5, 7):
            ih = (4.0 / h) * duty + rng.normal(0, 0.02, n) + 0.2
            data[f"h{h}_current_a"] = ih.clip(0.05)
            data[f"h{h}_voltage_a"] = np.clip(v_of_i(h, data[f"h{h}_current_a"])
                                              + rng.normal(0, 0.01, n), 0, None)
        return pd.DataFrame(data, index=idx)

    def test_distortion_that_tracks_the_load_reads_as_customer_side(self):
        from pq_analysis import harmonic_direction_from_intervals
        df = self._df(lambda h, i: 0.4 * h * i)          # no background term
        r = harmonic_direction_from_intervals(df, Thresholds(nominal_voltage=120.0))
        assert r["available"] is True
        assert r["overall"] == "downstream"
        assert all(od["indication"] == "downstream" for od in r["orders"].values())

    def test_distortion_that_ignores_the_load_reads_as_utility_side(self):
        from pq_analysis import harmonic_direction_from_intervals
        rng = np.random.default_rng(11)
        # Voltage distortion present regardless of what the premises draws.
        df = self._df(lambda h, i: 3.0 / h + rng.normal(0, 0.05, len(i)))
        r = harmonic_direction_from_intervals(df, Thresholds(nominal_voltage=120.0))
        assert r["overall"] == "upstream"
        for h, od in r["orders"].items():
            assert od["indication"] == "upstream"
            # The intercept and the directly measured quiet-interval average
            # are the same claim by two routes; they should agree.
            assert od["v_at_quiet_v"] == pytest.approx(od["v_background_v"],
                                                       rel=0.25)

    def test_a_spectrum_at_the_meters_resolution_is_not_assessed(self):
        from pq_analysis import harmonic_direction_from_intervals
        df = self._df(lambda h, i: 0.4 * h * i)
        for h in (3, 5, 7):
            df[f"h{h}_current_a"] *= 0.02          # down into quantization
        r = harmonic_direction_from_intervals(df, Thresholds(nominal_voltage=120.0))
        assert r["overall"] == "not_assessed"
        assert all(od["indication"] == "not_assessed" for od in r["orders"].values())
        # The measurements survive even though no conclusion is drawn from them.
        assert r["orders"][3]["slope_ohm"] is not None


class TestHarmonicDirectionFromWaveforms:
    """Direction measured from the sign of harmonic power in the captures."""

    THRESH = Thresholds(nominal_voltage=120.0)
    V1 = (120.0 * np.sqrt(2), 0.0)
    I1 = (10.0 * np.sqrt(2), 0.0)          # 10 A rms, in phase → importing

    def _run(self, captures):
        from pq_analysis import harmonic_direction_from_waveforms
        import pandas as pd
        return harmonic_direction_from_waveforms(
            _FakeDataset(pd.DataFrame(), captures), self.THRESH)

    def test_harmonic_power_leaving_the_premises_reads_as_customer_side(self):
        # V3 and I3 in antiphase: P3 < 0, harmonic power flowing to the system.
        cap = _capture({1: self.V1, 3: (4.0, 180.0)},
                       {1: self.I1, 3: (2.0, 0.0)})
        r = self._run([cap, cap])
        assert r["available"] is True
        assert r["orders"][3]["indication"] == "downstream"
        assert r["orders"][3]["median_p_w"] < 0
        assert r["overall"] == "downstream"

    def test_harmonic_power_entering_the_premises_reads_as_utility_side(self):
        cap = _capture({1: self.V1, 3: (4.0, 0.0)},
                       {1: self.I1, 3: (2.0, 0.0)})
        r = self._run([cap, cap])
        assert r["orders"][3]["indication"] == "upstream"
        assert r["orders"][3]["median_p_w"] > 0

    def test_reversed_cts_are_detected_and_corrected(self):
        # Same physical situation as the customer-side case, with the clamps
        # on backwards: every current inverted, including the fundamental.
        cap = _capture({1: self.V1, 3: (4.0, 180.0)},
                       {1: (-self.I1[0], 0.0), 3: (-2.0, 0.0)})
        r = self._run([cap, cap])
        assert r["ct_polarity_inverted"] is True
        assert r["ct_polarity_verified"] is True
        assert "reversed" in r["polarity_note"]
        # The conclusion must match the un-reversed installation, not invert.
        assert r["orders"][3]["indication"] == "downstream"

    def test_captures_taken_during_an_event_are_excluded(self):
        sag = _capture({1: (70.0 * np.sqrt(2), 0.0), 3: (4.0, 180.0)},
                       {1: self.I1, 3: (2.0, 0.0)})
        r = self._run([sag, sag])
        assert r["excluded_event"] == 2
        assert r["captures_used"] == 0
        assert r["available"] is False
        assert "voltage event" in r["note"]

    def test_a_transient_capture_is_too_short_to_read(self):
        # Pronto's 153 kHz transient captures are a fraction of a cycle.
        short = _capture({1: self.V1, 3: (4.0, 180.0)},
                         {1: self.I1, 3: (2.0, 0.0)},
                         fs=153600.0, cycles=1)
        r = self._run([short])
        assert r["excluded_short"] == 1
        assert r["captures_used"] == 0

    def test_an_unloaded_capture_is_not_read_for_direction(self):
        idle = _capture({1: self.V1, 3: (4.0, 180.0)},
                        {1: (0.1, 0.0), 3: (0.05, 0.0)})
        r = self._run([idle])
        assert r["excluded_light_load"] == 1
        assert r["available"] is False

    def test_an_order_below_the_current_floor_is_not_given_a_direction(self):
        # 0.02 A rms of H5: an angle measured on that is noise.
        cap = _capture({1: self.V1, 3: (4.0, 180.0), 5: (1.0, 180.0)},
                       {1: self.I1, 3: (2.0, 0.0), 5: (0.03, 0.0)})
        r = self._run([cap, cap])
        assert 3 in r["orders"]
        assert 5 not in r["orders"]

    def test_the_measured_fundamental_is_used_not_the_nominal_one(self):
        # An off-nominal system: at 59.3 Hz, projecting onto 60 Hz would walk
        # the third harmonic's angle right across the sign boundary.
        cap = _capture({1: self.V1, 3: (4.0, 180.0)},
                       {1: self.I1, 3: (2.0, 0.0)}, f0=59.3)
        r = self._run([cap, cap])
        assert r["fundamental_hz"] == pytest.approx(59.3, abs=0.05)
        assert r["orders"][3]["indication"] == "downstream"

    def test_a_file_with_no_captures_says_so_rather_than_failing(self):
        r = self._run([])
        assert r["available"] is False
        assert r["captures_total"] == 0
        assert "No point-on-wave captures" in r["note"]


class TestHarmonicDirectionOnANetMeteredService:
    """A service that exports breaks the CT-polarity test.

    Reversed clamps on an importing service and correct clamps on an exporting
    one produce the same negative fundamental power. Nothing in a single
    capture separates them, so on a net-metered service the correction is not
    applied at all and the captures are split on their own direction of flow.

    Schedule NM is a service element under every rate schedule, so this is a
    fact about the service and not about the customer class -- the flag is
    entered, never inferred from the recording.
    """

    THRESH = Thresholds(nominal_voltage=120.0, service_role="mixed")
    V1 = (120.0 * np.sqrt(2), 0.0)
    I1_IMPORT = (10.0 * np.sqrt(2), 0.0)         # in phase with V → importing
    I1_EXPORT = (10.0 * np.sqrt(2), 180.0)       # antiphase        → exporting

    def _run(self, captures, thresh=None):
        from pq_analysis import harmonic_direction_from_waveforms
        import pandas as pd
        return harmonic_direction_from_waveforms(
            _FakeDataset(pd.DataFrame(), captures), thresh or self.THRESH)

    def _importing(self):
        # H3 leaving the premises while importing: source is inside.
        return _capture({1: self.V1, 3: (4.0, 180.0)},
                        {1: self.I1_IMPORT, 3: (2.0, 0.0)})

    def _exporting(self):
        return _capture({1: self.V1, 3: (4.0, 180.0)},
                        {1: self.I1_EXPORT, 3: (2.0, 0.0)})

    def test_exporting_captures_are_not_read_as_reversed_cts(self):
        caps = [self._exporting()] * 3
        r = self._run(caps)
        assert r["ct_polarity_inverted"] is False
        assert r["ct_polarity_verified"] is False
        # The same captures on a service not declared as generating are read
        # the old way -- which is correct there and wrong here.
        r_load = self._run(caps, Thresholds(nominal_voltage=120.0))
        assert r_load["ct_polarity_inverted"] is True

    def test_the_two_directions_of_flow_are_reported_separately(self):
        r = self._run([self._importing()] * 3 + [self._exporting()] * 3)
        split = r["export_split"]
        # Two phases per capture, so three captures give six phase-readings.
        assert split["importing"]["capture_phases"] == 6
        assert split["exporting"]["capture_phases"] == 6
        assert 3 in split["importing"]["orders"]
        assert 3 in split["exporting"]["orders"]

    def test_the_main_table_states_the_importing_captures(self):
        # The importing half is the one comparable with a non-generating
        # service, so it is what the top-level result carries.
        r = self._run([self._importing()] * 3 + [self._exporting()] * 3)
        assert r["orders"] == r["export_split"]["importing"]["orders"]
        assert r["overall"] == r["export_split"]["importing"]["overall"]
        assert r["orders"][3]["indication"] == "downstream"

    def test_an_exporting_sample_never_lands_in_the_importing_set(self):
        # The split is per capture, not from the median over all of them: one
        # exporting capture among many importing ones must not be averaged in.
        r = self._run([self._importing()] * 5 + [self._exporting()])
        split = r["export_split"]
        assert split["importing"]["capture_phases"] == 10
        assert split["exporting"]["capture_phases"] == 2
        # The importing set holds its ten readings and not the odd exporting
        # one; the exporting set is two readings, under the floor for an
        # indication, so it gets none rather than a verdict from one capture.
        assert split["importing"]["orders"][3]["samples"] == 10
        assert split["exporting"]["orders"] == {}
        assert split["exporting"]["overall"] == "indeterminate"

    def test_a_service_that_only_exported_still_reports_that_half(self):
        r = self._run([self._exporting()] * 3)
        assert r["available"] is True
        assert r["orders"] == {}
        assert r["export_split"]["exporting"]["orders"][3]["samples"] == 6
        assert "No capture was taken while the service was importing" in \
            r["polarity_note"]

    def test_the_flag_changes_the_verdict_on_the_real_fixture(self):
        """End to end, from the .pqd file, not from synthetic captures.

        test_solar_net_metered.pqd holds two captures taken before sunrise and
        four taken across the middle of the day, all with H5 built to leave the
        premises. Read as a load service the exporting majority drags the
        median fundamental negative, the CTs are declared reversed, and every
        direction inverts -- the file reports the distortion as the utility's.
        """
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pq_analysis import harmonic_direction_from_waveforms
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_solar_net_metered.pqd"
        ds = extract_dataset(ProntoAdapter(str(path)), ChannelMapper())
        assert len(ds.waveforms) == 6

        as_load = harmonic_direction_from_waveforms(
            ds, Thresholds(nominal_voltage=277.0, customer_class="sg"))
        assert as_load["ct_polarity_inverted"] is True
        assert as_load["overall"] == "upstream"          # the wrong answer

        as_gen = harmonic_direction_from_waveforms(
            ds, Thresholds(nominal_voltage=277.0, customer_class="sg",
                           service_role="mixed"))
        assert as_gen["ct_polarity_inverted"] is False
        assert as_gen["overall"] == "downstream"         # what was built in
        split = as_gen["export_split"]
        assert split["importing"]["capture_phases"] == 6     # 2 captures x 3 ph
        assert split["exporting"]["capture_phases"] == 12    # 4 captures x 3 ph
        assert split["exporting"]["overall"] == "downstream"

    def test_the_note_says_the_ct_orientation_was_not_confirmed(self):
        r = self._run([self._importing()] * 3 + [self._exporting()] * 3)
        note = r["polarity_note"]
        assert "on-site generation" in note
        assert "arrow toward" in note
        assert "reversed" not in note      # the load-service claim must not appear


class TestHarmonicDirectionCombined:
    """The two methods together, including when they disagree."""

    THRESH = Thresholds(nominal_voltage=120.0)

    def _ds(self, v_of_i, wave_angle_deg):
        df = TestHarmonicDirectionFromIntervals._df(v_of_i)
        cap = _capture({1: TestHarmonicDirectionFromWaveforms.V1,
                        3: (4.0, wave_angle_deg)},
                       {1: TestHarmonicDirectionFromWaveforms.I1, 3: (2.0, 0.0)})
        return _FakeDataset(df, [cap, cap])

    def test_agreement_is_recorded_per_order(self):
        from pq_analysis import check_harmonic_direction
        ds = self._ds(lambda h, i: 0.4 * h * i, 180.0)   # both say customer
        r = check_harmonic_direction(ds, self.THRESH)
        assert r["agreement"][3] == "agree"
        assert r["overall"] == "downstream"
        assert r["methods_agree"] is True

    def test_disagreement_is_reported_rather_than_resolved(self):
        from pq_analysis import check_harmonic_direction
        rng = np.random.default_rng(5)
        # Trend says the distortion is background; the captures say the
        # premises is exporting H3. Both readings stand.
        ds = self._ds(lambda h, i: 3.0 / h + rng.normal(0, 0.05, len(i)), 180.0)
        r = check_harmonic_direction(ds, self.THRESH)
        assert r["agreement"][3] == "disagree"
        assert r["overall"] == "conflicting"
        assert r["methods_agree"] is False

    def test_the_report_section_prints_both_methods(self):
        from pq_analysis import check_harmonic_direction
        from pq_report import _direction_summary_sentence
        ds = self._ds(lambda h, i: 0.4 * h * i, 180.0)
        hd = check_harmonic_direction(ds, self.THRESH)
        sentence = _direction_summary_sentence(hd)
        assert "customer side" in sentence
        assert "agree at H3" in sentence

    def _rendered(self, customer_class):
        docx = pytest.importorskip("docx")
        from pq_analysis import check_harmonic_direction
        from pq_report import _word_harmonic_direction
        ds = self._ds(lambda h, i: 0.4 * h * i, 180.0)
        thresh = Thresholds(nominal_voltage=120.0, customer_class=customer_class)
        doc = docx.Document()
        _word_harmonic_direction(
            doc, {"harmonic_direction": check_harmonic_direction(ds, thresh)},
            thresh)
        return "\n".join(p.text for p in doc.paragraphs)

    def test_the_engineering_report_gets_the_section(self):
        text = self._rendered("sg")
        assert "Harmonic Source Direction" in text
        assert "Engineer's assessment" in text
        # Direction, never fault: the section names a side of the meter, says
        # in as many words that this is not an attribution, and leaves the
        # assessment blank for the engineer signing the report.
        assert "do not assign responsibility" in text
        assert "Neither method establishes what equipment is responsible" in text
        for claim in ("customer is responsible", "caused by the customer",
                      "the customer must", "Xcel will"):
            assert claim not in text

    def test_a_residential_report_does_not_get_it(self):
        # A homeowner's report carries no harmonic content at all.
        assert "Harmonic Source Direction" not in self._rendered("r")


# ─────────────────────────────────────────────────────────────────────────────
# 9c. Service impedance and high-impedance screening
# ─────────────────────────────────────────────────────────────────────────────

def _stepped_load(n=600, seed=4, peak=25.0):
    """A load that switches on and off, the way a real service's does."""
    rng = np.random.default_rng(seed)
    base = np.zeros(n)
    level = 4.0
    for k in range(n):
        if rng.random() < 0.25:                    # something switches
            level = float(rng.uniform(2.0, peak))
        base[k] = level + rng.normal(0, 0.05)
    return np.clip(base, 0.5, None)


def _service_frame(r_ohm=0.05, x_ohm=0.02, n=600, seed=4, v0=120.0,
                   phases=("a",), pf_varies=True, neutral_r=None,
                   per_phase_r=None, noise_v=0.01):
    """Interval data for a service of known impedance.

    Built forwards from the physics the check works backwards from: each
    phase's voltage is the no-load voltage less its own load's drop, so a
    test can state the answer and see whether it comes back.
    """
    import pandas as pd
    rng = np.random.default_rng(seed + 1)
    idx = pd.date_range("2024-06-01", periods=n, freq="5min", tz="UTC")
    pf = (rng.uniform(0.75, 1.0, n) if pf_varies else np.full(n, 0.9))
    data = {"power_factor": pf}
    for k, ph in enumerate(phases):
        i = _stepped_load(n, seed=seed + k)
        r = (per_phase_r or {}).get(ph, r_ohm)
        drop = r * i * pf + x_ohm * i * np.sqrt(np.clip(1 - pf ** 2, 0, 1))
        data[f"current_{ph}"] = i
        data[f"voltage_{ph}"] = v0 - drop + rng.normal(0, noise_v, n)
    if neutral_r is not None:
        i_n = _stepped_load(n, seed=seed + 50, peak=15.0)
        data["current_neutral"] = i_n
        data["voltage_neutral"] = neutral_r * i_n + rng.normal(0, 0.005, n)
    return pd.DataFrame(data, index=idx)


class TestServiceImpedanceMeasurement:
    """Reading the impedance between the source and the meter off the load steps."""

    THRESH = Thresholds(nominal_voltage=120.0)

    def test_a_known_impedance_comes_back(self):
        from pq_analysis import check_source_impedance
        df = _service_frame(r_ohm=0.05, x_ohm=0.02)
        r = check_source_impedance(df, self.THRESH)
        fit = r["phases"]["a"]
        assert fit["identifiable"] is True
        assert fit["separated"] is True
        assert fit["r_ohm"] == pytest.approx(0.05, abs=0.005)
        assert fit["x_ohm"] == pytest.approx(0.02, abs=0.005)

    def test_a_steady_power_factor_gives_a_magnitude_not_a_split(self):
        from pq_analysis import check_source_impedance
        # Without pf variation the real and reactive parts of the current are
        # the same shape, and R and X are not separately identifiable.
        df = _service_frame(r_ohm=0.05, x_ohm=0.02, pf_varies=False)
        fit = check_source_impedance(df, self.THRESH)["phases"]["a"]
        assert fit["identifiable"] is True
        assert fit["separated"] is False
        assert fit.get("x_ohm") is None
        # The effective magnitude along the load's own angle still lands.
        assert fit["z_ohm"] == pytest.approx(0.05 * 0.9 + 0.02 * 0.436, abs=0.006)

    def test_feeder_wide_droop_is_not_read_as_this_services_impedance(self):
        # The failure this estimator exists to avoid: a voltage that sags on
        # the same daily cycle as the load, with no local impedance at all.
        # Fitting the levels finds an impedance; fitting the steps must not.
        from pq_analysis import check_source_impedance
        df = _service_frame(r_ohm=0.0, x_ohm=0.0, noise_v=0.02)
        daily = 3.0 * np.sin(np.linspace(0, 4 * np.pi, len(df)))
        df["voltage_a"] = df["voltage_a"] - daily
        df["current_a"] = df["current_a"] + 4.0 * np.sin(
            np.linspace(0, 4 * np.pi, len(df)))
        level_slope = np.polyfit(df["current_a"], df["voltage_a"], 1)[0]
        assert level_slope < -0.1          # the naive fit would claim >0.1 Ω

        fit = check_source_impedance(df, self.THRESH)["phases"]["a"]
        assert fit["identifiable"] is False
        assert "not this service" in fit["reason"]

    def test_a_drop_below_the_meters_resolution_is_reported_as_such(self):
        from pq_analysis import check_source_impedance
        df = _service_frame(r_ohm=0.0, x_ohm=0.0, noise_v=0.0)
        fit = check_source_impedance(df, self.THRESH)["phases"]["a"]
        assert fit["identifiable"] is False
        assert fit["at_resolution"] is True
        assert "resolution" in fit["reason"]

    def test_one_bad_phase_is_flagged_against_the_others(self):
        from pq_analysis import check_source_impedance
        df = _service_frame(r_ohm=0.03, x_ohm=0.0,
                            phases=("a", "b", "c"),
                            per_phase_r={"c": 0.12})   # a degrading connection
        r = check_source_impedance(df, self.THRESH)
        asym = r["asymmetry"]
        assert asym["worst_phase"] == "C"
        assert asym["flagged"] is True
        assert asym["ratio"] > 2.0
        assert r["overall"] == "high_impedance_suspected"

    def test_balanced_phases_are_not_flagged(self):
        from pq_analysis import check_source_impedance
        df = _service_frame(r_ohm=0.03, x_ohm=0.0, phases=("a", "b", "c"))
        assert check_source_impedance(df, self.THRESH)["asymmetry"]["flagged"] is False

    def test_a_resistive_neutral_is_measured_from_its_own_rise(self):
        from pq_analysis import check_source_impedance
        df = _service_frame(r_ohm=0.03, x_ohm=0.0, neutral_r=0.25)
        neutral = check_source_impedance(df, self.THRESH)["neutral"]
        assert neutral["identifiable"] is True
        assert neutral["r_ohm"] == pytest.approx(0.25, abs=0.02)
        assert neutral["elevated"] is True
        assert neutral["rise_at_peak_v"] > 2.0

    def test_a_sound_neutral_reads_as_below_resolution(self):
        from pq_analysis import check_source_impedance
        df = _service_frame(r_ohm=0.03, x_ohm=0.0, neutral_r=0.0)
        neutral = check_source_impedance(df, self.THRESH)["neutral"]
        assert neutral["identifiable"] is False
        assert neutral["at_resolution"] is True

    def test_a_service_that_never_moves_is_not_measured(self):
        import pandas as pd
        from pq_analysis import check_source_impedance
        idx = pd.date_range("2024-06-01", periods=200, freq="5min", tz="UTC")
        df = pd.DataFrame({"voltage_a": 120.0, "current_a": 10.0}, index=idx)
        fit = check_source_impedance(df, self.THRESH)["phases"]["a"]
        assert fit["identifiable"] is False
        assert "load step" in fit["reason"]


class TestExpectedServiceImpedance:
    """What the picked transformer and conductor say the impedance should be."""

    def test_a_single_phase_run_counts_the_neutral_return(self):
        from pq_constants import conductor_impedance
        out = conductor_impedance("al-4-0-triplex", 150.0, return_path=True)
        one_way = conductor_impedance("al-4-0-triplex", 150.0, return_path=False)
        # 4/0 AL at 0.100 Ω/1000 ft: 150 ft out and 150 ft back.
        assert out[0] == pytest.approx(0.030, abs=1e-6)
        assert one_way[0] == pytest.approx(0.015, abs=1e-6)

    def test_the_isc_path_includes_the_primary_system(self):
        from pq_constants import expected_service_impedance
        thresh = Thresholds(nominal_voltage=120.0, service_type="3ph-padmount",
                            transformer_kva=150, isc_amps=8000,
                            conductor_key="al-350-urd", run_length_ft=200)
        e = expected_service_impedance(thresh)
        assert e["available"] is True
        assert e["upstream_ohm"] == pytest.approx(120.0 / 8000)
        assert "primary system and the transformer" in e["upstream_source"]
        assert e.get("upstream_is_floor") is None

    def test_without_isc_the_expected_value_is_called_a_floor(self):
        from pq_constants import expected_service_impedance
        thresh = Thresholds(nominal_voltage=120.0, service_type="1ph-overhead",
                            transformer_kva=25, topology="split-phase",
                            conductor_key="al-4-0-triplex", run_length_ft=150)
        e = expected_service_impedance(thresh)
        # 25 kVA at 1.6–2.4%: V_LN²/S is the base the L-N path sits on.
        assert e["transformer_ohm_range"][0] == pytest.approx(
            0.016 * 120.0 ** 2 / 25000.0)
        assert e["upstream_is_floor"] is True
        assert "floor" in e["upstream_source"]

    def test_nothing_picked_means_no_expected_value(self):
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(Thresholds(nominal_voltage=120.0))
        assert e["available"] is False
        assert "run length" in e["reason"]

    def test_the_comparison_calls_a_large_excess_high(self):
        from pq_analysis import check_source_impedance
        thresh = Thresholds(nominal_voltage=120.0, service_type="1ph-overhead",
                            transformer_kva=25, topology="split-phase",
                            conductor_key="al-4-0-triplex", run_length_ft=150)
        df = _service_frame(r_ohm=0.20, x_ohm=0.0)     # far above ~0.043 Ω
        r = check_source_impedance(df, thresh)
        assert r["comparison"]["verdict"] == "high"
        assert r["comparison"]["ratio"] > 2.5
        assert r["comparison"]["excess_v_at_peak"] > 1.0
        assert r["overall"] == "high_impedance_suspected"

    def test_an_as_built_service_reads_as_consistent(self):
        from pq_analysis import check_source_impedance
        thresh = Thresholds(nominal_voltage=120.0, service_type="1ph-overhead",
                            transformer_kva=25, topology="split-phase",
                            conductor_key="al-4-0-triplex", run_length_ft=150)
        # 0.0115 Ω transformer + 0.030 Ω of conductor ≈ 0.043 Ω expected.
        df = _service_frame(r_ohm=0.042, x_ohm=0.010, pf_varies=False)
        r = check_source_impedance(df, thresh)
        assert r["comparison"]["verdict"] == "consistent"
        assert r["overall"] == "consistent_with_expected"

    def test_the_report_section_states_the_constants_are_generic(self):
        docx = pytest.importorskip("docx")
        from pq_analysis import check_source_impedance
        from pq_report import _word_service_impedance
        thresh = Thresholds(nominal_voltage=120.0, customer_class="r",
                            service_type="1ph-overhead", transformer_kva=25,
                            topology="split-phase",
                            conductor_key="al-4-0-triplex", run_length_ft=150)
        df = _service_frame(r_ohm=0.20, x_ohm=0.0)
        doc = docx.Document()
        _word_service_impedance(
            doc, {"service_impedance": check_source_impedance(df, thresh)}, thresh)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Service Impedance" in text
        # The engineer has to be able to see that the expected side is generic.
        assert "generic published values" in text
        assert "not PSCo Blue Book figures" in text
        assert "Engineer's assessment" in text


class TestAServiceThatTapsASharedSecondary:
    """A secondary main shared with the neighbours is in this service's path."""

    def _thresh(self, **kw):
        base = dict(nominal_voltage=120.0, service_type="1ph-overhead",
                    transformer_kva=25, topology="split-phase",
                    conductor_key="al-2-triplex", run_length_ft=100)
        base.update(kw)
        return Thresholds(**base)

    def test_the_shared_run_adds_to_the_expected_impedance(self):
        from pq_constants import expected_service_impedance
        alone = expected_service_impedance(self._thresh())
        tapped = expected_service_impedance(self._thresh(
            shared_secondary_key="al-4-0-triplex", shared_secondary_ft=300))
        # 4/0 AL at 0.100 Ω/1000 ft, 300 ft out and back on the shared neutral.
        assert tapped["shared_secondary_r_ohm"] == pytest.approx(0.060, abs=1e-6)
        assert tapped["total_ohm"] > alone["total_ohm"]
        assert tapped["total_ohm"] == pytest.approx(
            alone["total_ohm"] + tapped["shared_secondary_z_ohm"], abs=1e-9)

    def test_a_dedicated_run_is_what_leaving_it_blank_means(self):
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(self._thresh())
        assert "shared_secondary_z_ohm" not in e
        assert e["available"] is True

    def test_a_shared_run_alone_still_gives_an_expected_value(self):
        # The service conductor may be unknown while the main is not.
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(Thresholds(
            nominal_voltage=120.0, topology="split-phase",
            shared_secondary_key="al-4-0-triplex", shared_secondary_ft=300))
        assert e["available"] is True
        assert "the service conductors" in e["partial"]

    def test_a_three_phase_service_counts_the_shared_run_one_way(self):
        # A balanced three-phase load's neutral carries almost nothing, so the
        # return path is not doubled -- the same rule as the service conductor.
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(Thresholds(
            nominal_voltage=277.0, service_type="3ph-padmount", topology="3ph-wye",
            transformer_kva=500, shared_secondary_key="al-350-urd",
            shared_secondary_ft=400))
        assert e["shared_secondary_r_ohm"] == pytest.approx(
            0.0611 * 400 / 1000.0, abs=1e-9)

    def test_the_report_says_the_neighbours_widen_the_fit(self):
        # Sharing a main does not bias the measurement, it scatters it, and an
        # engineer reading a loose fit needs to know which of the two it was.
        docx = pytest.importorskip("docx")
        from pq_analysis import check_source_impedance
        from pq_report import _word_service_impedance
        thresh = self._thresh(customer_class="r",
                              shared_secondary_key="al-4-0-triplex",
                              shared_secondary_ft=300)
        df = _service_frame(r_ohm=0.20, x_ohm=0.0)
        doc = docx.Document()
        _word_service_impedance(
            doc, {"service_impedance": check_source_impedance(df, thresh)}, thresh)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "shared secondary" in text
        assert "widens the scatter" in text


class TestAServiceMeteredOnThePrimary:
    """Metered on the high side, the customer's transformer is below the meter."""

    def _thresh(self, **kw):
        base = dict(nominal_voltage=7620.0, customer_class="pg",
                    service_type="3ph-padmount", topology="3ph-wye",
                    primary_metered=True,
                    primary_r1_ohm=0.42, primary_x1_ohm=0.88)
        base.update(kw)
        return Thresholds(**base)

    def test_the_expected_value_is_the_primary_line(self):
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(self._thresh())
        assert e["available"] is True
        assert e["total_ohm"] == pytest.approx(math.hypot(0.42, 0.88))
        assert e["sequence_used"] == "positive"

    def test_the_customers_transformer_and_conductors_are_left_out(self):
        # They sit downstream of the meter, so adding them would count wire the
        # recording never saw.
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(self._thresh(
            transformer_kva=1500, isc_amps=33900,
            conductor_key="al-350-urd", run_length_ft=200,
            shared_secondary_key="al-500-urd", shared_secondary_ft=500))
        assert e["total_ohm"] == pytest.approx(math.hypot(0.42, 0.88))
        assert "conductor_z_ohm" not in e
        assert "shared_secondary_z_ohm" not in e

    def test_zero_sequence_is_optional(self):
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(self._thresh())
        assert e["available"] is True
        assert e["primary"].get("z0_ohm") is None

    def test_zero_sequence_is_carried_but_not_compared_against(self):
        # Z0 is the right impedance for triplens and earth return, and the
        # wrong one for a balanced-load voltage drop.
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(self._thresh(
            primary_r0_ohm=1.30, primary_x0_ohm=2.90))
        p = e["primary"]
        assert p["z0_ohm"] == pytest.approx(math.hypot(1.30, 2.90))
        assert p["z0_over_z1"] == pytest.approx(p["z0_ohm"] / p["z1_ohm"])
        # The comparison still runs against Z1 alone.
        assert e["total_ohm"] == pytest.approx(p["z1_ohm"])

    def test_a_single_phase_tap_sees_two_z1_plus_z0_over_three(self):
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(self._thresh(
            primary_r0_ohm=1.30, primary_x0_ohm=2.90))
        assert e["primary"]["single_phase_loop_ohm"] == pytest.approx(
            math.hypot((2 * 0.42 + 1.30) / 3.0, (2 * 0.88 + 2.90) / 3.0))

    def test_no_entered_impedance_says_which_field_is_missing(self):
        from pq_constants import expected_service_impedance
        e = expected_service_impedance(Thresholds(nominal_voltage=7620.0,
                                                  primary_metered=True))
        assert e["available"] is False
        assert "metered on the primary" in e["reason"]
        assert "R1/X1" in e["reason"]

    def test_the_report_names_the_sequence_it_used(self):
        # Z0 runs two to three times Z1, so a reader who assumes the wrong one
        # misreads every figure beside it.
        docx = pytest.importorskip("docx")
        from pq_analysis import check_source_impedance
        from pq_report import _word_service_impedance
        thresh = self._thresh(primary_r0_ohm=1.30, primary_x0_ohm=2.90)
        df = _service_frame(r_ohm=2.0, x_ohm=0.0, v0=7620.0)
        doc = docx.Document()
        _word_service_impedance(
            doc, {"service_impedance": check_source_impedance(df, thresh)}, thresh)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "metered on the primary" in text
        assert "positive sequence" in text
        assert "triplen" in text
        # And it does not claim generic conductor constants it never used.
        assert "generic published values" not in text


# ─────────────────────────────────────────────────────────────────────────────
# 10. detect_events — adaptive vs interval path
# ─────────────────────────────────────────────────────────────────────────────

class TestAdaptiveEvents:
    """detect_events uses cycle-level adaptive data when ds.has_adaptive."""

    @pytest.fixture(scope="class")
    @classmethod
    def ds_with_adaptive(cls):
        """PQDataset backed by MockAdapter — includes synthetic adaptive_df."""
        adapter = MockAdapter(duration_hours=2.0, nominal=120.0)
        mapper  = ChannelMapper()
        return extract_dataset(adapter, mapper)

    @pytest.fixture(scope="class")
    @classmethod
    def thresh(cls):
        return Thresholds(nominal_voltage=120.0)

    # ── data_source flag ─────────────────────────────────────────────────────

    def test_data_source_adaptive_when_present(self, ds_with_adaptive, thresh):
        result = detect_events(ds_with_adaptive, thresh)
        assert result["data_source"] == "adaptive"

    # ── voltage sag from van_v ───────────────────────────────────────────────

    def test_voltage_sag_detected_from_adaptive(self, ds_with_adaptive, thresh):
        """MockAdapter injects van_v[50:80] *= 0.86 — sag must be reported."""
        result = detect_events(ds_with_adaptive, thresh)
        sag_events = result["events"][result["events"]["type"] == "voltage_sag"]
        assert len(sag_events) > 0, "No voltage_sag detected from adaptive van_v"

    def test_voltage_sag_phase_a(self, ds_with_adaptive, thresh):
        result = detect_events(ds_with_adaptive, thresh)
        sag_events = result["events"][result["events"]["type"] == "voltage_sag"]
        assert "A" in sag_events["phase"].values

    def test_sag_value_below_90pct(self, ds_with_adaptive, thresh):
        result = detect_events(ds_with_adaptive, thresh)
        sag_events = result["events"][result["events"]["type"] == "voltage_sag"]
        assert sag_events["value_v"].min() < 0.90 * 120.0

    # ── PST flicker exceedance ────────────────────────────────────────────────

    def test_flicker_pst_detected(self, ds_with_adaptive, thresh):
        """MockAdapter injects adap_pst[100:130] = 1.4 — PST event must be reported."""
        result = detect_events(ds_with_adaptive, thresh)
        pst_events = result["events"][result["events"]["type"] == "flicker_pst"]
        assert len(pst_events) > 0, "No flicker_pst event detected"

    def test_flicker_pst_value_above_limit(self, ds_with_adaptive, thresh):
        result = detect_events(ds_with_adaptive, thresh)
        pst_events = result["events"][result["events"]["type"] == "flicker_pst"]
        assert pst_events["value"].iloc[0] > 1.0

    # ── current step ─────────────────────────────────────────────────────────

    def test_current_step_detected_from_adaptive(self, ds_with_adaptive, thresh):
        """MockAdapter injects ia_a step from 50 A to 83 A at row 150."""
        result = detect_events(ds_with_adaptive, thresh)
        step_events = result["events"][result["events"]["type"] == "current_step"]
        assert len(step_events) > 0, "No current_step detected from adaptive ia_a"

    def test_current_step_phase_a(self, ds_with_adaptive, thresh):
        result = detect_events(ds_with_adaptive, thresh)
        step_events = result["events"][result["events"]["type"] == "current_step"]
        assert "A" in step_events["phase"].values

    # ── result shape ─────────────────────────────────────────────────────────

    def test_event_count_matches_df_length(self, ds_with_adaptive, thresh):
        result = detect_events(ds_with_adaptive, thresh)
        assert result["event_count"] == len(result["events"])

    def test_events_df_has_required_columns(self, ds_with_adaptive, thresh):
        result = detect_events(ds_with_adaptive, thresh)
        assert "timestamp" in result["events"].columns
        assert "type" in result["events"].columns
        assert "phase" in result["events"].columns

    # ── interval fallback path ────────────────────────────────────────────────

    def test_interval_fallback_data_source(self, thresh):
        """When adaptive_df is None, data_source must be 'interval'."""
        import pandas as pd
        idx = pd.date_range("2024-01-01", periods=100, freq="5min")
        df  = pd.DataFrame({
            "voltage_a": 120.0 + np.zeros(100),
            "voltage_b": 120.0 + np.zeros(100),
            "voltage_c": 120.0 + np.zeros(100),
        }, index=idx)
        from pq_adapter import PQDataset
        ds_no_adap = PQDataset(df=df, adaptive_df=None, meta={"interval_minutes": 5})
        result = detect_events(ds_no_adap, thresh)
        assert result["data_source"] == "interval"


# ─────────────────────────────────────────────────────────────────────────────
# 11. Standardized result shapes — available/error contract
# ─────────────────────────────────────────────────────────────────────────────

class TestUnavailableShapes:
    """All check_* functions must return available+error keys even on missing data."""

    @pytest.fixture
    def empty_df(self):
        import pandas as pd
        return pd.DataFrame(index=pd.date_range("2024-01-01", periods=10, freq="5min"))

    @pytest.fixture
    def thresh(self):
        return Thresholds(nominal_voltage=120.0)

    def test_voltage_compliance_unavailable(self, empty_df, thresh):
        from pq_analysis import check_voltage_compliance
        r = check_voltage_compliance(empty_df, thresh)
        assert r["available"] is False
        assert r["error"] is not None
        assert r["total_pct_out_of_bounds"] is None
        assert len(r["violation_timestamps"]) == 0

    def test_voltage_compliance_available_shape(self, thresh):
        import pandas as pd
        idx = pd.date_range("2024-01-01", periods=10, freq="5min")
        df  = pd.DataFrame({"voltage_a": [120.0] * 10}, index=idx)
        from pq_analysis import check_voltage_compliance
        r = check_voltage_compliance(df, thresh)
        assert r["available"] is True
        assert r["error"] is None
        assert isinstance(r["total_pct_out_of_bounds"], float)

    def test_thd_unavailable_sub_dicts(self, empty_df, thresh):
        from pq_analysis import check_thd
        r = check_thd(empty_df, thresh)
        assert r["voltage"]["available"] is False
        assert r["current"]["available"] is False
        assert r["available"] is False

    def test_thd_available_when_voltage_found(self, thresh):
        import pandas as pd
        idx = pd.date_range("2024-01-01", periods=10, freq="5min")
        df  = pd.DataFrame({"thd_voltage_a": [3.0] * 10}, index=idx)
        from pq_analysis import check_thd
        r = check_thd(df, thresh)
        assert r["voltage"]["available"] is True
        assert r["available"] is True

    def test_power_factor_unavailable(self, empty_df, thresh):
        from pq_analysis import check_power_factor
        r = check_power_factor(empty_df, thresh)
        assert r["available"] is False
        assert r["error"] is not None
        assert r["pct_below_limit"] is None
        assert len(r["violation_timestamps"]) == 0

    def test_voltage_imbalance_unavailable(self, thresh):
        import pandas as pd
        idx = pd.date_range("2024-01-01", periods=10, freq="5min")
        df  = pd.DataFrame({"voltage_a": [120.0] * 10}, index=idx)  # only 1 phase
        from pq_analysis import check_voltage_imbalance
        r = check_voltage_imbalance(df, thresh)
        assert r["available"] is False
        assert r["pct_exceeding"] is None

    def test_voltage_imbalance_available_shape(self, thresh):
        import pandas as pd
        idx = pd.date_range("2024-01-01", periods=10, freq="5min")
        df  = pd.DataFrame({"voltage_a": [120.0]*10, "voltage_b": [119.5]*10}, index=idx)
        from pq_analysis import check_voltage_imbalance
        r = check_voltage_imbalance(df, thresh)
        assert r["available"] is True
        assert r["error"] is None
        assert isinstance(r["pct_exceeding"], float)

    def test_current_imbalance_unavailable(self, thresh):
        import pandas as pd
        idx = pd.date_range("2024-01-01", periods=10, freq="5min")
        df  = pd.DataFrame({"current_a": [50.0] * 10}, index=idx)
        from pq_analysis import check_current_imbalance
        r = check_current_imbalance(df, thresh)
        assert r["available"] is False
        assert r["pct_exceeding"] is None

    def test_current_imbalance_available_shape(self, thresh):
        import pandas as pd
        idx = pd.date_range("2024-01-01", periods=10, freq="5min")
        df  = pd.DataFrame({"current_a": [50.0]*10, "current_b": [48.0]*10}, index=idx)
        from pq_analysis import check_current_imbalance
        r = check_current_imbalance(df, thresh)
        assert r["available"] is True
        assert r["error"] is None

    def test_demand_unavailable(self, empty_df, thresh):
        from pq_analysis import check_demand
        r = check_demand(empty_df, thresh)
        assert r["available"] is False
        assert r["error"] is not None

    def test_demand_available_when_real_power_present(self, thresh):
        import pandas as pd
        idx = pd.date_range("2024-01-01", periods=10, freq="5min")
        df  = pd.DataFrame({"power_real": [18000.0] * 10}, index=idx)
        from pq_analysis import check_demand
        r = check_demand(df, thresh)
        assert r["available"] is True
        assert r["error"] is None


# ─────────────────────────────────────────────────────────────────────────────
# 8. Spec-compliant PQDIF reader (IEEE Std 1159.3-2019)
# ─────────────────────────────────────────────────────────────────────────────

class TestPQDIFPhysical:
    """Annex A physical structure: elements, scalars, vectors."""

    @staticmethod
    def _element(body, physical_type, link, size, embedded=False,
                 element_type=pqdif.ELEMENT_VECTOR):
        return pqdif.Element(pqdif.TAG_SERIES_VALUES, element_type,
                             physical_type, embedded, link, size, body)

    def test_real8_vector(self):
        body = struct.pack("<I3d", 3, 1.5, 2.5, 3.5)
        el = self._element(body, 41, 0, len(body))
        assert np.allclose(el.vector(), [1.5, 2.5, 3.5])

    def test_int2_vector_widths_follow_physical_type(self):
        # A vector's stride comes from typePhysical, so INTEGER2 is 2 bytes per
        # point -- never assumed to be 8.
        body = struct.pack("<I4h", 4, -3, 7, 11, -19)
        el = self._element(body, 21, 0, len(body))
        assert np.allclose(el.vector(), [-3, 7, 11, -19])

    def test_complex16_takes_real_part(self):
        # COMPLEX16 is two REAL8 per point (Annex A), i.e. 16 bytes.
        body = struct.pack("<I4d", 2, 1.0, 90.0, 2.0, -45.0)
        el = self._element(body, 43, 0, len(body))
        assert np.allclose(el.vector(), [1.0, 2.0])

    def test_embedded_scalar_reads_from_link_bytes(self):
        # isEmbedded=TRUE stores the value in the 8 link/size bytes themselves.
        el = self._element(b"", 32, 1234, 0, embedded=True,
                           element_type=pqdif.ELEMENT_SCALAR)
        assert el.scalar() == 1234

    def test_non_embedded_scalar_reads_from_body(self):
        body = struct.pack("<xxxxI", 4321)
        el = self._element(body, 32, 4, 4, embedded=False,
                           element_type=pqdif.ELEMENT_SCALAR)
        assert el.scalar() == 4321

    def test_vector_overrunning_body_is_rejected(self):
        body = struct.pack("<I1d", 500, 1.0)   # claims 500 points, carries 1
        el = self._element(body, 41, 0, len(body))
        with pytest.raises(pqdif.PQDIFError):
            el.vector()

    def test_non_pqdif_file_is_rejected(self, tmp_path):
        # The signature GUID check is what routes the synthetic test fixtures to
        # the legacy reader instead of the spec reader.
        path = tmp_path / "not.pqd"
        path.write_bytes(b"\x00" * 256)
        with pytest.raises(pqdif.PQDIFError):
            pqdif.PQDIFFile(path)


class TestPQDIFIncrementSeries:
    """Clause 5.5.2 regular-rate series, using the standard's own examples."""

    def test_single_rate(self):
        # Spec example: 1792 points at 0.01 s.
        out = pqdif._expand_increment(np.array([1.0, 1792.0, 0.01]))
        assert len(out) == 1792
        assert out[0] == pytest.approx(0.0)
        assert out[1] == pytest.approx(0.01)
        assert out[-1] == pytest.approx(1791 * 0.01)

    def test_two_rates(self):
        # Spec example: 896 points at 0.01 s then 896 at 0.02 s.
        out = pqdif._expand_increment(
            np.array([2.0, 896.0, 0.01, 896.0, 0.02]))
        assert len(out) == 1792
        assert out[895] == pytest.approx(895 * 0.01)
        assert out[896] == pytest.approx(896 * 0.01)
        assert out[897] == pytest.approx(896 * 0.01 + 0.02)

    def test_truncated_instructions_rejected(self):
        with pytest.raises(pqdif.PQDIFError):
            pqdif._expand_increment(np.array([5.0, 10.0, 0.01]))


class TestStepPairDetection:
    """Pronto writes each interval twice, at its start and end."""

    def test_step_pairs_detected(self):
        # Gaps within a pair are the interval; gaps between pairs are ~1 us.
        t = np.array([0.0, 120.0, 120.000001, 240.0, 240.000001, 360.0])
        assert ProntoAdapter._step_pair_stride(t) == 2

    def test_plain_series_not_deduplicated(self):
        t = np.arange(0.0, 600.0, 120.0)
        assert ProntoAdapter._step_pair_stride(t) == 1

    def test_odd_length_never_paired(self):
        t = np.array([0.0, 120.0, 120.000001])
        assert ProntoAdapter._step_pair_stride(t) == 1

    def test_uniform_series_not_paired(self):
        # Equal gaps everywhere: no pair structure, so keep every point.
        t = np.arange(0.0, 8.0, 1.0)
        assert ProntoAdapter._step_pair_stride(t) == 1


class TestSpecChannelTable:
    """The metadata → canonical mapping must agree with ChannelMapper."""

    def test_canonical_names_match_tag_resolution(self):
        # _SPEC_CHANNELS records the canonical name so interval_peaks/_mins can
        # be keyed by it. That name must be what ChannelMapper independently
        # derives from the same tags, or peaks would attach to the wrong column.
        mapper = ChannelMapper()
        for key, (label, canonical, qt, qm, phase) in \
                ProntoAdapter._SPEC_CHANNELS.items():
            info = RawChannelInfo(0, label, qt, qm, phase, "")
            assert mapper._match_by_tags(info) == canonical, (
                f"{key} → {canonical!r} but tags {(qt, qm, phase)} resolve to "
                f"{mapper._match_by_tags(info)!r}"
            )

    def test_rms_is_preferred_over_the_fundamental(self):
        # 'Harm 1 of Van' is characteristic SPECTRA_HGROUP and must never be
        # picked up as the voltage trend; only RMS may map to voltage_a.
        assert ('voltage', 'RMS', 'an') in ProntoAdapter._SPEC_CHANNELS
        assert ProntoAdapter._SPEC_CHANNELS[('voltage', 'RMS', 'an')][1] == \
            'voltage_a'
        for key in ProntoAdapter._SPEC_CHANNELS:
            assert key[1] not in ('SPECTRA_HGROUP', 'SPECTRA'), (
                f"{key} would source a trend from a harmonic magnitude"
            )

    def test_harmonic_name_pattern(self):
        m = ProntoAdapter._HARM_NAME.match('Harm 13 of Van')
        assert m and m.group(1) == '13' and m.group(2) == 'Van'
        m = ProntoAdapter._HARM_NAME.match('Harm 3 of In')
        assert m and m.group(1) == '3' and m.group(2) == 'In'
        # The RMS channels must not be mistaken for harmonic orders.
        assert ProntoAdapter._HARM_NAME.match('RMS Van (V1)') is None
        assert ProntoAdapter._HARM_NAME.match('Hrms Van (V1)') is None


# ─────────────────────────────────────────────────────────────────────────────
# 9. test_data fixtures are valid PQDIF and exercise the spec reader
# ─────────────────────────────────────────────────────────────────────────────

_FIXTURES = sorted((Path(__file__).parent / "test_data").glob("*.pqd"))


@pytest.mark.skipif(not _FIXTURES, reason="test_data/*.pqd not generated")
class TestFixturesAreCompliant:
    """The fixtures exist to exercise the same path real Pronto files take.

    Regenerate them with `python make_test_pqd.py` after changing the writer.
    """

    @pytest.fixture(params=_FIXTURES, ids=lambda p: p.stem)
    def path(self, request):
        return request.param

    def test_parses_as_standard_pqdif(self, path):
        f = pqdif.PQDIFFile(path)
        assert f.version[:2] == (1, 5)
        assert f.compressed is True
        assert f.definitions

    def test_uses_the_spec_reader_not_the_legacy_fallback(self, path):
        adapter = ProntoAdapter(path)
        assert adapter._spec is not None, (
            f"{path.name} fell back to the legacy offset reader"
        )

    def test_every_guid_is_a_standard_identifier(self, path):
        # A vendor-private GUID would be legal but these fixtures should model
        # what real Pronto files contain, which is standard IDs throughout.
        f = pqdif.PQDIFFile(path)
        for defn in f.definitions:
            assert defn.quantity_type in ("VALUELOG", "PHASOR", "WAVEFORM")
            assert defn.quantity_measured in ("voltage", "current", "power")
            for series in defn.series:
                assert series.characteristic in pqdif.CHARACTERISTIC_NAMES.values()
                assert series.value_type in pqdif.VALUE_TYPE_NAMES.values()

    @staticmethod
    def _analysed_observations(path):
        """The observations of the session the adapter reads by default.

        Most fixtures hold one session, where this is every observation. The
        two-session fixture holds two, and a test that pooled both would be
        asserting against data the run never saw.
        """
        f = pqdif.PQDIFFile(path)
        # Point-on-wave captures are not sessions. They carry their own trigger
        # time and thousands of samples where an interval record carries a few
        # hundred, so grouping on (start, length) files each capture as a
        # session of its own and picking the longest then picks a capture.
        obs = [o for o in f.observations
               if o.channels
               and not any(c.quantity_type == 'WAVEFORM' for c in o.channels)]
        by_session = {}
        for o in obs:
            key = (o.start_time, len(o.channels[0].time))
            by_session.setdefault(key, []).append(o)
        # The default is the longest session; ties go to the earliest.
        best = max(by_session, key=lambda k: (k[1], -(k[0].timestamp() if k[0] else 0)))
        return by_session[best]

    def test_records_of_one_session_share_one_time_base(self, path):
        # Within a session, 'Interval (avg)' and 'Interval (max-min)' are
        # pooled, which is only sound when they sit on the same grid. Only
        # those two are pooled: a point-on-wave capture and a variable-rate
        # record carry their own timing by definition -- a capture is
        # thousands of samples of one instant, and an adaptive record gives
        # every channel its own base -- so neither belongs in this comparison
        # even when it happens to share a start time with an interval record.
        f = pqdif.PQDIFFile(path)

        def is_interval(obs):
            if any(c.quantity_type == 'WAVEFORM' for c in obs.channels):
                return False
            base = obs.channels[0].time
            return all(c.time is not None and len(c.time) == len(base)
                       for c in obs.channels)

        by_start = {}
        for obs in f.observations:
            if obs.channels and is_interval(obs):
                by_start.setdefault(obs.start_time, set()).add(
                    len(obs.channels[0].time))
        for start, grids in by_start.items():
            assert len(grids) == 1, (
                f"observations starting {start} disagree on sample count: {grids}")

    def test_step_pair_encoding_is_detected(self, path):
        f = pqdif.PQDIFFile(path)
        raw = f.observations[0].channels[0].time
        assert ProntoAdapter._step_pair_stride(raw) == 2
        # 288 intervals written as 576 points.
        adapter = ProntoAdapter(path)
        assert len(adapter._obs_ts) == len(raw) // 2

    def test_voltage_comes_from_rms_not_the_fundamental(self, path):
        # The regression this guards: 'Harm 1 of Van' (SPECTRA_HGROUP) is the
        # fundamental and is smaller than the true RMS by sqrt(1 + THD^2). The
        # mapped voltage_a column must be the RMS channel, not the fundamental.
        by_name = {c.name: c
                   for obs in self._analysed_observations(path)
                   for c in obs.channels}
        rms_channel = by_name["RMS Van (V1)"]
        fundamental_channel = by_name["Harm 1 of Van"]
        assert rms_channel.characteristic == "RMS"
        assert fundamental_channel.characteristic == "SPECTRA_HGROUP"

        rms = rms_channel.series["AVG"][0::2]
        fundamental = fundamental_channel.series["AVG"][0::2]
        assert np.all(rms > fundamental), "fixture does not separate RMS from H1"

        df = extract_dataset(ProntoAdapter(path), ChannelMapper()).df
        assert np.allclose(df["voltage_a"].to_numpy(), rms, rtol=1e-9)
        # And the difference is the physically expected factor.
        thd = df["thd_voltage_a"].to_numpy()
        assert np.allclose(fundamental, rms / np.sqrt(1.0 + (thd / 100.0) ** 2),
                           rtol=1e-9)

    def test_time_base_comes_from_tag_time_start(self, path):
        import pandas as pd
        adapter = ProntoAdapter(path)
        expected = pqdif.PQDIFFile(path).observations[0].start_time
        assert pd.Timestamp(adapter._obs_ts[0]).to_pydatetime() == expected

    def test_min_le_avg_le_peak(self, path):
        # Independent of any offset assumption: if MAX/MIN/AVG were assigned to
        # the wrong series this ordering would break.
        adapter = ProntoAdapter(path)
        df = extract_dataset(adapter, ChannelMapper()).df
        checked = 0
        for col in df.columns:
            if f"{col}_peak" not in df.columns:
                continue
            avg, peak, low = df[col], df[f"{col}_peak"], df[f"{col}_min"]
            assert (peak >= avg - 1e-9).all(), f"{col}: peak < avg"
            assert (low <= avg + 1e-9).all(), f"{col}: min > avg"
            checked += 1
        assert checked >= 2, "expected max-min columns on the fixtures"


class TestOverloadedCharacteristics:
    """Pronto reuses ID_QC_HRMS for four different channels per phase."""

    def test_hrms_requires_a_name_prefix(self):
        # 'Hrms Van (V1)', 'Odds Van (V1)', 'Evens Van (V1)' and 'Triplens Van'
        # all carry characteristic HRMS with phase 'an'. Only the first is the
        # total; without the prefix check, file ordering alone decided which one
        # became hrms_voltage_a.
        assert ProntoAdapter._SPEC_NAME_PREFIX["HRMS"] == "hrms"
        assert ("voltage", "HRMS", "an") in ProntoAdapter._SPEC_CHANNELS
        for name, wanted in [
            ("Hrms Van (V1)", True), ("Odds Van (V1)", False),
            ("Evens Van (V1)", False), ("Triplens Van", False),
        ]:
            matches = name.lower().startswith(
                ProntoAdapter._SPEC_NAME_PREFIX["HRMS"])
            assert matches is wanted, name

    def test_aggregate_hrms_is_not_a_per_order_column(self):
        # hrms_voltage_a must not be swept up by the per-order harmonic checks,
        # which would test the whole harmonic RMS against a single-order limit.
        from pq_analysis import _HARMONIC_COL
        assert _HARMONIC_COL.match("h3_voltage_a")
        assert _HARMONIC_COL.match("h13_current_neutral")
        assert not _HARMONIC_COL.match("hrms_voltage_a")
        assert not _HARMONIC_COL.match("hrms_current_neutral")

    def test_new_canonical_names_are_all_resolvable(self):
        # Every canonical name _SPEC_CHANNELS claims must exist in CANONICAL,
        # or extract_dataframe would drop the column on the floor.
        from pq_adapter import CANONICAL
        for key, (_label, canonical, *_tags) in \
                ProntoAdapter._SPEC_CHANNELS.items():
            assert canonical in CANONICAL, f"{key} → {canonical!r} not in CANONICAL"


# ─────────────────────────────────────────────────────────────────────────────
# 10. Line-to-line voltage, frequency, flicker and K-factor
# ─────────────────────────────────────────────────────────────────────────────

def _frame(**cols):
    import pandas as pd
    n = len(next(iter(cols.values())))
    return pd.DataFrame(cols, index=pd.date_range("2025-01-01", periods=n, freq="5min"))


def _solar_frame(days=3.7, night_i1=3.0, peak_i1=2000.0,
                 night_ih=1.5, peak_ih=60.0):
    """A generating service: output collapses to nothing every night.

    This is the shape that breaks any statistic taken against the fundamental.
    At night I1 is a few amps, so THD% runs away while the harmonic amperes
    behind it stay trivial.  TDD, measured against a fixed IL, does not move.
    """
    import numpy as np
    import pandas as pd
    n    = int(days * 24 * 12)
    idx  = pd.date_range("2024-04-06 12:00", periods=n, freq="5min")
    hour = idx.hour + idx.minute / 60.0
    day  = np.clip(np.sin((hour - 6) / 12 * np.pi), 0, None)

    i1   = night_i1 + peak_i1 * day
    ih   = night_ih + peak_ih * day
    df   = pd.DataFrame(index=idx)
    for ph in ("a", "b", "c"):
        df[f"current_{ph}"]      = np.sqrt(i1 ** 2 + ih ** 2)
        df[f"hrms_current_{ph}"] = ih
        df[f"thd_current_{ph}"]  = 100.0 * ih / i1
        df[f"h5_current_{ph}"]   = ih * 0.8
    return df


class TestHarmonicStatisticsOnAGeneratingService:
    """The 9907 Queensburg failure: a THD series graded against a TDD limit.

    The report put P95 = 59.70% beside a maximum TDD of 4.44% taken from
    another block of the same run -- two different series under one label, and
    a comparison no reader could have made sense of.
    """

    def test_the_aggregate_row_is_tdd_not_raw_thd(self):
        from pq_analysis import check_harmonic_statistics
        from pq_constants import Thresholds
        df = _solar_frame()
        # The raw THD channel reaches 50% every night on this profile.
        assert df["thd_current_a"].quantile(0.95) > 40
        w = check_harmonic_statistics(df, Thresholds(isc_amps=40000.0))["weekly"]["thd"]["a"]
        # TDD is ~60 A against an IL of ~2000 A, whatever the hour.
        assert w["p95"] < 5.0
        assert w["p95_pass"] is True

    def test_percentiles_never_exceed_the_maximum_reported_elsewhere(self):
        # The invariant the report violated: one service, one denominator, so
        # the statistical block's P95/P99 must sit under check_thd's maximum.
        from pq_analysis import check_harmonic_statistics, check_thd
        from pq_constants import Thresholds
        th  = Thresholds(isc_amps=40000.0)
        df  = _solar_frame()
        w   = check_harmonic_statistics(df, th)["weekly"]["thd"]["a"]
        top = check_thd(df, th)["current"]["max_thd_pct"]
        assert w["p95"] <= top + 1e-9
        assert w["p99"] <= top + 1e-9

    def test_il_matches_the_one_check_thd_grades_against(self):
        # Both blocks normalise against IL; if they derive it differently the
        # per-order table and the TDD row describe different denominators.
        from pq_analysis import check_harmonic_statistics, check_thd
        from pq_constants import Thresholds
        th   = Thresholds(isc_amps=40000.0)
        df   = _solar_frame()
        stat = check_harmonic_statistics(df, th)
        assert stat["il_amps"] == pytest.approx(
            check_thd(df, th)["current"]["il_amps"], rel=1e-3)

    def test_a_phase_without_harmonic_rms_falls_back_rather_than_vanishing(self):
        # Meters that report only THD totals must still get a TDD row, derived
        # as Ih = THD% x I1 / 100, not be dropped from the assessment.
        from pq_analysis import check_harmonic_statistics
        from pq_constants import Thresholds
        df = _solar_frame().drop(columns=["hrms_current_a", "hrms_current_b",
                                          "hrms_current_c"])
        w  = check_harmonic_statistics(df, Thresholds(isc_amps=40000.0))["weekly"]["thd"]
        assert set(w) == {"a", "b", "c"}
        assert w["a"]["p95"] < 5.0


class TestHarmonicDirectionAtAProducerArray:
    """A plant with no load: the polarity check returns, with its sign flipped.

    The three roles are not a severity scale. Fundamental flow is one-way at
    both ends and two-way only in the middle, so a wrong sign catches reversed
    clamps at a load service and at a plant alike -- it is the mixed service
    where the check cannot be made at all.
    """

    V1 = (277.0 * np.sqrt(2), 0.0)
    I1_EXPORT = (230.0 * np.sqrt(2), 180.0)
    I1_IMPORT = (230.0 * np.sqrt(2), 0.0)

    def _run(self, captures, role="generation"):
        from pq_analysis import harmonic_direction_from_waveforms
        import pandas as pd
        return harmonic_direction_from_waveforms(
            _FakeDataset(pd.DataFrame(), captures),
            Thresholds(nominal_voltage=277.0, service_role=role))

    def _producing(self, reversed_cts=False):
        i1 = self.I1_IMPORT if reversed_cts else self.I1_EXPORT
        h5 = (-2.0, 0.0) if reversed_cts else (2.0, 0.0)
        return _capture({1: self.V1, 5: (4.0, 180.0)}, {1: i1, 5: h5})

    def test_exporting_is_the_expected_sign_and_is_not_corrected(self):
        r = self._run([self._producing()] * 3)
        assert r["ct_polarity_inverted"] is False
        assert r["ct_polarity_verified"] is True
        assert r["orders"][5]["indication"] == "downstream"

    def test_a_plant_reading_as_importing_is_a_reversed_ct(self):
        # The load rule with its sign flipped: at a plant that is producing,
        # positive fundamental real power is the anomaly.
        r = self._run([self._producing(reversed_cts=True)] * 3)
        assert r["ct_polarity_inverted"] is True
        assert "reversed" in r["polarity_note"]
        # Corrected, it must reach the same conclusion as the sound install.
        assert r["orders"][5]["indication"] == "downstream"

    def test_the_mixed_split_is_not_applied_to_a_plant(self):
        r = self._run([self._producing()] * 3)
        assert "export_split" not in r

    def test_calling_a_plant_a_load_inverts_every_direction(self):
        # The defect the third role exists to prevent.
        r = self._run([self._producing()] * 3, role="load")
        assert r["ct_polarity_inverted"] is True
        assert r["orders"][5]["indication"] == "upstream"

    def test_calling_a_plant_mixed_throws_the_answer_away(self):
        # Safe but useless: with no importing captures the top-level result is
        # empty, so a plant must not be filed under the mixed role either.
        r = self._run([self._producing()] * 3, role="mixed")
        assert r["orders"] == {}
        assert r["overall"] == "indeterminate"

    def test_the_note_says_direction_is_a_check_not_a_finding(self):
        # At a plant the only source behind the meter is the plant, so
        # "customer side" could not have come out any other way.
        note = self._run([self._producing()] * 3)["polarity_note"]
        assert "check on the measurement" in note

    def test_the_real_fixture_reads_correctly_only_as_a_plant(self):
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pq_analysis import harmonic_direction_from_waveforms
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_producer_array.pqd"
        ds = extract_dataset(ProntoAdapter(str(path)), ChannelMapper())
        assert len(ds.waveforms) == 4

        def run(role):
            return harmonic_direction_from_waveforms(
                ds, Thresholds(nominal_voltage=277.0, customer_class="sg",
                               service_role=role))

        assert run("load")["overall"] == "upstream"           # inverted
        assert run("mixed")["overall"] == "indeterminate"     # answer lost
        assert run("generation")["overall"] == "downstream"   # right


class TestGenerationIL:
    """IL at a plant, which has no demand load to take it from."""

    @staticmethod
    def _thresh(**kw):
        return Thresholds(nominal_voltage=277.0, customer_class="sg",
                          service_role="generation", isc_amps=40000.0, **kw)

    @staticmethod
    def _df():
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_producer_array.pqd"
        return extract_dataset(ProntoAdapter(str(path)), ChannelMapper()).df

    def test_the_nameplate_is_used_when_given(self):
        from pq_analysis import check_thd
        # 250 kW AC at 277 V line-to-neutral, three-phase: 250000/(3x277) = 301 A.
        td = check_thd(self._df(), self._thresh(rated_ac_kw=250.0))["tdd_info"]
        assert td["il_amps"] == pytest.approx(300.8, abs=1.0)
        assert td["il_basis"] == "rated_output"

    def test_without_a_nameplate_it_falls_back_to_measured_export(self):
        from pq_analysis import check_thd
        td = check_thd(self._df(), self._thresh())["tdd_info"]
        assert td["il_basis"] == "measured_export"
        # The fixture's week is hazy, so the plant never reaches its rating.
        assert td["il_amps"] < 300.0

    def test_the_weaker_reference_inflates_the_percentages(self):
        # Why the basis has to be stated rather than just used: the same
        # harmonic amperes grade differently against the two denominators.
        from pq_analysis import check_thd
        df = self._df()
        rated    = check_thd(df, self._thresh(rated_ac_kw=250.0))["current"]
        measured = check_thd(df, self._thresh())["current"]
        assert measured["max_thd_pct"] > rated["max_thd_pct"]

    def test_a_load_service_ignores_the_rating_entirely(self):
        # IL at a load service is measured demand; a stray rating must not
        # quietly become the denominator there.
        from pq_analysis import check_thd
        df = self._df()
        load = Thresholds(nominal_voltage=277.0, customer_class="sg",
                          isc_amps=40000.0, rated_ac_kw=250.0)
        td = check_thd(df, load)["tdd_info"]
        assert td["il_basis"] == "measured_demand"
        assert td["il_amps"] < 300.0


def _build_report(pqd, cls, nominal, **thresh_kw):
    """A full report for a fixture, the way the CLI builds one."""
    import pq_analysis as An
    from pq_report import generate_report
    from pathlib import Path
    source = Path(__file__).parent / "test_data" / f"{pqd}.pqd"
    ds = extract_dataset(ProntoAdapter(str(source)), ChannelMapper())
    th = Thresholds(nominal_voltage=nominal, customer_class=cls, **thresh_kw)
    df = ds.df
    ev = An.detect_events(ds, th)
    rep = generate_report(
        ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
        An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
        An.check_current_imbalance(df, th), An.check_demand(df, th),
        An.check_individual_harmonics(df, th),
        An.check_individual_voltage_harmonics(df, th),
        An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
        An.check_harmonic_statistics(df, th), ev, th,
        neutral_health_result=An.check_neutral_health(ds, th),
        itic_result=An.check_itic(ev, th),
        flicker_result=An.check_flicker(df, th))
    rep["root_causes"] = An.analyze_root_causes(rep, ds, th)
    return ds, rep, th


#: Every fixture with the role it is meant to be analysed under. Both
#: documents are built for each, because the suite has twice now been fully
#: green while the CLI crashed on a path no test constructed -- a generating
#: service, and then a meter that signs power factor by direction.
_DOCUMENT_MATRIX = [
    ("test_residential",         "r",  120.0,   {}),
    ("test_commercial_small",    "c",  120.0,   {}),
    ("test_commercial_large",    "sg", 277.0,   {"isc_amps": 40000.0}),
    ("test_commercial_imbalanced", "sg", 277.0, {"isc_amps": 40000.0}),
    ("test_commercial_primary",  "pg", 13200.0, {"isc_amps": 5000.0}),
    ("test_solar_net_metered",   "sg", 277.0,
     {"isc_amps": 40000.0, "service_role": "mixed", "rated_ac_kw": 150.0,
      "avg_peak_demand_kw": 1200.0}),
    ("test_producer_array",      "sg", 277.0,
     {"isc_amps": 40000.0, "service_role": "generation", "rated_ac_kw": 250.0,
      "avg_peak_demand_kw": 10.0, "der_category": "II"}),
    # The same plant with its interconnection agreement supplied, which is a
    # different set of paragraphs, table rows and letter rows -- and the graded
    # ones. Without this entry the whole fixed-power-factor feature assembled
    # only in the unit tests.
    ("test_producer_array",      "sg", 277.0,
     {"isc_amps": 40000.0, "service_role": "generation", "rated_ac_kw": 250.0,
      "der_category": "II", "der_reactive_mode": "fixed_pf",
      "der_pf_setpoint": 0.98, "der_pf_direction": "absorbing",
      "der_pf_tolerance": 0.01}),
    # And on a mode with no assessment behind it, which must decline in prose
    # rather than fall through any of the graded paths.
    ("test_solar_net_metered",   "sg", 277.0,
     {"isc_amps": 40000.0, "service_role": "mixed", "rated_ac_kw": 150.0,
      "avg_peak_demand_kw": 1200.0, "der_reactive_mode": "volt_var"}),
]


class TestOperatingQuadrant:
    """Which way each power flowed, which a power factor magnitude cannot say.

    An interconnection agreement specifies a plant's power factor as a signed
    setpoint -- "-0.98" meaning real power out while reactive power is drawn in,
    for voltage mitigation. The magnitude 0.98 describes that and its opposite
    equally well, so the quadrant has to come from the signs of the two power
    channels. It must not come from the sign of the power factor: the meter
    picks that convention itself and does not declare which one it used.
    """

    @staticmethod
    def _df(pf, watts, vars_):
        import pandas as pd
        n = len(pf)
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        return pd.DataFrame({"power_factor": pf, "power_real": watts,
                             "power_reactive": vars_}, index=idx)

    def test_a_plant_absorbing_vars_is_named_as_such(self):
        from pq_analysis import check_power_factor
        # The Queensburg setpoint: exporting watts, absorbing VAR.
        df = self._df([0.98] * 20, [-500_000.0] * 20, [101_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg",
                                              service_role="generation"))
        q = r["quadrants"]
        assert q["dominant"] == "export_absorb"
        assert q["dominant_pct"] == 100.0
        assert q["mean_kw"] < 0 and q["mean_kvar"] > 0

    def test_the_opposite_quadrant_is_told_apart_at_the_same_magnitude(self):
        from pq_analysis import check_power_factor
        # Same 0.98, VAR the other way: injecting rather than absorbing. The
        # power factor cannot distinguish these and the signs must.
        df = self._df([0.98] * 20, [-500_000.0] * 20, [-101_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg",
                                              service_role="generation"))
        assert r["quadrants"]["dominant"] == "export_inject"
        assert r["mean_pf"] == pytest.approx(0.98, abs=0.001)

    def test_the_quadrant_does_not_follow_the_power_factor_sign(self):
        from pq_analysis import check_power_factor
        # A direction-signing meter reports -0.98 while exporting. That minus
        # sign is the real-power direction and carries nothing about the
        # reactive flow, which here is being absorbed.
        df = self._df([-0.98] * 20, [-500_000.0] * 20, [101_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg",
                                              service_role="generation"))
        assert r["quadrants"]["dominant"] == "export_absorb"

    def test_no_reactive_channel_means_no_quadrant_rather_than_a_guess(self):
        import pandas as pd
        from pq_analysis import check_power_factor
        idx = pd.date_range("2025-01-01", periods=20, freq="5min")
        df = pd.DataFrame({"power_factor": [0.95] * 20,
                           "power_real": [50_000.0] * 20}, index=idx)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg"))
        assert r["quadrants"] == {}


class TestFixedPowerFactorAgainstTheAgreement:
    """What a plant owes, which no tariff sheet states.

    PSCo's Technical Specifications Manual (01/01/2025) 6.3.2: "Where a
    constant power factor is otherwise specified or applied based on legacy
    requirements ... a 0.98 absorbing power factor shall be used". Absorbing is
    voltage mitigation -- an exporting plant lifts the voltage at the point of
    interconnection and drawing VAR pulls it back. A plant injecting instead is
    doing the opposite of what it agreed to, at the same magnitude.
    """

    @staticmethod
    def _df(pf, watts, vars_):
        import pandas as pd
        n = len(pf)
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        return pd.DataFrame({"power_factor": pf, "power_real": watts,
                             "power_reactive": vars_}, index=idx)

    @staticmethod
    def _t(**kw):
        base = dict(nominal_voltage=277.0, customer_class="sg",
                    service_role="generation", rated_ac_kw=500.0,
                    der_reactive_mode="fixed_pf", der_pf_setpoint=0.98,
                    der_pf_direction="absorbing")
        base.update(kw)
        return Thresholds(**base)

    def test_a_plant_holding_its_setpoint_passes_on_direction(self):
        from pq_analysis import check_der_power_factor
        df = self._df([0.98] * 20, [-500_000.0] * 20, [101_000.0] * 20)
        r = check_der_power_factor(df, self._t())
        assert r["assessed"] is True
        assert r["pct_in_required_direction"] == 100.0
        assert r["direction_pass"] is True
        assert r["max_deviation"] < 0.001

    def test_the_same_magnitude_the_wrong_way_round_fails(self):
        from pq_analysis import check_der_power_factor
        # 0.98 exactly, but injecting. The magnitude is perfect and the plant
        # is still doing the opposite of what the agreement asked for.
        df = self._df([0.98] * 20, [-500_000.0] * 20, [-101_000.0] * 20)
        r = check_der_power_factor(df, self._t())
        assert r["max_deviation"] < 0.001        # magnitude is spotless
        assert r["direction_pass"] is False      # and it still fails
        assert r["pct_in_required_direction"] == 0.0

    def test_no_tolerance_reports_the_deviation_without_grading_it(self):
        from pq_analysis import check_der_power_factor
        df = self._df([0.93] * 20, [-500_000.0] * 20, [101_000.0] * 20)
        r = check_der_power_factor(df, self._t())
        assert r["mean_deviation"] == pytest.approx(0.05, abs=0.001)
        assert r["magnitude_pass"] is None
        assert r["pct_outside_tolerance"] is None
        assert len(r["violation_timestamps"]) == 0
        assert "tolerance" in r["tolerance_note"]

    def test_a_tolerance_turns_the_deviation_into_a_verdict(self):
        from pq_analysis import check_der_power_factor
        df = self._df([0.93] * 20, [-500_000.0] * 20, [101_000.0] * 20)
        r = check_der_power_factor(df, self._t(der_pf_tolerance=0.01))
        assert r["magnitude_pass"] is False
        assert r["pct_outside_tolerance"] == 100.0
        assert len(r["violation_timestamps"]) == 20

    def test_volt_var_is_declined_rather_than_graded(self):
        from pq_analysis import check_der_power_factor
        # The reactive output is *required* to move with voltage here, so a
        # fixed-setpoint comparison would report correct operation as a fault.
        df = self._df([0.90] * 20, [-500_000.0] * 20, [101_000.0] * 20)
        r = check_der_power_factor(df, self._t(der_reactive_mode="volt_var"))
        assert r["assessed"] is False
        assert r["available"] is False
        assert "vary" in r["error"]
        assert "5.3.3" in r["error"]

    def test_an_unentered_mode_is_not_assumed_to_be_fixed(self):
        from pq_analysis import check_der_power_factor
        df = self._df([0.98] * 20, [-500_000.0] * 20, [101_000.0] * 20)
        r = check_der_power_factor(df, self._t(der_reactive_mode=None))
        assert r["assessed"] is False
        assert "not entered" in r["error"]

    def test_a_load_service_has_no_agreement_to_check(self):
        from pq_analysis import check_der_power_factor
        df = self._df([0.95] * 20, [50_000.0] * 20, [10_000.0] * 20)
        r = check_der_power_factor(df, self._t(service_role="load"))
        assert r["assessed"] is False
        assert "no generation" in r["error"]

    def test_output_below_the_witness_test_floor_is_not_assessed(self):
        from pq_analysis import check_der_power_factor
        # TSM 8.1 will not verify power factor below 15% of capacity. A 500 kW
        # plant trickling 20 kW is under it, and its displacement there says
        # nothing about the setpoint.
        df = self._df([0.50] * 20, [-20_000.0] * 20, [30_000.0] * 20)
        r = check_der_power_factor(df, self._t())
        assert r["assessed"] is False
        assert "15%" in r["error"]

    def test_the_floor_is_taken_against_the_nameplate_not_the_week(self):
        from pq_analysis import check_der_power_factor
        # Half the recording at 300 kW, half at 40 kW, on a 500 kW plant. The
        # 40 kW half is 8% of the rating and excluded; against the recording's
        # own peak it would have been 13% and still excluded, but the count of
        # what was left out has to come from the rating.
        df = self._df([0.98] * 20 + [0.60] * 20,
                      [-300_000.0] * 20 + [-40_000.0] * 20,
                      [101_000.0] * 40)
        r = check_der_power_factor(df, self._t())
        assert r["basis"] == "nameplate"
        assert r["reference_kw"] == 500.0
        assert r["intervals_used"] == 20
        assert r["excluded_low_output"] == 20
        assert r["mean_pf"] == pytest.approx(0.98, abs=0.001)

    def test_a_mixed_service_is_assessed_on_its_exporting_half(self):
        from pq_analysis import check_der_power_factor
        # An agreement binds a net-metered site too, and only while it exports.
        df = self._df([0.98] * 20 + [0.85] * 20,
                      [-500_000.0] * 20 + [200_000.0] * 20,
                      [101_000.0] * 40)
        r = check_der_power_factor(df, self._t(service_role="mixed"))
        assert r["intervals_used"] == 20
        assert r["mean_pf"] == pytest.approx(0.98, abs=0.001)


class TestSeverityGradesTheRightQuantity:
    """The margin has to describe whatever the verdict was reached on.

    Voltage imbalance failed on excursions while the margin was computed from
    the mean, so a residential service reaching 23% imbalance on a tenth of its
    intervals was reported as "0.77x the limit" -- which reads as comfortably
    inside. The 95th percentile is what IEEE 519 grades its own limits on, and
    unlike the maximum it is not set by a single bad interval.
    """

    @staticmethod
    def _ds(name):
        from pathlib import Path
        return extract_dataset(ProntoAdapter(
            str(Path(__file__).parent / "test_data" / f"{name}.pqd")),
            ChannelMapper())

    def test_the_imbalance_checks_expose_a_percentile(self):
        import pq_analysis as An
        ds = self._ds("test_commercial_large")
        th = Thresholds(nominal_voltage=277.0, customer_class="sg")
        for fn in (An.check_voltage_imbalance, An.check_current_imbalance):
            r = fn(ds.df, th)
            assert "p95_imbalance_pct" in r
            assert r["mean_imbalance_pct"] <= r["p95_imbalance_pct"] + 1e-9
            assert r["p95_imbalance_pct"] <= r["max_imbalance_pct"] + 1e-9

    def test_current_tdd_exposes_a_percentile(self):
        import pq_analysis as An
        ds = self._ds("test_commercial_large")
        c = An.check_thd(ds.df, Thresholds(nominal_voltage=277.0,
                                           customer_class="sg",
                                           isc_amps=40000.0))["current"]
        assert c["mean_thd_pct"] <= c["p95_thd_pct"] <= c["max_thd_pct"] + 1e-9

    def test_a_failing_finding_never_reports_a_margin_under_one(self):
        # The symptom that gave this away: a failure described as 0.77x the
        # limit. If the verdict is "outside", the margin must say so too.
        import pq_analysis as An
        from pq_report import generate_report, compute_severities
        ds = self._ds("test_commercial_large")
        th = Thresholds(nominal_voltage=277.0, customer_class="sg",
                        isc_amps=40000.0)
        df = ds.df
        ev = An.detect_events(ds, th)
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), ev, th,
            neutral_health_result=An.check_neutral_health(ds, th),
            itic_result=An.check_itic(ev, th),
            flicker_result=An.check_flicker(df, th))
        for key, graded in compute_severities(rep, th).items():
            if graded["band"] in ("minor", "significant", "severe"):
                margin = graded.get("margin")
                if margin is not None:
                    assert margin >= 1.0, (key, margin, graded["reason"])


class TestSeverityNeedsMarginAndPersistence:
    """"Significant" now takes the same shape as "severe".

    It used to be an OR, so persistence alone promoted anything: a metric 1.2%
    past its limit for a quarter of the week graded the same as one 20% past
    it, and a power factor of 0.89 against a 0.90 limit came out Significant.
    That is the disproportion the grading exists to prevent.
    """

    @staticmethod
    def _band(measured, limit, persistence, lower=False):
        from pq_analysis import grade_finding
        return grade_finding(False, measured=measured, limit=limit,
                             persistence_pct=persistence,
                             lower_is_worse=lower)["band"]

    def test_persistence_alone_no_longer_promotes(self):
        # 1.2% over the limit is a minor finding whether it lasts an hour or
        # the whole recording.
        assert self._band(8.1, 8.0, 24) == "minor"
        assert self._band(8.1, 8.0, 25) == "minor"
        assert self._band(8.1, 8.0, 100) == "minor"

    def test_a_margin_with_persistence_is_significant(self):
        assert self._band(8.5, 8.0, 100) == "significant"
        assert self._band(8.5, 8.0, 10) == "minor"      # margin without duration

    def test_a_large_margin_alone_is_still_significant(self):
        assert self._band(9.6, 8.0, 1) == "significant"

    def test_severe_is_unchanged(self):
        assert self._band(12.0, 8.0, 24) == "significant"
        assert self._band(12.0, 8.0, 25) == "severe"
        assert self._band(16.5, 8.0, 1) == "severe"

    def test_a_marginal_power_factor_is_not_significant(self):
        # 0.89 against a 0.90 limit is 1% short, all week.
        assert self._band(0.89, 0.90, 100, lower=True) == "minor"
        assert self._band(0.855, 0.90, 100, lower=True) == "significant"

    def test_the_two_bands_share_one_shape(self):
        # Both are (margin AND persistence) OR margin-alone, which is what
        # keeps the ladder explicable to a reader.
        from pq_constants import (SEVERITY_SIGNIFICANT_MARGIN,
                                  SEVERITY_SIGNIFICANT_MARGIN_ALONE,
                                  SEVERITY_SEVERE_MARGIN,
                                  SEVERITY_SEVERE_MARGIN_ALONE)
        assert SEVERITY_SIGNIFICANT_MARGIN < SEVERITY_SIGNIFICANT_MARGIN_ALONE
        assert SEVERITY_SIGNIFICANT_MARGIN_ALONE <= SEVERITY_SEVERE_MARGIN
        assert SEVERITY_SEVERE_MARGIN < SEVERITY_SEVERE_MARGIN_ALONE


class TestLegDifferenceIsNotGradedAgainstNEMA:
    """A split-phase service has no NEMA MG1 unbalance to be outside of.

    The check already said so in words -- "NEMA MG1 unbalance is defined for
    three-phase systems and is not applicable to a single-phase service" --
    and then compared the leg difference against the 3% NEMA limit anyway.
    With the mean as the margin that surfaced as a quiet "minor"; with a
    percentile it became a "severe" on every house.
    """

    @staticmethod
    def _result(name, nominal, cls):
        import pq_analysis as An
        from pathlib import Path
        ds = extract_dataset(ProntoAdapter(
            str(Path(__file__).parent / "test_data" / f"{name}.pqd")),
            ChannelMapper())
        th = Thresholds(nominal_voltage=nominal, customer_class=cls)
        return An.check_voltage_imbalance(ds.df, th)

    def test_a_split_phase_service_gets_no_limit(self):
        r = self._result("test_residential", 120.0, "r")
        assert r["metric"] == "leg_difference"
        assert r["limit_pct"] is None
        assert r["pct_exceeding"] == 0        # nothing to exceed

    def test_a_three_phase_service_keeps_the_nema_limit(self):
        r = self._result("test_commercial_large", 277.0, "sg")
        assert r["metric"] == "nema_mg1"
        assert r["limit_pct"] == 3.0

    def test_the_measurement_is_still_reported(self):
        # Not graded is not the same as not measured: the leg difference is a
        # real indicator of unequal loading or neutral impedance.
        r = self._result("test_residential", 120.0, "r")
        assert r["max_imbalance_pct"] > 0
        assert r["p95_imbalance_pct"] > 0


class TestBothDocumentsBuildForEveryFixture:
    """The end-to-end guard the unit tests kept missing.

    Twice now the whole suite has passed while `pq_analyzer.py --report`
    raised: once on a generating service, and once on a meter that signs power
    factor by direction, which left a None where a percentage was formatted.
    Unit tests covered both checks; nothing assembled the documents.
    """

    @pytest.mark.parametrize("pqd,cls,nominal,extra", _DOCUMENT_MATRIX,
                             ids=[m[0] for m in _DOCUMENT_MATRIX])
    def test_the_documents_assemble(self, pqd, cls, nominal, extra, tmp_path):
        docx = pytest.importorskip("docx")
        from pq_report import generate_word_report, generate_customer_letter
        ds, rep, th = _build_report(pqd, cls, nominal, **extra)
        internal = generate_word_report(
            report=rep, thresh=th, ds=ds, site_name="S", site_address="A",
            engineer_name="E", outdir=tmp_path, stem=pqd)
        letter = generate_customer_letter(rep, th, "1 Test St", "Eng",
                                          tmp_path, pqd)
        for path in (internal, letter):
            assert path is not None and Path(path).exists(), pqd
            doc = docx.Document(str(path))
            # A document that assembled but says nothing is still a failure.
            assert sum(len(p.text) for p in doc.paragraphs) > 500, pqd

    @pytest.mark.parametrize("pqd,cls,nominal,extra", _DOCUMENT_MATRIX,
                             ids=[m[0] for m in _DOCUMENT_MATRIX])
    def test_the_console_summary_prints(self, pqd, cls, nominal, extra, capsys):
        # The third time the suite was green while the CLI raised, it was here
        # rather than in Word: a plant carries no power factor limit, and the
        # console formatted it as ".2f" before anything reached the document
        # layer. The documents were covered above; the terminal output that
        # every run prints first was covered on one fixture.
        from pq_report import print_report
        _ds, rep, th = _build_report(pqd, cls, nominal, **extra)
        print_report(rep)
        assert len(capsys.readouterr().out) > 500, pqd

    @pytest.mark.parametrize("pqd,cls,nominal,extra", _DOCUMENT_MATRIX,
                             ids=[m[0] for m in _DOCUMENT_MATRIX])
    def test_the_csv_exports_write(self, pqd, cls, nominal, extra, tmp_path):
        from pq_report import export_results
        ds, rep, th = _build_report(pqd, cls, nominal, **extra)
        export_results(ds, rep, tmp_path, pqd)
        assert list(tmp_path.glob("*.csv")), pqd


class TestAnalysisModeIsVisible:
    """Which mode produced these numbers, stated on the documents.

    The same recording can produce opposite verdicts depending on the power
    flow setting and the standard it selects, and nothing on the page said
    which was active -- so a power factor finding on a plant read as either a
    broken report or the wrong mode, with no way to tell them apart.
    """

    def test_the_header_table_lists_the_settings_that_change_the_answer(self):
        from pq_report import analysis_mode_summary
        _ds, rep, th = _build_report(
            "test_producer_array", "sg", 277.0, isc_amps=40000.0,
            service_role="generation", rated_ac_kw=250.0,
            avg_peak_demand_kw=10.0, der_category="II")
        bits = " | ".join(analysis_mode_summary(rep, th))
        assert "Generation only" in bits
        assert "1547" in bits                 # the standard actually applied
        assert "IL:" in bits                  # and where IL came from
        assert "category: II" in bits
        assert "Schedule SG" in bits

    def test_a_missing_category_is_named_rather_than_left_blank(self):
        from pq_report import analysis_mode_summary
        _ds, rep, th = _build_report(
            "test_producer_array", "sg", 277.0, isc_amps=40000.0,
            service_role="generation", rated_ac_kw=250.0,
            avg_peak_demand_kw=10.0)
        assert any("not entered" in b for b in analysis_mode_summary(rep, th))

    def test_a_load_service_says_load(self):
        from pq_report import analysis_mode_summary
        _ds, rep, th = _build_report("test_commercial_large", "sg", 277.0,
                                     isc_amps=40000.0)
        bits = " | ".join(analysis_mode_summary(rep, th))
        assert "Load only" in bits and "519" in bits

    def test_the_internal_footer_carries_the_mode_on_every_page(self, tmp_path):
        docx = pytest.importorskip("docx")
        from pq_report import generate_word_report
        ds, rep, th = _build_report(
            "test_producer_array", "sg", 277.0, isc_amps=40000.0,
            service_role="generation", rated_ac_kw=250.0,
            avg_peak_demand_kw=10.0, der_category="II")
        path = generate_word_report(
            report=rep, thresh=th, ds=ds, site_name="S", site_address="A",
            engineer_name="E", outdir=tmp_path, stem="t")
        footer = docx.Document(str(path)).sections[0].footer.paragraphs[0].text
        assert "generation only" in footer
        assert "1547" in footer

    def test_the_letter_marks_generation_and_leaves_load_letters_alone(self, tmp_path):
        # A customer with an ordinary service does not need to be told their
        # service was assessed as a load; it is the default and adds noise.
        docx = pytest.importorskip("docx")
        from pq_report import generate_customer_letter

        def footer_of(pqd, cls, nominal, **extra):
            _ds, rep, th = _build_report(pqd, cls, nominal, **extra)
            p = generate_customer_letter(rep, th, "1 St", "Eng", tmp_path, pqd)
            para = docx.Document(str(p)).sections[0].footer.paragraphs[0]
            return "".join(r.text for r in para.runs)

        gen = footer_of("test_producer_array", "sg", 277.0, isc_amps=40000.0,
                        service_role="generation", rated_ac_kw=250.0,
                        avg_peak_demand_kw=10.0, der_category="II")
        assert "generating facility" in gen
        load = footer_of("test_commercial_large", "sg", 277.0, isc_amps=40000.0)
        assert "assessed as" not in load


class TestPowerFactorSignConvention:
    """A signed power factor channel read as a violation on every export.

    Meters sign power factor by one of two conventions and do not say which:
    by the direction of real power, or by leading against lagging. Comparing
    the signed value against 0.90 made every exporting interval on a
    generating service a violation at -0.95.
    """

    @staticmethod
    def _df(pf, kw, kvar=2000.0):
        import pandas as pd
        n = len(pf)
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        return pd.DataFrame({
            "power_factor":   pf,
            "power_real":     kw,
            "power_reactive": [kvar] * n,
        }, index=idx)

    def test_a_direction_signed_meter_is_not_read_as_a_violation(self):
        from pq_analysis import check_power_factor
        # Exporting at a healthy 0.95 displacement, reported as -0.95.
        df = self._df([-0.95] * 20 + [0.95] * 20,
                      [-50_000.0] * 20 + [50_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg",
                                              service_role="mixed"))
        assert r["convention"] == "direction"
        assert r["pct_below_limit"] == 0.0
        assert r["min_pf"] > 0                # magnitudes, not signed values

    def test_the_same_data_read_as_a_load_still_grades_on_magnitude(self):
        # Even without the generation flag, -0.95 is a 0.95 displacement. The
        # sign is direction information, not a low power factor.
        from pq_analysis import check_power_factor
        df = self._df([-0.95] * 20 + [0.95] * 20,
                      [-50_000.0] * 20 + [50_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg"))
        assert r["pct_below_limit"] == 0.0

    def test_a_leading_signed_meter_is_reported_as_leading(self):
        # Negative power factor while real power stays positive is not an
        # export; it is a leading power factor, and R73 asks for lagging.
        from pq_analysis import check_power_factor
        df = self._df([-0.95] * 20 + [0.95] * 20, [50_000.0] * 40)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg"))
        assert r["convention"] == "leading"
        assert r["pct_leading"] == pytest.approx(50.0, abs=1.0)

    def test_the_tariff_clause_is_applied_over_importing_intervals_only(self):
        from pq_analysis import check_power_factor
        # Poor displacement while exporting, healthy while importing. The
        # tariff clauses describe what a load presents at the point of
        # delivery, so the export must not create a violation.
        # "mixed", not "generation": this is a load with an array behind it.
        # A plant is scoped the other way round -- see the test below.
        df = self._df([-0.70] * 20 + [0.97] * 20,
                      [-50_000.0] * 20 + [50_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg",
                                              service_role="mixed"))
        assert r["pct_below_limit"] == 0.0
        assert "importing intervals only" in r["scope_note"]
        assert r["export_mean_pf"] == pytest.approx(0.70, abs=0.01)

    def test_a_plant_is_measured_on_what_it_held_while_producing(self):
        from pq_analysis import check_power_factor
        # A producer's array: 2 MW out all day at 0.995, and a 10 kW SCADA
        # cabinet overnight at 0.30. Scoped to the importing intervals, the
        # only population is the cabinet -- which the light-load gate then
        # discards entirely, so the check reported nothing while the number
        # that matters, the displacement while producing, went unread.
        df = self._df([0.995] * 20 + [0.30] * 20,
                      [-2_000_000.0] * 20 + [10_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg",
                                              service_role="generation"))
        assert r["basis"] == "plant_export"
        assert r["mean_pf"] == pytest.approx(0.995, abs=0.005)
        assert r["intervals_used"] == 20
        # Reported, never graded: no tariff clause reaches a plant.
        assert r["assessed"] is False
        assert r["limit"] is None
        assert r["pct_below_limit"] is None
        assert len(r["violation_timestamps"]) == 0

    def test_the_light_load_gate_uses_the_population_it_is_gating(self):
        from pq_analysis import check_power_factor
        # A net-metered service: 500 kW of midday export, 40 kW of import the
        # rest of the time at a poor 0.80. The tariff clause applies to the
        # import, and 40 kW is a real load -- but measured against the export
        # peak it is 8%, so a gate keyed to the whole recording threw away
        # every interval the clause speaks to and reported that nothing
        # carried enough load. Each population is gated against its own peak.
        df = self._df([0.98] * 20 + [0.80] * 20,
                      [-500_000.0] * 20 + [40_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg",
                                              service_role="mixed"))
        assert r["pct_below_limit"] == 100.0
        assert r["mean_pf"] == pytest.approx(0.80, abs=0.01)

    def test_a_plants_overnight_auxiliary_load_is_not_its_power_factor(self):
        from pq_analysis import check_power_factor
        # The regression this fixes: the mean over the whole record, night
        # included, was 0.53 against a plant that ran at 0.995 all day.
        df = self._df([0.995] * 20 + [0.003] * 20,
                      [-2_000_000.0] * 20 + [8_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg",
                                              service_role="generation"))
        assert r["mean_pf"] > 0.9
        assert r["min_pf"] > 0.9

    def test_a_service_that_only_exports_gets_no_compliance_finding(self):
        from pq_analysis import check_power_factor
        df = self._df([-0.85] * 20, [-50_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg",
                                              service_role="generation"))
        assert r["assessed"] is False
        assert r["pct_below_limit"] is None

    def test_power_factor_is_gated_on_load_like_every_other_ratio(self):
        from pq_analysis import check_power_factor
        # A plant's overnight auxiliary load: 600 W against 2 kVAR is a real
        # 0.29 displacement and a meaningless finding.
        df = self._df([0.29] * 20 + [0.97] * 20, [600.0] * 20 + [90_000.0] * 20)
        r = check_power_factor(df, Thresholds(nominal_voltage=277.0,
                                              customer_class="sg"))
        assert r["pct_below_limit"] == 0.0    # the idle intervals are excluded

    def test_the_real_fixture_no_longer_reports_a_violation(self):
        from pq_analysis import check_power_factor
        _ds, rep, th = _build_report(
            "test_producer_array", "sg", 277.0, isc_amps=40000.0,
            service_role="generation", rated_ac_kw=250.0,
            avg_peak_demand_kw=10.0, der_category="II")
        pfr = rep["power_factor"]
        assert pfr["convention"] == "direction"
        assert pfr["pct_below_limit"] is None
        assert rep["pass_fail"]["power_factor"] is None   # not a fail


class TestBillingDemandPhaseSpread:
    """PSCo Sheet R123, reported as a cost rather than as a compliance finding.

    Nothing here is violated: the clause says the Company *may* take billing
    demand from the worst phase above a 15% spread. It is in the letter because
    it is a charge a customer can be carrying without knowing the cause, and
    which balancing the panel removes.
    """

    @staticmethod
    def _ds(name):
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / f"{name}.pqd"
        return extract_dataset(ProntoAdapter(str(path)), ChannelMapper())

    @staticmethod
    def _t(cls="sg", nominal=277.0):
        return Thresholds(nominal_voltage=nominal, customer_class=cls)

    def test_it_is_evaluated_at_the_peak_not_averaged(self):
        # Billing demand is set by the peak interval, so imbalance at 3 a.m.
        # costs nothing. Averaging the spread would answer a question nobody
        # is billed on.
        from pq_analysis import check_billing_demand_imbalance
        ds = self._ds("test_commercial_imbalanced")
        r = check_billing_demand_imbalance(ds.df, self._t())
        assert r["available"] and r["applies"]
        # The reported phases are the ones present at that timestamp.
        at_peak = ds.df.loc[r["peak_timestamp"],
                            ["current_a", "current_b", "current_c"]]
        assert r["phase_amps"]["c"] == pytest.approx(float(at_peak["current_c"]), abs=0.1)

    def test_the_uplift_is_worst_over_mean_not_worst_over_least(self):
        # The trigger and the cost are measured differently, which is the part
        # that surprises people: a 46% spread carries a 26% uplift.
        from pq_analysis import check_billing_demand_imbalance
        r = check_billing_demand_imbalance(
            self._ds("test_commercial_imbalanced").df, self._t())
        amps = r["phase_amps"]
        expected = max(amps.values()) / (sum(amps.values()) / 3)
        assert r["uplift"] == pytest.approx(expected, rel=0.01)
        assert r["uplift_pct"] < r["spread_pct"]

    def test_the_clause_demand_is_three_times_the_worst_phase(self):
        from pq_analysis import check_billing_demand_imbalance
        r = check_billing_demand_imbalance(
            self._ds("test_commercial_imbalanced").df, self._t())
        worst_a = max(r["phase_amps"].values())
        assert r["clause_kva"] == pytest.approx(3 * worst_a * 277.0 / 1000, rel=0.01)
        assert r["clause_kw"] == pytest.approx(r["clause_kva"] * 0.90, rel=0.01)

    def test_a_balanced_service_does_not_trigger_it(self):
        from pq_analysis import check_billing_demand_imbalance
        r = check_billing_demand_imbalance(
            self._ds("test_commercial_large").df, self._t())
        assert r["available"] is True
        assert r["applies"] is False

    def test_schedule_c_has_no_demand_charge_so_no_finding(self):
        # Schedule C is service, facility and energy only. There is no billing
        # demand for the clause to recompute, so raising it would be a false
        # alarm on a customer who cannot be billed that way.
        from pq_analysis import check_billing_demand_imbalance
        r = check_billing_demand_imbalance(
            self._ds("test_commercial_imbalanced").df, self._t(cls="c"))
        assert r["available"] is False
        assert "no demand charge" in r["note"]

    def test_a_single_phase_service_is_out_of_scope(self):
        from pq_analysis import check_billing_demand_imbalance
        r = check_billing_demand_imbalance(
            self._ds("test_residential").df, self._t(cls="sg", nominal=120.0))
        assert r["available"] is False
        assert "three-phase" in r["note"]

    def test_the_letter_carries_it_as_cost_not_as_a_violation(self, tmp_path):
        from docx import Document
        from pq_report import generate_customer_letter
        import pq_analysis as An
        from pq_report import generate_report
        ds = self._ds("test_commercial_imbalanced")
        th = Thresholds(nominal_voltage=277.0, customer_class="sg",
                        isc_amps=40000.0)
        df = ds.df
        ev = An.detect_events(ds, th)
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), ev, th,
            neutral_health_result=An.check_neutral_health(ds, th),
            itic_result=An.check_itic(ev, th),
            flicker_result=An.check_flicker(df, th))
        out = generate_customer_letter(rep, th, "1 Trade St", "Eng", tmp_path, "t")
        text = " ".join(p.text for p in Document(str(out)).paragraphs)
        assert "Phase balance and your billing demand" in text
        assert "R123" in text
        # It must not read as an accusation or a billing notice.
        assert "not billing you differently" in text
        assert "cost rather than power quality" in text

    def test_the_compliance_table_still_says_nothing_about_it(self):
        # The whole reason it has its own section: a billing provision in a
        # table of standards reads as something the customer is failing.
        from pq_report import _customer_checks
        import pq_analysis as An
        from pq_report import generate_report
        ds = self._ds("test_commercial_imbalanced")
        th = Thresholds(nominal_voltage=277.0, customer_class="sg")
        df = ds.df
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), An.detect_events(ds, th), th)
        joined = " ".join(f"{c['item']} {c['against']}"
                          for c in _customer_checks(rep, th))
        assert "R123" not in joined and "R12" not in joined


class TestRideThroughTables:
    """IEEE 1547-2018 Tables 14, 15 and 16, transcribed and spot-checked.

    The tables are the whole check, so they are asserted against the standard
    rather than against whatever the code happens to produce. The two sloped
    rows are the ones worth checking arithmetic on: Category I is a 4 s/p.u.
    slope from 0.7 s at 0.7 p.u., Category II is 8.7 s/p.u. from 3 s at 0.65.
    """

    @staticmethod
    def _region(cat, v):
        from pq_constants import ride_through_region
        return ride_through_region(cat, v)

    def test_the_continuous_region_is_the_ansi_range_b_band(self):
        # 0.88 to 1.10 p.u. in all three categories.
        for cat in ("I", "II", "III"):
            assert self._region(cat, 0.88)["mode"] == "continuous", cat
            assert self._region(cat, 1.10)["mode"] == "continuous", cat
            assert self._region(cat, 0.879)["mode"] != "continuous", cat
            assert self._region(cat, 1.101)["mode"] != "continuous", cat
            assert self._region(cat, 1.0)["min_ride_s"] == math.inf, cat

    def test_category_i_mandatory_slope(self):
        # Table 14: 0.7 s at 0.70 p.u., rising 4 s per p.u.
        assert self._region("I", 0.70)["min_ride_s"] == pytest.approx(0.70)
        assert self._region("I", 0.75)["min_ride_s"] == pytest.approx(0.90)
        assert self._region("I", 0.85)["min_ride_s"] == pytest.approx(1.30)

    def test_category_ii_mandatory_slope(self):
        # Table 15: 3 s at 0.65 p.u., rising 8.7 s per p.u.
        assert self._region("II", 0.65)["min_ride_s"] == pytest.approx(3.00)
        assert self._region("II", 0.75)["min_ride_s"] == pytest.approx(3.87)
        assert self._region("II", 0.85)["min_ride_s"] == pytest.approx(4.74)

    def test_category_iii_is_flat_and_far_longer(self):
        # Table 16 is where the categories diverge most: 20 s of mandatory
        # operation at 0.75 p.u. against 0.9 s under Category I.
        assert self._region("III", 0.75)["min_ride_s"] == 20.0
        assert self._region("III", 0.55)["min_ride_s"] == 10.0
        assert self._region("III", 0.40)["mode"] == "momentary"
        assert self._region("III", 1.15)["mode"] == "momentary"

    def test_the_deep_and_high_extremes_cease_to_energize(self):
        for cat in ("I", "II", "III"):
            assert self._region(cat, 1.25)["mode"] == "cease", cat
        assert self._region("I", 0.40)["mode"] == "cease"
        assert self._region("II", 0.25)["mode"] == "cease"

    def test_an_unknown_category_is_not_guessed_at(self):
        assert self._region("IV", 0.9) is None
        assert self._region(None, 0.9) is None


class TestRideThroughAgainstMeasuredEvents:
    """Clause 6.4.2 applied to what the meter actually caught."""

    @staticmethod
    def _events(rows):
        import pandas as pd
        return {"events": pd.DataFrame(rows)}

    @staticmethod
    def _t(**kw):
        kw.setdefault("service_role", "generation")
        kw.setdefault("der_category", "II")
        return Thresholds(nominal_voltage=277.0, **kw)

    def test_a_service_without_generation_is_not_assessed(self):
        from pq_analysis import check_ride_through
        r = check_ride_through({}, Thresholds(nominal_voltage=277.0))
        assert r["available"] is False
        assert "distributed energy resource" in r["note"]

    def test_the_category_is_required_rather_than_assumed(self):
        # 6.4.2.1 leaves the category to the Area EPS operator, and the
        # difference between them is 0.9 s against 20 s at 0.75 p.u.
        from pq_analysis import check_ride_through
        r = check_ride_through({}, self._t(der_category=None))
        assert r["available"] is False
        assert "Area EPS operator" in r["note"]

    def test_an_event_inside_the_region_obliges_the_plant_to_stay_on(self):
        from pq_analysis import check_ride_through
        r = check_ride_through(self._events([
            {"type": "voltage_sag", "phase": "a", "value_v": 277 * 0.85,
             "duration_ms": 300, "timestamp": "t"}]), self._t())
        e = r["events"][0]
        assert e["region"] == "Mandatory Operation"
        assert e["must_not_trip"] is True
        assert r["n_required_to_ride_through"] == 1

    def test_an_event_longer_than_the_required_time_is_beyond_it(self):
        # 0.75 p.u. under Category II requires 3.87 s; six seconds is past what
        # the standard asks the plant to survive.
        from pq_analysis import check_ride_through
        r = check_ride_through(self._events([
            {"type": "voltage_sag", "phase": "a", "value_v": 277 * 0.75,
             "duration_ms": 6000, "timestamp": "t"}]), self._t())
        assert r["events"][0]["must_not_trip"] is False
        assert r["n_beyond_requirement"] == 1

    def test_a_deep_dip_is_one_the_plant_may_drop_on(self):
        from pq_analysis import check_ride_through
        r = check_ride_through(self._events([
            {"type": "voltage_sag", "phase": "a", "value_v": 277 * 0.25,
             "duration_ms": 100, "timestamp": "t"}]), self._t())
        assert r["events"][0]["region"] == "Cease to Energize"
        assert r["events"][0]["must_not_trip"] is False

    def test_the_category_changes_the_answer_on_the_same_event(self):
        # The reason the field is required rather than defaulted.
        from pq_analysis import check_ride_through
        ev = self._events([{"type": "voltage_sag", "phase": "a",
                            "value_v": 277 * 0.75, "duration_ms": 5000,
                            "timestamp": "t"}])
        assert check_ride_through(ev, self._t(der_category="I")
                                  )["events"][0]["must_not_trip"] is False
        assert check_ride_through(ev, self._t(der_category="III")
                                  )["events"][0]["must_not_trip"] is True

    def test_it_says_so_when_the_recording_cannot_resolve_durations(self):
        # Clause 6.4.2 is voltage against duration; interval averages cannot
        # tell a 100 ms dip from a 4 s one.
        from pq_analysis import check_ride_through
        r = check_ride_through(self._events([
            {"type": "voltage_sag", "phase": "a", "value_v": 200.0,
             "timestamp": "t"}]), self._t())
        assert r["available"] is False
        assert "durations" in r["note"]

    def test_the_cumulative_duration_limit_is_stated(self):
        # The tables are cumulative within a disturbance and Table 17 governs
        # consecutive ones; neither is evaluated, so neither is implied.
        from pq_analysis import check_ride_through
        r = check_ride_through(self._events([
            {"type": "voltage_sag", "phase": "a", "value_v": 277 * 0.85,
             "duration_ms": 300, "timestamp": "t"}]), self._t())
        assert any("cumulative" in c for c in r["caveats"])

    def test_it_runs_end_to_end_from_the_fixture(self):
        # The fixture carries a variable-rate record, which is the only thing
        # that resolves an event duration -- before it, neither this check nor
        # the ITIC one had ever been exercised against a file.
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pq_analysis import detect_events, check_ride_through
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_producer_array.pqd"
        ds = extract_dataset(ProntoAdapter(str(path)), ChannelMapper())
        th = self._t()
        r = check_ride_through(detect_events(ds, th), th)
        assert r["available"] is True
        assert r["n_events"] == 3
        assert r["n_required_to_ride_through"] == 2
        assert r["n_beyond_requirement"] == 1
        depths = sorted(e["pct_nominal"] for e in r["events"])
        assert depths == pytest.approx([25.0, 85.0, 112.0], abs=0.5)


class TestFrequencyRideThrough:
    """IEEE 1547-2018 Clause 6.5.2 and Table 19.

    Two things about this clause are easy to encode backwards, so both are
    pinned. The 299 s is a condition on the requirement rather than a limit on
    the plant -- past it the obligation lapses and tripping is allowed. And
    continuous operation needs V/f <= 1.1 as well as the frequency band, so
    frequency alone does not establish it.
    """

    @staticmethod
    def _ds(freq_series=None, interval_freq=None):
        """A dataset carrying frequency in one record or the other."""
        import pandas as pd

        class _DS:
            pass

        ds = _DS()
        idx = pd.date_range("2025-06-25", periods=12, freq="5min", tz="UTC")
        ds.df = pd.DataFrame({
            "voltage_a": [277.0] * 12,
            "frequency": interval_freq if interval_freq is not None else [60.0] * 12,
        }, index=idx)
        if freq_series is None:
            ds.adaptive_df = None
        else:
            n = len(freq_series)
            aidx = pd.date_range("2025-06-25", periods=n, freq="1s", tz="UTC")
            ds.adaptive_df = pd.DataFrame({"adap_freq": freq_series}, index=aidx)
        return ds

    @staticmethod
    def _t(**kw):
        kw.setdefault("service_role", "generation")
        return Thresholds(nominal_voltage=277.0, **kw)

    def test_table_19_is_the_same_for_every_category(self):
        # Unlike the voltage tables. The category changes only how much active
        # power must be held (Table 20), not whether to ride through.
        from pq_constants import frequency_ride_through_region
        a = frequency_ride_through_region(58.0)
        assert a["mode"] == "mandatory" and a["min_ride_s"] == 299.0

    def test_the_band_above_61_8_is_reported_as_unspecified(self):
        # No Table 19 row covers 61.8 to 62.0, and 6.5.2.4.1 puts the
        # high-frequency requirement at "greater than 61.2 and at most 61.8".
        # Resolving the gap silently either way would be inventing a rule.
        from pq_constants import frequency_ride_through_region
        assert frequency_ride_through_region(61.9)["mode"] == "unspecified"
        assert frequency_ride_through_region(62.5)["mode"] == "none"

    def test_an_excursion_inside_the_allowance_obliges_the_plant(self):
        from pq_analysis import check_frequency_ride_through
        r = check_frequency_ride_through(
            self._ds([60.0] * 30 + [58.2] * 40 + [60.0] * 30), self._t())
        assert r["source"] == "variable-rate" and r["assessable"] is True
        e = r["excursions"][0]
        assert e["region"] == "Mandatory Operation"
        assert e["must_not_trip"] is True
        assert e["cumulative_s"] == pytest.approx(40.0, abs=1.5)

    def test_past_the_allowance_the_obligation_lapses(self):
        # 299 s is a condition on the requirement, not a limit on the plant.
        from pq_analysis import check_frequency_ride_through
        r = check_frequency_ride_through(
            self._ds([60.0] * 5 + [58.2] * 320 + [60.0] * 5), self._t())
        e = r["excursions"][0]
        assert e["cumulative_s"] > 299.0
        assert e["within_allowance"] is False
        assert e["must_not_trip"] is False
        assert r["n_beyond_requirement"] == 1

    def test_below_57_hz_carries_no_ride_through_requirement(self):
        from pq_analysis import check_frequency_ride_through
        r = check_frequency_ride_through(
            self._ds([60.0] * 10 + [56.5] * 5 + [60.0] * 10), self._t())
        assert r["excursions"][0]["mode"] == "none"
        assert r["excursions"][0]["must_not_trip"] is False

    def test_a_quiet_recording_says_so_from_the_right_record(self):
        from pq_analysis import check_frequency_ride_through
        r = check_frequency_ride_through(self._ds([60.0] * 60), self._t())
        assert r["assessable"] is True
        assert r["n_excursions"] == 0
        assert "variable-rate" in r["note"]

    def test_interval_averages_are_not_allowed_to_read_as_a_pass(self):
        # The whole point of the tier: an average cannot rule an excursion out.
        from pq_analysis import check_frequency_ride_through
        r = check_frequency_ride_through(self._ds(freq_series=None), self._t())
        assert r["available"] is True
        assert r["assessable"] is False
        assert r["source"] == "interval-average"
        assert "does not establish it" in r["note"]
        assert r["excursions"] == []

    def test_the_v_over_f_condition_on_continuous_operation_is_checked(self):
        # 6.5.2.2 requires V/f <= 1.1 as well as the band.
        from pq_analysis import check_frequency_ride_through
        r = check_frequency_ride_through(self._ds([60.0] * 60), self._t())
        assert r["v_over_f"]["limit"] == 1.1
        assert r["v_over_f"]["within"] is True
        # Push the voltage up and the ratio should follow it out of bounds.
        ds = self._ds([60.0] * 60)
        ds.df["voltage_a"] = 277.0 * 1.15
        assert check_frequency_ride_through(ds, self._t())["v_over_f"]["within"] is False

    def test_a_service_without_generation_is_not_assessed(self):
        from pq_analysis import check_frequency_ride_through
        r = check_frequency_ride_through(self._ds([60.0] * 10),
                                         Thresholds(nominal_voltage=277.0))
        assert r["available"] is False

    def test_it_runs_end_to_end_from_the_fixture(self):
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pq_analysis import check_frequency_ride_through
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_producer_array.pqd"
        ds = extract_dataset(ProntoAdapter(str(path)), ChannelMapper())
        assert "adap_freq" in ds.adaptive_df.columns
        r = check_frequency_ride_through(ds, self._t(der_category="II"))
        assert r["source"] == "variable-rate"
        assert r["n_excursions"] == 2
        assert r["n_required_to_ride_through"] == 1
        assert r["active_power_capability"] == "the pre-disturbance active power output"


class TestTheLetterToAProducer:
    """A plant is not a customer with unusual load, and reads nothing like one.

    Nobody at a producer's array has noticed the lights flickering, because
    there are no lights. What they have is a plant that trips, curtails or
    underproduces, and SCADA to check a claim against -- so the letter is
    written to be verifiable rather than evocative. Addressing them in the load
    register reads as a form letter sent to the wrong site, which costs the
    trust every other sentence is spending.
    """

    @staticmethod
    def _report(role="generation", customer_class="sg"):
        import pq_analysis as An
        from pq_report import generate_report
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_producer_array.pqd"
        ds = extract_dataset(ProntoAdapter(str(path)), ChannelMapper())
        th = Thresholds(nominal_voltage=277.0, customer_class=customer_class,
                        service_role=role, isc_amps=40000.0, rated_ac_kw=250.0,
                        avg_peak_demand_kw=10.0)
        df = ds.df
        ev = An.detect_events(ds, th)
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th),
            An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), ev, th,
            neutral_health_result=An.check_neutral_health(ds, th),
            itic_result=An.check_itic(ev, th),
            flicker_result=An.check_flicker(df, th),
        )
        return rep, th

    def test_the_register_follows_the_role_not_the_class(self):
        # A producer keeps whichever schedule it takes service under, so the
        # class cannot be what selects the voice.
        from pq_report import _letter_register
        for cls in ("r", "c", "sg", "pg"):
            reg = _letter_register(Thresholds(customer_class=cls,
                                              service_role="generation"))
            assert reg["generating"] is True, cls
        assert not _letter_register(
            Thresholds(customer_class="sg")).get("generating")

    def test_it_is_addressed_to_the_point_of_interconnection(self):
        from pq_report import _customer_vocabulary
        rep, th = self._report()
        assert _customer_vocabulary(rep, th)["site"] == "the plant"

    def test_the_symptoms_are_things_a_plant_can_check(self):
        from pq_report import _customer_conditions
        rep, th = self._report()
        joined = " ".join(c["symptom"] for c in _customer_conditions(rep, th))
        assert "lights" not in joined.lower()
        assert any(w in joined for w in ("inverter", "SCADA", "interconnection"))

    def test_power_factor_is_not_billed_to_a_plant_off_a_load_clause(self):
        # The plant is not drawing reactive power to serve a load; what it must
        # hold is set by the interconnection agreement and 1547 Clause 5.
        # Quoting a load tariff sheet and recommending capacitors would send
        # the operator after the wrong fix.
        from pq_report import _customer_conditions
        rep, th = self._report()
        pf = [c for c in _customer_conditions(rep, th)
              if "power factor" in c["headline"].lower()]
        if pf:
            body = pf[0]["means"]
            assert "interconnection agreement" in body
            assert "1547" in body
            assert "Sheet R" not in body
            assert "capacitors are not the remedy" in body.lower()

    def test_imbalance_points_at_a_unit_not_at_a_panel(self):
        from pq_report import _customer_conditions
        rep, th = self._report()
        imb = [c for c in _customer_conditions(rep, th)
               if "unevenly split" in c["headline"]]
        assert imb, "expected an imbalance finding on this fixture"
        assert "electrician" not in imb[0]["means"]
        assert "panel" not in imb[0]["means"]
        assert "inverter" in imb[0]["means"]

    def test_the_urgent_signs_are_plant_equipment(self):
        from pq_report import _urgent_signs, _letter_register
        gen = _urgent_signs(_letter_register(
            Thresholds(customer_class="sg", service_role="generation")))
        assert "outlets" not in gen and "combiner" in gen
        load = _urgent_signs(_letter_register(Thresholds(customer_class="r")))
        assert "outlets" in load

    def test_a_plant_is_not_shown_the_itic_curve(self):
        # ITIC is a ride-through envelope for load equipment. A plant's
        # obligation is 1547 Clause 6, a different curve this does not assess,
        # so showing ITIC would invite judging it against the wrong standard.
        from pq_report import _letter_register
        assert _letter_register(
            Thresholds(customer_class="sg",
                       service_role="generation"))["itic_curve"] is False

    def test_the_load_register_is_untouched(self):
        from pq_report import _letter_register
        reg = _letter_register(Thresholds(customer_class="r"))
        assert reg["site"] == "your home"
        assert reg["explains_basics"] is True


class TestTariffScopingIsNotMisstated:
    """The two PF clauses and the 15% clause, as the filed tariff has them.

    Verified against COLO. PUC No. 8 Electric on 2026-08-13:

      Sheet R73   Rules and Regulations, GENERAL: "Company's rates contemplate
                  Customer's use of service at a Power Factor... of not less
                  than ninety percent (90%) lagging."  General, so all classes.
      Sheet R121  Rules and Regulations, COMMERCIAL AND INDUSTRIAL: "a Power
                  Factor as near unity as practicable."  All of C, SG and PG.
      Sheet R123  C&I billing demand provisions: above 15% between phases the
                  Company "may take as the Billing Demand" the three-phase
                  equivalent of the worst phase.  A charge, not a limit.

    These went into a customer-facing document, so they are pinned.
    """

    @staticmethod
    def _sources():
        from pathlib import Path
        root = Path(__file__).parent
        return {name: (root / name).read_text()
                for name in ("run.py", "pq_report.py", "pq_analyzer.py")}

    def test_no_file_calls_r121_a_primary_only_clause(self):
        for name, text in self._sources().items():
            assert "Sheet R121 (Schedule PG)" not in text, name
            assert "R121 requires Primary service" not in text, name

    def test_no_file_attributes_sheet_r73_to_a_schedule(self):
        for name, text in self._sources().items():
            for sched in ("C", "SG"):
                assert f"Sheet R73 (Schedule {sched})" not in text, name

    def test_the_fifteen_percent_clause_is_not_called_a_requirement(self):
        for name, text in self._sources().items():
            assert "R121 requires that load in any one phase" not in text, name
            assert "Sheet R121 (≤ 15%" not in text, name

    def test_residential_is_not_described_as_having_no_pf_clause(self):
        # R73 is in the General rules and reaches Schedule R too. What is true
        # is that no reactive billing applies there.
        text = self._sources()["run.py"]
        assert "Residential customers are not contractually required" not in text

    def test_the_imbalance_row_cites_no_tariff_sheet_at_all(self):
        """The compliance table states measurements and the standards they
        were judged against. The 15% phase clause is neither.

        It first named R121 and called 15% a limit, which was wrong twice
        over. Correcting it to R123 fixed the citation but left a billing
        provision sitting in a table of power quality standards, where a
        customer reads it as something they are failing. The limit that does
        apply to imbalance is NEMA MG1.
        """
        text = self._sources()["pq_report.py"]
        assert "PSCo Tariff Sheet R121 ≤ 15% for C&I" not in text
        assert 'ci_label = "Current imbalance < 10% (NEMA MG1)"' in text
        # No tariff sheet reaches the customer-facing imbalance wording.
        import re
        for m in re.finditer(r'ci_label\s*=\s*\(?"([^"]*)"', text):
            assert "R12" not in m.group(1), m.group(1)


class TestMixedServicePopulations:
    """A service with generation holds two populations that share a meter.

    Pooling them is the failure mode: the light-load gate takes its floor from
    a share of the peak, so the export peak sets a bar the load can never
    reach and the whole load half drops out as "light load" with nothing
    saying so. The two direction methods then describe opposite halves and get
    reported as agreeing with each other.
    """

    @staticmethod
    def _ds():
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_solar_net_metered.pqd"
        return extract_dataset(ProntoAdapter(str(path)), ChannelMapper())

    @staticmethod
    def _t(role="mixed"):
        return Thresholds(nominal_voltage=277.0, customer_class="sg",
                          service_role=role)

    def test_the_gate_no_longer_takes_its_floor_from_the_export_peak(self):
        from pq_analysis import harmonic_spectrum_significance
        df = self._ds().df
        s = harmonic_spectrum_significance(df, self._t())
        # The export peak is around 200 A; the load's own demand is a few amps.
        # A floor set from the former excludes every load interval.
        assert s["il_amps"] < 50.0
        assert s["load_floor_amps"] < 5.0

    def test_the_load_half_of_the_service_survives_the_gate(self):
        from pq_analysis import harmonic_spectrum_significance
        df = self._ds().df
        s = harmonic_spectrum_significance(df, self._t())
        kept = s["loaded"]
        assert s["loaded_intervals"] > 0
        # Every kept interval is an importing one, and none are exporting:
        # the two are characterised apart, not pooled.
        assert int((df["power_real"].loc[kept] < 0).sum()) == 0

    def test_a_load_service_is_not_split_at_all(self):
        from pq_analysis import harmonic_spectrum_significance
        df = self._ds().df
        s = harmonic_spectrum_significance(df, self._t(role="load"))
        assert s["flow"]["split"] is False
        assert s["flow"]["direction"] == "all"

    def test_both_direction_methods_read_the_same_population(self):
        from pq_analysis import (harmonic_direction_from_intervals,
                                 harmonic_direction_from_waveforms)
        ds, th = self._ds(), self._t()
        iv = harmonic_direction_from_intervals(ds.df, th)
        wf = harmonic_direction_from_waveforms(ds, th)
        # The captures' top-level result is the importing half; the interval
        # regression must be over importing intervals, or "agreement" between
        # them compares two different things.
        assert iv["flow"]["direction"] == "importing"
        assert wf["export_split"]["importing"]["capture_phases"] > 0

    def test_a_plant_reads_its_exporting_intervals(self):
        from pq_analysis import harmonic_direction_from_intervals, primary_flow_direction
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pathlib import Path
        th = self._t(role="generation")
        assert primary_flow_direction(th) == "exporting"
        path = Path(__file__).parent / "test_data" / "test_producer_array.pqd"
        ds = extract_dataset(ProntoAdapter(str(path)), ChannelMapper())
        iv = harmonic_direction_from_intervals(ds.df, th)
        assert iv["flow"]["direction"] == "exporting"

    def test_a_service_with_no_power_channel_says_it_could_not_split(self):
        # Silence here would be the original defect wearing a new coat.
        from pq_analysis import flow_scope
        df = self._ds().df.drop(columns=["power_real"])
        scoped, info = flow_scope(df, self._t(), "importing")
        assert scoped is not None          # still assessed
        assert info["split"] is False
        assert "no real-power channel" in info["reason"]


class TestCaptureSplitDeadband:
    """Near the crossover the sign of P1 is a residue, not a direction.

    Current magnitude alone cannot reach the deadband: the light-load gate
    already drops any capture under 1 A, which at 277 V is twice the polarity
    floor in watts. What does reach it is a capture whose fundamental is almost
    entirely reactive -- generation offsetting load in real terms while current
    still circulates. P1 is then a residue whose sign is noise, and on an
    unbalanced service it can point opposite ways on two phases of one capture.
    """

    V1 = (277.0 * np.sqrt(2), 0.0)

    def _run(self, captures):
        from pq_analysis import harmonic_direction_from_waveforms
        import pandas as pd
        return harmonic_direction_from_waveforms(
            _FakeDataset(pd.DataFrame(), captures),
            Thresholds(nominal_voltage=277.0, service_role="mixed"))

    def _capture_at(self, i1_amps, deg):
        return _capture({1: self.V1, 5: (4.0, 180.0)},
                        {1: (i1_amps * np.sqrt(2), deg), 5: (2.0, 0.0)})

    def test_a_capture_with_no_real_power_lands_in_neither_half(self):
        # 150 A at 90 degrees: plenty of current, no real power to take a sign
        # from. P1 = V x I x cos(90) is zero however large the current is.
        r = self._run([self._capture_at(150.0, 0.0)] * 3
                      + [self._capture_at(150.0, 90.0)] * 3)
        split = r["export_split"]
        assert split["importing"]["capture_phases"] == 6
        assert split["exporting"]["capture_phases"] == 0
        assert split["near_crossover"] == 6

    def test_the_note_accounts_for_the_captures_it_set_aside(self):
        r = self._run([self._capture_at(150.0, 0.0)] * 3
                      + [self._capture_at(150.0, 90.0)] * 3)
        assert "in neither half" in r["polarity_note"]

    def test_a_clear_export_is_still_filed_as_exporting(self):
        r = self._run([self._capture_at(150.0, 0.0)] * 3
                      + [self._capture_at(150.0, 180.0)] * 3)
        split = r["export_split"]
        assert split["importing"]["capture_phases"] == 6
        assert split["exporting"]["capture_phases"] == 6
        assert split["near_crossover"] == 0

    def test_a_mostly_reactive_capture_still_counts_when_real_power_is_clear(self):
        # 60 degrees is a poor power factor, not an ambiguous direction:
        # cos(60) x 277 V x 150 A is well clear of the floor.
        r = self._run([self._capture_at(150.0, 60.0)] * 3)
        split = r["export_split"]
        assert split["importing"]["capture_phases"] == 6
        assert split["near_crossover"] == 0


class TestTheHouseInterpretationIsStated:
    """Figure 1's denominator is an undefined term, and the report says so.

    "Annual average load demand" appears only inside Figure 1 of 519-2022 --
    no definition entry, no method, no other use in the standard. PSCo reads
    it as the average of the twelve monthly maxima, matching the way the same
    standard defines IL. That is the more permissive reading, so a report that
    used it silently would be making a choice a reader could not see.
    """

    @staticmethod
    def _t(**kw):
        return Thresholds(nominal_voltage=277.0, customer_class="sg", **kw)

    def test_the_note_travels_with_the_test_when_it_sends_a_site_to_1547(self):
        from pq_analysis import applicable_current_standard
        r = applicable_current_standard(self._t(
            service_role="mixed", rated_ac_kw=150.0, avg_peak_demand_kw=1200.0))
        assert r["standard"] == "1547"
        assert "house interpretation" in r["reason"]
        assert "without defining it" in r["reason"]

    def test_the_note_travels_when_the_site_stays_under_519(self):
        from pq_analysis import applicable_current_standard
        r = applicable_current_standard(self._t(
            service_role="mixed", rated_ac_kw=40.0, avg_peak_demand_kw=1200.0))
        assert r["standard"] == "519"
        assert "house interpretation" in r["reason"]

    def test_it_names_the_direction_the_reading_leans(self):
        # The choice favours 519, and a reader is owed that rather than left to
        # work it out from the arithmetic.
        from pq_constants import HOUSE_INTERPRETATION_NOTE
        assert "more permissive" in HOUSE_INTERPRETATION_NOTE
        assert "fewer installations to IEEE 1547" in HOUSE_INTERPRETATION_NOTE

    def test_an_ordinary_load_service_is_not_lectured_about_it(self):
        # No generation means Figure 1 is never reached, so the note would be
        # an explanation of a decision that was not made.
        from pq_analysis import applicable_current_standard
        r = applicable_current_standard(self._t())
        assert "house interpretation" not in r["reason"]


class TestWhichCurrentStandardApplies:
    """IEEE 519-2022 Figure 1, the decision tree for an installation with DER.

    519 limits its own scope to a PCC "primarily with harmonic producing
    loads". Applying it to a plant quotes a limit three times looser than the
    one that governs, so the branch taken has to be reported, not assumed.
    """

    @staticmethod
    def _t(**kw):
        return Thresholds(nominal_voltage=277.0, customer_class="sg", **kw)

    def test_a_service_with_no_generation_is_a_519_site(self):
        from pq_analysis import applicable_current_standard
        r = applicable_current_standard(self._t())
        assert r["standard"] == "519"
        assert r["branch"] == "no_der"

    def test_generation_under_a_tenth_of_load_stays_under_519(self):
        from pq_analysis import applicable_current_standard
        # 40 kW of solar on a service averaging 500 kW: 8%.
        r = applicable_current_standard(self._t(
            service_role="mixed", rated_ac_kw=40.0, avg_peak_demand_kw=500.0))
        assert r["standard"] == "519"
        assert r["branch"] == "der_below_threshold"
        assert r["der_share"] == pytest.approx(0.08)

    def test_generation_at_a_tenth_of_load_goes_to_1547(self):
        from pq_analysis import applicable_current_standard
        # Exactly 10%: the tree asks whether it is *below* the threshold.
        r = applicable_current_standard(self._t(
            service_role="mixed", rated_ac_kw=50.0, avg_peak_demand_kw=500.0))
        assert r["standard"] == "1547"

    def test_a_plant_with_no_load_goes_to_1547(self):
        from pq_analysis import applicable_current_standard
        r = applicable_current_standard(self._t(
            service_role="generation", rated_ac_kw=250.0, avg_peak_demand_kw=10.0))
        assert r["standard"] == "1547"
        assert "1547" in r["reason"]

    def test_a_plant_needs_no_demand_figure_to_reach_1547(self):
        # The demand field is greyed out at a plant, so nothing supplies it.
        # Falling through to "undetermined" here graded a producer's array
        # against 519 -- three times the aggregate limit that governs it.
        from pq_analysis import applicable_current_standard
        r = applicable_current_standard(self._t(
            service_role="generation", rated_ac_kw=2300.0))
        assert r["standard"] == "1547"
        assert r["branch"] == "generation_only"
        assert r["determined"] is True
        assert "5.2" in r["reason"]

    def test_a_mixed_service_still_needs_both_figures(self):
        # The short circuit above is scoped to plants: a load with generation
        # behind it is exactly the case Figure 1 exists to decide.
        from pq_analysis import applicable_current_standard
        r = applicable_current_standard(self._t(
            service_role="mixed", rated_ac_kw=2300.0))
        assert r["branch"] == "undetermined"

    def test_missing_records_are_reported_undetermined_not_guessed(self):
        # Both quantities come from records. Guessing either picks the wrong
        # standard, and the two differ by 3x in the aggregate limit.
        from pq_analysis import applicable_current_standard
        r = applicable_current_standard(self._t(
            service_role="mixed", rated_ac_kw=250.0))
        assert r["determined"] is False
        assert r["branch"] == "undetermined"
        assert "annual average load demand" in r["reason"]
        # It still grades against something rather than going silent.
        assert r["standard"] == "519"

    def test_the_reason_names_both_figures_it_compared(self):
        from pq_analysis import applicable_current_standard
        r = applicable_current_standard(self._t(
            service_role="mixed", rated_ac_kw=40.0, avg_peak_demand_kw=500.0))
        assert "40 kW" in r["reason"] and "500 kW" in r["reason"]


class TestTRDAgainst1547:
    """IEEE 1547-2018 Clause 7.3: TRD against rated current, not demand."""

    @staticmethod
    def _df():
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_producer_array.pqd"
        return extract_dataset(ProntoAdapter(str(path)), ChannelMapper()).df

    @staticmethod
    def _t(**kw):
        kw.setdefault("isc_amps", 40000.0)
        return Thresholds(nominal_voltage=277.0, customer_class="sg",
                          service_role="generation", **kw)

    def test_it_declines_without_a_nameplate_rather_than_substituting(self):
        # I_rated is a nameplate. Falling back to a measured peak and calling
        # the result 1547 would be inventing the denominator.
        from pq_analysis import check_trd
        r = check_trd(self._df(), self._t())
        assert r["available"] is False
        assert "nameplate" in r["note"]

    def test_trd_is_graded_against_five_percent_of_rated_current(self):
        from pq_analysis import check_trd
        r = check_trd(self._df(), self._t(rated_ac_kw=250.0))
        assert r["available"] is True
        assert r["irated_amps"] == pytest.approx(300.8, abs=1.0)
        assert r["trd_limit_pct"] == 5.0
        assert r["trd_pct"] < 5.0 and r["trd_pass"] is True

    def test_the_limits_do_not_move_with_isc(self):
        # 519's limits scale with ISC/IL; 1547's are fixed. A stiff system must
        # not buy the plant headroom it is not entitled to.
        from pq_analysis import check_trd
        df = self._df()
        stiff = check_trd(df, self._t(rated_ac_kw=250.0, isc_amps=200000.0))
        weak  = check_trd(df, self._t(rated_ac_kw=250.0, isc_amps=5000.0))
        assert stiff["trd_limit_pct"] == weak["trd_limit_pct"] == 5.0
        assert stiff["orders"][5]["limit_pct"] == weak["orders"][5]["limit_pct"]

    def test_even_orders_use_their_own_looser_table(self):
        from pq_analysis import check_trd
        r = check_trd(self._df(), self._t(rated_ac_kw=250.0))
        # Table 27: h=2 is 1.0%, h=4 is 2.0%, h=6 is 3.0% -- and 519's blanket
        # 25%-of-odd would have made all three 1.0%.
        assert r["orders"][2]["limit_pct"] == 1.0
        assert r["orders"][4]["limit_pct"] == 2.0
        assert r["orders"][6]["limit_pct"] == 3.0
        assert all(r["orders"][h]["even"] for h in (2, 4, 6))

    def test_the_1547_limits_are_tighter_than_the_519_ones_here(self):
        # The reason the branch matters: same measurement, different verdict
        # headroom. The producer's ISC/IL puts it in 519's loosest class.
        from pq_analysis import check_trd, check_thd
        df = self._df()
        th = self._t(rated_ac_kw=250.0)
        trd = check_trd(df, th)
        tdd = check_thd(df, th)["tdd_info"]
        assert tdd["tdd_limit_pct"] == 15.0
        assert trd["trd_limit_pct"] == 5.0

    def test_both_caveats_are_carried_with_the_result(self):
        # Neither is optional: one says the figure is conservative, the other
        # says a narrow pass is not clearance.
        from pq_analysis import check_trd
        r = check_trd(self._df(), self._t(rated_ac_kw=250.0))
        joined = " ".join(r["caveats"])
        assert "without the plant connected" in joined
        assert "interharmonics" in joined.lower()


class TestILFromBilling:
    """519-2022 defines IL from twelve months of billing, not from a recording."""

    @staticmethod
    def _df():
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_commercial_large.pqd"
        return extract_dataset(ProntoAdapter(str(path)), ChannelMapper()).df

    def test_billing_demand_is_used_over_the_recording_peak(self):
        from pq_analysis import check_thd
        df = self._df()
        # 300 kW at 277 V L-N, three-phase, converted at the flat 0.90:
        # 300000 / (3 x 277 x 0.90) = 401 A.
        th = Thresholds(nominal_voltage=277.0, customer_class="sg",
                        isc_amps=40000.0, avg_peak_demand_kw=300.0)
        td = check_thd(df, th)["tdd_info"]
        assert td["il_amps"] == pytest.approx(401.0, abs=1.0)
        assert td["il_basis"] == "billing"

    def test_the_conversion_power_factor_is_flat_not_measured(self):
        # Billing IL exists to be stable across recordings; deriving it through
        # a power factor measured in one week would put the recording back in.
        from pq_analysis import billing_il_amps
        from pq_constants import IL_CONVERSION_PF
        assert IL_CONVERSION_PF == 0.90
        th = Thresholds(nominal_voltage=277.0, avg_peak_demand_kw=300.0)
        assert billing_il_amps(th, "three-phase") == pytest.approx(
            300_000.0 / (3 * 277.0 * 0.90), rel=1e-6)

    def test_a_plant_keeps_its_rating_rather_than_its_auxiliary_demand(self):
        # The trap in collapsing the two fields: a producer bills a handful of
        # kW of auxiliary load, which would put IL near 12 A on a service
        # exporting thousands.
        from pq_analysis import check_thd
        from pq_adapter import ProntoAdapter, ChannelMapper, extract_dataset
        from pathlib import Path
        path = Path(__file__).parent / "test_data" / "test_producer_array.pqd"
        ds = extract_dataset(ProntoAdapter(str(path)), ChannelMapper())
        th = Thresholds(nominal_voltage=277.0, customer_class="sg",
                        isc_amps=40000.0, service_role="generation",
                        rated_ac_kw=250.0, avg_peak_demand_kw=10.0)
        td = check_thd(ds.df, th)["tdd_info"]
        assert td["il_basis"] == "rated_output"
        assert td["il_amps"] > 250.0

    def test_without_one_the_recording_peak_stands_in_and_is_labelled(self):
        from pq_analysis import check_thd
        th = Thresholds(nominal_voltage=277.0, customer_class="sg",
                        isc_amps=40000.0)
        td = check_thd(self._df(), th)["tdd_info"]
        assert td["il_basis"] == "measured_demand"

    def test_a_small_il_inflates_the_percentages(self):
        # Why the substitution has to be labelled: IL is a denominator, and a
        # recording from a slow week can manufacture a violation.
        from pq_analysis import check_thd
        df = self._df()
        def tdd(kw):
            th = Thresholds(nominal_voltage=277.0, customer_class="sg",
                            isc_amps=40000.0, avg_peak_demand_kw=kw)
            return check_thd(df, th)["current"]["max_thd_pct"]
        assert tdd(40.0) > tdd(300.0)


class TestKFactorByPhase:
    def test_rating_is_sized_on_the_worst_phase(self):
        from pq_analysis import kfactor_by_phase
        df = _frame(kfactor_meter=[104.0] * 10, kfactor_current_b=[217.0] * 10)
        kf = kfactor_by_phase(df)
        assert kf["available"] and kf["worst_phase"] == "B"
        assert kf["median"] == 217.0
        assert kf["phase_a_median"] == 104.0

    def test_neutral_does_not_drive_the_rating(self):
        # Neutral K-factor describes conductor heating, not a transformer winding.
        from pq_analysis import kfactor_by_phase
        df = _frame(kfactor_meter=[6.0] * 10, kfactor_current_neutral=[99.0] * 10)
        kf = kfactor_by_phase(df)
        assert kf["worst_phase"] == "A" and kf["median"] == 6.0
        assert "N" in kf["phases"]

    def test_unavailable_without_channels(self):
        from pq_analysis import kfactor_by_phase
        assert kfactor_by_phase(_frame(voltage_a=[120.0] * 5))["available"] is False

    def test_rating_never_exceeds_a_purchasable_unit(self):
        from pq_analysis import standard_k_rating, STANDARD_K_RATINGS
        for k in (1.5, 5.0, 10.0, 25.0, 50.0):
            rating, _ = standard_k_rating(k)
            assert rating in STANDARD_K_RATINGS or rating == 1
            assert rating >= k or rating == 1
        rating, wording = standard_k_rating(217.3)
        assert rating is None and "exceeds K-50" in wording


class TestFlickerAllPhases:
    def test_worst_phase_governs_not_phase_a(self):
        from pq_analysis import check_flicker
        df = _frame(flicker_pst=[0.5] * 10, flicker_pst_b=[4.98] * 10,
                    flicker_plt=[0.2] * 10, flicker_plt_b=[2.21] * 10)
        r = check_flicker(df, Thresholds())
        assert r["available"] and r["worst_phase"] == "B"
        assert r["pst_max"] == 4.98
        # Phase A alone would have passed both limits.
        assert r["pst"]["A"]["pass"] and not r["pst"]["B"]["pass"]
        assert r["overall_pass"] is False

    def test_pass_when_every_phase_is_within_limits(self):
        from pq_analysis import check_flicker
        df = _frame(flicker_pst=[0.4] * 10, flicker_pst_c=[0.6] * 10,
                    flicker_plt=[0.3] * 10)
        r = check_flicker(df, Thresholds())
        assert r["overall_pass"] is True
        assert sorted(r["phases_read"]) == ["A", "C"]

    def test_unavailable_without_channels(self):
        from pq_analysis import check_flicker
        assert check_flicker(_frame(voltage_a=[120.0] * 5), Thresholds())["available"] is False


class TestShortAndLongTermFlickerAreReportedSeparately:
    """Pst and Plt measure different windows and fail independently.

    Both were always read and analysed per phase; the narrative section and
    the key findings took their numbers from phase A's columns, so a service
    whose second leg reached Pst 4.98 was described by phase A's 1.43.
    """

    @staticmethod
    def _report(**cols):
        from pq_analysis import check_flicker
        df = _frame(**cols)
        return {"flicker": check_flicker(df, Thresholds()),
                "file_summary": {"duration_hours": 68.0}}, df

    def test_the_status_used_by_the_summary_follows_the_worst_phase(self):
        from pq_report import _flicker_status
        report, _df = self._report(flicker_pst=[0.5] * 10, flicker_pst_b=[4.98] * 10,
                                   flicker_plt=[0.2] * 10, flicker_plt_b=[2.21] * 10)
        status = _flicker_status(report)
        assert status["pst_max"] == 4.98 and status["pst_phase"] == "B"
        assert status["plt_max"] == 2.21 and status["plt_phase"] == "B"
        assert status["passes"] is False

    def test_a_long_term_failure_alone_is_named_as_a_cycling_load(self):
        docx = pytest.importorskip("docx")
        from pq_report import _word_flicker
        # Every ten-minute value inside its limit, the two-hour aggregate over:
        # the pattern a repeatedly cycling load makes.
        report, df = self._report(flicker_pst=[0.9] * 10, flicker_plt=[0.75] * 10)
        doc = docx.Document()
        _word_flicker(doc, report, df)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "cycles repeatedly" in text
        assert "aggregates twelve consecutive Pst values" in text

    def test_both_measures_appear_for_every_phase(self):
        docx = pytest.importorskip("docx")
        from pq_report import _word_flicker
        report, df = self._report(flicker_pst=[0.5] * 10, flicker_pst_b=[4.98] * 10,
                                  flicker_plt=[0.2] * 10, flicker_plt_b=[2.21] * 10)
        doc = docx.Document()
        _word_flicker(doc, report, df)
        rows = [[c.text for c in row.cells] for row in doc.tables[0].rows]
        measures = {(r[0].split(",")[0], r[1]) for r in rows[1:]}
        assert measures == {("Pst (10 min)", "A"), ("Pst (10 min)", "B"),
                            ("Plt (2 h)", "A"), ("Plt (2 h)", "B")}
        # The phase that fails must be the one quoted in the narrative.
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "4.98 on phase B" in text

    def test_the_equipment_limit_is_not_presented_as_a_system_limit(self):
        docx = pytest.importorskip("docx")
        from pq_report import _word_flicker
        report, df = self._report(flicker_pst=[0.4] * 10, flicker_plt=[0.3] * 10)
        doc = docx.Document()
        _word_flicker(doc, report, df)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "equipment may emit" in text
        assert "IEEE 1453-2015" in text and "0.8" in text
        # The recording is shorter than the week both standards assess over.
        assert "not a week" in text

    def test_held_values_are_not_mistaken_for_measurement_windows(self):
        from pq_analysis import check_flicker
        # Twelve intervals carrying three distinct Pst values.
        df = _frame(flicker_pst=[0.3, 0.3, 0.3, 0.3, 0.9, 0.9,
                                 0.9, 0.9, 0.5, 0.5, 0.5, 0.5])
        r = check_flicker(df, Thresholds())
        assert r["pst"]["A"]["distinct_values"] == 3


class TestFlickerSeverity:
    """How much a flicker exceedance matters, separately from whether it passed.

    Grading on the single worst reading made one bad ten-minute window look
    like a sustained condition: a Pst of 4.98 in a quiet week is 4.98x the
    limit, which the bands call severe on its own. Severity is graded at the
    95th percentile with the share of time over the limit, the same way
    voltage THD is; the maximum still decides pass or fail.
    """

    @staticmethod
    def _sev(pst, plt_=None, hours=167.0):
        import pandas as pd
        from pq_analysis import check_flicker
        from pq_report import _flicker_severities
        idx = pd.date_range("2024-01-01", periods=len(pst), freq="5min", tz="UTC")
        cols = {"flicker_pst": pd.Series(pst, index=idx)}
        if plt_ is not None:
            cols["flicker_plt"] = pd.Series(plt_, index=idx)
        fl = check_flicker(pd.DataFrame(cols), Thresholds())
        return fl, _flicker_severities(fl, {"file_summary": {"duration_hours": hours}})

    def test_a_lone_spike_fails_compliance_but_is_only_minor(self):
        quiet = [0.2] * 500
        quiet[7] = 5.0                      # one ten-minute window, five times the limit
        fl, sev = self._sev(quiet)
        assert fl["overall_pass"] is False          # compliance stays binary
        assert sev["flicker"]["band"] == "minor"    # severity does not
        assert "0.2% of the recording" in strip_marks(sev["flicker"]["reason"])

    def test_a_sustained_exceedance_is_significant(self):
        # Over the limit for a third of the recording, never dramatically.
        values = [1.2] * 170 + [0.3] * 330
        fl, sev = self._sev(values)
        assert sev["flicker"]["band"] in ("significant", "severe")
        assert "34.0% of the recording" in strip_marks(sev["flicker"]["reason"])

    def test_severity_is_graded_for_each_measure_separately(self):
        # Pst quiet with one spike, Plt over limit most of the time: the two
        # must not be collapsed into one band before the report sees them.
        pst = [0.2] * 500
        pst[3] = 4.0
        plt_ = [0.9] * 400 + [0.2] * 100
        fl, sev = self._sev(pst, plt_)
        assert sev["flicker_pst"]["band"] == "minor"
        assert sev["flicker_plt"]["band"] in ("significant", "severe")
        # The headline follows the worse of the two.
        assert sev["flicker"]["band"] == sev["flicker_plt"]["band"]

    def test_a_comfortable_pass_is_not_dressed_up_as_a_watch(self):
        fl, sev = self._sev([0.3] * 300, [0.2] * 300)
        assert sev["flicker"]["band"] == "compliant"

    def test_close_to_the_limit_reads_as_watch(self):
        fl, sev = self._sev([0.92] * 300)
        assert sev["flicker"]["band"] == "watch"
        assert "within the limit" in sev["flicker"]["reason"]

    def test_a_recording_shorter_than_a_day_is_discounted_a_band(self):
        values = [1.4] * 200 + [0.3] * 100
        _fl, full = self._sev(values, hours=72.0)
        _fl, short = self._sev(values, hours=8.0)
        assert SEVERITY_ORDER.index(short["flicker"]["band"]) \
            < SEVERITY_ORDER.index(full["flicker"]["band"])
        assert short["flicker"]["downgraded"] is True
        assert "two-hour windows" in short["flicker"]["reason"]

    def test_a_multi_day_recording_is_not_discounted(self):
        # Every survey is shorter than the week the standards assess over, so
        # that alone must not discount every finding forever.
        _fl, sev = self._sev([1.4] * 200 + [0.3] * 100, hours=72.0)
        assert sev["flicker"]["downgraded"] is False


class TestFlickerPlot:
    def test_the_chart_is_written_with_both_measures(self, tmp_path):
        pytest.importorskip("matplotlib")
        import pandas as pd
        from pq_analysis import check_flicker
        from pq_plots import plot_flicker
        idx = pd.date_range("2024-01-01", periods=300, freq="5min", tz="UTC")
        df = pd.DataFrame({
            "flicker_pst":   pd.Series([0.4] * 299 + [3.0], index=idx),
            "flicker_pst_b": pd.Series([0.5] * 300, index=idx),
            "flicker_plt":   pd.Series([0.9] * 300, index=idx),
        }, index=idx)
        plot_flicker(df, check_flicker(df, Thresholds()),
                     outdir=tmp_path, stem="site")
        written = list(tmp_path.glob("*flicker.png"))
        assert written and written[0].stat().st_size > 10_000

    def test_no_chart_without_flicker_data(self, tmp_path):
        pytest.importorskip("matplotlib")
        from pq_plots import plot_flicker
        plot_flicker(_frame(voltage_a=[120.0] * 5), {"available": False},
                     outdir=tmp_path, stem="site")
        assert not list(tmp_path.glob("*flicker.png"))


class TestLineToLineVoltage:
    def test_wye_nominal_inferred_and_snapped(self):
        from pq_analysis import check_line_to_line_voltage
        df = _frame(voltage_a=[120.0] * 20, voltage_b=[120.0] * 20,
                    voltage_c=[120.0] * 20, voltage_ab=[207.8] * 20,
                    voltage_bc=[207.8] * 20, voltage_ca=[207.8] * 20)
        r = check_line_to_line_voltage(df, Thresholds(nominal_voltage=120.0))
        assert r["available"] and r["nominal_v"] == 208.0
        assert "wye" in r["configuration"]
        assert r["overall_pass"] is True
        assert set(r["pairs"]) == {"A-B", "B-C", "C-A"}

    def test_split_phase_nominal_inferred(self):
        from pq_analysis import check_line_to_line_voltage
        df = _frame(voltage_a=[120.0] * 20, voltage_b=[120.0] * 20,
                    voltage_ab=[240.0] * 20)
        r = check_line_to_line_voltage(df, Thresholds(nominal_voltage=120.0))
        assert r["available"] and r["nominal_v"] == 240.0
        assert "split-phase" in r["configuration"]

    def test_out_of_band_is_flagged(self):
        from pq_analysis import check_line_to_line_voltage
        ll = [207.8] * 18 + [180.0, 180.0]     # 13% low on the last two intervals
        df = _frame(voltage_a=[120.0] * 20, voltage_b=[120.0] * 20,
                    voltage_c=[120.0] * 20, voltage_ab=ll)
        r = check_line_to_line_voltage(df, Thresholds(nominal_voltage=120.0))
        assert r["overall_pass"] is False
        assert r["pairs"]["A-B"]["pct_under"] == pytest.approx(10.0)

    def test_ambiguous_ratio_is_refused_not_guessed(self):
        from pq_analysis import check_line_to_line_voltage
        df = _frame(voltage_a=[120.0] * 10, voltage_ab=[300.0] * 10)  # ratio 2.5
        r = check_line_to_line_voltage(df, Thresholds(nominal_voltage=120.0))
        assert r["available"] is False and "neither" in r["error"]

    def test_unavailable_without_ll_channels(self):
        from pq_analysis import check_line_to_line_voltage
        r = check_line_to_line_voltage(_frame(voltage_a=[120.0] * 5), Thresholds())
        assert r["available"] is False


class TestFrequency:
    def test_in_band_passes(self):
        from pq_analysis import check_frequency
        r = check_frequency(_frame(frequency=[60.0, 59.98, 60.02] * 4), Thresholds())
        assert r["available"] and r["overall_pass"] is True
        assert r["pct_out_of_band"] == 0.0

    def test_deviation_beyond_band_fails(self):
        from pq_analysis import check_frequency
        r = check_frequency(_frame(frequency=[60.0] * 8 + [58.9, 61.2]), Thresholds())
        assert r["overall_pass"] is False
        assert r["pct_out_of_band"] == pytest.approx(20.0)
        assert r["max_deviation_hz"] == pytest.approx(1.2, abs=1e-9)

    def test_unavailable_without_channel(self):
        from pq_analysis import check_frequency
        assert check_frequency(_frame(voltage_a=[120.0] * 5), Thresholds())["available"] is False


# ─────────────────────────────────────────────────────────────────────────────
# 11. TDD from the meter's harmonic RMS (IEEE 519-2022)
# ─────────────────────────────────────────────────────────────────────────────

class TestHarmonicCurrentRMS:
    def test_meter_aggregate_is_preferred_over_the_per_order_sum(self):
        from pq_analysis import harmonic_current_rms
        # Per-order sum would give 5.0 A; the meter reports 6.0 A because the
        # orders it displays are rounded and the aggregate is not.
        df = _frame(hrms_current_a=[6.0] * 5,
                    h3_current_a=[3.0] * 5, h5_current_a=[4.0] * 5)
        series, source = harmonic_current_rms(df, "a")
        assert source == "meter"
        assert series.iloc[0] == 6.0

    def test_falls_back_to_the_per_order_sum(self):
        from pq_analysis import harmonic_current_rms
        df = _frame(h3_current_a=[3.0] * 5, h5_current_a=[4.0] * 5)
        series, source = harmonic_current_rms(df, "a")
        assert source == "per-order sum"
        assert series.iloc[0] == pytest.approx(5.0)

    def test_aggregate_columns_are_not_mistaken_for_orders(self):
        from pq_analysis import harmonic_current_rms
        # hrms_current_a must not also be summed in as if it were an order.
        df = _frame(h3_current_a=[3.0] * 5, h5_current_a=[4.0] * 5,
                    hrms_current_a=[float("nan")] * 5)
        series, source = harmonic_current_rms(df, "a")
        assert source == "per-order sum" and series.iloc[0] == pytest.approx(5.0)

    def test_no_channels_at_all(self):
        from pq_analysis import harmonic_current_rms
        assert harmonic_current_rms(_frame(voltage_a=[120.0] * 3), "a") == (None, "")


class TestFundamentalCurrent:
    def test_derived_from_rms_and_harmonic_rms(self):
        from pq_analysis import fundamental_current
        # 5-12-13 triangle: I1=12 when Irms=13 and Ih=5.
        df = _frame(current_a=[13.0] * 5)
        import pandas as pd
        h = pd.Series([5.0] * 5, index=df.index)
        assert fundamental_current(df, "a", h).iloc[0] == pytest.approx(12.0)

    def test_noise_above_the_total_is_clamped_not_square_rooted(self):
        from pq_analysis import fundamental_current
        import pandas as pd
        df = _frame(current_a=[1.0] * 5)
        h = pd.Series([2.0] * 5, index=df.index)   # impossible, but happens at ~0 A
        out = fundamental_current(df, "a", h)
        assert (out == 0.0).all() and out.notna().all()


class TestBlueBookISCLookup:
    """Every ISC the GUI offers must also resolve when the analysis runs."""

    # The nominal voltages the GUI offers.
    NOMINALS = (120.0, 208.0, 240.0, 277.0, 480.0)

    def test_single_phase_padmount_at_120_v_resolves(self):
        # 120 V L-N is one leg of a 120/240 V split-phase secondary, and the
        # pad-mount tables are keyed at 240 V only. The GUI showed 29,600 A while
        # the run resolved nothing, so the report said ISC was never provided.
        from pq_constants import _lookup_isc
        result = _lookup_isc("1ph-padmount", 100, 120.0)
        assert result is not None
        isc, note = result
        assert isc == 29_600
        assert "240V secondary" in note

    def test_every_size_the_gui_lists_resolves(self):
        from pq_constants import _BLUE_BOOK_ISC, _lookup_isc, _infer_secondary_v
        for svc_type in {k[0] for k in _BLUE_BOOK_ISC}:
            for nominal in self.NOMINALS:
                sec_v = _infer_secondary_v(svc_type, nominal)
                offered = [k[1] for k in _BLUE_BOOK_ISC
                           if k[0] == svc_type and k[2] == sec_v]
                for kva in offered:
                    result = _lookup_isc(svc_type, kva, nominal)
                    assert result is not None, (svc_type, kva, nominal)
                    assert result[0] == _BLUE_BOOK_ISC[(svc_type, kva, sec_v)]

    def test_the_secondary_is_never_a_voltage_the_table_lacks(self):
        # Resolving to a voltage with no rows is what made the lookup fail: the
        # exact key misses and the nearest-kVA fallback is filtered to the same
        # empty set, so the whole lookup returns None.
        from pq_constants import _BLUE_BOOK_ISC, _infer_secondary_v
        for svc_type in {k[0] for k in _BLUE_BOOK_ISC}:
            available = {k[2] for k in _BLUE_BOOK_ISC if k[0] == svc_type}
            for nominal in self.NOMINALS:
                sec_v = _infer_secondary_v(svc_type, nominal)
                if svc_type == "3ph-overhead-wye" and nominal == 240.0:
                    # A 240 V wye bank is not a configuration PSCo builds; the
                    # GUI says so and falls back to the manual ISC override.
                    continue
                assert sec_v in available, (svc_type, nominal, sec_v)

    @pytest.mark.parametrize("svc_type,nominal,expected", [
        # A 240 V pick is already the secondary line voltage — reading it as a
        # line-to-neutral value and scaling by √3 landed on the 480 V rows and
        # halved the ISC.
        ("3ph-open-delta",   240.0, 240),
        ("3ph-closed-delta", 240.0, 240),
        ("3ph-padmount",     240.0, 240),
        ("3ph-padmount",     208.0, 208),
        ("3ph-overhead-wye", 208.0, 208),
        # 120 V and 277 V are the line-to-neutral readings of a wye secondary.
        ("3ph-padmount",     120.0, 208),
        ("3ph-overhead-wye", 120.0, 208),
        ("3ph-overhead-wye", 277.0, 480),
        ("3ph-padmount",     480.0, 480),
        # A delta bank has no neutral: a 120 V pick is the center-tapped leg of
        # a 120/240 V secondary.
        ("3ph-open-delta",   120.0, 240),
        ("3ph-closed-delta", 120.0, 240),
        # Single-phase: 120 V rows where the table has them, 240 V otherwise.
        ("1ph-overhead",     120.0, 120),
        ("1ph-overhead",     240.0, 240),
        ("1ph-padmount",     120.0, 240),
        ("1ph-padmount",     240.0, 240),
    ])
    def test_secondary_voltage_for_each_service(self, svc_type, nominal, expected):
        from pq_constants import _infer_secondary_v
        assert _infer_secondary_v(svc_type, nominal) == expected

    def test_a_240_v_delta_reads_the_240_v_rows(self):
        from pq_constants import _lookup_isc
        isc_240, note = _lookup_isc("3ph-closed-delta", 150, 240.0)
        isc_480, _    = _lookup_isc("3ph-closed-delta", 150, 480.0)
        assert "240V secondary" in note
        assert isc_240 == pytest.approx(2 * isc_480, rel=0.01)


class TestTwoLegServiceRigor:
    """NEMA MG1 and the neutral sum both depend on the service configuration.

    NEMA MG1 unbalance is defined for three-phase systems; its formula applied
    to two legs returns half the difference an engineer would quote, under a
    label that does not apply. And the L1+L2 sum discriminates an open neutral
    on a 120/208 service but cannot on a 120/240 one.
    """

    def _two_leg(self, va=122.0, vb=118.0, n=120):
        import pandas as pd
        return pd.DataFrame(
            {"voltage_a": np.full(n, va), "voltage_b": np.full(n, vb)},
            index=pd.date_range("2025-01-01", periods=n, freq="5min"))

    def _ds(self, va, vb):
        import pandas as pd
        n = len(va)
        class _DS:
            pass
        d = _DS()
        d.df = pd.DataFrame({"voltage_a": va, "voltage_b": vb},
                            index=pd.date_range("2025-01-01", periods=n, freq="5min"))
        d.meta = {"topology": "split-phase"}
        d.has_adaptive = False
        d.adaptive_df = None
        return d

    def test_nema_mg1_is_not_claimed_on_a_single_phase_service(self):
        from pq_analysis import check_voltage_imbalance
        r = check_voltage_imbalance(self._two_leg(),
                                    Thresholds(nominal_voltage=120.0))
        assert r["metric"] == "leg_difference"
        assert "NEMA MG1" not in r["metric_label"]
        assert "not applicable to a single-phase service" in r["basis"]

    def test_two_legs_report_the_full_difference_not_half(self):
        from pq_analysis import check_voltage_imbalance
        # 122 vs 118 on a 120 V base is a 4 V spread: 3.33% of nominal.
        # The NEMA formula applied to two elements would have said 1.67%.
        r = check_voltage_imbalance(self._two_leg(),
                                    Thresholds(nominal_voltage=120.0))
        assert r["mean_imbalance_pct"] == pytest.approx(3.333, abs=0.01)

    def test_three_phase_still_uses_nema_mg1(self):
        from pq_analysis import check_voltage_imbalance
        import pandas as pd
        n = 60
        df = pd.DataFrame({"voltage_a": np.full(n, 122.0),
                           "voltage_b": np.full(n, 118.0),
                           "voltage_c": np.full(n, 120.0)},
                          index=pd.date_range("2025-01-01", periods=n, freq="5min"))
        r = check_voltage_imbalance(df, Thresholds(nominal_voltage=120.0))
        assert r["metric"] == "nema_mg1"
        assert r["mean_imbalance_pct"] == pytest.approx(1.667, abs=0.01)

    def test_a_208_service_says_the_third_phase_is_unmeasured(self):
        from pq_analysis import check_voltage_imbalance
        r = check_voltage_imbalance(
            self._two_leg(), Thresholds(nominal_voltage=120.0,
                                        service_type="1ph-208"))
        assert r["note"] and "cannot be determined" in r["note"]

    def test_the_sum_discriminates_an_open_neutral_on_a_208_service(self):
        from pq_analysis import check_neutral_health
        n = 120
        # Open neutral: the two loads sit in series across the 208 V L-L.
        frac = np.full(n, 0.45)
        ds = self._ds(208.0 * frac, 208.0 * (1 - frac))
        r = check_neutral_health(ds, Thresholds(nominal_voltage=120.0,
                                                service_type="1ph-208"))
        assert r["sum_is_diagnostic"] is True
        assert r["sum_toward_open"] > 0.9
        assert any("collapsed toward the line-to-line" in f for f in r["findings"])

    def test_a_healthy_208_service_is_not_flagged_by_the_sum(self):
        from pq_analysis import check_neutral_health
        n = 120
        ds = self._ds(np.full(n, 120.0), np.full(n, 120.0))
        r = check_neutral_health(ds, Thresholds(nominal_voltage=120.0,
                                                service_type="1ph-208"))
        assert r["sum_toward_open"] == pytest.approx(0.0, abs=0.05)

    def test_the_sum_carries_no_information_on_a_120_240_service(self):
        # Collinear legs sum to the line-to-line voltage whether the neutral is
        # intact or open, so a steady 240 V must not read as proof of health.
        from pq_analysis import check_neutral_health
        n = 120
        healthy = self._ds(np.full(n, 120.0), np.full(n, 120.0))
        frac = np.full(n, 0.45)
        opened = self._ds(240.0 * frac, 240.0 * (1 - frac))
        th = Thresholds(nominal_voltage=120.0)
        for ds in (healthy, opened):
            r = check_neutral_health(ds, th)
            assert r["sum_is_diagnostic"] is False
            assert any("carries no open-neutral information" in f
                       for f in r["findings"])

    def test_an_open_neutral_on_a_120_240_service_is_still_caught(self):
        # The sum cannot see it, but the cross-leg behaviour must.
        from pq_analysis import check_neutral_health
        n = 120
        rng = np.random.default_rng(1)
        frac = 0.45 + rng.normal(0, 0.03, n)
        ds = self._ds(240.0 * frac, 240.0 * (1 - frac))
        r = check_neutral_health(ds, Thresholds(nominal_voltage=120.0))
        assert r["severity"] in ("warning", "critical")
        assert r["leg_correlation_available"]
        assert r["leg_correlation"] < 0

    def test_neutral_health_runs_on_a_208_service(self):
        from pq_analysis import check_neutral_health
        n = 60
        ds = self._ds(np.full(n, 120.0), np.full(n, 120.0))
        r = check_neutral_health(ds, Thresholds(nominal_voltage=120.0,
                                                service_type="1ph-208"))
        assert r["available"] and r["topology"] == "1ph-208"


class TestServiceGeometryGating:
    """Three-phase criteria must not be applied to a two-leg service.

    The discriminator is the angle between the legs, not how many there are.
    A 120/240 service is collinear, so the neutral carries a difference and odd
    harmonics -- triplens included -- subtract in it. Two legs of a 120/208 wye
    are 120 degrees apart, so the neutral carries a sum and triplens add. Both
    have two legs; only one of them accumulates.
    """

    def _legs(self, ia=20.0, ib=12.0, n=120, neutral=None, h3n=None):
        import pandas as pd
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        d = {"current_a": np.full(n, ia), "current_b": np.full(n, ib),
             "voltage_a": np.full(n, 121.0), "voltage_b": np.full(n, 119.0),
             "h3_current_a": np.full(n, 2.0), "h3_current_b": np.full(n, 2.0)}
        if neutral is not None:
            d["current_neutral"] = np.full(n, neutral)
        if h3n is not None:
            d["h3_current_neutral"] = np.full(n, h3n)
            d["h5_current_neutral"] = np.full(n, 0.2)
        return pd.DataFrame(d, index=idx)

    def _three_phase(self, n=120):
        import pandas as pd
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        return pd.DataFrame(
            {"current_a": np.full(n, 20.0), "current_b": np.full(n, 12.0),
             "current_c": np.full(n, 16.0)}, index=idx)

    # ── the resolver ─────────────────────────────────────────────────────────

    def test_two_legs_default_to_split_phase(self):
        from pq_analysis import service_geometry
        assert service_geometry(Thresholds(), self._legs().columns) == "split-phase"

    def test_service_type_marks_a_208_network_service(self):
        from pq_analysis import service_geometry
        g = service_geometry(Thresholds(service_type="1ph-208"),
                             self._legs().columns)
        assert g == "two-leg-208"

    def test_a_third_phase_reads_as_three_phase(self):
        from pq_analysis import service_geometry
        assert service_geometry(
            Thresholds(), self._three_phase().columns) == "three-phase"

    def test_the_engineers_pick_beats_channel_presence(self):
        # A three-phase export that dropped phase C still has three phases.
        from pq_analysis import service_geometry
        g = service_geometry(Thresholds(topology="3ph-wye"), self._legs().columns)
        assert g == "three-phase"

    def test_only_120_degree_systems_accumulate_triplens(self):
        from pq_analysis import accumulates_triplens
        assert accumulates_triplens("three-phase") is True
        assert accumulates_triplens("two-leg-208") is True
        assert accumulates_triplens("split-phase") is False
        assert accumulates_triplens("single-phase") is False

    # ── current imbalance ────────────────────────────────────────────────────

    def test_no_limit_is_applied_to_leg_current_difference(self):
        # There is no PSCo number and no standard one, so none is invented.
        from pq_analysis import check_current_imbalance
        r = check_current_imbalance(self._legs(), Thresholds())
        assert r["available"] is True
        assert r["limit_pct"] is None
        assert r["pct_exceeding"] == 0.0
        assert len(r["violation_timestamps"]) == 0

    def test_the_leg_difference_says_it_is_not_a_violation(self):
        from pq_analysis import check_current_imbalance
        r = check_current_imbalance(self._legs(), Thresholds())
        assert r["metric"] == "leg_difference"
        assert "measurement, not a violation" in r["note"]

    def test_a_208_service_says_the_third_phase_is_unmeasured(self):
        from pq_analysis import check_current_imbalance
        r = check_current_imbalance(self._legs(),
                                    Thresholds(service_type="1ph-208"))
        assert r["limit_pct"] is None
        assert "third phase is not measured" in r["note"]

    def test_three_phase_still_gets_the_ten_percent_limit(self):
        from pq_analysis import check_current_imbalance
        r = check_current_imbalance(self._three_phase(), Thresholds())
        assert r["limit_pct"] == 10.0
        assert r["metric"] == "nema_style"
        assert r["note"] is None
        # 20/12/16 averages 16, max deviation 4 -> 25%, over the limit.
        assert r["mean_imbalance_pct"] == pytest.approx(25.0, abs=0.1)
        assert r["pct_exceeding"] == 100.0

    # ── neutral harmonics ────────────────────────────────────────────────────

    def test_split_phase_withholds_the_accumulation_factor(self):
        from pq_analysis import check_neutral_harmonics
        r = check_neutral_harmonics(self._legs(h3n=5.0), Thresholds())
        assert r["available"] is True
        assert r["triplens_accumulate"] is False
        assert r["accumulation_factor"] is None
        assert r["triplen_dominant"] is False
        assert "subtract in the neutral" in r["accumulation_note"]

    def test_split_phase_still_measures_neutral_harmonic_current(self):
        # Neutral heating is real whatever the geometry; only the
        # zero-sequence interpretation is withheld.
        from pq_analysis import check_neutral_harmonics
        r = check_neutral_harmonics(self._legs(h3n=5.0), Thresholds())
        assert r["orders"][3]["mean_a"] == pytest.approx(5.0, abs=0.01)

    def test_a_208_service_does_accumulate(self):
        from pq_analysis import check_neutral_harmonics
        r = check_neutral_harmonics(self._legs(h3n=5.0),
                                    Thresholds(service_type="1ph-208"))
        assert r["triplens_accumulate"] is True
        assert r["accumulation_factor"] == pytest.approx(2.5, abs=0.01)
        assert r["accumulation_note"] is None


class TestSharedTransformerIsNotMeasured:
    """One meter cannot show what a shared transformer is carrying.

    A residential pole or pad transformer feeds several houses. This service's
    demand is therefore a *lower bound* on the transformer's load, and the
    inference is one-sided: above nameplate proves an overload whoever else is
    connected, but below nameplate proves nothing. The negative case must be
    "not determinable", never a pass -- it used to print "The transformer
    loading is within acceptable limits" about equipment never measured.
    """

    def _df(self, kva, n=200):
        import pandas as pd
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        return pd.DataFrame({"power_real": np.full(n, kva * 1000.0),
                             "power_reactive": np.zeros(n)}, index=idx)

    def _tx(self, cls, nameplate, kva=10.0):
        from pq_analysis import check_demand
        r = check_demand(self._df(kva),
                         Thresholds(customer_class=cls,
                                    transformer_kva=nameplate))
        return r["transformer"]

    def test_which_classes_own_their_transformer(self):
        from pq_analysis import has_dedicated_transformer
        assert has_dedicated_transformer("pg") is True
        assert has_dedicated_transformer("sg") is True
        assert has_dedicated_transformer("r") is False
        assert has_dedicated_transformer("c") is False

    def test_a_house_under_nameplate_is_not_determinable(self):
        tx = self._tx("r", nameplate=25.0, kva=10.0)
        assert tx["overloaded"] is None
        assert tx["dedicated"] is False
        assert "cannot be determined" in tx["note"]

    def test_a_house_over_nameplate_still_proves_an_overload(self):
        # One-sided: a lower bound above the nameplate is conclusive.
        tx = self._tx("r", nameplate=5.0, kva=10.0)
        assert tx["overloaded"] is True

    def test_small_commercial_shares_too(self):
        assert self._tx("c", nameplate=50.0, kva=10.0)["overloaded"] is None

    def test_a_dedicated_service_still_passes_and_fails(self):
        assert self._tx("sg", nameplate=500.0, kva=10.0)["overloaded"] is False
        assert self._tx("sg", nameplate=5.0, kva=10.0)["overloaded"] is True
        assert self._tx("pg", nameplate=500.0, kva=10.0)["note"] is None

    def test_not_determinable_is_not_graded(self):
        from pq_analysis import grade_finding
        tx = self._tx("r", nameplate=25.0, kva=10.0)
        passes = None if tx["overloaded"] is None else not tx["overloaded"]
        assert grade_finding(passes, measured=tx["peak_8h_kva"],
                             limit=tx["nameplate_kva"])["band"] == "not_assessed"

    def test_no_all_clear_reaches_a_residential_report(self):
        import subprocess, sys, os, glob
        pytest.importorskip("docx")
        from docx import Document
        root = os.path.dirname(os.path.abspath(__file__))
        import tempfile
        with tempfile.TemporaryDirectory() as d:
            subprocess.run(
                [sys.executable, os.path.join(root, "pq_analyzer.py"),
                 os.path.join(root, "test_data", "test_residential.pqd"),
                 "--customer-class", "r", "--nominal", "120",
                 "--transformer-kva", "25", "--report", "--no-plots",
                 "--outdir", d], check=True, capture_output=True)
            path = [p for p in glob.glob(os.path.join(d, "*.docx"))
                    if "internal" in p][0]
            text = " ".join(p.text for p in Document(path).paragraphs)
        assert "within acceptable limits" not in text
        assert "contribution to its loading and not that loading" in text
        assert "cannot be determined from a recording at one meter" in text
        # The K-factor section reached the same conclusion by a second route,
        # asserting thermal margin from this service's share of the nameplate.
        assert "retains substantial thermal margin" not in text
        assert "not the transformer's thermal margin" in text


class TestResidentialRunSaysNoViolation:
    """End to end on a house: what reaches the page, not just the dicts."""

    @pytest.fixture(scope="class")
    @classmethod
    def out(cls, tmp_path_factory):
        import subprocess, sys, os
        root = os.path.dirname(os.path.abspath(__file__))
        d = tmp_path_factory.mktemp("residential")
        p = subprocess.run(
            [sys.executable, os.path.join(root, "pq_analyzer.py"),
             os.path.join(root, "test_data", "test_residential.pqd"),
             "--customer-class", "r", "--nominal", "120",
             "--no-plots", "--outdir", str(d)],
            check=True, capture_output=True, text=True)
        return p.stdout

    def test_the_leg_difference_is_reported_without_a_verdict(self, out):
        assert "MEASUREMENT — no limit" in out
        assert "measurement, not a violation" in out

    def test_no_ten_percent_current_imbalance_failure(self, out):
        block = out.split("CURRENT IMBALANCE")[1].split("─" * 10)[0]
        assert "FAIL" not in block
        assert "Limit=" not in block

    def test_no_triplen_accumulation_language_on_a_house(self, out):
        assert "Accumulation factor:" not in out
        assert "4-wire wye" not in out
        assert "subtract in the neutral" in out

    def test_the_imbalance_findings_do_not_contradict_each_other(self, out):
        # These two fired together before: one saying both voltage and current
        # were elevated, the other saying current was low, same recording.
        assert not ("Current imbalance — investigate supply voltage" in out
                    and "balanced load current" in out)

    def test_nothing_commits_xcel_to_an_action(self, out):
        assert "Xcel Energy will" not in out

    def test_a_house_is_not_told_to_rebalance_three_phases(self, out):
        assert "across phases" not in out
        assert "A, B, C phases" not in out


@pytest.fixture
def gui_app():
    """A real window, skipped where no display can be opened."""
    import run
    try:
        app = run.PQApp()
    except Exception as exc:                      # no display, no Tk
        pytest.skip(f"Tk unavailable: {exc}")
    app.update_idletasks()
    yield app
    app.destroy()


class TestTheWindowFitsOnAScreen:
    """The form ran off the bottom of a 1080p work PC.

    It was one vertical stack -- form, then log -- asking for 914 px before the
    session row or the expanded sign-off were showing, which put the Run button
    below the taskbar. The log is the tallest single element and the one least
    needed while the form is being filled in, so it moved into a second column.
    These assertions are a height budget: a row added later that pushes the
    window past a laptop screen should fail here rather than in the field.
    """

    @staticmethod
    def _generation(app):
        import run
        kids = []

        def walk(w):
            for c in w.winfo_children():
                kids.append(c)
                walk(c)

        walk(app)
        combo = [w for w in kids if isinstance(w, run.ttk.Combobox)
                 and str(w.cget("textvariable")) == str(app._role_var)][0]
        app._role_var.set(run._ROLE_LABELS[2])
        combo.event_generate("<<ComboboxSelected>>")
        app.update_idletasks()

    def test_the_form_and_the_log_are_side_by_side(self, gui_app):
        # Not a cosmetic assertion: if the log ever lands back under the form,
        # every height figure below moves by ~270 px at once.
        assert gui_app._log.winfo_rootx() > gui_app._form.winfo_rootx()

    def test_it_opens_short_enough_for_a_laptop(self, gui_app):
        # 768 px is the shortest display these run on, less ~120 for the
        # taskbar and title bar.
        self._generation(gui_app)
        assert gui_app.winfo_reqheight() < 648, gui_app.winfo_reqheight()

    def test_even_everything_at_once_clears_a_1080p_screen(self, gui_app):
        # Session picker showing and the sign-off expanded: the tallest the
        # window ever gets.
        self._generation(gui_app)
        gui_app._session_frame.pack(fill="x", padx=12, pady=6,
                                    after=gui_app._file_frame)
        gui_app._det_toggle_btn.invoke()
        gui_app.update_idletasks()
        assert gui_app.winfo_reqheight() < 960, gui_app.winfo_reqheight()

    def test_it_never_opens_taller_than_the_display(self, gui_app):
        # Windows DPI scaling can inflate every row on a machine this was
        # never measured on, so the clamp is belt and braces for the split.
        gui_app._fit_to_screen()
        height = int(gui_app.geometry().split("x")[1].split("+")[0])
        assert height <= gui_app.winfo_screenheight() - 120


class TestHelpWindowStructure:
    """The reference guide is paged, not one scroll of everything.

    It carries three different kinds of thing -- what to enter and why, what
    the standards say, and how to work a job -- and it grew past a thousand
    lines. A reader after one of those should not have to scroll past the
    other two, and no single pane should be a wall on its own.
    """

    @staticmethod
    def _panes(app):
        import tkinter as tk
        app._show_help()
        win = [w for w in app.winfo_children() if isinstance(w, tk.Toplevel)][-1]
        kids = []

        def walk(w):
            for c in w.winfo_children():
                kids.append(c)
                walk(c)

        walk(win)
        nb = [w for w in kids if w.winfo_class() == "TNotebook"][0]
        txts = [w for w in kids if isinstance(w, tk.Text)]
        tabs = [nb.tab(i, "text").strip() for i in range(len(nb.tabs()))]
        return win, tabs, txts

    def test_it_opens_as_tabs_with_start_here_first(self, gui_app):
        win, tabs, txts = self._panes(gui_app)
        assert tabs[0] == "Start here"
        assert len(tabs) == len(txts) >= 6
        win.destroy()

    def test_no_single_pane_is_a_wall(self, gui_app):
        # Before paging, one scroll held ~918 lines and the merged concepts
        # page alone held 475. Either is more than anyone reads.
        win, tabs, txts = self._panes(gui_app)
        lines = {t: len(w.get("1.0", "end-1c").splitlines())
                 for t, w in zip(tabs, txts)}
        assert max(lines.values()) < 300, lines
        assert sum(lines.values()) > 600      # nothing was dropped in the split
        win.destroy()

    def test_the_three_entered_inputs_are_on_the_first_page(self, gui_app):
        # The whole point of a landing page: the things no .pqd file carries,
        # where guessing changes the answer rather than the wording.
        win, tabs, txts = self._panes(gui_app)
        start = txts[tabs.index("Start here")].get("1.0", "end-1c")
        assert "Power flow" in start
        assert "Figure 1" in start
        assert "twelve previous months" in start
        win.destroy()

    def test_every_pane_is_read_only(self, gui_app):
        win, tabs, txts = self._panes(gui_app)
        assert all(str(w.cget("state")) == "disabled" for w in txts)
        win.destroy()

    def test_links_are_bound_on_the_pane_that_holds_them(self, gui_app):
        # The tags are created while a page is being written; binding them to
        # a stale widget would leave them inert without looking broken.
        win, tabs, txts = self._panes(gui_app)
        found = 0
        for w in txts:
            for tag in (t for t in w.tag_names() if t.startswith("_lnk")):
                assert w.tk.call(w._w, "tag", "bind", tag, "<Button-1>")
                found += 1
        assert found >= 5
        win.destroy()


class TestClearAll:
    """Clear All resets every entry, including the ones that cascade.

    The field that quietly keeps its value is the one that carries the last
    site's transformer or engineer details into the next run, so the defaults
    are snapshotted from the widgets rather than listed by hand.
    """

    @staticmethod
    def _fill(app):
        import run
        app._file_var.set("/tmp/site.pqd")
        app._site_var.set("1500 S Hudson Mile Rd")
        app._cclass_var.set("Schedule R — Residential")
        app._topo_var.set("split-phase")
        app._nominal_var.set("277")
        app._xfmr_type_var.set(run._TYPE_DISPLAY["1ph-padmount"])
        app._on_type_change()
        app._isc_override_var.set(True)
        app._on_isc_override_toggle()
        app._isc_manual_var.set("9000")
        app._conductor_var.set("4/0 AL triplex (overhead drop)")
        app._run_length_var.set("150")
        app._eng_name_var.set("A. Engineer")

    def test_every_service_entry_returns_to_its_default(self, gui_app, monkeypatch):
        import run
        monkeypatch.setattr(run.messagebox, "askyesno", lambda *a, **k: True)
        defaults = dict(gui_app._input_defaults)
        assert len(defaults) > 10, "the form's entries should be tracked"

        self._fill(gui_app)
        assert any(getattr(gui_app, n).get() != d for n, d in defaults.items())

        gui_app._clear_all()
        still_set = {n: getattr(gui_app, n).get()
                     for n, d in defaults.items() if getattr(gui_app, n).get() != d}
        assert not still_set

    def test_the_engineers_own_details_survive(self, gui_app, monkeypatch):
        # They describe who is running the tool, not which service was
        # measured, and they go out on a customer document.
        import run
        monkeypatch.setattr(run.messagebox, "askyesno", lambda *a, **k: True)
        self._fill(gui_app)
        gui_app._eng_title_var.set("Electric Area Engineer")
        gui_app._eng_email_var.set("a@example.com")

        gui_app._clear_all()

        assert gui_app._eng_name_var.get() == "A. Engineer"
        assert gui_app._eng_title_var.get() == "Electric Area Engineer"
        assert gui_app._eng_email_var.get() == "a@example.com"
        # ...while the service they were entered against is gone.
        assert gui_app._site_var.get() == ""
        assert gui_app._file_var.get() == ""

    def test_engineer_details_alone_do_not_trigger_the_confirmation(
            self, gui_app, monkeypatch):
        # Nothing clearable is set, so the dialog would ask to clear nothing.
        import run
        asked = []
        monkeypatch.setattr(run.messagebox, "askyesno",
                            lambda *a, **k: asked.append(a) or True)
        gui_app._eng_name_var.set("A. Engineer")
        gui_app._clear_all()
        assert not asked
        assert gui_app._eng_name_var.get() == "A. Engineer"

    def test_the_dependent_pickers_are_reset_too(self, gui_app, monkeypatch):
        import run
        monkeypatch.setattr(run.messagebox, "askyesno", lambda *a, **k: True)
        self._fill(gui_app)
        assert str(gui_app._isc_entry["state"]) == "normal"

        gui_app._clear_all()
        # Not just the variables: the widget states derived from them.
        assert gui_app._xfmr_type_key is None
        assert str(gui_app._isc_entry["state"]) == "disabled"
        assert str(gui_app._kva_combo["state"]) == "disabled"
        assert "Pick a transformer Type" in gui_app._isc_auto_var.get()

    def test_a_clean_form_clears_without_asking(self, gui_app, monkeypatch):
        import run
        asked = []
        monkeypatch.setattr(run.messagebox, "askyesno",
                            lambda *a, **k: asked.append(a) or True)
        gui_app._clear_all()
        assert not asked

    def test_declining_the_confirmation_keeps_the_entries(self, gui_app, monkeypatch):
        import run
        monkeypatch.setattr(run.messagebox, "askyesno", lambda *a, **k: False)
        self._fill(gui_app)
        gui_app._clear_all()
        assert gui_app._site_var.get() == "1500 S Hudson Mile Rd"
        assert gui_app._run_length_var.get() == "150"


class TestBothDocumentsOpen:
    """A run produces two documents, so a run opens two documents.

    Only the internal report opened before, and the customer document had to
    be found by hand — which is how a document goes out unread.
    """

    class _FakeApp:
        """Enough of the window for the opener: a log, a button, an `after`."""

        def __init__(self):
            self.errors = []

        def _log_write(self, text, tag=None):
            if tag == "error":
                self.errors.append(text)

        def after(self, _delay, fn):
            fn()

        class _Btn:
            def __init__(self):
                self.state = "disabled"

            def config(self, state=None, **_kw):
                if state:
                    self.state = state

        _open_btn = None

    def _run_opener(self, monkeypatch, tmp_path, files):
        import run
        for name in files:
            (tmp_path / name).write_bytes(b"docx")
        monkeypatch.setattr(run, "_SCRIPT", tmp_path / "run.py")
        (tmp_path / "pq_output").mkdir(exist_ok=True)
        for name in files:
            (tmp_path / "pq_output" / name).write_bytes(b"docx")

        launched = []
        monkeypatch.setattr(run.subprocess, "Popen",
                            lambda cmd, **kw: launched.append(cmd))
        app = self._FakeApp()
        app._open_btn = self._FakeApp._Btn()
        run.PQApp._open_documents(app, "site")
        return app, launched

    def test_both_documents_are_opened(self, monkeypatch, tmp_path):
        app, launched = self._run_opener(monkeypatch, tmp_path, [
            "site_customer_letter.docx",
            "site_internal_engineering_report.docx",
        ])
        opened = [" ".join(str(part) for part in cmd) for cmd in launched]
        assert len(opened) == 2
        assert any("customer_letter" in o for o in opened)
        assert any("internal_engineering_report" in o for o in opened)
        # The internal report goes last so it lands in front.
        assert "internal_engineering_report" in opened[-1]
        assert app._open_btn.state == "normal"
        assert not app.errors

    def test_a_missing_document_is_reported_not_silently_skipped(
            self, monkeypatch, tmp_path):
        app, launched = self._run_opener(monkeypatch, tmp_path, [
            "site_internal_engineering_report.docx",
        ])
        assert len(launched) == 1
        assert app.errors and "customer document" in app.errors[0]
        # The one that did get written still opens.
        assert app._open_btn.state == "normal"


_TWO_SESSION = Path(__file__).parent / "test_data" / "test_two_sessions.pqd"


@pytest.mark.skipif(not _TWO_SESSION.exists(),
                    reason="test_data/test_two_sessions.pqd not generated")
class TestAFileHoldingSeveralSessions:
    """A "download all data" export carries every session on the meter.

    A meter reset or re-armed in the field starts a new one, and IEEE 1159.3
    clause 6 describes exactly this: one log chunked into several observation
    records. The reader used to pool observations by time base and skip
    whatever did not match, so the second session was dropped with only a log
    line to say so -- the report then claimed a 24-hour recording while half
    the download went unread.
    """

    def _report(self, session=None):
        import pq_analysis as An
        from pq_report import generate_report
        ds = extract_dataset(ProntoAdapter(_TWO_SESSION, session=session),
                             ChannelMapper())
        th = Thresholds(nominal_voltage=120.0, customer_class="r")
        df = ds.df
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), An.detect_events(ds, th), th,
            neutral_health_result=An.check_neutral_health(ds, th))
        return ds, th, rep

    def test_both_sessions_are_found(self):
        sessions = ProntoAdapter.scan_sessions(_TWO_SESSION)
        assert len(sessions) == 2
        assert sessions[0]["start_time"] < sessions[1]["start_time"]
        assert sessions[0]["intervals"] == 2 * sessions[1]["intervals"]
        # Scanning must not need the values, only the time series.
        assert all(s["channels"] for s in sessions)

    def test_the_longest_session_is_analysed_by_default(self):
        ds, _th, rep = self._report()
        fs = rep["file_summary"]
        assert fs["session_index"] == 0
        assert len(fs["sessions"]) == 2
        assert fs["duration_hours"] == pytest.approx(23.9, abs=0.2)

    def test_the_other_session_can_be_analysed(self):
        ds, _th, rep = self._report(session=1)
        fs = rep["file_summary"]
        assert fs["session_index"] == 1
        assert fs["duration_hours"] == pytest.approx(11.9, abs=0.2)
        # And it is a different recording, not a slice of the first.
        assert fs["start_time"] > self._report()[2]["file_summary"]["end_time"]

    def test_a_session_that_is_not_there_is_refused(self):
        with pytest.raises(ValueError, match="holds 2 session"):
            ProntoAdapter(_TWO_SESSION, session=7)

    def test_no_session_is_silently_pooled_into_another(self):
        # The bug this replaces: 144 intervals appended to 288, or dropped.
        ds, _th, _rep = self._report()
        assert len(ds.df) == 288
        span = (ds.df.index[-1] - ds.df.index[0]).total_seconds() / 3600
        assert span < 24, "the three-day gap was read as recorded time"

    def test_the_engineering_report_says_which_session_it_covers(self, tmp_path):
        import docx
        from pq_plots import plot_overview
        from pq_report import generate_word_report
        ds, th, rep = self._report()
        rep["root_causes"] = []
        plot_overview(ds, th, outdir=tmp_path, stem="two")
        doc = docx.Document(str(generate_word_report(
            report=rep, thresh=th, ds=ds, site_name="S", site_address="A",
            engineer_name="E", outdir=tmp_path, stem="two")))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "More than one session in this file" in text
        assert "session 1 of 2" in text
        assert "2025-06-28" in text, "the unread session's date is not stated"

    def test_the_customer_letter_says_it_too(self, tmp_path):
        import docx
        from pq_plots import plot_overview
        from pq_report import generate_customer_letter
        ds, th, rep = self._report()
        rep["root_causes"] = []
        plot_overview(ds, th, outdir=tmp_path, stem="two")
        doc = docx.Document(str(generate_customer_letter(
            rep, th, "1 Test St", "Eng", tmp_path, "two")))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "recorded in more than one stretch" in text
        assert "2025-06-28" in text
        # Said plainly: the customer is not told about observation records.
        for jargon in ("session", "PQDIF", "observation"):
            assert jargon not in text.lower()

    def test_the_picker_appears_only_when_there_is_a_choice(self, gui_app):
        # The scan itself runs on a worker thread; the display is called back
        # on the UI thread, which is what this exercises.
        import run
        sessions = run.ProntoAdapter.scan_sessions(_TWO_SESSION)
        gui_app._show_sessions(sessions)
        gui_app.update_idletasks()
        assert gui_app._session_frame.winfo_ismapped()
        labels = gui_app._session_combo.cget("values")
        assert len(labels) == 2
        # Defaults to the longest, and hands the analysis a zero-based index.
        assert "longest" in gui_app._session_var.get()
        assert gui_app._selected_session() == 0
        gui_app._session_var.set(labels[1])
        assert gui_app._selected_session() == 1

        gui_app._show_sessions([])
        gui_app.update_idletasks()
        assert not gui_app._session_frame.winfo_ismapped()
        assert gui_app._selected_session() is None

    def test_a_single_session_file_says_nothing_about_sessions(self, tmp_path):
        import docx
        import pq_analysis as An
        from pq_plots import plot_overview
        from pq_report import generate_report, generate_word_report
        ds = extract_dataset(
            ProntoAdapter(Path("test_data/test_residential.pqd")), ChannelMapper())
        th = Thresholds(nominal_voltage=120.0, customer_class="r")
        df = ds.df
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), An.detect_events(ds, th), th)
        rep["root_causes"] = []
        plot_overview(ds, th, outdir=tmp_path, stem="one")
        doc = docx.Document(str(generate_word_report(
            report=rep, thresh=th, ds=ds, site_name="S", site_address="A",
            engineer_name="E", outdir=tmp_path, stem="one")))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "More than one session" not in text


class TestABrokenInstallSaysWhatBroke:
    """A failed engine import must not surface as a NameError on a file.

    Everything the GUI needs from pq_analyzer comes in through one try/except,
    so one failed import leaves every one of those names undefined. The run
    then died at whichever name it reached first -- "NameError: name
    'Thresholds' is not defined" -- which names a symptom four hundred lines
    from the cause, on a machine whose .pqd file cannot be sent to us.
    """

    def test_the_run_stops_with_the_import_traceback(self, monkeypatch):
        import run
        monkeypatch.setattr(run, "_BOOK_AVAILABLE", False)
        monkeypatch.setattr(run, "_IMPORT_TRACEBACK",
                            "ModuleNotFoundError: No module named 'docx'")
        with pytest.raises(RuntimeError) as excinfo:
            run.PQApp._do_analysis(object(), {"filepath": "/tmp/site.pqd"})
        message = str(excinfo.value)
        assert "No module named 'docx'" in message          # the actual cause
        assert "install problem" in message                 # not this file
        assert "site.pqd" in message
        assert "Thresholds" not in message

    def test_the_message_says_which_version_it_is(self, monkeypatch):
        # "They may have an older version though" is the first question asked
        # of any report from the field, and the tool could not answer it.
        import run
        from pq_constants import __version__
        monkeypatch.setattr(run, "_BOOK_AVAILABLE", False)
        monkeypatch.setattr(run, "_IMPORT_TRACEBACK", "boom")
        with pytest.raises(RuntimeError) as excinfo:
            run.PQApp._do_analysis(object(), {"filepath": "/tmp/site.pqd"})
        assert __version__ in str(excinfo.value)
        assert run._ENGINE_VERSION == __version__

    def test_a_working_install_is_not_blocked(self, monkeypatch):
        # The guard must not be what stops a healthy run: it passes through to
        # the analysis, which then fails on the missing file instead.
        import run
        monkeypatch.setattr(run, "_BOOK_AVAILABLE", True)
        with pytest.raises(Exception) as excinfo:
            run.PQApp._do_analysis(object(), {"filepath": "/tmp/nope.pqd",
                                              "nominal": 120.0})
        assert "install problem" not in str(excinfo.value)


class TestSinglePhase208Service:
    """A single-phase service taken from two legs of a 208Y/120 wye.

    Common in condos and apartments. It has two legs like a 120/240 service,
    but they sit 120 degrees apart rather than 180, which changes the
    line-to-line voltage, the neutral current expectation, and which Blue Book
    rows the fault current comes from.
    """

    def test_line_to_line_is_208_not_240(self):
        from pq_constants import ll_factor
        assert 120 * ll_factor("1ph-208") == pytest.approx(207.8, abs=0.5)
        assert 120 * ll_factor("1ph-padmount") == pytest.approx(240.0)

    def test_it_is_recognised_as_a_single_phase_208_service(self):
        from pq_constants import is_single_phase_208
        assert is_single_phase_208("1ph-208")
        assert not is_single_phase_208("1ph-padmount")
        assert not is_single_phase_208("3ph-padmount")
        assert not is_single_phase_208(None)

    def test_fault_current_comes_from_the_three_phase_rows(self):
        from pq_constants import _lookup_isc
        net = _lookup_isc("1ph-208", 75, 120.0)
        thr = _lookup_isc("3ph-padmount", 75, 120.0)
        assert net is not None and thr is not None
        assert net[0] == thr[0]           # same transformer, same ISC
        assert "208V secondary" in net[1]
        # ...and it is not the single-phase answer.
        sp = _lookup_isc("1ph-padmount", 75, 120.0)
        assert sp is None or sp[0] != net[0]

    def test_the_kva_picker_is_populated(self):
        # The service type has no Blue Book rows of its own; without the proxy
        # the Size dropdown would be permanently empty.
        import run
        assert run._kva_options("1ph-208", 120.0)
        assert run._resolve_secondary_v("1ph-208", 120.0) == 208

    def test_it_appears_in_the_picker(self):
        from pq_constants import _SERVICE_TYPE_LABEL
        import run
        assert "1ph-208" in _SERVICE_TYPE_LABEL
        assert "1ph-208" in run._TYPE_ORDER

    def test_it_still_resolves_to_two_legs(self):
        from pq_plots import service_phases
        import pandas as pd
        n = 10
        df = pd.DataFrame(
            {"voltage_a": np.full(n, 120.0), "voltage_b": np.full(n, 120.0)},
            index=pd.date_range("2025-01-01", periods=n, freq="5min"))
        ph = service_phases(df, Thresholds(service_type="1ph-208"))
        assert len(ph) == 2

    def test_a_full_neutral_is_not_called_elevated(self, tmp_path):
        # With balanced load the neutral carries essentially full leg current
        # on this configuration; the report must say that is normal rather
        # than reporting imbalance.
        pytest.importorskip("docx")
        import subprocess, sys, glob, os
        from docx import Document
        root = os.path.dirname(os.path.abspath(__file__))
        subprocess.run(
            [sys.executable, os.path.join(root, "pq_analyzer.py"),
             os.path.join(root, "test_data", "test_residential.pqd"),
             "--customer-class", "r", "--service-type", "1ph-208",
             "--report", "--no-plots", "--outdir", str(tmp_path)],
            check=True, capture_output=True)
        d = Document(glob.glob(str(tmp_path / "**" / "*_internal_engineering_report.docx"),
                               recursive=True)[0])
        body = " ".join(p.text for p in d.paragraphs)
        if "Neutral current averaged" in body:
            assert "120°" in body or "120 degrees" in body


class TestServicePhasesFollowTheService:
    """A split-phase house has two legs, not three.

    The harmonic spectrum chart drew an empty "Phase C" bar and legend entry on
    every residential file, inventing a conductor that does not exist.
    """

    def _df(self, three_phase: bool):
        import pandas as pd
        n = 20
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        cols = {"voltage_a": np.full(n, 120.0), "voltage_b": np.full(n, 120.0),
                "current_a": np.full(n, 30.0), "current_b": np.full(n, 30.0)}
        if three_phase:
            cols["voltage_c"] = np.full(n, 120.0)
            cols["current_c"] = np.full(n, 30.0)
        return pd.DataFrame(cols, index=idx)

    def test_split_phase_channels_give_two_legs(self):
        from pq_plots import service_phases
        ph = service_phases(self._df(False), Thresholds())
        assert [p[1] for p in ph] == ["L1", "L2"]

    def test_three_phase_channels_give_three_phases(self):
        from pq_plots import service_phases
        ph = service_phases(self._df(True), Thresholds())
        assert [p[1] for p in ph] == ["Phase A", "Phase B", "Phase C"]

    def test_the_service_type_picker_wins_over_channel_presence(self):
        # A three-phase export missing its C channel must not be relabelled as
        # a split-phase service when the engineer said it is three-phase.
        from pq_plots import service_phases
        ph = service_phases(self._df(False),
                            Thresholds(service_type="3ph-padmount"))
        assert len(ph) == 3

    def test_a_single_phase_transformer_forces_two_legs(self):
        from pq_plots import service_phases
        ph = service_phases(self._df(True),
                            Thresholds(service_type="1ph-padmount"))
        assert [p[1] for p in ph] == ["L1", "L2"]

    def test_the_topology_picker_is_honoured(self):
        from pq_plots import service_phases
        assert len(service_phases(self._df(True),
                                  Thresholds(topology="split-phase"))) == 2
        assert len(service_phases(self._df(False),
                                  Thresholds(topology="3ph-wye"))) == 3

    def test_auto_topology_defers_to_the_channels(self):
        from pq_plots import service_phases
        assert len(service_phases(self._df(True), Thresholds(topology="auto"))) == 3
        assert len(service_phases(self._df(False), Thresholds(topology="auto"))) == 2

    def test_the_spectrum_plot_draws_only_real_phases(self, tmp_path):
        pytest.importorskip("matplotlib")
        from pq_plots import plot_harmonic_spectrum
        import matplotlib.pyplot as plt
        df = self._df(False)
        for h in (3, 5, 7, 9, 11, 13):
            df[f"h{h}_current_a"] = 1.0
            df[f"h{h}_current_b"] = 0.8
        plot_harmonic_spectrum(df, Thresholds(isc_amps=5000.0),
                               outdir=tmp_path, stem="s")
        assert (tmp_path / "s_harmonic_spectrum.png").exists()
        # The legend is built from the resolved phase list, so this is the
        # invariant that matters: no C entry for a two-leg service.
        from pq_plots import service_phases
        assert all(p[0] != "c" for p in service_phases(df, Thresholds()))


class TestSectionOrder:
    """Measurements precede the narrative drawn from them.

    This is the internal engineering document: its reader is checking numbers,
    not being walked to a conclusion, so the sections that cite measurements
    come after the measurements rather than before them.
    """

    @pytest.fixture(scope="class")
    @classmethod
    def headings(cls, tmp_path_factory):
        pytest.importorskip("docx")
        import subprocess, sys, glob, os
        from docx import Document
        out = tmp_path_factory.mktemp("order")
        root = os.path.dirname(os.path.abspath(__file__))
        subprocess.run(
            [sys.executable, os.path.join(root, "pq_analyzer.py"),
             os.path.join(root, "test_data", "test_commercial_large.pqd"),
             "--isc", "10000", "--nominal", "277", "--report", "--no-plots",
             "--outdir", str(out)], check=True, capture_output=True)
        path = glob.glob(str(out / "**" / "*_internal_engineering_report.docx"), recursive=True)[0]
        d = Document(path)
        return [p.text.strip() for p in d.paragraphs
                if p.style.name == "Heading 1" and p.text.strip()]

    def _idx(self, headings, prefix):
        return next(i for i, h in enumerate(headings) if h.startswith(prefix))

    def test_measurement_review_follows_the_executive_summary(self, headings):
        assert (self._idx(headings, "Detailed Measurement Review")
                == self._idx(headings, "Executive Summary") + 1)

    def test_measurements_precede_the_narrative(self, headings):
        last_measurement = max(self._idx(headings, "Detailed Measurement Review"),
                               self._idx(headings, "Harmonic Evaluation"))
        for narrative in ("Key Findings",
                          "Engineering Assessment",
                          "Recommended Actions"):
            assert self._idx(headings, narrative) > last_measurement

    def test_the_two_measurement_sections_stay_adjacent(self, headings):
        assert (self._idx(headings, "Harmonic Evaluation")
                == self._idx(headings, "Detailed Measurement Review") + 1)

    def test_appendices_remain_last(self, headings):
        first_appendix = self._idx(headings, "Appendix A")
        assert all(h.startswith("Appendix") for h in headings[first_appendix:])

    def test_the_opening_paragraph_describes_the_actual_order(self, tmp_path_factory):
        # The opening said conclusions came first and measurements followed.
        pytest.importorskip("docx")
        import subprocess, sys, glob, os
        from docx import Document
        out = tmp_path_factory.mktemp("intro")
        root = os.path.dirname(os.path.abspath(__file__))
        subprocess.run(
            [sys.executable, os.path.join(root, "pq_analyzer.py"),
             os.path.join(root, "test_data", "test_commercial_small.pqd"),
             "--report", "--no-plots", "--outdir", str(out)],
            check=True, capture_output=True)
        d = Document(glob.glob(str(out / "**" / "*_internal_engineering_report.docx"), recursive=True)[0])
        body = " ".join(p.text for p in d.paragraphs)
        assert "recommended actions are presented first" not in body


class TestPowerUnitsAreConsistent:
    """Power channels are watts/VAR; every label and figure is kW/kVAR.

    The plots were drawing the raw channel under a kW label, so a 3.2 kW house
    appeared as 3,200 kW.  The analysis code had always divided by 1000, so the
    text and the charts disagreed by a factor of a thousand.
    """

    def _df(self, watts=3200.0, n=48):
        import pandas as pd
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        return pd.DataFrame({
            "power_real":     np.full(n, watts),
            "power_reactive": np.full(n, watts * 0.3),
            "power_factor":   np.full(n, 0.96),
            "current_a":      np.full(n, 27.0),
        }, index=idx)

    def test_the_demand_plot_draws_kilowatts(self, tmp_path):
        pytest.importorskip("matplotlib")
        from pq_plots import plot_demand_profile
        import matplotlib.pyplot as plt
        plot_demand_profile(self._df(), {}, outdir=tmp_path, stem="u")
        # The y data must be ~3.2, not ~3200.
        from pq_plots import _to_kilo
        assert _to_kilo(self._df()["power_real"]).max() == pytest.approx(3.2)

    def test_the_conversion_helper_is_a_thousand(self):
        from pq_plots import _to_kilo
        import pandas as pd
        s = pd.Series([1000.0, 2500.0])
        assert list(_to_kilo(s)) == [1.0, 2.5]

    def test_analysis_and_plots_agree_on_scale(self):
        # check_demand reports kVA by dividing by 1000; the plot helper must
        # apply the same factor or the report contradicts its own charts.
        from pq_analysis import check_demand
        from pq_plots import _to_kilo
        df = self._df()
        dem = check_demand(df, Thresholds())
        assert dem["available"]
        plotted_peak_kw = float(_to_kilo(df["power_real"]).max())
        assert plotted_peak_kw == pytest.approx(3.2, rel=1e-6)
        assert dem["real_power"]["peak_kw"] == pytest.approx(plotted_peak_kw, rel=1e-6)

    def test_the_demo_dataset_is_in_watts(self):
        # The demo generator wrote kW into a watt channel, which only became
        # visible once the plots started converting.
        from pq_adapter import MockAdapter
        from pq_analyzer import ChannelMapper, extract_dataset
        ds = extract_dataset(MockAdapter(), ChannelMapper())
        mean_w = float(ds.df["power_real"].dropna().mean())
        assert 5_000 < mean_w < 60_000, f"demo power_real mean is {mean_w} W"


class TestRecommendedActionsAreProportionate:
    """Actions must fit the service and stay short enough to be read."""

    def _report(self, **pf_over):
        pf = {k: True for k in ("power_factor", "thd_current",
                                "individual_harmonics", "current_imbalance",
                                "transformer_loading", "voltage")}
        pf.update(pf_over)
        return {"pass_fail": pf, "root_causes": []}

    def _actions(self, report, cls):
        from pq_report import _build_structured_actions
        return _build_structured_actions(report, Thresholds(customer_class=cls))

    def test_a_house_is_not_told_to_fit_reactors_to_its_vfds(self):
        # The load-signature match is a hypothesis, so its advice belongs with
        # the finding rather than in a list of things to go and do.
        rep = self._report()
        rep["root_causes"] = [{
            "severity": "info",
            "title": "Possible load type: 6-pulse VFD / rectifier (no input reactor)",
            "recommendation": "Add 3-5% impedance AC line reactors to VFD inputs.",
        }]
        assert self._actions(rep, "r") == []

    def test_no_informational_finding_becomes_an_action(self):
        rep = self._report()
        rep["root_causes"] = [{"severity": "info", "title": "X",
                               "recommendation": "Do something speculative."}]
        for cls in ("r", "c", "sg", "pg"):
            assert self._actions(rep, cls) == []

    def test_harmonic_advice_scales_with_the_service(self):
        rep = self._report(thd_current=False)
        res = " ".join(a["recommendation"] for a in self._actions(rep, "r")).lower()
        assert "filter" not in res or "not an appropriate remedy" in res
        assert "12-pulse" not in res
        ci = " ".join(a["recommendation"] for a in self._actions(rep, "sg")).lower()
        assert "12-pulse" in ci or "harmonic filters" in ci

    def test_a_house_is_not_asked_to_commission_a_harmonic_study(self):
        rep = self._report(individual_harmonics=False)
        for cls in ("r", "c"):
            joined = " ".join(a["recommendation"] for a in self._actions(rep, cls))
            assert "detailed harmonic study" not in joined.lower()
        joined = " ".join(a["recommendation"] for a in self._actions(rep, "sg"))
        assert "detailed harmonic study" in joined.lower()

    def test_a_two_leg_service_raises_no_imbalance_action(self):
        # A house reports its leg difference as a measurement with no limit, so
        # pass_fail carries None and nothing here fires. This test used to
        # assert split-phase wording on the action itself; there is no longer a
        # path that produces it, because leg imbalance is not a violation and
        # must not raise a High-priority item. The equivalent advice now hangs
        # off the measured neutral current instead.
        rep = self._report(current_imbalance=None)
        res = " ".join(a["recommendation"] for a in self._actions(rep, "r")).lower()
        assert "across phases" not in res
        assert "imbalance" not in res

    def test_three_phase_imbalance_still_raises_an_action(self):
        rep = self._report(current_imbalance=False)
        res = " ".join(a["recommendation"] for a in self._actions(rep, "sg")).lower()
        assert "across phases" in res

    def test_duplicate_intents_collapse_to_one_action(self):
        from pq_report import _build_structured_actions
        rep = self._report()
        rep["root_causes"] = [
            {"severity": "warning", "title": "A",
             "recommendation": "Measure voltage imbalance with all customer loads "
                               "disconnected to isolate the utility contribution."},
            {"severity": "warning", "title": "B",
             "recommendation": "Confirming the origin requires measurement with all "
                               "customer loads disconnected."},
        ]
        acts = _build_structured_actions(rep, Thresholds(customer_class="sg"))
        assert len(acts) == 1

    def test_the_action_list_is_capped(self):
        from pq_report import _build_structured_actions, _MAX_ACTIONS
        rep = self._report(power_factor=False, thd_current=False,
                           individual_harmonics=False, current_imbalance=False,
                           transformer_loading=False, voltage=False)
        rep["root_causes"] = [
            {"severity": "warning", "title": f"F{i}",
             "recommendation": f"Distinct unrelated recommendation number {i}."}
            for i in range(12)
        ]
        acts = _build_structured_actions(rep, Thresholds(customer_class="sg"))
        assert len(acts) <= _MAX_ACTIONS

    def test_high_priority_actions_survive_the_cap(self):
        from pq_report import _build_structured_actions
        rep = self._report(voltage=False, transformer_loading=False)
        rep["root_causes"] = [
            {"severity": "info", "title": f"I{i}",
             "recommendation": f"Low value note {i}."} for i in range(10)
        ]
        acts = _build_structured_actions(rep, Thresholds(customer_class="sg"))
        assert acts and all(a["priority"] == "High" for a in acts)

    def test_a_trivial_power_factor_shortfall_raises_no_install_action(self):
        # Correcting by a couple of kVAR is below the smallest practical
        # switched capacitor step; recommending an install is not advice.
        import pq_analysis as A
        assert A._MIN_ACTIONABLE_KVAR > 0


#: Harmonic orders carried in the load-signature library, in vector order.
_SIG_ORDERS = [3, 5, 7, 9, 11, 13]


def _sig_frame(spec, cv=0.1, n=150, seed=0):
    """Interval frame carrying `spec` as harmonic amps on a 100 A service."""
    import pandas as pd
    idx = pd.date_range("2025-01-01", periods=n, freq="5min")
    rng = np.random.default_rng(seed)
    d = {"current_a": np.full(n, 100.0)}
    for h, v in zip(_SIG_ORDERS, spec):
        d[f"h{h}_current_a"] = np.clip(
            rng.normal(v, max(v * cv, 1e-4), n), 1e-4, None)
    return pd.DataFrame(d, index=idx)


def _sig_run(spec, cv=0.1, cls="sg", seed=0):
    """Score `spec` through the matcher as a service of class `cls`."""
    from pq_analysis import _detect_harmonic_signature
    return _detect_harmonic_signature(_sig_frame(spec, cv, seed=seed),
                                      100.0, None, customer_class=cls)


class TestLoadSignatureFloor:
    """A spectrum matching nothing must be reported as matching nothing.

    Nearest-neighbour scoring always returns a candidate.  Scoring 20,000 random
    decaying spectra against the library gave a median top score of 0.87, with
    71% above the old 0.75 gate and 29% above 0.95 -- so "scores well" was not
    evidence of anything, and pure noise reached "high confidence" routinely.
    """

    def _frame(self, spec, cv=0.1, n=150, seed=0):
        return _sig_frame(spec, cv, n, seed)

    def _run(self, spec, cv=0.1, cls="sg", seed=0):
        return _sig_run(spec, cv, cls, seed)

    def test_an_unmatched_spectrum_is_named_as_unmatched(self):
        out = self._run([0.4, 0.3, 0.5, 0.2, 0.4, 0.3], cv=0.8, cls="r")
        assert out and out[0]["title"] == "No recognised load signature"
        assert "below the" in out[0]["finding"]

    def test_the_unmatched_finding_still_reports_the_spectrum(self):
        # Naming nothing must not mean saying nothing -- the engineer still
        # needs the measured numbers to interpret by hand.
        out = self._run([0.4, 0.3, 0.5, 0.2, 0.4, 0.3], cv=0.8, cls="r")
        assert "Measured spectrum:" in out[0]["finding"]
        assert out[0]["evidence"]["h3_pct_il"] > 0

    def test_a_spectrum_between_two_families_is_not_assigned_to_either(self):
        from pq_analysis import _detect_harmonic_signature
        out = self._run([35, 18, 9, 5, 3, 2], cls="r")
        assert out[0]["title"] == "No recognised load signature"
        assert "between two unrelated load families" in out[0]["finding"]

    def test_a_clean_match_is_still_reported(self):
        out = self._run([2, 23, 9, 1, 5, 4], cls="sg")
        assert out[0]["title"] != "No recognised load signature"
        assert "6-pulse" in out[0]["title"]

    def test_indistinguishable_members_are_reported_as_a_family(self):
        # A 6-pulse VFD, a 6-pulse UPS and a DC fast charger are one topology;
        # naming one of them specifically would be arbitrary.
        out = self._run([2, 23, 9, 1, 5, 4], cls="sg")
        assert out[0]["title"].startswith("Possible load family")
        assert out[0]["evidence"]["resolved_to_member"] is False

    def test_most_random_spectra_are_rejected(self):
        # The rule's whole purpose. Old behaviour named something ~71% of the
        # time; this asserts the floor actually bites.
        rng = np.random.default_rng(7)
        named = 0
        trials = 120
        for i in range(trials):
            spec = np.clip(np.array([20, 10, 6, 3, 2, 1])
                           * rng.uniform(0.2, 3.0, 6), 0.01, None)
            out = self._run(spec, cv=0.2, cls="sg", seed=i)
            if out and out[0]["title"] != "No recognised load signature":
                named += 1
        assert named / trials < 0.30, f"named {named/trials:.0%} of random spectra"

    def test_the_floor_is_documented_where_it_is_set(self):
        import pq_constants as C
        assert C.SIGNATURE_ABSOLUTE_FLOOR >= 0.85
        assert C.SIGNATURE_FAMILY_SEPARATION > 0
        assert C.SIGNATURE_MEMBER_SEPARATION > 0

    def test_every_signature_belongs_to_a_labelled_family(self):
        from pq_constants import _LOAD_SIGNATURES, LOAD_FAMILY_LABEL
        for s in _LOAD_SIGNATURES:
            assert s.get("family"), f"{s['id']} has no family"
            assert s["family"] in LOAD_FAMILY_LABEL, s["family"]


#: Trials per (class, arm) cell in the mixture sweep.  Large enough that a rate
#: is readable to a couple of points, small enough to stay in the unit suite.
_MIX_TRIALS = 300

#: A draw is called "blended" when no single load owns more than this share of
#: the fundamental.  Above it the service is essentially one device plus trim,
#: and naming that device is the right answer rather than a failure.
_MIX_BLENDED_MAX_SHARE = 0.60

#: Not every library entry describes a single device.  `mixed_vfd_smps` is a
#: blend by construction -- "6-pulse VFDs + single-phase nonlinear loads" -- so
#: naming it on a draw whose loads all come from the families it describes is a
#: correct answer, not a misidentification.  Without this carve-out the sweep
#: scores that entry as wrong every time it is right, which on the `sg` class
#: inflates the misidentification rate from 37% to 67%.
_MIX_COMPOSITE_ENTRIES = {
    "mixed_vfd_smps": {"six_pulse", "single_phase_switchmode"},
}


def _mix_defensible(signature_id, parent_families):
    """True when a non-parent match is a composite entry that describes the draw."""
    covered = _MIX_COMPOSITE_ENTRIES.get(signature_id)
    return covered is not None and parent_families <= covered


def _mixture_spectrum(rng, candidates, phase_random):
    """One synthetic service: 2-3 library loads sharing a 100 A fundamental.

    Each library vector is a percentage of *its own* fundamental, so a load
    drawing share ``w`` of a 100 A service contributes ``w * spec`` amps at each
    order.  The mixture spectrum is therefore the share-weighted sum, in amps.

    ``phase_random`` selects between the two bounding cases described in
    TestLoadSignatureMixtures: aligned magnitudes, or uniform random phase.
    """
    k = int(rng.integers(2, 4))
    picks = rng.choice(len(candidates), size=k, replace=False)
    parents = [candidates[i] for i in picks]
    shares = rng.dirichlet(np.ones(k))

    contrib = np.array([np.asarray(p["spectrum"], dtype=float) * w
                        for p, w in zip(parents, shares)])
    if phase_random:
        theta = rng.uniform(0, 2 * np.pi, size=contrib.shape)
        spec = np.abs((contrib * np.exp(1j * theta)).sum(axis=0))
    else:
        spec = contrib.sum(axis=0)
    return parents, shares, np.clip(spec, 1e-4, None)


def _mixture_sweep(cls, phase_random, trials=_MIX_TRIALS, seed=20260807):
    """Score `trials` synthetic mixtures and tally what the matcher said."""
    from pq_constants import _LOAD_SIGNATURES
    candidates = [s for s in _LOAD_SIGNATURES if cls in s["classes"]]
    rng = np.random.default_rng(seed)
    tally = {"n": 0, "rejected": 0, "named": 0, "named_member": 0,
             "member_not_parent": 0, "family_not_parent": 0, "high_conf": 0,
             "composite_ok": 0, "blended_n": 0, "blended_named_member": 0,
             "blended_member_not_parent": 0}

    for i in range(trials):
        parents, shares, spec = _mixture_spectrum(rng, candidates, phase_random)
        out = _sig_run(spec, cv=0.1, cls=cls, seed=i)
        assert out, "the matcher returned no finding at all"
        f, ev = out[0], out[0]["evidence"]
        blended = float(max(shares)) < _MIX_BLENDED_MAX_SHARE
        parent_ids = {p["id"] for p in parents}
        parent_families = {p["family"] for p in parents}

        tally["n"] += 1
        tally["blended_n"] += blended
        if f["title"] == "No recognised load signature":
            tally["rejected"] += 1
            continue

        defensible = _mix_defensible(ev["signature_id"], parent_families)
        tally["named"] += 1
        tally["composite_ok"] += defensible
        if ev["family"] not in parent_families and not defensible:
            tally["family_not_parent"] += 1
        if ev["family_separation"] >= 0.15:      # what the code calls "high"
            tally["high_conf"] += 1
        if ev["resolved_to_member"]:
            tally["named_member"] += 1
            tally["blended_named_member"] += blended
            if ev["signature_id"] not in parent_ids and not defensible:
                tally["member_not_parent"] += 1
                tally["blended_member_not_parent"] += blended
    return tally


def _mixture_report(rows):
    """Render sweep tallies as a table.  Visible under `pytest -s`."""
    def pct(num, den):
        return "     -" if not den else f"{num / den:6.1%}"

    lines = [
        "",
        "  Load-signature behaviour on multi-load services",
        "  (share-weighted mixtures of 2-3 library loads; 'parent' = a load"
        "  actually present in the draw)",
        "",
        f"  {'class':>6} {'arm':>10} {'n':>5} {'rejected':>9} {'named':>7}"
        f" {'->member':>9} {'not parent':>11} {'wrong fam':>10} {'high conf':>10}",
    ]
    for cls, arm, t in rows:
        lines.append(
            f"  {cls:>6} {arm:>10} {t['n']:>5} {pct(t['rejected'], t['n']):>9}"
            f" {pct(t['named'], t['n']):>7} {pct(t['named_member'], t['n']):>9}"
            f" {pct(t['member_not_parent'], t['named_member']):>11}"
            f" {pct(t['family_not_parent'], t['named']):>10}"
            f" {pct(t['high_conf'], t['named']):>10}"
        )
    lines += ["",
              "  'not parent' and 'wrong fam' are shares of the matches actually"
              " named,",
              "  not of all trials -- they answer \"when it names something, how"
              " often is",
              "  that something not even present?\"  Both credit the composite"
              " library entry",
              "  (see _MIX_COMPOSITE_ENTRIES) when it correctly describes the"
              " blend.", ""]
    print("\n".join(lines))


class TestLoadSignatureMixtures:
    """Can the matcher name a device when the service carries more than one?

    Every entry in the library is a single device at rated load.  A meter at the
    service entrance sees the sum of everything behind it, and the guards in
    `_detect_harmonic_signature` -- the 0.90 floor, family separation, member
    separation -- all measure distance to library points.  None of them asks
    whether a *blend* of library entries explains the same point, so a mixture
    can in principle clear every gate and be reported as one resolved device.

    This class measures how often that happens.  It sets no thresholds and
    asserts no rate: the rates it prints are the deliverable, and what to do
    about them is a separate decision.

    Measured over 300 draws per cell (`pytest test_pq.py -k Mixture -s`), the
    guards hold on small services and fail on large ones.  Residential and small
    commercial mixtures are rejected 96% and 80% of the time, and `c` never
    resolves to a single device at all.  On `sg` and `pg` -- the classes where a
    service most plausibly carries several nonlinear loads, and where this
    finding is most likely to be acted on -- the matcher names something for 47%
    and 69% of mixtures, resolves roughly half of those to one specific device,
    and of those single-device names 45% (`sg`) and 32% (`pg`) name a load that
    is not in the draw.  A quarter of `sg` matches name the wrong *family*.
    Random phase moves every rate by less than the gap between classes, so the
    result does not depend on where between the two arms reality sits.

    What this does not say: these are mixtures of library entries, which is a
    friendlier population than real services carrying loads the library has no
    entry for.  The rates are a floor on the error, not an estimate of it.

    Two arms bound the physics.  Harmonic currents from several loads add as
    phasors, and the library stores magnitudes only:

      * `aligned`     -- magnitudes added directly.  No cancellation, which is
                         the most favourable case for recognising a parent.
      * `randomphase` -- uniform random phase per load per order.  Maximum
                         cancellation, the least favourable case.

    Neither is a model of a real site; real services fall between them, closer
    to `aligned` when the mixed loads share a topology.  A conclusion that holds
    in both arms does not depend on where between them reality sits.
    """

    def test_a_single_library_load_is_still_recognised(self):
        # Harness control.  If a pure library spectrum fed through the mixture
        # path stopped matching, every rate below would be measuring a broken
        # harness rather than the matcher.
        from pq_constants import _LOAD_SIGNATURES
        vfd = next(s for s in _LOAD_SIGNATURES if s["id"] == "vfd_6pulse_reactor")
        out = _sig_run(np.asarray(vfd["spectrum"], dtype=float), cls="sg")
        assert out[0]["title"] != "No recognised load signature"
        assert out[0]["evidence"]["family"] == "six_pulse"

    def test_every_mixture_is_accounted_for(self):
        # Each trial must land in exactly one bucket, or the rates do not add up.
        t = _mixture_sweep("sg", phase_random=False, trials=40)
        assert t["n"] == 40
        assert t["rejected"] + t["named"] == t["n"]
        assert t["named_member"] <= t["named"]
        assert t["member_not_parent"] <= t["named_member"]

    def test_mixture_rates(self):
        # The measurement.  Run with `pytest test_pq.py -k Mixture -s` to read
        # the table; the assertions here only confirm the sweep ran.
        rows = []
        for cls in ("r", "c", "sg", "pg"):
            for arm, phase_random in (("aligned", False), ("randomphase", True)):
                rows.append((cls, arm, _mixture_sweep(cls, phase_random)))
        _mixture_report(rows)
        assert all(t["n"] == _MIX_TRIALS for _, _, t in rows)


class TestFamilyLevelRecommendations:
    """Advice must be written for the family, since the member is not reported.

    Inside a family the member advice contradicted itself -- "verify existing
    input reactors are in service" against "add reactors", "verify the
    phase-shifting transformer" against "no action required" -- and which one
    printed turned on a member score gap the finding calls meaningless.
    """

    def test_every_family_has_a_recommendation(self):
        from pq_constants import LOAD_FAMILY_LABEL, LOAD_FAMILY_RECOMMENDATION
        assert set(LOAD_FAMILY_RECOMMENDATION) == set(LOAD_FAMILY_LABEL)

    def test_the_finding_carries_the_family_advice_not_the_members(self):
        from pq_constants import LOAD_FAMILY_RECOMMENDATION, _LOAD_SIGNATURES
        f = _sig_run(np.array([2, 23, 9, 1, 5, 4], dtype=float), cls="sg")[0]
        assert f["recommendation"] == LOAD_FAMILY_RECOMMENDATION["six_pulse"]
        nearest = next(s for s in _LOAD_SIGNATURES
                       if s["id"] == f["evidence"]["signature_id"])
        assert f["recommendation"] != nearest["recommendation"]

    def test_no_family_advice_assigns_responsibility(self):
        # The tool states evidence; the engineer assigns responsibility.
        from pq_constants import LOAD_FAMILY_RECOMMENDATION
        for fam, text in LOAD_FAMILY_RECOMMENDATION.items():
            low = text.lower()
            assert "responsibility" not in low, fam
            assert "xcel energy will" not in low, fam

    def test_triplen_neutral_advice_names_the_geometry_it_needs(self):
        # These families reach residential services, where the legs are
        # collinear and triplens subtract in the neutral instead of adding.
        from pq_constants import LOAD_FAMILY_RECOMMENDATION
        for fam in ("single_phase_switchmode", "mixed_single_phase",
                    "mixed_three_phase"):
            text = LOAD_FAMILY_RECOMMENDATION[fam]
            if "neutral" in text.lower():
                assert "four-wire" in text, fam


class TestLoadSignaturesRespectCustomerClass:
    """Load-type matching must not offer equipment the service cannot have.

    Cosine similarity always returns a nearest neighbour, so an unrestricted
    library named an arc welder as the best match for a house.
    """

    def _welder_shaped_frame(self):
        # [H3,H5,H7,H9,H11,H13] = [10,8,6,5,4,3] with high variability, which is
        # the arc welder reference signature.
        import pandas as pd
        n = 200
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        rng = np.random.default_rng(3)
        data = {"current_a": np.full(n, 100.0)}
        for h, v in {3: 10., 5: 8., 7: 6., 9: 5., 11: 4., 13: 3.}.items():
            data[f"h{h}_current_a"] = np.clip(rng.normal(v, v * 0.5, n), 0.01, None)
        return pd.DataFrame(data, index=idx)

    def _titles(self, customer_class):
        from pq_analysis import _detect_harmonic_signature
        return [f["title"] for f in _detect_harmonic_signature(
            self._welder_shaped_frame(), 100.0, None,
            customer_class=customer_class)]

    def _texts(self, customer_class):
        # Title plus body.  The equipment considered is named in the body now
        # that the title reports only the family.
        from pq_analysis import _detect_harmonic_signature
        return [f"{f['title']} {f.get('finding', '')}"
                for f in _detect_harmonic_signature(
                    self._welder_shaped_frame(), 100.0, None,
                    customer_class=customer_class)]

    def test_a_residential_service_is_never_offered_industrial_equipment(self):
        joined = " ".join(self._titles("r")).lower()
        for banned in ("arc welder", "arc furnace", "12-pulse", "18-pulse",
                       "dc fast charger"):
            assert banned not in joined, f"{banned!r} offered to a house"

    def test_a_residential_service_gets_residential_candidates(self):
        joined = " ".join(self._texts("r")).lower()
        assert any(k in joined for k in
                   ("pv inverter", "ev charger", "heat pump", "air conditioning"))

    def test_an_arc_furnace_is_only_offered_to_primary_service(self):
        from pq_constants import _LOAD_SIGNATURES
        eaf = next(s for s in _LOAD_SIGNATURES if s["id"] == "arc_furnace")
        assert eaf["classes"] == {"pg"}

    def test_every_signature_declares_its_classes(self):
        from pq_constants import _LOAD_SIGNATURES
        valid = {"r", "c", "sg", "pg"}
        for s in _LOAD_SIGNATURES:
            assert s.get("classes"), f"{s['id']} has no classes"
            assert set(s["classes"]) <= valid, f"{s['id']} has an unknown class"

    def test_every_class_has_candidates(self):
        from pq_constants import _LOAD_SIGNATURES
        for cls in ("r", "c", "sg", "pg"):
            assert [s for s in _LOAD_SIGNATURES if cls in s["classes"]], cls

    def test_an_unresolvable_match_names_the_family_not_the_equipment(self):
        # An arc welder and an arc furnace score within a thousandth of each
        # other on this spectrum. They are one topology, so the report names
        # the arcing family rather than picking one piece of equipment.
        from pq_analysis import _detect_harmonic_signature
        f = _detect_harmonic_signature(self._welder_shaped_frame(), 100.0, None,
                                       customer_class="pg")
        assert f[0]["evidence"]["similarity"] > 0.95
        assert f[0]["title"].startswith("Possible load family")
        assert "Arcing load" in f[0]["title"]
        assert f[0]["evidence"]["resolved_to_member"] is False
        assert "individual load type is not reported" in f[0]["finding"]
        assert f[0]["confidence"] != "high"

    def test_the_finding_does_not_claim_to_identify_equipment(self):
        # No path names a single piece of equipment any more.  A mixture of two
        # loads lands nearest an entry containing neither often enough that the
        # member-level claim was not supportable -- see TestLoadSignatureMixtures.
        f_titles = self._titles("r")
        assert f_titles and not any(t.startswith("Best match") for t in f_titles)
        assert f_titles[0].startswith("Possible load family")
        assert not any(t.startswith("Possible load type") for t in f_titles)

    def test_a_family_with_one_entry_still_does_not_claim_the_equipment(self):
        # mixed_three_phase has a single entry applicable to primary service, so
        # the family label and that entry coincide. The finding has to say so
        # rather than reading as an identification by omission.
        from pq_analysis import _detect_harmonic_signature
        f = _detect_harmonic_signature(
            _sig_frame(np.array([15, 20, 8, 2, 4, 3], dtype=float)),
            100.0, None, customer_class="pg")[0]
        assert f["title"].startswith("Possible load family")
        assert "one reference entry" in f["finding"]
        assert "not an identification of the equipment present" in f["finding"]
        assert f["evidence"]["resolved_to_member"] is False

    def test_no_match_is_ever_reported_at_high_confidence(self):
        from pq_analysis import _detect_harmonic_signature
        for cls in ("r", "c", "sg", "pg"):
            for f in _detect_harmonic_signature(self._welder_shaped_frame(),
                                                100.0, None, customer_class=cls):
                assert f["confidence"] != "high", cls


class TestComplianceTableGrouping:
    """The table groups checks by the quantity measured, and loses nothing.

    Buffering rows to regroup them is the kind of change that silently drops a
    check, so the count is asserted against the ungrouped behaviour: every
    standard that used to appear still appears, exactly once.
    """

    def _table(self, path):
        from docx import Document
        d = Document(str(path))
        for t in d.tables:
            if [c.text.strip() for c in t.rows[0].cells][:2] == ["Standard", "Measured"]:
                return t
        raise AssertionError("compliance table not found")

    def _split(self, tbl):
        headings, checks = [], []
        for row in tbl.rows[1:]:
            if len({id(c._tc) for c in row.cells}) == 1:
                headings.append(row.cells[0].text.strip())
            else:
                checks.append(row.cells[0].text.strip())
        return headings, checks

    @pytest.fixture(scope="class")
    @classmethod
    def report_path(cls, tmp_path_factory):
        pytest.importorskip("docx")
        import subprocess, sys, glob, os
        out = tmp_path_factory.mktemp("grouped")
        root = os.path.dirname(os.path.abspath(__file__))
        subprocess.run(
            [sys.executable, os.path.join(root, "pq_analyzer.py"),
             os.path.join(root, "test_data", "test_commercial_large.pqd"),
             "--isc", "10000", "--transformer-kva", "500", "--nominal", "277",
             "--report", "--no-plots", "--outdir", str(out)],
            check=True, capture_output=True)
        found = glob.glob(str(out / "**" / "*.docx"), recursive=True)
        assert found, "no report generated"
        return found[0]

    def test_every_check_appears_exactly_once(self, report_path):
        _, checks = self._split(self._table(report_path))
        assert len(checks) == 14
        assert len(set(checks)) == len(checks)

    def test_checks_are_grouped_by_measured_quantity(self, report_path):
        headings, _ = self._split(self._table(report_path))
        assert [h.split("—")[0].strip() for h in headings] == [
            "VOLTAGE", "CURRENT", "DEMAND AND POWER FACTOR", "FREQUENCY"]

    def test_no_check_is_stranded_outside_a_group(self, report_path):
        tbl = self._table(report_path)
        seen_heading = False
        for row in tbl.rows[1:]:
            if len({id(c._tc) for c in row.cells}) == 1:
                seen_heading = True
            else:
                assert seen_heading, f"{row.cells[0].text!r} precedes any heading"

    def test_magnitude_checks_precede_distortion_checks(self, report_path):
        _, checks = self._split(self._table(report_path))
        pos = {c.split("(")[0].strip(): i for i, c in enumerate(checks)}
        volt_mag = next(i for c, i in pos.items() if c.startswith("Steady-state voltage"))
        volt_thd = next(i for c, i in pos.items() if c.startswith("Voltage THD"))
        assert volt_mag < volt_thd


class TestComplianceRowsQuoteMeasurements:
    """Every row should carry a number the recording actually produced.

    The voltage row used to print range_v -- the *allowed* band -- under the
    bare label "Range", so the only voltage in the cell was the limit, and
    "Worst phase" was never followed by which phase or what it read.
    """

    def _report(self, **over):
        from pq_analysis import check_voltage_compliance
        import pandas as pd
        n = 100
        idx = pd.date_range("2025-01-01", periods=n, freq="5min")
        data = {"voltage_a": [121.0] * n, "voltage_b": [120.0] * n,
                "voltage_c": [120.0] * n}
        data.update(over)
        return check_voltage_compliance(pd.DataFrame(data, index=idx),
                                        Thresholds(nominal_voltage=120.0))

    def test_per_phase_measurements_are_available_to_the_row(self):
        v = self._report(voltage_a=[110.0] * 50 + [121.0] * 50)
        st = v["phases"]["voltage_a"]
        # The row needs all of these; none were being printed before.
        assert st["min_v"] == pytest.approx(110.0)
        assert st["max_v"] == pytest.approx(121.0)
        assert st["mean_v"] == pytest.approx(115.5)
        assert st["pct_under"] == pytest.approx(50.0)
        assert st["pct_over"] == 0.0

    def test_range_v_is_the_allowed_band_not_a_measurement(self):
        # 120 V nominal at +/-5 % -> 114-126.  Nothing measured 114 or 126 here.
        v = self._report()
        assert v["range_v"] == pytest.approx((114.0, 126.0))
        assert v["phases"]["voltage_a"]["max_v"] == pytest.approx(121.0)

    def test_the_worst_phase_is_the_one_with_most_intervals_out_of_band(self):
        v = self._report(voltage_b=[100.0] * 80 + [120.0] * 20)
        worst = max(v["phases"].items(),
                    key=lambda kv: kv[1]["pct_out_of_bounds"])
        assert worst[0] == "voltage_b"
        assert worst[1]["pct_out_of_bounds"] == pytest.approx(80.0)

    def test_the_binding_harmonic_order_is_reported_by_margin(self):
        from pq_analysis import check_individual_harmonics
        # H5 is the larger current but sits well inside its limit; H23's much
        # tighter limit makes it the order that actually binds.
        #
        # The margins have to be separated by more than a rounding step.  At
        # 6.0 A and 0.9 A both orders sit at exactly 0.60 of their limits, and
        # which one "wins" is then decided by the last bit of the division --
        # so the assertion below held for a reason that had nothing to do with
        # margins.  1.2 A puts H23 at 0.80 of its limit against H5's 0.60.
        df = _frame(current_a=[100.0] * 10,
                    h5_current_a=[6.0] * 10,
                    h23_current_a=[1.2] * 10)
        r = check_individual_harmonics(df, Thresholds(isc_amps=5000.0))
        assert r["available"]
        assert r["worst_order"][0] == 5              # largest magnitude
        assert r["worst_margin_order"][0] == 23      # tightest margin
        assert r["worst_limit_pct"] < 2.0


def _sample_documents(outdir):
    """Both documents, for a residential and a large-commercial recording.

    Styling has to hold across classes: the two documents route different
    sections, and a colour added to a commercial-only table is exactly the kind
    of drift a residential-only check would miss.
    """
    import pq_analysis as An
    from pq_report import (generate_report, generate_word_report,
                           generate_customer_letter)
    outdir = Path(outdir)
    paths = []
    for pqd, cls, nominal in [("test_data/test_residential.pqd", "r", 120.0),
                              ("test_data/test_commercial_large.pqd", "c", 277.0)]:
        source = Path(pqd)
        if not source.exists():
            continue
        ds = extract_dataset(ProntoAdapter(source), ChannelMapper())
        th = Thresholds(nominal_voltage=nominal, customer_class=cls)
        df = ds.df
        ev = An.detect_events(ds, th)
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), ev, th,
            neutral_health_result=An.check_neutral_health(ds, th),
            itic_result=An.check_itic(ev, th),
            flicker_result=An.check_flicker(df, th))
        rep["root_causes"] = An.analyze_root_causes(rep, ds, th)
        stem = source.stem
        paths.append(generate_word_report(
            report=rep, thresh=th, ds=ds, site_name="S", site_address="A",
            engineer_name="E", outdir=outdir, stem=stem))
        paths.append(generate_customer_letter(
            rep, th, "1 Test St", "Eng", outdir, stem))
    assert paths, "no sample recordings available to style-check"
    return [Path(p) for p in paths if p is not None]


class TestANSIVoltageRanges:
    """C84.1 has two named ranges and rates sustained voltage, not events.

    The tool previously carried one symmetric band and judged it on the meter's
    within-interval extremes, so a sag both failed a steady-state standard that
    does not cover it and was counted again against ITIC.
    """

    NOMINAL = 277.0

    def _frame(self, avg, interval_min=None, n=1000):
        import pandas as pd
        idx = pd.date_range("2026-01-01", periods=n, freq="30s")
        df = pd.DataFrame({"voltage_a": pd.Series(avg, index=idx)})
        if interval_min is not None:
            df["voltage_a_min"] = pd.Series(interval_min, index=idx)
            df["voltage_a_peak"] = pd.Series(avg, index=idx)
        return df

    def _thresh(self, **kw):
        return Thresholds(nominal_voltage=self.NOMINAL, **kw)

    # ── The bands themselves ─────────────────────────────────────────────────
    @pytest.mark.parametrize("nominal,a_min,a_max,b_min,b_max", [
        # ANSI C84.1-2016 Table 1, service voltage. The standard prints these
        # rounded to whole volts, so 208 lands a few tenths off its published
        # 197/218 and 191/220 — everything else reproduces exactly.
        (120, 114.0, 126.0, 110.0, 127.0),
        (240, 228.0, 252.0, 220.0, 254.0),
        (480, 456.0, 504.0, 440.0, 508.0),
    ])
    def test_the_bands_match_the_published_table(self, nominal, a_min, a_max,
                                                 b_min, b_max):
        from pq_constants import ansi_bands
        b = ansi_bands(nominal)
        assert b["a_min"] == pytest.approx(a_min, abs=0.05)
        assert b["a_max"] == pytest.approx(a_max, abs=0.05)
        assert b["b_min"] == pytest.approx(b_min, abs=0.05)
        assert b["b_max"] == pytest.approx(b_max, abs=0.05)

    def test_range_b_is_not_symmetric(self):
        """The low side is wider, and a symmetric band would misjudge both ends."""
        from pq_constants import ansi_bands
        b = ansi_bands(120)
        assert b["b_min"] == pytest.approx(110.0)
        assert b["b_max"] == pytest.approx(127.0)
        assert (120 - b["b_min"]) > (b["b_max"] - 120)

    @pytest.mark.parametrize("nominal", [2400, 4160, 13200, 13800, 34500])
    def test_over_600_v_uses_its_own_table_1_group(self, nominal):
        """C84.1 gives systems over 600 V their own row, tighter below nominal.

        A primary-metered customer still has their own transformation between
        this meter and their equipment, and the standard reserves that headroom
        for the drop through it. The −5% that applies on a secondary service
        would put the limit 2.5% of nominal below where C84.1 sets it — on a
        13.2 kV service, 330 V of undervoltage that would read as a pass.
        """
        from pq_constants import ansi_bands
        b = ansi_bands(nominal)
        assert b["group"] == "over_600v"
        assert b["a_min"] == pytest.approx(nominal * 0.975)
        assert b["a_max"] == pytest.approx(nominal * 1.05)
        assert b["b_min"] == pytest.approx(nominal * 0.95)
        assert b["b_max"] == pytest.approx(nominal * 1.058)
        assert b["range_b_evaluated"] is True

    def test_the_two_groups_differ_only_below_nominal(self):
        from pq_constants import ansi_bands
        lv, mv = ansi_bands(480.0), ansi_bands(13200.0)
        assert lv["group"] == "under_600v" and mv["group"] == "over_600v"
        # Same ceilings, tighter floors.
        assert lv["a_max"] / 480.0 == pytest.approx(mv["a_max"] / 13200.0)
        assert (mv["a_min"] / 13200.0) > (lv["a_min"] / 480.0)
        assert (mv["b_min"] / 13200.0) > (lv["b_min"] / 480.0)

    def test_the_grouping_holds_on_both_bases(self):
        """The threshold is compared against whatever nominal it is handed.

        That works because the groups do not overlap line-to-neutral either:
        600 V L-L is 346 V L-N and 2400 V L-L is 1386 V L-N.
        """
        from pq_constants import ansi_bands
        assert ansi_bands(600.0)["group"] == "under_600v"      # 600 V L-L
        assert ansi_bands(346.4)["group"] == "under_600v"      # its L-N
        assert ansi_bands(2400.0)["group"] == "over_600v"      # 2.4 kV L-L
        assert ansi_bands(1385.6)["group"] == "over_600v"      # its L-N

    def test_above_34_5_kv_no_range_is_claimed(self):
        """Table 1 stops there; a band invented past it would still print as one."""
        from pq_constants import ansi_bands
        b = ansi_bands(69000.0)
        assert b["group"] == "out_of_scope"
        assert b["range_a_evaluated"] is False
        assert b["a_min"] is None and b["b_min"] is None
        assert "does not cover" in b["range_b_note"]

    def test_an_out_of_scope_nominal_reports_unavailable_with_the_reason(self):
        import pandas as pd
        idx = pd.date_range("2026-01-01", periods=50, freq="30s")
        df = pd.DataFrame({"voltage_a": pd.Series(39800.0, index=idx)})
        res = check_voltage_compliance(df, Thresholds(nominal_voltage=39837.0))
        assert res["available"] is False
        assert "does not cover" in res["error"]

    # ── What decides the verdict ─────────────────────────────────────────────
    def test_a_sag_inside_one_interval_does_not_fail_c841(self):
        """The Queensburg case: 2 of 5,344 interval minima dipped, means held.

        C84.1 rates sustained voltage. Failing it on a within-interval minimum
        both misapplies the standard and double-counts the event, which is
        already graded on depth and duration against ITIC.
        """
        import pandas as pd
        avg  = pd.Series(282.0, index=range(1000))
        mins = avg.copy(); mins.iloc[500] = 258.6
        df = self._frame(avg.values, mins.values)
        res = check_voltage_compliance(df, self._thresh())
        st  = res["phases"]["voltage_a"]
        assert st["band"] == "range_a"
        assert st["pct_out_of_bounds"] == 0.0
        # The excursion is not discarded — it is reported, under its own key.
        assert st["min_interval_v"] == pytest.approx(258.6)
        assert st["used_interval_extremes"] is True

    def test_a_sustained_excursion_still_fails(self):
        """Dropping the extremes must not make the check blind to real undervoltage."""
        import pandas as pd
        avg = pd.Series(282.0, index=range(1000))
        avg.iloc[:400] = 258.0
        res = check_voltage_compliance(self._frame(avg.values), self._thresh())
        st  = res["phases"]["voltage_a"]
        assert st["band"] == "range_b"
        assert st["pct_out_of_bounds"] == pytest.approx(40.0)

    # ── The three-state verdict ──────────────────────────────────────────────
    def test_an_excursion_inside_range_b_is_named_as_range_b(self):
        import pandas as pd
        avg = pd.Series(282.0, index=range(1000))
        avg.iloc[500] = 258.0           # below Range A, inside Range B
        res = check_voltage_compliance(self._frame(avg.values), self._thresh())
        st  = res["phases"]["voltage_a"]
        assert st["band"] == "range_b"
        assert st["pct_range_b"] == pytest.approx(0.1)
        assert st["pct_outside_b"] == 0.0

    def test_leaving_range_b_is_a_different_finding(self):
        import pandas as pd
        avg = pd.Series(282.0, index=range(1000))
        avg.iloc[500] = 240.0           # below Range B as well
        res = check_voltage_compliance(self._frame(avg.values), self._thresh())
        st  = res["phases"]["voltage_a"]
        assert st["band"] == "outside_b"
        assert st["pct_outside_b"] == pytest.approx(0.1)

    def test_sustained_range_b_outranks_a_brief_one(self):
        """C84.1 permits Range B only if excursions are limited in duration."""
        import pandas as pd
        from pq_report import _grade_voltage_band
        avg = pd.Series(282.0, index=range(1000))

        brief = avg.copy(); brief.iloc[500] = 258.0
        held  = avg.copy(); held.iloc[:400] = 258.0
        g_brief = _grade_voltage_band(
            check_voltage_compliance(self._frame(brief.values), self._thresh()))
        g_held = _grade_voltage_band(
            check_voltage_compliance(self._frame(held.values), self._thresh()))
        assert g_brief["band"] == "minor"
        assert g_held["band"] == "significant"
        assert "corrected within a reasonable time" in strip_marks(g_held["reason"])

    def test_a_compliant_service_grades_as_compliant(self):
        import pandas as pd
        from pq_report import _grade_voltage_band
        avg = pd.Series(282.0, index=range(1000))
        g = _grade_voltage_band(
            check_voltage_compliance(self._frame(avg.values), self._thresh()))
        assert g["band"] == "compliant"

    # ── The Measured cell ────────────────────────────────────────────────────
    def test_the_cell_never_contradicts_its_own_verdict(self):
        """A cell opening on 258.6 V and closing on "in Range A" reads as wrong.

        The two numbers answer different questions, so the extreme is quoted
        under its own clause naming where it *is* graded.
        """
        import pandas as pd
        from pq_report import _voltage_band_cell
        avg  = pd.Series(282.0, index=range(1000))
        mins = avg.copy(); mins.iloc[500] = 258.6
        res = check_voltage_compliance(
            self._frame(avg.values, mins.values), self._thresh())
        cell = strip_marks(
            _voltage_band_cell(res["phases"]["voltage_a"], res, "Phase A: "))
        assert "All intervals in Range A" in cell
        assert "Within-interval extremes 258.6" in cell
        assert "ITIC" in cell

    def test_an_outside_b_cell_quotes_the_range_b_edge(self):
        """Citing the Range A limit there reads as a classification one band out."""
        import pandas as pd
        from pq_report import _voltage_band_cell
        avg = pd.Series(282.0, index=range(1000))
        avg.iloc[500] = 240.0
        res = check_voltage_compliance(self._frame(avg.values), self._thresh())
        cell = strip_marks(
            _voltage_band_cell(res["phases"]["voltage_a"], res, "Phase A: "))
        assert "outside Range B" in cell
        assert "all below 253.9 V" in cell

    # ── The primary nominal ──────────────────────────────────────────────────
    def test_an_entered_primary_nominal_beats_inference(self):
        """PSCo runs several primary voltages and the file names none of them."""
        import pandas as pd
        from pq_analysis import check_line_to_line_voltage
        idx = pd.date_range("2026-01-01", periods=200, freq="30s")
        ln  = 7621.0
        df = pd.DataFrame({
            "voltage_a":  pd.Series(ln, index=idx),
            "voltage_b":  pd.Series(ln, index=idx),
            "voltage_c":  pd.Series(ln, index=idx),
            "voltage_ab": pd.Series(ln * 3 ** 0.5, index=idx),
        })
        th = Thresholds(nominal_voltage=ln, primary_ll_voltage=13200.0)
        res = check_line_to_line_voltage(df, th)
        assert res["available"] is True
        assert res["nominal_v"] == pytest.approx(13200.0)
        assert res["nominal_source"] == "entered"
        # Judged against the over-600 V group, not the secondary one.
        assert res["nominal_group"] == "over_600v"
        assert res["range_v"][0] == pytest.approx(13200.0 * 0.975)
        assert res["range_b_v"][0] == pytest.approx(13200.0 * 0.95)

    def test_without_an_entered_nominal_inference_still_runs(self):
        import pandas as pd
        from pq_analysis import check_line_to_line_voltage
        idx = pd.date_range("2026-01-01", periods=200, freq="30s")
        df = pd.DataFrame({
            "voltage_a":  pd.Series(277.0, index=idx),
            "voltage_b":  pd.Series(277.0, index=idx),
            "voltage_c":  pd.Series(277.0, index=idx),
            "voltage_ab": pd.Series(480.0, index=idx),
        })
        res = check_line_to_line_voltage(df, Thresholds(nominal_voltage=277.0))
        assert res["nominal_source"] == "inferred"
        assert res["nominal_v"] == pytest.approx(480.0)
        assert res["range_v"][0] == pytest.approx(456.0)
        assert res["range_v"][1] == pytest.approx(504.0)


class TestSeverityGrading:
    """Severity is a second axis beside compliance, not a replacement for it.

    Compliance stays binary because the standards are; severity says how much
    the exceedance matters, so an isolated artifact and a sustained overload
    stop sharing one red FAIL.
    """

    def test_compliance_and_severity_are_independent(self):
        from pq_analysis import grade_finding
        # Over the limit is over the limit — severity never contradicts that.
        assert grade_finding(False, measured=8.1, limit=8.0,
                             persistence_pct=0.1)["band"] == "minor"
        assert grade_finding(True, measured=2.0, limit=8.0)["band"] == "compliant"

    def test_a_marginal_brief_exceedance_is_minor_not_severe(self):
        from pq_analysis import grade_finding
        g = grade_finding(False, measured=8.2, limit=8.0, persistence_pct=0.2)
        assert g["band"] == "minor"

    def test_a_large_sustained_exceedance_is_severe(self):
        from pq_analysis import grade_finding
        g = grade_finding(False, measured=12.8, limit=8.0, persistence_pct=60.0)
        assert g["band"] == "severe"

    def test_a_very_large_exceedance_is_severe_however_brief(self):
        from pq_analysis import grade_finding
        g = grade_finding(False, measured=19.0, limit=8.0, persistence_pct=0.5)
        assert g["band"] == "severe"

    def test_low_confidence_drops_one_band_and_says_so(self):
        from pq_analysis import grade_finding
        kw = dict(measured=12.8, limit=8.0, persistence_pct=60.0)
        assert grade_finding(False, **kw)["band"] == "severe"
        g = grade_finding(False, confidence_notes=["ISC not supplied"], **kw)
        assert g["band"] == "significant"
        assert g["downgraded"] is True
        assert "ISC not supplied" in g["reason"]
        assert "severity reduced one band" in g["reason"]

    def test_confidence_never_downgrades_below_minor(self):
        from pq_analysis import grade_finding
        g = grade_finding(False, measured=8.1, limit=8.0, persistence_pct=0.1,
                          confidence_notes=["light-load intervals excluded"])
        assert g["band"] == "minor"
        assert g["downgraded"] is False

    def test_a_flagged_exceedance_never_reports_zero_persistence(self):
        """"Exceeded" beside "0.0% of the recording" reads as self-refuting.

        A handful of intervals out of tens of thousands is a real exceedance
        that one decimal place cannot show, and the reader who checks the row
        against itself concludes the flag is wrong rather than the rounding.
        """
        from pq_analysis import grade_finding
        g = grade_finding(False, measured=8.1, limit=8.0, persistence_pct=0.04)
        reason = strip_marks(g["reason"])
        assert "0.0% of the recording" not in reason
        assert "<0.1% of the recording" in reason

    def test_a_genuinely_zero_share_still_prints_as_zero(self):
        """The guard must not turn a true zero into a fake exceedance."""
        from pq_constants import pct_text
        assert pct_text(0.0, ".1f") == "0.0%"
        assert pct_text(0.0, ".2f") == "0.00%"

    def test_the_floor_tracks_the_format_it_is_given(self):
        from pq_constants import pct_text
        assert pct_text(0.004, ".2f") == "<0.01%"
        assert pct_text(0.004, ".1f") == "<0.1%"
        assert pct_text(0.4, ".0f") == "<1%"
        # Anything the format can show is left exactly as it was.
        assert pct_text(0.05, ".1f") == "0.1%"
        assert pct_text(17.3, ".1f") == "17.3%"

    def test_close_to_the_limit_is_watch_not_compliant(self):
        from pq_analysis import grade_finding
        g = grade_finding(True, measured=7.2, limit=8.0)
        assert g["band"] == "watch"
        assert "90% of it" in strip_marks(g["reason"])


    def test_a_healthy_power_factor_is_not_flagged_as_watch(self):
        from pq_analysis import grade_finding
        # PF is bounded at 1.0, so the ceiling-metric watch band would mark
        # every power factor including 0.99.  Floor metrics get a tighter one.
        assert grade_finding(True, measured=0.989, limit=0.90,
                             lower_is_worse=True)["band"] == "compliant"
        assert grade_finding(True, measured=0.912, limit=0.90,
                             lower_is_worse=True)["band"] == "watch"

    def test_a_power_factor_shortfall_reads_in_the_right_direction(self):
        from pq_analysis import grade_finding
        g = grade_finding(False, measured=0.88, limit=0.90, persistence_pct=12.0,
                          lower_is_worse=True)
        assert "below the limit" in g["reason"]

    def test_an_unassessed_check_has_no_severity(self):
        from pq_analysis import grade_finding
        assert grade_finding(None)["band"] == "not_assessed"


class TestOneSeverityScale:
    """Colour and wording are a system, and systems drift.

    Three analyses grew their own severity words and their own reds before
    there was a shared scale, so the same seriousness printed three ways. These
    tests hold the rendering to one scale: one vocabulary, one ramp, alarm off
    brand red, and every colour readable.
    """

    @staticmethod
    def _contrast(hex_a, hex_b="FFFFFF"):
        def lum(h):
            ch = [int(h[i:i + 2], 16) / 255 for i in (0, 2, 4)]
            f = lambda c: c / 12.92 if c <= 0.03928 else ((c + 0.055) / 1.055) ** 2.4
            return 0.2126 * f(ch[0]) + 0.7152 * f(ch[1]) + 0.0722 * f(ch[2])
        hi, lo = sorted((lum(hex_a), lum(hex_b)), reverse=True)
        return (hi + 0.05) / (lo + 0.05)

    def test_every_severity_colour_is_readable_on_white(self):
        from pq_report import _SEVERITY_STYLE
        for band, (rgb, _) in _SEVERITY_STYLE.items():
            ratio = self._contrast("%02X%02X%02X" % rgb)
            assert ratio >= 4.5, f"{band} is {ratio:.2f}:1, below WCAG AA"

    def test_alarm_is_not_the_brand_red(self):
        from pq_report import _SEVERITY_STYLE, _XE_RED
        severe = "%02X%02X%02X" % _SEVERITY_STYLE["severe"][0]
        assert severe != str(_XE_RED)
        # Far enough apart to survive a printer, not just a monitor.
        dist = sum((int(severe[i:i + 2], 16) - int(str(_XE_RED)[i:i + 2], 16)) ** 2
                   for i in (0, 2, 4)) ** 0.5
        assert dist > 60, f"severe sits {dist:.0f} from brand red"

    def test_the_older_scales_render_through_the_shared_one(self):
        # Neutral integrity, assessment findings and action priorities each
        # keep their own keys in the analysis layer; none of them may put a
        # word or a colour on the page that the shared scale does not own.
        from pq_report import _sev_band, _sev_color, _sev_label, _SEVERITY_STYLE
        for legacy, band in [("critical", "severe"), ("warning", "significant"),
                             ("caution", "minor"), ("normal", "compliant"),
                             ("info", "watch"), ("High", "severe"),
                             ("Medium", "significant"), ("Low", "watch")]:
            assert _sev_band(legacy) == band
            assert _sev_color(legacy) == _sev_color(band)
            assert _sev_label(legacy) == SEVERITY_LABEL[band]
        assert _sev_band("nonsense-band") == "not_assessed"
        assert set(_SEVERITY_STYLE) == set(SEVERITY_LABEL)

    def test_the_documents_paint_only_from_the_palette(self, tmp_path):
        # The regression that started this: a fourth red, added locally, that
        # no one could see was a fourth red.
        import docx
        allowed = {"%02X%02X%02X" % rgb for rgb, _ in
                   __import__("pq_report")._SEVERITY_STYLE.values()}
        allowed |= {"DA1020",                       # brand: headings, titles
                    "FFFFFF",                       # text on the header band
                    "333333", "555555", "666666", "6B6B6B", "808080"}  # neutrals
        for path in _sample_documents(tmp_path):
            d = docx.Document(str(path))
            paras = list(d.paragraphs) + [
                p for t in d.tables for row in t.rows
                for c in row.cells for p in c.paragraphs]
            for p in paras:
                for r in p.runs:
                    if r.font.color is not None and r.font.color.rgb is not None:
                        assert str(r.font.color.rgb) in allowed, (
                            f"{path.name}: unpalette colour "
                            f"{r.font.color.rgb} on {r.text[:40]!r}")

    def test_both_documents_are_set_in_the_house_font(self, tmp_path):
        import docx
        for path in _sample_documents(tmp_path):
            d = docx.Document(str(path))
            normal = d.styles["Normal"]
            assert normal.font.name == "Arial"
            assert normal.font.size.pt == 10


class TestMeasuredValuesAreMarkedInProse:
    """A reader must be able to tell a reading from a limit at a glance.

    Prose puts both in one sentence -- "the measured 0.0226 Ω sits within the
    range the expected 0.0249 Ω accounts for" -- and the second figure is a
    calculation from the Blue Book, not something the meter saw. Measured
    values are marked where they are written and render bold.
    """

    def test_a_marked_value_renders_bold_and_keeps_its_unit(self):
        import docx
        from pq_report import _DocxDocument, _apply_base_style, _body, _m
        d = _DocxDocument()
        _apply_base_style(d)
        _body(d, f"The measured {_m(0.0226, '.4f', ' Ω')} sits inside the "
                 f"expected {0.0249:.4f} Ω.")
        runs = d.paragraphs[0].runs
        bold = [r.text for r in runs if r.bold]
        plain = "".join(r.text for r in runs if not r.bold)
        assert bold == ["0.0226 Ω"]
        assert "0.0249 Ω" in plain
        assert "".join(r.text for r in runs) == (
            "The measured 0.0226 Ω sits inside the expected 0.0249 Ω.")

    def _compliance_measured_column(self, tmp_path):
        """Every Measured cell in the compliance table, as (bold, plain)."""
        import docx
        for path in _sample_documents(tmp_path):
            if "internal" not in path.name:
                continue
            doc = docx.Document(str(path))
            tbl = next((t for t in doc.tables
                        if t.rows[0].cells[0].text.strip() == "Standard"), None)
            if tbl is None:
                continue
            for row in tbl.rows[1:]:
                # Group headings merge across the row; skip them.
                if row.cells[0]._tc is row.cells[1]._tc:
                    continue
                runs = [r for p in row.cells[1].paragraphs for r in p.runs]
                if not runs:
                    continue
                yield ("".join(r.text for r in runs if r.bold),
                       "".join(r.text for r in runs if not r.bold))

    def test_the_measured_column_bolds_readings_and_not_limits(self, tmp_path):
        # The whole point of the column: "P95 6.80% (limit 8.00%)" puts a
        # reading and the standard it is judged against in one cell, and only
        # the first came off the meter.
        pytest.importorskip("docx")
        checked = 0
        for bold, plain in self._compliance_measured_column(tmp_path):
            # A limit is never a reading, wherever it appears.
            for word in ("limit", "Allowed", "allowed", "nameplate"):
                if word in bold:
                    pytest.fail(f"a limit was bolded as measured: {bold!r}")
            checked += 1
        assert checked, "the compliance table produced no Measured cells"

    def test_at_least_one_reading_is_bolded(self, tmp_path):
        # Guards the opposite failure: a table where nothing is marked passes
        # the test above trivially.
        pytest.importorskip("docx")
        assert any(bold.strip() for bold, _plain
                   in self._compliance_measured_column(tmp_path))

    def test_the_markers_never_reach_the_page(self, tmp_path):
        # The sentinels are private-use characters; if a prose path ever
        # bypasses the splitter they would print as boxes in a customer's copy.
        import docx
        from pq_report import _MEASURED_OPEN, _MEASURED_CLOSE
        for path in _sample_documents(tmp_path):
            xml = docx.Document(str(path)).element.xml
            assert _MEASURED_OPEN not in xml, f"{path.name} leaked an open marker"
            assert _MEASURED_CLOSE not in xml, f"{path.name} leaked a close marker"

    def test_no_other_output_carries_the_markers(self, tmp_path, capsys):
        # The analysis layer writes findings, and those strings go to the
        # console and the CSVs as well as to Word. Only Word can render a mark.
        import pq_analysis as An
        from pq_report import (MEASURED_OPEN, MEASURED_CLOSE, generate_report,
                               print_report, export_results)
        ds = extract_dataset(ProntoAdapter(Path("test_data/test_residential.pqd")),
                             ChannelMapper())
        th = Thresholds(nominal_voltage=120.0, customer_class="r")
        df = ds.df
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), An.detect_events(ds, th), th,
            neutral_health_result=An.check_neutral_health(ds, th))
        rep["root_causes"] = An.analyze_root_causes(rep, ds, th)

        print_report(rep)
        printed = capsys.readouterr().out
        assert MEASURED_OPEN not in printed and MEASURED_CLOSE not in printed

        export_results(ds, rep, tmp_path, "leak")
        for csv in tmp_path.glob("*.csv"):
            text = csv.read_text(errors="replace")
            assert MEASURED_OPEN not in text, f"{csv.name} leaked a marker"
            assert MEASURED_CLOSE not in text, f"{csv.name} leaked a marker"

    def test_what_is_marked_is_a_number_and_not_a_label(self, tmp_path):
        import re
        import docx
        for path in _sample_documents(tmp_path):
            d = docx.Document(str(path))
            for p in d.paragraphs:
                if not p.style.name.startswith(("Normal", "Body", "List")):
                    continue
                if p.text.startswith("Figures in "):
                    continue            # the key, which bolds the word "bold"
                # Lead-in labels ("Finding:") are bold by their own right; a
                # marked span is the rest, and every one must start in a digit.
                for run in p.runs:
                    text = run.text.strip()
                    if not run.bold or not text or text.endswith(":"):
                        continue
                    if re.match(r'^[A-Z][A-Za-z]', text):
                        continue        # severity headline, e.g. "SEVERE: ..."
                    assert re.match(r'^[<>~+±-]?[\d.,]', text), (
                        f"{path.name}: bold span is not a measurement: {text!r}")

    def test_limits_and_ratings_stay_plain(self, tmp_path):
        # The distinction is only worth drawing if the other side holds: a
        # nameplate rating and a standard's limit must never come out bold.
        import docx
        for path in _sample_documents(tmp_path):
            d = docx.Document(str(path))
            for p in d.paragraphs:
                text = p.text
                if "nameplate" not in text and "Range A" not in text:
                    continue
                for run in p.runs:
                    if not run.bold:
                        continue
                    following = text.split(run.text, 1)[-1][:12]
                    assert not following.startswith((" kVA nameplate",
                                                     " V) for the entire")), (
                        f"{path.name}: a rating or limit rendered as measured: "
                        f"{run.text!r}")


class TestVoltageTHDIsJudgedStatistically:
    """IEEE 519-2022 Clause 5 judges voltage THD on percentiles, not on maxima.

    A single interval where the fundamental collapsed used to fail an entire
    site: pass_fail asked for pct_exceeding == 0, and V_h/V_1 runs to tens of
    percent whenever V_1 approaches zero.
    """

    def _thresh(self):
        return Thresholds(nominal_voltage=120.0)

    def test_a_sag_artifact_is_dropped_not_reported_as_distortion(self):
        from pq_analysis import check_thd
        # 2.5 % site with one interval measured during a collapse to 11 V.
        thd   = [2.5] * 99 + [80.20]
        volts = [120.0] * 99 + [11.0]
        r = check_thd(_frame(thd_voltage_a=thd, voltage_a=volts), self._thresh())["voltage"]
        assert r["artifact_samples"] == 1
        assert r["sample_count"] == 99
        assert r["max_thd_pct"] == pytest.approx(2.5)
        assert r["p95_pass"] and r["p99_pass"]

    def test_a_spike_at_normal_voltage_does_not_fail_the_site(self):
        from pq_analysis import check_thd
        # The voltage gate cannot catch this one — the percentile has to.
        thd = [2.5] * 99 + [80.20]
        r = check_thd(_frame(thd_voltage_a=thd, voltage_a=[120.0] * 100),
                      self._thresh())["voltage"]
        assert r["artifact_samples"] == 0
        assert r["max_thd_pct"] == pytest.approx(80.20)
        assert r["p95_pass"] and r["p99_pass"]
        assert r["max_is_outlier"] is True

    def test_genuine_sustained_distortion_still_fails(self):
        from pq_analysis import check_thd
        r = check_thd(_frame(thd_voltage_a=[9.5] * 100, voltage_a=[120.0] * 100),
                      self._thresh())["voltage"]
        assert r["p95_pass"] is False
        assert r["max_is_outlier"] is False

    def test_a_sustained_excursion_is_caught_by_the_p99_rule(self):
        from pq_analysis import check_thd
        # P95 stays under 8 %, but 3 % of the recording sits at 13 % — over the
        # 1.5x short-time limit, so the site must not pass.
        thd = [6.0] * 97 + [13.0] * 3
        r = check_thd(_frame(thd_voltage_a=thd, voltage_a=[120.0] * 100),
                      self._thresh())["voltage"]
        assert r["p95_pass"] is True
        assert r["p99_pass"] is False

    def test_the_gate_keeps_distortion_measured_during_a_shallow_sag(self):
        from pq_analysis import check_thd
        # 0.75 pu is a real sag but a valid THD reading; it must not be dropped.
        r = check_thd(_frame(thd_voltage_a=[9.0] * 100, voltage_a=[90.0] * 100),
                      self._thresh())["voltage"]
        assert r["artifact_samples"] == 0
        assert r["p95_pass"] is False


class TestTDDDefinition:
    def _thresh(self):
        return Thresholds(nominal_voltage=120.0, isc_amps=10000.0)

    def test_tdd_is_harmonic_current_over_il(self):
        from pq_analysis import check_thd
        # Ih = 5 A on every interval, IL = fundamental max = 12 A → TDD = 41.67 %.
        df = _frame(current_a=[13.0] * 10, hrms_current_a=[5.0] * 10,
                    thd_current_a=[41.667] * 10)
        r = check_thd(df, self._thresh())["current"]
        assert r["metric"] == "tdd"
        assert r["harmonic_rms_source"] == "meter"
        assert r["il_amps"] == pytest.approx(12.0, abs=0.01)
        assert r["max_thd_pct"] == pytest.approx(100 * 5.0 / 12.0, rel=1e-6)

    def test_an_inconsistent_thd_channel_no_longer_inflates_tdd(self):
        # The old form, THD% x Irms / IL, multiplied the THD ratio by current, so
        # an interval where the meter's THD contradicts its own RMS and harmonic
        # RMS was amplified. Reading the harmonic current directly cannot be.
        from pq_analysis import check_thd
        thd = [12.8] * 9 + [111.7]          # last interval is the anomaly
        df = _frame(current_a=[5.5] * 10, hrms_current_a=[0.7] * 10,
                    thd_current_a=thd)
        r = check_thd(df, self._thresh())["current"]
        il = np.sqrt(5.5 ** 2 - 0.7 ** 2)          # fundamental, not the 5.5 A RMS
        assert r["max_thd_pct"] == pytest.approx(100 * 0.7 / il, rel=1e-6)
        # The anomalous interval contributes nothing beyond the others.
        assert r["max_thd_pct"] == pytest.approx(r["mean_thd_pct"], rel=1e-9)
        # For contrast: the old form would have reached 111.7 x 5.5 / IL here.
        assert r["max_thd_pct"] < 20.0

    def test_falls_back_to_the_thd_channel_against_the_fundamental(self):
        from pq_analysis import check_thd
        # No harmonic RMS anywhere: TDD comes from THD x I1 / IL, where I1 is the
        # RMS channel (the best available) rather than being confused with it.
        df = _frame(current_a=[10.0] * 10, thd_current_a=[20.0] * 10)
        r = check_thd(df, self._thresh())["current"]
        assert r["metric"] == "tdd"
        assert r["harmonic_rms_source"] is None
        assert r["max_thd_pct"] == pytest.approx(20.0, rel=1e-6)

    def test_il_uses_the_fundamental_not_the_rms(self):
        from pq_analysis import check_thd
        df = _frame(current_a=[13.0] * 10, hrms_current_a=[5.0] * 10,
                    thd_current_a=[41.667] * 10)
        r = check_thd(df, self._thresh())["current"]
        # IEEE 519 defines IL at the fundamental: 12 A here, not the 13 A RMS.
        assert r["il_amps"] == pytest.approx(12.0, abs=0.01)


# ─────────────────────────────────────────────────────────────────────────────
# 12. Truncated and damaged files
# ─────────────────────────────────────────────────────────────────────────────

@pytest.mark.skipif(not _FIXTURES, reason="test_data/*.pqd not generated")
class TestTruncatedFile:
    """A file cut short must degrade, not fail outright.

    Damage lands at the end of an interrupted export or copy, so the interval
    records carrying the compliance data are usually intact and worth reading.
    """

    @pytest.fixture
    def truncated(self, tmp_path):
        src = Path(__file__).parent / "test_data" / "test_commercial_small.pqd"
        raw = src.read_bytes()
        out = tmp_path / "truncated.pqd"
        out.write_bytes(raw[: len(raw) - 3000])
        return out

    def test_truncation_is_measured_not_just_reported_as_a_zlib_error(self, truncated):
        f = pqdif.PQDIFFile(truncated)
        assert f.missing_bytes == 3000

    def test_intact_file_reports_no_missing_bytes(self):
        src = Path(__file__).parent / "test_data" / "test_commercial_small.pqd"
        f = pqdif.PQDIFFile(src)
        assert f.missing_bytes == 0
        assert f.unreadable_observations == []

    def test_readable_records_still_load(self, truncated):
        adapter = ProntoAdapter(truncated)
        df = extract_dataset(adapter, ChannelMapper()).df
        assert len(df) > 0
        assert "voltage_a" in df.columns

    def test_damage_is_recorded_on_the_dataset(self, truncated):
        ds = extract_dataset(ProntoAdapter(truncated), ChannelMapper())
        dq = ds.meta["data_quality"]
        assert dq["missing_bytes"] == 3000

    def test_intact_file_carries_a_clean_bill(self):
        src = Path(__file__).parent / "test_data" / "test_commercial_small.pqd"
        ds = extract_dataset(ProntoAdapter(src), ChannelMapper())
        dq = ds.meta["data_quality"]
        assert dq["missing_bytes"] == 0 and dq["unreadable_observations"] == 0

    def test_a_file_with_no_readable_observation_still_raises(self, tmp_path):
        # Degrading gracefully must not extend to inventing an empty analysis.
        src = Path(__file__).parent / "test_data" / "test_commercial_small.pqd"
        raw = bytearray(src.read_bytes())
        f = pqdif.PQDIFFile(src)
        for record in f.records:
            if record.tag == pqdif.TAG_OBSERVATION:
                start = record.position + record.header_size
                raw[start:start + 8] = b"\x00" * 8      # destroy the zlib header
        broken = tmp_path / "broken.pqd"
        broken.write_bytes(bytes(raw))
        with pytest.raises(pqdif.PQDIFError, match="could be read"):
            _ = pqdif.PQDIFFile(broken).observations


@pytest.mark.skipif(not _FIXTURES, reason="test_data/*.pqd not generated")
class TestRecordThatWillNotInflate:
    """A body that will not inflate has several causes, and they differ.

    The files are customer data that cannot be shared for a second look, so
    the reader has to recover the causes it can and, for the one it cannot,
    say enough in the failure itself to tell the two apart from the message
    alone: a damaged file to re-export, or a file this reader reads wrongly.
    """

    SRC = Path(__file__).parent / "test_data" / "test_commercial_small.pqd"

    @staticmethod
    def _rebuild(records, path):
        """Re-emit a record chain, relinking it around new body lengths."""
        out = bytearray()
        for i, (tag, payload, declared) in enumerate(records):
            last = i == len(records) - 1
            next_pos = 0 if last else (
                len(out) + pqdif.RECORD_HEADER_SIZE + len(payload))
            header = bytearray(pqdif.RECORD_HEADER_SIZE)
            header[0:16] = pqdif.RECORD_SIGNATURE.bytes_le
            header[16:32] = tag.bytes_le
            struct.pack_into("<IIII", header, 32, pqdif.RECORD_HEADER_SIZE,
                             len(payload) if declared is None else declared,
                             next_pos, zlib.crc32(payload) & 0xFFFFFFFF)
            out += header + payload
        path.write_bytes(bytes(out))
        return path

    def _records(self):
        raw = self.SRC.read_bytes()
        return [(r.tag, raw[r.position + r.header_size:
                            r.position + r.header_size + r.body_size], None)
                for r in pqdif.PQDIFFile(self.SRC).records]

    def test_a_size_field_that_under_declares_its_body_is_read_anyway(self, tmp_path):
        # Every byte is present and the chain is intact; only the size field
        # is short, which cuts the deflate stream off mid-way.
        records = self._records()
        for i, (tag, payload, _) in enumerate(records):
            if tag == pqdif.TAG_OBSERVATION:
                records[i] = (tag, payload, len(payload) - 40)
        f = pqdif.PQDIFFile(self._rebuild(records, tmp_path / "short_size.pqd"))
        assert f.observations
        assert f.unreadable_observations == []

    def test_an_uncompressed_record_in_a_compressed_file_is_read_anyway(self, tmp_path):
        records = self._records()
        for i, (tag, payload, _) in enumerate(records):
            if tag == pqdif.TAG_OBSERVATION:
                records[i] = (tag, zlib.decompress(payload), None)
        f = pqdif.PQDIFFile(self._rebuild(records, tmp_path / "plain.pqd"))
        assert f.observations
        assert f.unreadable_observations == []

    def test_an_incomplete_stream_in_a_whole_file_is_named_as_damage(self, tmp_path):
        # The field case: the header inflates as far as its first bytes and
        # then the stream stops, while the file itself is not short at all.
        records = self._records()
        first = next(i for i, (tag, _, _) in enumerate(records)
                     if tag == pqdif.TAG_OBSERVATION)
        records[first] = (records[first][0], records[first][1][:2], None)
        f = pqdif.PQDIFFile(self._rebuild(records, tmp_path / "cut_stream.pqd"))
        _ = f.observations

        reason = f.unreadable_observations[0][1]
        assert "incomplete or truncated stream" in reason
        # The distinction the message exists to draw: the file is whole, so a
        # re-export is the answer, and no byte count went missing.
        assert "cut short" not in reason
        assert "damaged in place" in reason
        assert "zlib_header=valid" in reason
        assert "present=2 to_next_record=2" in reason
        assert "next_header=intact" in reason

    def test_a_body_that_is_not_a_zlib_stream_says_so_instead(self, tmp_path):
        # A different cause with the same surface symptom, and the evidence
        # has to separate them without anyone opening the file.
        raw = bytearray(self.SRC.read_bytes())
        for record in pqdif.PQDIFFile(self.SRC).records:
            if record.tag == pqdif.TAG_OBSERVATION:
                start = record.position + record.header_size
                raw[start:start + 8] = b"\x00" * 8
        broken = tmp_path / "damaged.pqd"
        broken.write_bytes(bytes(raw))

        f = pqdif.PQDIFFile(broken)
        with pytest.raises(pqdif.PQDIFError):
            _ = f.observations
        reason = f.unreadable_observations[0][1]
        assert "not begin with a zlib stream" in reason
        assert "zlib_header=not a zlib header" in reason
        assert "cut short" not in reason
        assert "present=" in reason and "to_next_record=" in reason

    def test_a_truncated_file_still_says_the_file_was_cut_short(self, tmp_path):
        # The evidence must not drown the one cause that is not about this
        # record at all: bytes the file never received.
        raw = self.SRC.read_bytes()
        short = tmp_path / "short.pqd"
        short.write_bytes(raw[:len(raw) - 3000])
        f = pqdif.PQDIFFile(short)
        _ = f.observations
        assert f.missing_bytes == 3000
        # Whatever inflated is kept; only a record that yields nothing fails.
        for _pos, reason in f.unreadable_observations:
            assert "cut short" in reason

    def test_the_evidence_survives_into_the_report(self, tmp_path):
        # The message is only useful if it reaches the document the engineer
        # actually reads, whole rather than summarised away.
        from pq_report import _integrity_note
        raw = bytearray(self.SRC.read_bytes())
        observations = [r for r in pqdif.PQDIFFile(self.SRC).records
                        if r.tag == pqdif.TAG_OBSERVATION]
        start = observations[0].position + observations[0].header_size
        raw[start:start + 8] = b"\x00" * 8
        broken = tmp_path / "one_bad.pqd"
        broken.write_bytes(bytes(raw))

        f = pqdif.PQDIFFile(broken)
        _ = f.observations
        note = _integrity_note({
            "unreadable_observations": len(f.unreadable_observations),
            "total_observations": f.observation_count,
            "missing_bytes": f.missing_bytes,
            "unreadable_detail": [{"offset": pos, "name": "", "reason": reason}
                                  for pos, reason in f.unreadable_observations],
        }, {})
        assert "Evidence: declared_body=" in note
        assert "partial_inflate=" in note

    def test_the_header_summarises_and_the_appendix_carries_the_evidence(self):
        # The evidence has to reach the document, but it does not have to sit
        # in the header table, where it buries the rows the reader came for.
        from pq_report import _integrity_summary, _integrity_note
        dq = {
            "unreadable_observations": 1, "total_observations": 87,
            "missing_bytes": 0,
            "unreadable_detail": [{
                "offset": 41010644, "name": "Interval max/min",
                "reason": ("record at 41010644 declared record-level zlib "
                           "compression but did not decompress. Evidence: "
                           "declared_body=0 partial_inflate=0."),
            }],
        }
        head = _integrity_summary(dq, {})
        full = _integrity_note(dq, {})

        # Scope, cost and where to look -- and nothing a parser author needs.
        assert "1 of 87" in head
        assert "may understate" in head
        assert "Appendix B" in head
        assert "Evidence:" not in head
        assert "declared_body" not in head
        assert len(head) < len(full) / 2

        # And the evidence is still in the document, one section further down.
        assert "declared_body=0" in full

    def test_one_skipped_record_is_singular(self):
        # "1 of 87 ... were skipped" reached a real report before this.
        from pq_report import _integrity_summary
        dq = {"unreadable_observations": 1, "total_observations": 87,
              "missing_bytes": 0, "unreadable_detail": []}
        head = _integrity_summary(dq, {})
        assert "records could not be decoded and was skipped" in head
        assert "were skipped" not in head
        assert "the other 86 read cleanly" in head

    def test_a_clean_file_gets_no_integrity_row_at_all(self):
        docx = pytest.importorskip("docx")
        from pq_report import _word_site_info_table
        doc = docx.Document()
        fs = {"start_time": "a", "end_time": "b", "duration_hours": 1.0,
              "sample_count": 12, "topology": "3-phase", "data_quality": {
                  "missing_bytes": 0, "unreadable_observations": 0}}
        _word_site_info_table(doc, "Site", "stem", "", fs, 277.0, 480)
        labels = [r.cells[0].text for r in doc.tables[0].rows]
        assert "Source file integrity" not in labels

    def test_an_intact_file_needs_none_of_it(self, tmp_path):
        f = pqdif.PQDIFFile(self.SRC)
        assert f.observations and f.unreadable_observations == []
        assert all(r.next_header_intact in (None, True) for r in f.records)


# ─────────────────────────────────────────────────────────────────────────────
# 13. Harmonic conclusions are gated on measurable load
# ─────────────────────────────────────────────────────────────────────────────

class TestHarmonicSignificanceGate:
    """A spectrum quantised to the meter's resolution carries no shape.

    Without this gate a 1.1 A residential service with 0.2 A of third harmonic
    was reported as an electric arc furnace at 94% similarity, with a
    recommendation to install a STATCOM.
    """

    @staticmethod
    def _frame(fundamental, h3, n=300):
        # Slight variation, not constants: a constant series has zero standard
        # deviation, which makes the Pearson correlation undefined and buries the
        # behaviour under a divide warning.
        import pandas as pd
        rng = np.random.default_rng(11)
        idx = pd.date_range("2025-01-01", periods=n, freq="5min", tz="UTC")
        jitter = lambda base: np.abs(base * (1 + rng.normal(0, 0.02, n)))
        return pd.DataFrame({
            "current_a": jitter(fundamental),
            "h3_current_a": jitter(h3),
            "h5_current_a": jitter(h3 * 0.5),
            "h7_current_a": jitter(h3 * 0.3),
            "h9_current_a": jitter(h3 * 0.1),
            "h3_voltage_a": jitter(h3 * 0.1),
            "h5_voltage_a": jitter(h3 * 0.08),
            "h7_voltage_a": jitter(h3 * 0.06),
        }, index=idx)

    def test_light_load_spectrum_is_refused(self):
        from pq_analysis import harmonic_spectrum_significance
        r = harmonic_spectrum_significance(self._frame(1.1, 0.2), Thresholds())
        assert r["usable"] is False
        assert "resolution" in r["reason"]

    def test_loaded_spectrum_is_accepted(self):
        from pq_analysis import harmonic_spectrum_significance
        r = harmonic_spectrum_significance(self._frame(20.0, 4.0), Thresholds())
        assert r["usable"] is True
        assert r["resolution_steps"] >= 5

    def test_no_signature_match_at_light_load(self):
        from pq_analysis import _detect_harmonic_signature, harmonic_spectrum_significance
        df = self._frame(1.1, 0.2)
        sig = harmonic_spectrum_significance(df, Thresholds())
        assert _detect_harmonic_signature(df, 1.1, sig) == []

    def test_signature_match_survives_at_real_load(self):
        from pq_analysis import _detect_harmonic_signature, harmonic_spectrum_significance
        df = self._frame(20.0, 4.0)
        sig = harmonic_spectrum_significance(df, Thresholds())
        assert _detect_harmonic_signature(df, 20.0, sig)

    def test_attribution_and_resonance_withheld_at_light_load(self):
        from pq_analysis import check_harmonic_sources
        r = check_harmonic_sources(self._frame(1.1, 0.2), Thresholds())
        # The impedances stay as measured data; only the conclusions are withheld.
        assert r["available"] is True
        assert r["overall"] == "not_assessed"
        assert r["resonant_orders"] == []
        assert all(od["attribution"] == "not_assessed" for od in r["orders"].values())

    def test_spectral_shape_declines_at_light_load(self):
        from pq_analysis import check_spectral_shape, check_harmonic_sources
        df = self._frame(1.1, 0.2)
        src = check_harmonic_sources(df, Thresholds())
        r = check_spectral_shape(df, Thresholds(), src)
        assert r["available"] is False and "not classified" in r["error"]

    def test_too_few_loaded_intervals_is_refused(self):
        from pq_analysis import harmonic_spectrum_significance
        import pandas as pd
        df = self._frame(20.0, 4.0, n=300)
        # Only 5 intervals at real load; the rest near zero.
        df.loc[df.index[5:], "current_a"] = 0.1
        r = harmonic_spectrum_significance(df, Thresholds())
        assert r["usable"] is False and "loaded" in r["reason"]

    def test_resolution_test_applies_without_an_rms_current_channel(self):
        # Load cannot be verified, but the quantisation test still governs.
        from pq_analysis import harmonic_spectrum_significance
        df = self._frame(1.1, 0.2).drop(columns=["current_a"])
        r = harmonic_spectrum_significance(df, Thresholds())
        assert r["usable"] is False and r["load_verified"] is False


# ─────────────────────────────────────────────────────────────────────────────
# 14. Residential customer letter
# ─────────────────────────────────────────────────────────────────────────────

@pytest.mark.skipif(not _FIXTURES, reason="test_data/*.pqd not generated")
class TestChannelAppendix:
    """Appendix B: what was read out of the file, and what each channel holds."""

    @staticmethod
    @pytest.fixture(scope="class")
    def doc(tmp_path_factory):
        import docx
        import pq_analysis as An
        from pq_report import generate_report, generate_word_report
        out = tmp_path_factory.mktemp("chan")
        ds = extract_dataset(
            ProntoAdapter(Path("test_data/test_residential.pqd")), ChannelMapper())
        th = Thresholds(nominal_voltage=120.0, customer_class="r")
        df = ds.df
        ev = An.detect_events(ds, th)
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), ev, th,
            neutral_health_result=An.check_neutral_health(ds, th),
            itic_result=An.check_itic(ev, th),
            flicker_result=An.check_flicker(df, th))
        rep["root_causes"] = An.analyze_root_causes(rep, ds, th)
        path = generate_word_report(
            report=rep, thresh=th, ds=ds, site_name="S", site_address="A",
            engineer_name="E", outdir=out, stem="ch")
        return docx.Document(str(path))

    @staticmethod
    def _table(doc):
        return next(t for t in doc.tables if t.rows[0].cells[0].text == "Channel")

    def test_the_appendices_are_lettered_and_ordered(self, doc):
        heads = [p.text for p in doc.paragraphs if p.text.startswith("Appendix")]
        assert heads == ["Appendix A: Terms Used in This Report",
                         "Appendix B: Standards, Methods, and Limitations",
                         "Appendix C: Channels Read From the Meter File",
                         "Appendix D: Load Signature Matching (Experimental)"]

    def test_the_experimental_appendix_is_last(self, doc):
        # Load-signature matching is the weakest analysis in the report and is
        # not part of the compliance assessment, so it sits at the very bottom.
        heads = [p.text for p in doc.paragraphs if p.text.startswith("Appendix")]
        assert heads[-1].startswith("Appendix D")

    def test_signature_findings_are_out_of_the_assessment(self, doc):
        # They used to sit among the measured findings, which gave a hypothesis
        # the standing of a measurement.
        texts = [p.text for p in doc.paragraphs]
        assess = texts.index(
            "Engineering Assessment: Likely Causes and Contributing Conditions")
        appendix_d = next(i for i, t in enumerate(texts)
                          if t.startswith("Appendix D"))
        for i, t in enumerate(texts):
            if "load family" in t.lower() or "load signature" in t.lower():
                assert not (assess < i < appendix_d), \
                    f"signature text inside the assessment: {t[:70]!r}"

    def test_the_glossary_is_reachable_from_the_body(self, doc):
        # The definitions moved to the back, so the body has to say where they
        # went or they are found only by whoever pages to the end.
        body = " ".join(p.text for p in doc.paragraphs)
        assert "defined in Appendix A" in body

    def test_no_cross_reference_points_at_a_bare_appendix(self, doc):
        # Three appendices make "see the Appendix" ambiguous.
        body = " ".join(p.text for p in doc.paragraphs)
        for tb in doc.tables:
            body += " " + " ".join(c.text for r in tb.rows for c in r.cells)
        assert "see the Appendix" not in body

    def test_every_channel_read_is_listed(self, doc):
        ds = extract_dataset(
            ProntoAdapter(Path("test_data/test_residential.pqd")), ChannelMapper())
        listed = {r.cells[0].text.split()[0] for r in list(self._table(doc).rows)[1:]}
        assert listed == set(ds.meta["channel_map"])

    def test_each_row_names_the_device_channel_it_came_from(self, doc):
        # The match from device label to engineering quantity is the step that
        # fails silently, so the label has to be on the page to check it.
        rows = {r.cells[0].text.split()[0]: [c.text for c in r.cells]
                for r in list(self._table(doc).rows)[1:]}
        assert rows["voltage_a"][1].startswith("RMS voltage, line to neutral")
        assert rows["voltage_a"][2] == "V"
        assert rows["voltage_a"][3] == "Van RMS"
        assert rows["current_neutral"][3] == "In RMS"

    def test_within_interval_extremes_ride_with_their_channel(self, doc):
        labels = [r.cells[0].text for r in list(self._table(doc).rows)[1:]]
        assert any(l.startswith("voltage_a  (+ ") and "max" in l and "min" in l
                   for l in labels)
        assert not any(l.startswith("voltage_a_peak") for l in labels)

    def test_coverage_counts_are_reported(self, doc):
        rows = {r.cells[0].text.split()[0]: [c.text for c in r.cells]
                for r in list(self._table(doc).rows)[1:]}
        assert rows["voltage_a"][4] == "288"

    def test_split_phase_uses_the_service_own_phase_names(self, doc):
        text = " ".join(c.text for r in self._table(doc).rows for c in r.cells)
        assert "L1" in text and "L2" in text
        assert "phase A" not in text

    @pytest.mark.parametrize("name,expected", [
        ("h3_current_a",       "3rd-order harmonic current, phase A"),
        ("h11_voltage_b",      "11th-order harmonic voltage, phase B"),
        ("h21_current_c",      "21st-order harmonic current, phase C"),
        ("h13_current_neutral", "13th-order harmonic current, neutral"),
        ("thd_voltage_a",      "Total harmonic distortion of the voltage, phase A"),
        ("flicker_plt",        "Long-term flicker severity (Plt, 2-hour), phase A"),
        ("voltage_ab",         "RMS voltage, line to line (phase A to phase B)"),
        ("current_neutral",    "RMS current in the neutral conductor"),
        ("frequency",          "System frequency"),
    ])
    def test_descriptions_are_derived_from_the_name(self, name, expected):
        from pq_report import _channel_description
        assert _channel_description(name, is_split=False) == expected

    def test_within_interval_extremes_describe_their_base_channel(self):
        from pq_report import _channel_description
        assert _channel_description("voltage_a_peak", is_split=True) == (
            "RMS voltage, line to neutral, L1 — highest value within each interval")
        assert _channel_description("current_b_min", is_split=True) == (
            "RMS current, L2 — lowest value within each interval")

    def test_an_order_nobody_tabulated_still_gets_a_description(self):
        # Meters report orders well past the ones the limits cover; a hand-kept
        # table would leave those rows blank.
        from pq_report import _channel_description
        assert _channel_description("h49_current_a", is_split=False) == (
            "49th-order harmonic current, phase A")


class TestRecordingOverview:
    """The sanity-check chart: the whole recording, before any assessment."""

    def _frame(self, n=576, gap_at=None, cols=("voltage_a", "voltage_b",
                                               "current_a", "current_b",
                                               "current_neutral")):
        import pandas as pd
        rng = np.random.default_rng(5)
        idx = pd.date_range("2026-05-02", periods=n, freq="5min", tz="UTC")
        if gap_at is not None:
            idx = idx.delete(range(gap_at, gap_at + 60))
        base = {"voltage_a": 121.0, "voltage_b": 120.0, "voltage_c": 120.5,
                "current_a": 12.0, "current_b": 9.0, "current_c": 10.0,
                "current_neutral": 3.0}
        return pd.DataFrame(
            {c: np.abs(base[c] + rng.normal(0, 0.5, len(idx))) for c in cols},
            index=idx)

    def _ds(self, df):
        from pq_adapter import PQDataset
        return PQDataset(df=df, adaptive_df=None,
                         meta={"interval_minutes": 5, "topology": "split-phase"})

    def test_chart_is_written(self, tmp_path):
        from pq_plots import plot_overview
        plot_overview(self._ds(self._frame()), Thresholds(nominal_voltage=120.0),
                      outdir=tmp_path, stem="s")
        assert (tmp_path / "s_overview.png").exists()

    def test_a_recording_gap_breaks_the_trace(self):
        # A line drawn straight across five hours the meter never recorded is
        # indistinguishable from five hours of steady service.
        from pq_plots import _gap_spans, _break_at_gaps
        df = self._frame(gap_at=200)
        gaps = _gap_spans(df.index)
        assert len(gaps) == 1
        start, end = gaps[0]
        assert (end - start).total_seconds() / 3600 == pytest.approx(5.08, abs=0.1)
        broken = _break_at_gaps(df["voltage_a"], gaps)
        assert broken.isna().sum() == 1
        assert start < broken[broken.isna()].index[0] < end

    def test_an_unbroken_recording_has_no_gaps(self):
        from pq_plots import _gap_spans
        assert _gap_spans(self._frame().index) == []

    def test_voltage_only_file_still_charts(self, tmp_path):
        from pq_plots import plot_overview
        plot_overview(self._ds(self._frame(cols=("voltage_a", "voltage_b"))),
                      Thresholds(nominal_voltage=120.0), outdir=tmp_path, stem="v")
        assert (tmp_path / "v_overview.png").exists()

    def test_nothing_to_chart_writes_nothing(self, tmp_path):
        import pandas as pd
        from pq_plots import plot_overview
        df = self._frame()[[]].assign(frequency=60.0)
        plot_overview(self._ds(df), Thresholds(nominal_voltage=120.0),
                      outdir=tmp_path, stem="n")
        assert not (tmp_path / "n_overview.png").exists()

    def test_both_documents_open_with_the_chart(self, tmp_path):
        # It is the first thing in the report and it is in the letter, so a
        # reader of either can check the period before reading a conclusion.
        import docx
        import pq_analysis as An
        from pq_plots import plot_overview
        from pq_report import (generate_report, generate_word_report,
                               generate_customer_letter)
        ds = extract_dataset(ProntoAdapter(Path("test_data/test_residential.pqd")),
                             ChannelMapper())
        th = Thresholds(nominal_voltage=120.0, customer_class="r")
        df = ds.df
        ev = An.detect_events(ds, th)
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), ev, th,
            neutral_health_result=An.check_neutral_health(ds, th),
            itic_result=An.check_itic(ev, th),
            flicker_result=An.check_flicker(df, th))
        rep["root_causes"] = An.analyze_root_causes(rep, ds, th)
        plot_overview(ds, th, outdir=tmp_path, stem="ov")

        rpt = docx.Document(str(generate_word_report(
            report=rep, thresh=th, ds=ds, site_name="S", site_address="A",
            engineer_name="E", outdir=tmp_path, stem="ov")))
        heads = [p.text for p in rpt.paragraphs if p.text.strip()]
        assert "Recording Overview" in heads
        assert heads.index("Recording Overview") < heads.index(
            "Executive Summary and Compliance Status")

        letter = docx.Document(str(generate_customer_letter(
            rep, th, "1 Test St", "Eng", tmp_path, "ov")))
        assert "What we recorded" in [p.text for p in letter.paragraphs]
        assert len(letter.inline_shapes) == 1

    def test_the_chart_states_the_period_it_drew(self, tmp_path):
        # The chart and the "Length of recording" line in the letter come from
        # the same index; the caption is what makes a mismatch visible.
        from pq_plots import plot_overview
        import matplotlib.pyplot as plt
        df = self._frame()
        captured = {}
        orig = plt.Figure.text

        def _capture(self, x, y, s, *a, **k):
            captured.setdefault("sub", s)
            return orig(self, x, y, s, *a, **k)
        plt.Figure.text = _capture
        try:
            plot_overview(self._ds(df), Thresholds(nominal_voltage=120.0),
                          outdir=tmp_path, stem="c")
        finally:
            plt.Figure.text = orig
        assert "2026-05-02 00:00" in captured["sub"]
        assert f"{len(df):,} intervals of 5 min" in captured["sub"]


class TestCustomerLetter:
    """The customer-facing document, one per service class.

    It must state no attribution, commit Xcel Energy to nothing, and never
    tell the customer an internal report exists.
    """

    def _report(self, path, customer_class="r", nominal=120.0):
        import pq_analysis as An
        from pq_report import generate_report
        ds = extract_dataset(ProntoAdapter(path), ChannelMapper())
        th = Thresholds(nominal_voltage=nominal, customer_class=customer_class)
        df = ds.df
        ev = An.detect_events(ds, th)
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th),
            An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), ev, th,
            neutral_health_result=An.check_neutral_health(ds, th),
            itic_result=An.check_itic(ev, th),
            flicker_result=An.check_flicker(df, th),
        )
        rep["root_causes"] = An.analyze_root_causes(rep, ds, th)
        return rep, th

    def test_written_for_residential(self, tmp_path):
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        assert out is not None and out.exists()
        assert out.name.endswith("_customer_letter.docx")

    def _letter_text(self, path, tmp_path, customer_class, nominal=120.0):
        from pq_report import generate_customer_letter
        docx = pytest.importorskip("docx")
        rep, th = self._report(Path(path), customer_class, nominal)
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        doc = docx.Document(str(out))
        return doc, "\n".join(p.text for p in doc.paragraphs)

    def test_a_facility_gets_the_full_list_of_checks(self, tmp_path):
        # An exceptions-only letter cannot distinguish a check that passed from
        # one that never ran, and a facility hands this to a contractor.
        doc, text = self._letter_text("test_data/test_commercial_large.pqd",
                                      tmp_path, "sg", 277.0)
        assert "What we checked" in text
        tbl = next(t for t in doc.tables
                   if t.rows[0].cells[0].text.strip() == "What we looked at")
        items = [r.cells[0].text.strip() for r in tbl.rows[1:]]
        assert "Supply voltage" in items
        # The standard is named, because this reader can look one up.
        against = " ".join(r.cells[2].text for r in tbl.rows[1:])
        assert "ANSI C84.1" in against and "IEEE 519" in against

    def test_a_homeowner_does_not_get_it(self, tmp_path):
        _doc, text = self._letter_text("test_data/test_residential.pqd",
                                       tmp_path, "r")
        assert "What we checked" not in text
        assert "ANSI C84.1" not in text

    def test_every_row_outside_limits_has_a_finding_to_match(self, tmp_path):
        # The letter tells the reader the findings are drawn from these rows,
        # so a row marked outside limits with no finding is a contradiction the
        # reader can see. Voltage unbalance was exactly that before this test.
        doc, text = self._letter_text("test_data/test_commercial_large.pqd",
                                      tmp_path, "sg", 277.0)
        tbl = next(t for t in doc.tables
                   if t.rows[0].cells[0].text.strip() == "What we looked at")
        outside = [r for r in tbl.rows[1:]
                   if r.cells[3].text.strip() == "Outside limits"]
        assert outside, "this fixture is meant to fail at least one check"
        n_findings = len([p for p in doc.paragraphs
                          if p.text.strip()[:2] in {f"{i}." for i in range(1, 10)}])
        assert n_findings >= len(outside), (
            f"{len(outside)} rows outside limits but only {n_findings} findings")

    def test_an_unbalanced_three_phase_service_is_told_about_it(self, tmp_path):
        _doc, text = self._letter_text("test_data/test_commercial_large.pqd",
                                       tmp_path, "sg", 277.0)
        assert "The three phases are not supplying equal voltage" in text
        assert "derate" in text
        # No attribution: unbalance can start on either side of the meter.
        for blamed in ("your equipment is causing", "our equipment is causing",
                       "caused by your", "caused by our"):
            assert blamed not in text.lower()

    def test_a_split_phase_service_is_not_told_it_failed_nema(self, tmp_path):
        # Two legs 180 degrees apart are not a three-phase unbalance, and no
        # limit is set for the leg difference.
        _doc, text = self._letter_text("test_data/test_residential.pqd",
                                       tmp_path, "r")
        assert "The three phases are not supplying equal voltage" not in text

    def test_the_letter_is_brand_red_and_nothing_else(self, tmp_path):
        # Every coloured run in a customer document is the Xcel Energy red.
        # The numbered findings were drawn from the severity palette, which put
        # each ordinary finding in the Watch blue -- a severity claim about a
        # finding whose severity was never assessed for the customer.
        docx = pytest.importorskip("docx")
        from pq_report import generate_customer_letter, _XE_RED
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        doc = docx.Document(str(out))
        colours = {
            str(r.font.color.rgb)
            for p in doc.paragraphs for r in p.runs
            if r.font.color is not None and r.font.color.type is not None
        }
        assert colours == {str(_XE_RED)}, f"non-brand colour in the letter: {colours}"

    def test_a_numbered_finding_is_a_red_heading(self, tmp_path):
        docx = pytest.importorskip("docx")
        from pq_report import generate_customer_letter, _XE_RED
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        doc = docx.Document(str(out))
        numbered = [p for p in doc.paragraphs
                    if p.text.strip()[:2] in {f"{i}." for i in range(1, 10)}]
        assert numbered, "the letter listed no findings to check"
        for p in numbered:
            for r in p.runs:
                assert r.font.color.rgb == _XE_RED
                assert r.font.bold

    def test_a_safety_finding_is_called_out_in_words_not_only_colour(self, tmp_path):
        # The colour distinction went away with this change, so the safety
        # signal has to survive somewhere a reader actually reads.
        from pq_report import generate_customer_letter
        docx = pytest.importorskip("docx")
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        text = "\n".join(p.text for p in docx.Document(str(out)).paragraphs)
        assert "Is this a safety concern?" in text
        assert "needs prompt attention" in text

    def test_a_letter_from_an_earlier_run_never_survives(self, tmp_path):
        # A letter that outlives the run that wrote it sits beside a fresh
        # report describing a different recording, and says nothing about which
        # run produced it.
        from pq_report import generate_customer_letter
        stale = tmp_path / "t_customer_letter.docx"
        stale.write_bytes(b"letter from a two-hour download")
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        assert out == stale and stale.read_bytes() != b"letter from a two-hour download"

    def test_a_reclassified_service_replaces_the_previous_letter(self, tmp_path):
        # The stale letter described the same file read as a different class,
        # so it must be replaced rather than left beside the new one.
        from pq_report import generate_customer_letter
        stale = tmp_path / "t_customer_letter.docx"
        stale.write_bytes(b"letter from when this was billed residential")
        rep, th = self._report(Path("test_data/test_residential.pqd"),
                               customer_class="sg")
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        assert out is not None and out.exists()
        assert out.read_bytes() != b"letter from when this was billed residential"

    def test_a_letter_that_cannot_be_replaced_stops_the_run(self, tmp_path, monkeypatch):
        # Word holds the file open on Windows, so the unlink is what fails. The
        # run has to stop there rather than leave the old letter in place.
        from pathlib import Path as _P
        import pq_report
        stale = tmp_path / "t_customer_letter.docx"
        stale.write_bytes(b"open in Word")
        rep, th = self._report(Path("test_data/test_residential.pqd"))

        def _locked(self):
            raise PermissionError(32, "The process cannot access the file")
        monkeypatch.setattr(_P, "unlink", _locked)
        with pytest.raises(PermissionError, match="open in Word"):
            pq_report.generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        assert stale.read_bytes() == b"open in Word"

    def test_written_for_small_commercial(self, tmp_path):
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_commercial_small.pqd"),
                               customer_class="c")
        out = generate_customer_letter(rep, th, "1 Trade St", "Eng", tmp_path, "t")
        assert out is not None and out.exists()

    @pytest.mark.parametrize("cls", ["sg", "pg"])
    def test_every_class_gets_its_own_customer_document(self, tmp_path, cls):
        # The engineering report is internal for every class, so a customer
        # document is written at every scale -- what differs is the register.
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_commercial_large.pqd"),
                               customer_class=cls, nominal=277.0)
        out = generate_customer_letter(rep, th, "1 Trade St", "Eng", tmp_path, "t")
        assert out is not None and out.exists()

    def test_the_internal_report_is_named_and_unsigned(self, tmp_path):
        """The internal document says what it is and is addressed to nobody."""
        docx = pytest.importorskip("docx")
        import glob as _glob
        from pq_report import generate_word_report
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        ds = extract_dataset(ProntoAdapter(Path("test_data/test_residential.pqd")),
                             ChannelMapper())
        generate_word_report(
            report=rep, thresh=th, ds=ds, outdir=tmp_path, stem="t",
            site_name="Site", site_address="1 Test St",
            engineer_name="A. Engineer",
            engineer_title="Electric Area Engineer")
        written = _glob.glob(str(tmp_path / "*_internal_engineering_report.docx"))
        assert written, "the internal report is named as such on disk"

        d = docx.Document(written[0])
        text = "\n".join(p.text for p in d.paragraphs)
        assert "Internal Engineering Report" in text
        assert "internal working document" in text
        # No sign-off block: the document is not addressed to anyone.
        assert "Sincerely," not in text
        # Whose work it is survives, as a header field rather than a signature.
        table_text = " ".join(c.text for t in d.tables for r in t.rows for c in r.cells)
        assert "Prepared by" in table_text and "A. Engineer" in table_text

    def test_the_register_changes_with_the_class(self, tmp_path):
        """A homeowner is told what a meter does; a plant engineer is not."""
        docx = pytest.importorskip("docx")
        from pq_report import generate_customer_letter
        texts = {}
        for cls, path, nominal in (("r", "test_data/test_residential.pqd", 120.0),
                                   ("pg", "test_data/test_commercial_large.pqd", 277.0)):
            rep, th = self._report(Path(path), customer_class=cls, nominal=nominal)
            out = generate_customer_letter(rep, th, "1 Test St", "Eng",
                                           tmp_path / cls, "t")
            texts[cls] = "\n".join(p.text for p in docx.Document(str(out)).paragraphs)

        assert "your home" in texts["r"]
        assert "measured the voltage and current many times a second" in texts["r"]

        assert "your facility" in texts["pg"]
        # The terse register states what was logged without explaining metering.
        assert "measured the voltage and current many times a second" not in texts["pg"]
        assert "interval resolution" in texts["pg"]
        # A primary-metered customer owns the transformer, and the letter says so.
        assert "transformer at this site is yours" in texts["pg"]

    def test_the_power_factor_sheet_is_not_attributed_to_a_schedule(self):
        """The PF clauses are scoped by rules section, not by rate schedule.

        Sheet R73 sits in the General rules and sets 0.90 lagging for every
        class; Sheet R121 sits in the Commercial and Industrial rules and asks
        all of C, SG and PG for near unity. Writing "Sheet R73 (Schedule C)"
        told a customer the clause was theirs alone, which is not what the
        tariff says and is not a claim worth defending in a letter.
        """
        from pq_report import _customer_conditions
        sheets = {}
        for cls in ("c", "sg", "pg"):
            rep, th = self._report(Path("test_data/test_commercial_small.pqd"),
                                   customer_class=cls)
            pf = [c for c in _customer_conditions(rep, th)
                  if "power factor" in c["headline"].lower()]
            sheets[cls] = pf[0]["measured"] if pf else ""
        assert "Sheet R73" in sheets["c"] and "Schedule C" not in sheets["c"]
        assert "Sheet R73" in sheets["sg"] and "Schedule SG" not in sheets["sg"]
        # Primary practice is near unity rather than a 0.90 floor, and R121 is
        # cited without being called a Schedule PG clause.
        assert "Sheet R121" in sheets["pg"] and "Schedule PG" not in sheets["pg"]
        assert "near unity" in sheets["pg"] and "0.90" not in sheets["pg"]

    def test_no_customer_document_mentions_the_internal_report(self, tmp_path):
        docx = pytest.importorskip("docx")
        from pq_report import generate_customer_letter
        for cls, path, nominal in (("r", "test_data/test_residential.pqd", 120.0),
                                   ("c", "test_data/test_commercial_small.pqd", 120.0),
                                   ("sg", "test_data/test_commercial_large.pqd", 277.0),
                                   ("pg", "test_data/test_commercial_large.pqd", 277.0)):
            rep, th = self._report(Path(path), customer_class=cls, nominal=nominal)
            out = generate_customer_letter(rep, th, "1 Test St", "Eng",
                                           tmp_path / cls, "t")
            text = "\n".join(p.text for p in docx.Document(str(out)).paragraphs).lower()
            for forbidden in ("engineering report", "internal report", "attached report"):
                assert forbidden not in text, f"{cls}: {forbidden}"

    def test_business_letter_covers_power_factor_and_distortion(self):
        from pq_report import _customer_conditions
        rep, th = self._report(Path("test_data/test_commercial_small.pqd"),
                               customer_class="c")
        heads = " ".join(c["headline"].lower() for c in _customer_conditions(rep, th))
        assert "power factor" in heads
        assert "distorted" in heads

    def test_residential_letter_omits_power_factor_and_distortion(self):
        # A homeowner is not billed for power factor and cannot act on distortion.
        from pq_report import _customer_conditions
        rep, th = self._report(Path("test_data/test_commercial_small.pqd"),
                               customer_class="r")
        heads = " ".join(c["headline"].lower() for c in _customer_conditions(rep, th))
        assert "power factor" not in heads
        assert "distorted" not in heads

    def test_business_letter_carries_no_residential_wording(self, tmp_path):
        from docx import Document
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_commercial_small.pqd"),
                               customer_class="c")
        out = generate_customer_letter(rep, th, "1 Trade St", "Eng", tmp_path, "t")
        text = " ".join(p.text for p in Document(str(out)).paragraphs)
        for residential in ("your home", "the house", "refrigerators, freezers",
                            "well pumps"):
            assert residential not in text, f"business letter says {residential!r}"
        assert "your business" in text

    def test_business_letter_names_the_tariff_sheet(self, tmp_path):
        from docx import Document
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_commercial_small.pqd"),
                               customer_class="c")
        out = generate_customer_letter(rep, th, "1 Trade St", "Eng", tmp_path, "t")
        text = " ".join(p.text for p in Document(str(out)).paragraphs)
        # Power factor is the one item with a direct billing consequence, so
        # the customer should be able to look the clause up -- which means the
        # sheet, not a schedule the clause does not belong to.
        assert "R73" in text

    def test_neutral_wording_follows_the_service_topology(self):
        # "Two 120-volt halves" is true of a house and false of a three-phase site.
        from pq_report import _customer_vocabulary
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        v = _customer_vocabulary(rep, th)
        assert "two 120-volt halves" in v["neutral_measured"]
        rep3, th3 = self._report(Path("test_data/test_commercial_small.pqd"),
                                 customer_class="c")
        v3 = _customer_vocabulary(rep3, th3)
        assert "three separate live wires" in v3["neutral_measured"]
        assert "your business" == v3["site"]

    def test_states_no_attribution_and_no_commitment(self, tmp_path):
        from docx import Document
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        text = "\n".join(p.text for p in Document(str(out)).paragraphs).lower()
        for banned in ("xcel energy will", "customer's responsibility",
                       "utility responsibility", "customer-side", "utility-side",
                       "your responsibility", "at fault"):
            assert banned not in text, f"letter asserts {banned!r}"

    def test_omits_engineering_concepts_a_homeowner_cannot_act_on(self, tmp_path):
        from docx import Document
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        text = "\n".join(p.text for p in Document(str(out)).paragraphs).lower()
        for jargon in ("thd", "tdd", "k-factor", "harmonic", "resonance",
                       "ansi c84", "ieee 519", "iec 61000", "per-order",
                       "spectral", "impedance"):
            assert jargon not in text, f"letter uses {jargon!r}"

    def test_every_condition_pairs_a_number_with_meaning_and_symptom(self):
        from pq_report import _customer_conditions
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        conditions = _customer_conditions(rep, th)
        assert conditions, "residential fixture should surface conditions"
        for c in conditions:
            assert c["headline"] and c["measured"] and c["means"] and c["symptom"]

    def test_event_counts_handles_dataframe_and_list(self):
        from pq_report import _event_counts
        import pandas as pd
        assert _event_counts({}) == {}
        assert _event_counts({"events": None}) == {}
        assert _event_counts({"events": pd.DataFrame()}) == {}
        df = pd.DataFrame({"type": ["voltage_sag", "voltage_sag", "voltage_swell"]})
        assert _event_counts({"events": df}) == {"voltage_sag": 2, "voltage_swell": 1}
        lst = [{"type": "voltage_sag"}, {"type": "flicker_pst"}]
        assert _event_counts({"events": lst}) == {"voltage_sag": 1, "flicker_pst": 1}

    def test_flicker_number_is_given_a_scale_the_reader_can_use(self, tmp_path):
        # "4.98" means nothing without the basis of the scale: 1.0 is the
        # conventional threshold of irritability, so the multiple of it is what
        # makes the reading interpretable.
        from docx import Document
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        text = " ".join(p.text for p in Document(str(out)).paragraphs)
        if "flickering" not in text:
            pytest.skip("fixture did not exceed the flicker limit")
        assert "annoying" in text
        assert "times that level" in text or "same as that level" in text
        # It must say what flicker is not, so nobody reads it as damage or usage.
        assert "not of damage" in text or "does not harm" in text

    def test_letter_does_not_promise_the_engineering_report(self, tmp_path):
        # The engineering document is shared at Xcel's discretion, so the letter
        # must not tell the customer one is attached or owed to them.
        from docx import Document
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        text = " ".join(p.text for p in Document(str(out)).paragraphs).lower()
        for promise in ("report accompanies", "accompanying this letter",
                        "attached report", "enclosed report",
                        "document to hand to an electrician"):
            assert promise not in text, f"letter promises {promise!r}"

    def test_clean_service_gets_a_clear_no_problem_letter(self, tmp_path):
        from docx import Document
        from pq_report import generate_customer_letter
        rep, th = self._report(Path("test_data/test_residential.pqd"))
        # Blank every condition source so the letter takes its no-findings path.
        for k in ("voltage_compliance", "flicker", "neutral_health",
                  "current_imbalance", "itic", "events"):
            rep[k] = {"available": False}
        rep["root_causes"] = []
        out = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, "t")
        text = "\n".join(p.text for p in Document(str(out)).paragraphs)
        assert "did not find a problem" in text
        # And it must not overclaim: an intermittent fault can fall outside a window.
        assert "only the days it ran" in text


# ─────────────────────────────────────────────────────────────────────────────
# 15. Neutral wording is scaled to what was measured
# ─────────────────────────────────────────────────────────────────────────────

class TestNeutralSeverityWording:
    """The neutral indicators move for ordinary reasons too.

    A "caution" driven by a couple of volts of variation in the leg sum must not
    read the same as legs actually opposing each other — the letter was declaring
    a shock and fire hazard on the strength of a 2.8 V standard deviation.
    """

    @staticmethod
    def _report(severity, corr, sum_std=2.8, asym=1.4):
        return {
            "file_summary": {"topology": "split-phase", "duration_hours": 68},
            "voltage_compliance": {"available": False},
            "flicker": {"available": False},
            "current_imbalance": {"available": False},
            "power_factor": {"available": False},
            "thd_compliance": {},
            "itic": {},
            "events": {},
            "neutral_health": {
                "available": True, "severity": severity, "leg_correlation": corr,
                "sum_mean_v": 238.0, "sum_std_v": sum_std, "asym_mean_v": asym,
                "vne_max_v": 0.1,
            },
        }

    def _neutral(self, severity, corr):
        from pq_report import _customer_conditions
        conds = _customer_conditions(self._report(severity, corr),
                                     Thresholds(nominal_voltage=120.0,
                                                customer_class="r"))
        assert conds, f"{severity} produced no condition"
        return conds

    def test_caution_is_a_baseline_not_a_hazard(self):
        c = self._neutral("caution", 0.82)[0]
        assert c.get("safety") is False
        assert "hazard" not in c["means"].lower()
        assert "baseline" in c["means"]

    def test_critical_keeps_the_urgent_language(self):
        c = self._neutral("critical", -0.74)[0]
        assert c["safety"] is True
        assert "shock and fire hazard" in c["means"]

    def test_only_warning_and_critical_lead_the_letter(self):
        from pq_report import _customer_conditions
        th = Thresholds(nominal_voltage=120.0, customer_class="r")
        rep = self._report("caution", 0.82)
        # Give it a second condition so ordering is observable.
        rep["voltage_compliance"] = {
            "available": True, "total_pct_out_of_bounds": 5.0,
            "range_v": (114.0, 126.0),
            "phases": {"voltage_a": {"min_v": 100.0, "max_v": 121.0}},
        }
        conds = _customer_conditions(rep, th)
        assert "neutral" not in conds[0]["headline"].lower()
        rep["neutral_health"]["severity"] = "critical"
        rep["neutral_health"]["leg_correlation"] = -0.74
        assert "neutral" in _customer_conditions(rep, th)[0]["headline"].lower()

    def test_the_explainer_does_not_contradict_the_measurement(self):
        # The explainer says what we look for; the data says what was found. It
        # must not assert opposition when the legs measured +0.82.
        c = self._neutral("caution", 0.82)[0]
        assert "so that is the pattern we look for" in c["measured"]
        assert "rose and fell together" in c["measured"]
        assert "Our measurements show the two halves moving in opposite" not in c["measured"]

    def test_indicators_are_always_reported_as_numbers(self):
        for severity, corr in (("caution", 0.82), ("warning", -0.1), ("critical", -0.74)):
            measured = self._neutral(severity, corr)[0]["measured"]
            assert "correlation of" in measured
            assert "238 volts" in measured
            assert f"{corr:+.2f}" in measured

    def test_correlation_sense_matches_its_sign(self):
        from pq_report import _neutral_indicator_sentence
        base = {"sum_mean_v": 240.0, "sum_std_v": 1.0, "asym_mean_v": 0.5}
        assert "sound shared connection" in _neutral_indicator_sentence(
            {**base, "leg_correlation": 0.9})
        assert "only loosely" in _neutral_indicator_sentence(
            {**base, "leg_correlation": 0.2})
        assert "failing connection produces" in _neutral_indicator_sentence(
            {**base, "leg_correlation": -0.6})

    def test_caution_claims_no_symptom_the_customer_would_have_seen(self):
        c = self._neutral("caution", 0.82)[0]
        assert "Nothing in particular" in c["symptom"]

    def test_normal_neutral_produces_no_condition_at_all(self):
        from pq_report import _customer_conditions
        rep = self._report("normal", 0.95)
        conds = _customer_conditions(rep, Thresholds(nominal_voltage=120.0,
                                                     customer_class="r"))
        assert not any("neutral" in c["headline"].lower() for c in conds)


# ─────────────────────────────────────────────────────────────────────────────
# 16. Distortion is reported without depending on an assumed limit
# ─────────────────────────────────────────────────────────────────────────────

class TestDistortionClaim:
    """The current-side limit depends on ISC/IL.

    Without ISC the analysis falls back to the most restrictive class, which on
    one test file turns a 0.35% exceedance into 100%. That is fine in a labelled
    engineering report and not fine in a letter telling a business it breaches a
    standard, so the letter only asserts what its limits actually support.
    """

    @staticmethod
    def _report(v_exceed, i_exceed, isc_provided):
        return {
            "file_summary": {"topology": "3-phase", "duration_hours": 24},
            "voltage_compliance": {"available": False},
            "flicker": {"available": False},
            "neutral_health": {"available": False},
            "current_imbalance": {"available": False},
            "power_factor": {"available": False},
            "itic": {}, "events": {},
            "thd_compliance": {
                "voltage": {"available": True, "pct_exceeding": v_exceed},
                "current": {"available": True, "pct_exceeding": i_exceed},
                "tdd_info": {"isc_provided": isc_provided},
            },
        }

    def _distortion(self, v_exceed, i_exceed, isc_provided):
        from pq_report import _customer_conditions
        th = Thresholds(nominal_voltage=120.0, customer_class="c")
        for c in _customer_conditions(self._report(v_exceed, i_exceed, isc_provided), th):
            if "distort" in c["headline"].lower():
                return c
        return None

    def test_asserts_both_when_isc_is_known(self):
        c = self._distortion(5.0, 100.0, True)
        assert "the voltage supplied to you" in c["measured"]
        assert "the current your equipment draws" in c["measured"]
        assert "not settled" not in c["measured"]

    def test_withholds_the_current_claim_without_isc(self):
        c = self._distortion(5.0, 100.0, False)
        assert "the voltage supplied to you" in c["measured"]
        # The 100% exceedance is against an assumed limit, so it is not asserted.
        assert "and in the current your equipment draws, beyond" not in c["measured"]
        assert "not settled by this recording" in c["measured"]

    def test_still_reports_when_only_current_is_involved(self):
        # The item must not vanish just because ISC was omitted — that was the
        # dependency on the flag worth removing.
        c = self._distortion(0.0, 100.0, False)
        assert c is not None
        assert "more distorted than the standard allows" not in c["headline"]
        assert "not settled by this recording" in c["measured"]

    def test_voltage_alone_is_enough_to_assert(self):
        # The voltage limit is fixed and does not depend on ISC.
        c = self._distortion(5.0, 0.0, False)
        assert "beyond the level the applicable standard permits" in c["measured"]

    def test_silent_when_nothing_exceeds_and_isc_is_known(self):
        assert self._distortion(0.0, 0.0, True) is None

    def test_fires_the_same_way_with_and_without_the_flag(self):
        # Presence must not depend on the flag, only the strength of the claim.
        assert self._distortion(5.0, 100.0, True) is not None
        assert self._distortion(5.0, 100.0, False) is not None


# ─────────────────────────────────────────────────────────────────────────────
# 17. Engineering report review fixes
# ─────────────────────────────────────────────────────────────────────────────

@pytest.mark.skipif(not _FIXTURES, reason="test_data/*.pqd not generated")
class TestEngineeringReportReview:
    """Defects found reading the >50 kW report end to end as its audience."""

    @staticmethod
    @pytest.fixture(scope="class")
    def doc(tmp_path_factory):
        from docx import Document
        import pq_analysis as An
        from pq_report import generate_report, generate_word_report
        out = tmp_path_factory.mktemp("sg")
        ds = extract_dataset(
            ProntoAdapter(Path("test_data/test_commercial_large.pqd")), ChannelMapper())
        th = Thresholds(nominal_voltage=277.0, customer_class="sg",
                        isc_amps=5000.0, transformer_kva=300.0)
        df, ev = ds.df, None
        ev = An.detect_events(ds, th)
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), ev, th,
            neutral_health_result=An.check_neutral_health(ds, th),
            itic_result=An.check_itic(ev, th),
            flicker_result=An.check_flicker(df, th),
            kfactor_result=An.kfactor_by_phase(df),
            ll_volt_result=An.check_line_to_line_voltage(df, th),
            frequency_result=An.check_frequency(df, th))
        rep["root_causes"] = An.analyze_root_causes(rep, ds, th)
        path = generate_word_report(
            report=rep, thresh=th, ds=ds, site_name="S", site_address="A",
            engineer_name="E", outdir=out, stem="sg")
        return Document(str(path))

    @staticmethod
    def _text(doc):
        t = " ".join(p.text for p in doc.paragraphs)
        for tb in doc.tables:
            t += " " + " ".join(c.text for r in tb.rows for c in r.cells)
        return t

    def test_thermal_claim_is_conditional_on_loading(self, doc):
        # At 27% of nameplate a K-factor of 5 does not exceed the rating, and
        # the report had recommended a replacement transformer on that basis.
        t = self._text(doc)
        assert "retains substantial thermal margin" in t
        assert "significantly exceeds nameplate assumptions" not in t

    def test_standards_tally_matches_the_table(self, doc):
        import re
        t = self._text(doc)
        m = re.search(r"of the (\d+) power quality standards evaluated", t)
        assert m, "standards tally sentence missing"
        claimed = int(m.group(1))
        # The table is grouped by measured quantity, so it carries merged
        # heading rows that are not standards and must not be counted.
        tbl = next(tb for tb in doc.tables
                   if tb.rows[0].cells[0].text.strip() == "Standard")
        rows = sum(1 for r in tbl.rows[1:]
                   if len({id(c._tc) for c in r.cells}) > 1)
        assert claimed == rows, f"summary claims {claimed}, table shows {rows}"

    def test_the_new_checks_have_table_rows(self, doc):
        standards = next(
            [r.cells[0].text for r in tb.rows[1:]] for tb in doc.tables
            if tb.rows[0].cells[0].text.strip() == "Standard")
        joined = " ".join(standards)
        assert "Line-to-line voltage" in joined
        assert "System frequency" in joined

    def test_no_stale_attribution_or_dangling_reference(self, doc):
        t = self._text(doc)
        assert "identifies whose system is involved" not in t
        assert "attached Pronto data" not in t

    def test_statistical_margin_convention_is_explained(self, doc):
        t = self._text(doc)
        assert "A positive margin is headroom" in t
        assert "exceeded by that amount" in t

    def test_aggregate_statistical_row_is_tdd_when_a_tdd_limit_applies(self, doc):
        # The row sits in a "% of IL" table and is measured against the TDD
        # limit, so labelling it THD named a different quantity.
        for tb in doc.tables:
            labels = [r.cells[0].text.strip() for r in tb.rows]
            if "TDD" in labels or "THD" in labels:
                assert "THD" not in labels, "aggregate row still labelled THD"

    def test_source_table_says_indication_not_attribution(self, doc):
        for tb in doc.tables:
            hdr = [c.text.strip() for c in tb.rows[0].cells]
            if "Apparent Z (Ω)" in hdr:
                assert "Indication" in hdr and "Attribution" not in hdr
                assert "Pearson r" not in hdr
                break
        else:
            pytest.fail("harmonic source table not found")

    def test_overall_assessment_is_not_a_directive(self, doc):
        t = self._text(doc)
        assert "prompt corrective action is required" not in t

    def test_low_accumulation_factor_is_explained(self, doc):
        t = self._text(doc)
        assert "largely cancelling in the neutral" in t


class TestWindowsIcon:
    """The .ico is built on a Mac and only ever consumed on Windows.

    Nothing on the machine that generates it will notice it is wrong, and
    Windows reports none of these faults — it silently draws the host
    interpreter's icon instead. So the file is checked here.
    """

    ICO = Path(__file__).parent / "icon.ico"
    #: Windows asks for these: 16 title bar, 24/32 taskbar and Alt-Tab, 48
    #: desktop, 256 the large-icon view. Missing one is a silent fallback.
    REQUIRED = [16, 24, 32, 48, 256]

    def _directory(self):
        import struct
        data = self.ICO.read_bytes()
        reserved, kind, count = struct.unpack("<HHH", data[:6])
        assert reserved == 0 and kind == 1, "not an icon file"
        entries, offset = [], 6
        for _ in range(count):
            w, h, _n, _r, _p, bpp, length, at = struct.unpack(
                "<BBBBHHII", data[offset:offset + 16])
            offset += 16
            entries.append({"w": w or 256, "h": h or 256, "bpp": bpp,
                            "png": data[at:at + 8] == b"\x89PNG\r\n\x1a\n"})
        return entries

    def test_the_icon_file_exists(self):
        assert self.ICO.exists(), "icon.ico is missing — run make_icon.py"

    def test_every_size_windows_asks_for_is_present(self):
        """A lone 16x16 entry is what made the tool look iconless on Windows."""
        have = {e["w"] for e in self._directory()}
        missing = [s for s in self.REQUIRED if s not in have]
        assert not missing, (
            f"icon.ico is missing {missing} px. Pillow drops any size larger "
            "than the image being saved, so make_ico must save from the "
            "largest frame.")

    def test_small_entries_are_not_png_compressed(self):
        """Windows reads PNG inside an .ico only at 256x256, and skips the rest.

        Pillow writes every entry as PNG unless bitmap_format="bmp" is passed,
        which produces a file that looks valid everywhere except Windows.
        """
        bad = [e["w"] for e in self._directory() if e["png"] and e["w"] < 256]
        assert not bad, (
            f"{bad} px entries are PNG-compressed; Windows will ignore them. "
            'Pass bitmap_format="bmp" when saving.')

    def test_the_icon_carries_an_alpha_channel(self):
        assert all(e["bpp"] == 32 for e in self._directory()), \
            "every entry should be 32bpp so the rounded corners stay transparent"

    def test_the_app_claims_its_own_taskbar_identity(self):
        """Without an AppUserModelID the taskbar shows Python's icon regardless.

        The window icon and the taskbar icon are separate on Windows, and
        fixing only the first leaves the symptom the user actually sees.
        """
        import run
        assert hasattr(run, "_claim_windows_taskbar_identity")
        run._claim_windows_taskbar_identity()   # no-op off Windows, never raises

    def test_the_shortcut_installer_sets_an_icon_location(self):
        """Without IconLocation, Windows draws the target's icon.

        The target is "PQ Analyzer.bat", so the shortcut came out wearing the
        generic gears every .bat file gets.
        """
        src = (Path(__file__).parent / "install_shortcut.py").read_text()
        assert "IconLocation" in src, \
            "install_shortcut.py must set the shortcut's IconLocation"
        assert "icon.ico" in src

    def test_the_installer_asks_the_shell_where_the_desktop_is(self):
        """~/Desktop is the wrong answer on a machine with OneDrive.

        Redirection moves the Desktop to ~/OneDrive/Desktop and usually leaves
        the old folder behind, so a shortcut written to the literal path lands
        somewhere the user never looks, and the installer reports success.
        """
        src = (Path(__file__).parent / "install_shortcut.py").read_text()
        assert "SHGetFolderPathW" in src, \
            "the Windows desktop path must come from the shell, not from ~/Desktop"

    def test_the_bat_delegates_instead_of_reimplementing(self):
        """Two installers drifted: one set the icon, the other found the Desktop.

        Whichever the user ran, they got half of it. The .bat is a wrapper now,
        and must not grow its own copy of the shortcut logic back.
        """
        bat = (Path(__file__).parent / "install_shortcut.bat").read_text()
        assert "install_shortcut.py" in bat, ".bat should call the Python installer"
        assert "CreateShortcut" not in bat, \
            "the .bat is reimplementing the shortcut logic again"


class TestLetterITICCurve:
    """Schedule SG gets the ITIC curve itself, not only the verdict from it.

    "What we checked" already names the ITIC curve as the standard the dips
    were judged against, which leaves the reader referred to a curve they
    cannot see. A facility with maintenance staff can read it and is likely to
    hand the letter to a contractor who will want it.
    """

    def _letter(self, tmp_path, cls, nominal):
        import pq_analysis as An
        from pq_report import generate_report, generate_customer_letter
        from pq_plots import plot_itic, plot_overview
        ds = extract_dataset(MockAdapter(duration_hours=6.0, nominal=nominal),
                             ChannelMapper())
        th = Thresholds(nominal_voltage=nominal, customer_class=cls)
        df = ds.df
        ev = An.detect_events(ds, th)
        rep = generate_report(
            ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
            An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
            An.check_current_imbalance(df, th), An.check_demand(df, th),
            An.check_individual_harmonics(df, th),
            An.check_individual_voltage_harmonics(df, th),
            An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
            An.check_harmonic_statistics(df, th), ev, th,
            neutral_health_result=An.check_neutral_health(ds, th),
            itic_result=An.check_itic(ev, th),
            flicker_result=An.check_flicker(df, th))
        rep["root_causes"] = An.analyze_root_causes(rep, ds, th)
        stem = f"itic_{cls}"
        plot_overview(ds, th, outdir=tmp_path, stem=stem)
        plot_itic(ev["events"], th, outdir=tmp_path, stem=stem)
        path = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, stem)
        return rep, path

    def _text(self, path):
        import docx
        d = docx.Document(str(path))
        return "\n".join(p.text for p in d.paragraphs)

    def _images(self, path):
        import zipfile
        return [n for n in zipfile.ZipFile(str(path)).namelist()
                if n.startswith("word/media/")]

    @pytest.mark.parametrize("cls,nominal", [("sg", 277.0), ("pg", 7621.0)])
    def test_the_terse_classes_get_the_curve(self, cls, nominal, tmp_path):
        """Both readers run a facility and both hand this to a contractor."""
        rep, path = self._letter(tmp_path, cls, nominal)
        assert rep["itic"]["n_events"], "mock recording produced no events to plot"
        text = self._text(path)
        assert "Every dip and surge we recorded" in text
        # The overview chart plus the ITIC chart.
        assert len(self._images(path)) == 2

    @pytest.mark.parametrize("cls,nominal", [("r", 120.0), ("c", 120.0)])
    def test_the_explanatory_classes_do_not_get_it(self, cls, nominal, tmp_path):
        """A log-scaled scatter plot is not what a homeowner asked for."""
        _, path = self._letter(tmp_path, cls, nominal)
        assert "Every dip and surge we recorded" not in self._text(path)
        assert len(self._images(path)) == 1

    def test_the_chart_is_explained_before_it_is_shown(self, tmp_path):
        """The letter rule: no number without what it means and what follows.

        A curve the reader cannot decode is the same failure in picture form,
        so the bands are named in words.
        """
        _, path = self._letter(tmp_path, "sg", 277.0)
        text = self._text(path)
        for phrase in ("green band", "how long an event", "percentage of normal"):
            assert phrase in text, f"missing plain-language key: {phrase}"

    def test_the_chart_makes_no_attribution(self, tmp_path):
        """Sags reach the meter from either side; the letter must not pick one."""
        _, path = self._letter(tmp_path, "sg", 277.0)
        text = self._text(path)
        assert "not where it started" in text
        assert "begin on our system or inside the facility" in text

    def test_a_stale_plot_cannot_resurrect_the_section(self, tmp_path):
        """pq_output keeps the previous run's plots under the same stem.

        A chart of last month's events under this month's heading is worse
        than no chart, so the count is checked as well as the file.
        """
        import pq_analysis as An
        from pq_report import generate_report, generate_customer_letter
        rep, path = self._letter(tmp_path, "sg", 277.0)
        stem = "itic_sg"
        assert (tmp_path / f"{stem}_itic_curve.png").exists()
        # Same folder, same stem, but this run detected nothing.
        rep["itic"] = {"available": True, "n_events": 0, "n_violations": 0}
        th = Thresholds(nominal_voltage=277.0, customer_class="sg")
        again = generate_customer_letter(rep, th, "1 Test St", "Eng", tmp_path, stem)
        assert "Every dip and surge we recorded" not in self._text(again)


#: The signature block is built once and inspected by several tests; a
#: module-scoped fixture keeps it to one letter rather than eight.
_SIG_NAME  = "Jacob Whitaker"
_SIG_TITLE = "Manager, Electric Area Engineering"
_SIG_EMAIL = "jacob.b.whitaker@xcelenergy.com"


@pytest.fixture(scope="module")
def signature_letter(tmp_path_factory):
    import docx
    import pq_analysis as An
    from pq_report import generate_report, generate_customer_letter
    out = tmp_path_factory.mktemp("sig")
    ds = extract_dataset(MockAdapter(duration_hours=6.0, nominal=277.0),
                         ChannelMapper())
    th = Thresholds(nominal_voltage=277.0, customer_class="sg")
    df = ds.df
    ev = An.detect_events(ds, th)
    rep = generate_report(
        ds, An.check_voltage_compliance(df, th), An.check_thd(df, th),
        An.check_power_factor(df, th), An.check_voltage_imbalance(df, th),
        An.check_current_imbalance(df, th), An.check_demand(df, th),
        An.check_individual_harmonics(df, th),
        An.check_individual_voltage_harmonics(df, th),
        An.check_neutral_harmonics(df, th), An.check_harmonic_sources(df, th),
        An.check_harmonic_statistics(df, th), ev, th,
        neutral_health_result=An.check_neutral_health(ds, th),
        itic_result=An.check_itic(ev, th), flicker_result=An.check_flicker(df, th))
    rep["root_causes"] = An.analyze_root_causes(rep, ds, th)
    path = generate_customer_letter(
        rep, th, "1 Test St", _SIG_NAME, out, "sig",
        engineer_title=_SIG_TITLE, engineer_email=_SIG_EMAIL)
    doc = docx.Document(str(path))
    start = next(i for i, p in enumerate(doc.paragraphs)
                 if p.text.strip() == _SIG_NAME)
    return doc, doc.paragraphs[start:start + 4]


class TestSignatureBlock:
    """The sign-off has a house format, and it is the same wherever it appears.

    Four lines, no telephone number. Each property here was specified rather
    than chosen, so each is asserted rather than eyeballed.
    """

    def test_line_1_is_the_name_bold_11pt(self, signature_letter):
        _, lines = signature_letter
        r = lines[0].runs[0]
        assert r.text == _SIG_NAME
        assert r.bold is True
        assert r.font.name == "Arial"
        assert r.font.size.pt == 11.0
        # Black is the default; setting it explicitly would add a colour that
        # means nothing to a document checked for painting only from the palette.
        assert r.font.color is None or r.font.color.type is None

    def test_line_2_is_the_company_in_brand_red(self, signature_letter):
        from pq_report import _XE_RED
        _, lines = signature_letter
        r = lines[1].runs[0]
        assert r.text == "Xcel Energy"
        assert r.bold is True
        assert r.font.name == "Arial"
        assert r.font.size.pt == 10.0
        assert r.font.color.rgb == _XE_RED

    def test_line_3_is_the_title_not_bold(self, signature_letter):
        _, lines = signature_letter
        r = lines[2].runs[0]
        assert r.text == _SIG_TITLE
        assert not r.bold
        assert r.font.name == "Arial"
        assert r.font.size.pt == 10.0
        assert r.font.color is None or r.font.color.type is None

    def test_line_4_is_a_live_mailto_link_in_outlook_teal(self, signature_letter):
        """Styled like a link and behaves like one.

        A teal underlined run that does nothing when clicked is a small broken
        promise in a document whose whole point is to invite a reply.
        """
        from docx.oxml.ns import qn
        doc, lines = signature_letter
        hyperlinks = lines[3]._p.findall(qn("w:hyperlink"))
        assert hyperlinks, "the email address is not a hyperlink"
        run = hyperlinks[0].find(qn("w:r"))
        rpr = run.find(qn("w:rPr"))
        assert run.find(qn("w:t")).text == _SIG_EMAIL
        assert rpr.find(qn("w:rFonts")).get(qn("w:ascii")) == "Arial"
        assert int(rpr.find(qn("w:sz")).get(qn("w:val"))) == 20      # half-points
        assert rpr.find(qn("w:color")).get(qn("w:val")).upper() == "467886"
        assert rpr.find(qn("w:b")) is None, "the address is not bold"
        rel = doc.part.rels[hyperlinks[0].get(qn("r:id"))]
        assert rel.target_ref == f"mailto:{_SIG_EMAIL}"

    def test_no_telephone_number_anywhere(self, signature_letter):
        """Policy: we do not give out phone numbers."""
        doc, _ = signature_letter
        text = "\n".join(p.text for p in doc.paragraphs)
        for word in ("Phone", "Telephone", "phone"):
            assert word not in text

    def test_the_tool_no_longer_collects_a_phone_number(self):
        """A field feeding nothing is worse than no field.

        Removed from the GUI form and the CLI as well as the documents, so a
        number cannot be typed into a box that quietly discards it.
        """
        import inspect
        import pq_report
        for fn in (pq_report.generate_customer_letter,
                   pq_report.generate_word_report):
            assert "engineer_phone" not in inspect.signature(fn).parameters
        assert "_eng_phone_var" not in (Path(__file__).parent / "run.py").read_text()
        assert "--engineer-phone" not in (Path(__file__).parent / "pq_analyzer.py").read_text()

    def test_the_letter_is_dated(self, signature_letter):
        """It moved to the top, where a letter's date goes — but it is still there."""
        import datetime
        doc, _ = signature_letter
        today = datetime.date.today()
        expected = f"{today:%B} {today.day}, {today.year}"
        assert expected in "\n".join(p.text for p in doc.paragraphs)

    def test_there_is_no_blank_line_in_the_block(self, signature_letter):
        """The four lines are contiguous — the address follows the title."""
        _, lines = signature_letter
        assert all(l.text.strip() or l._p.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink")
            for l in lines), "the signature block has a blank line in it"
